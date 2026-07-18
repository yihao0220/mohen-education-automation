from __future__ import annotations

from hashlib import sha256
import json
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml import OxmlElement
import pytest

from tools.trim_future_biology_answers import (
    discover_answer_documents,
    scan_answer_batch,
    trim_answer_batch,
    trim_answer_document,
)


REAL_BATCH = Path("/Users/xiaosheng/Documents/墨痕教育/未来-高二-生物")


def _sha256(path: Path) -> str:
    return sha256(path.read_bytes()).hexdigest()


def _add_field_marker(document: Document) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    instruction = OxmlElement("w:instrText")
    instruction.text = ' INCLUDEPICTURE "E:\\\\生物\\\\课时对点练新.TIF" \\* MERGEFORMAT '
    run._element.append(instruction)


def _write_field_fixture(path: Path) -> None:
    document = Document()
    document.add_paragraph("前部教材知识")
    _add_field_marker(document)
    document.add_paragraph("[分值：100分]")
    document.add_paragraph("1．题目")
    document.add_paragraph("答案　A")
    document.add_paragraph("解析　第一题解析。")
    document.save(path)


def _write_text_fixture(path: Path) -> None:
    document = Document()
    document.add_paragraph("前部教材知识")
    document.add_paragraph("课时对点练　[分值：100分]")
    document.add_paragraph("1．题目")
    document.add_paragraph("答案　B")
    document.add_paragraph("解析　第二题解析。")
    document.save(path)


def _paragraph_texts(path: Path) -> list[str]:
    return [paragraph.text for paragraph in Document(path).paragraphs if paragraph.text.strip()]


def test_trim_field_marker_removes_marker_and_front_content(tmp_path: Path) -> None:
    source = tmp_path / "field.docx"
    output = tmp_path / "out" / "field.docx"
    _write_field_fixture(source)
    source_hash = _sha256(source)

    result = trim_answer_document(source, output)

    assert result.marker_kind == "field_image"
    assert _sha256(source) == source_hash
    assert _paragraph_texts(output) == [
        "[分值：100分]",
        "1．题目",
        "答案　A",
        "解析　第一题解析。",
    ]
    assert result.answer_count == 1
    assert result.analysis_count == 1


def test_trim_text_marker_keeps_score_suffix_in_same_paragraph(tmp_path: Path) -> None:
    source = tmp_path / "text.docx"
    output = tmp_path / "out" / "text.docx"
    _write_text_fixture(source)

    result = trim_answer_document(source, output)

    assert result.marker_kind == "visible_text"
    assert _paragraph_texts(output) == [
        "[分值：100分]",
        "1．题目",
        "答案　B",
        "解析　第二题解析。",
    ]


def test_batch_preflights_all_docs_and_skips_other_document_families(tmp_path: Path) -> None:
    root = tmp_path / "未来-高二-生物"
    elective_one = root / "选必一答案"
    elective_two = root / "选必二答案"
    elective_one.mkdir(parents=True)
    elective_two.mkdir(parents=True)
    _write_field_fixture(elective_one / "第一课.docx")
    _write_text_fixture(elective_two / "第二课.docx")
    Document().save(elective_two / "章末检测试卷.docx")
    (elective_one / ".~锁文件.docx").write_bytes(b"lock")

    scans = scan_answer_batch(root)
    output = root / "答案" / "按课时截取"
    manifest_path = trim_answer_batch(root, output_dir=output)
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))

    assert len(scans) == 3
    assert manifest["source_document_count"] == 3
    assert manifest["matched_document_count"] == 2
    assert manifest["skipped_document_count"] == 1
    assert (output / "选必一答案" / "第一课.docx").is_file()
    assert (output / "选必二答案" / "第二课.docx").is_file()
    assert not (output / "选必二答案" / "章末检测试卷.docx").exists()


def test_ambiguous_or_nested_marker_stops_before_output(tmp_path: Path) -> None:
    root = tmp_path / "未来-高二-生物"
    answers = root / "选必一答案"
    answers.mkdir(parents=True)
    (root / "选必二答案").mkdir()

    multiple = Document()
    multiple.add_paragraph("课时对点练")
    multiple.add_paragraph("课时对点练")
    multiple.save(answers / "重复标记.docx")

    with pytest.raises(ValueError, match="标记数量异常"):
        trim_answer_batch(root, output_dir=root / "答案" / "按课时截取")
    assert not (root / "答案" / "按课时截取").exists()

    (answers / "重复标记.docx").unlink()
    nested = Document()
    table = nested.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "课时对点练"
    nested.save(answers / "表格标记.docx")

    with pytest.raises(ValueError, match="必须位于正文顶层段落"):
        trim_answer_batch(root, output_dir=root / "答案" / "按课时截取")
    assert not (root / "答案" / "按课时截取").exists()


def test_trim_preserves_all_package_members_except_document_xml(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "out" / "source.docx"
    _write_field_fixture(source)

    trim_answer_document(source, output)

    with ZipFile(source) as source_zip, ZipFile(output) as output_zip:
        assert source_zip.namelist() == output_zip.namelist()
        for member in source_zip.namelist():
            if member == "word/document.xml":
                continue
            assert output_zip.read(member) == source_zip.read(member), member


def test_real_batch_has_expected_two_marker_forms() -> None:
    if not REAL_BATCH.is_dir():
        pytest.skip("当前机器没有未来高二生物真实答案批次")

    sources = discover_answer_documents(REAL_BATCH)
    scans = scan_answer_batch(REAL_BATCH)
    matched = [scan for scan in scans if scan.matched]

    assert len(sources) == 73
    assert len(matched) == 42
    assert sum(scan.marker_kind == "field_image" for scan in matched) == 24
    assert sum(scan.marker_kind == "visible_text" for scan in matched) == 18
    assert sum(not scan.matched for scan in scans) == 31
