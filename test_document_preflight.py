from __future__ import annotations

import hashlib
import json
from pathlib import Path

import pytest
from docx import Document

from shared_core.document_preflight import (
    build_preflight_bundle,
    write_preflight_artifacts,
)


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _build_minimal_question_doc(path: Path) -> None:
    doc = Document()
    doc.add_heading("数学练习", level=1)
    doc.add_paragraph("1．观察下表，选择正确答案。")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "数量"
    table.cell(0, 1).text = "3"
    table.cell(1, 0).text = "总价"
    table.cell(1, 1).text = "12"
    doc.add_paragraph("A．2元")
    doc.add_paragraph("B．4元")
    doc.add_paragraph("2．直接写出结果。")
    doc.add_paragraph("3+4=____")
    doc.save(path)


def _build_role_question_doc(path: Path) -> None:
    doc = Document()
    doc.add_heading("阶段练习", level=0)
    doc.add_heading("第一部分 现代文阅读", level=1)
    doc.add_paragraph("阅读下面的材料，根据要求完成题目。")
    doc.add_paragraph("1．下列说法正确的一项是（　　）")
    doc.add_paragraph("A．第一项")
    internal_heading = doc.add_paragraph()
    internal_heading.add_run("（一）高考实战").bold = True
    doc.add_paragraph("2．请概括材料的主要内容。")
    doc.add_paragraph("材料内容继续说明相关背景和条件。")
    doc.add_paragraph("（3）（a+b）×（c+d）=a×c+a×d+b×c+b×d")
    doc.add_paragraph("3．完成下一道题。")
    doc.add_paragraph("补写出下列句子中的空缺部分。(20分)")
    doc.add_paragraph("4．写出最终答案。")
    doc.add_paragraph("二、拓展训练")
    doc.add_paragraph("5．完成拓展题。")
    doc.save(path)


def _stable_bundle_view(bundle: dict) -> dict:
    view = json.loads(json.dumps(bundle, ensure_ascii=False))
    view["profile"]["source"]["path"] = "<SOURCE>"
    view["profile"]["source"]["size_bytes"] = "<SIZE>"
    view["profile"]["source"]["sha256"] = "<SHA256>"
    view["plan"]["source_sha256"] = "<SHA256>"
    return view


def test_build_preflight_bundle_is_read_only_and_maps_native_table(tmp_path: Path) -> None:
    source = tmp_path / "数学练习.docx"
    _build_minimal_question_doc(source)
    before = _sha256(source)

    bundle = build_preflight_bundle(source, include_docling=False)

    assert _sha256(source) == before
    assert bundle["profile"]["source"]["sha256"] == before
    assert bundle["profile"]["source"]["hash_verified_unchanged"] is True
    assert bundle["profile"]["native"]["table_count"] == 1
    assert bundle["profile"]["native"]["heading_count"] == 1
    assert bundle["plan"]["execution_enabled"] is False
    assert [action["question_ids"] for action in bundle["plan"]["actions"]] == [["1"], ["2"]]
    assert bundle["plan"]["actions"][0]["source_ref"]["table_indexes"] == [1]


def test_json_is_authoritative_and_markdown_is_generated(tmp_path: Path) -> None:
    source = tmp_path / "数学练习.docx"
    output_dir = tmp_path / "artifacts"
    _build_minimal_question_doc(source)

    paths = write_preflight_artifacts(source, output_dir, include_docling=False)

    profile = json.loads(paths["profile_json"].read_text(encoding="utf-8"))
    plan = json.loads(paths["plan_json"].read_text(encoding="utf-8"))
    profile_md = paths["profile_md"].read_text(encoding="utf-8")
    plan_md = paths["plan_md"].read_text(encoding="utf-8")
    assert profile["source"]["name"] == "数学练习.docx"
    assert plan["source_sha256"] == profile["source"]["sha256"]
    assert "# 文档画像：数学练习.docx" in profile_md
    assert "# 动作计划：数学练习.docx" in plan_md
    assert "## 段落角色证据" in profile_md
    assert "自动排除题内标题：否" in profile_md
    assert "仅供预检，不允许执行 WPS" in plan_md


def test_minimal_bundle_golden_json(tmp_path: Path, file_regression) -> None:
    source = tmp_path / "数学练习.docx"
    _build_minimal_question_doc(source)
    bundle = build_preflight_bundle(source, include_docling=False)

    file_regression.check(
        json.dumps(_stable_bundle_view(bundle), ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
        extension=".json",
    )


def test_docling_sidecar_compares_structure_without_changing_source(tmp_path: Path) -> None:
    pytest.importorskip("docling")
    source = tmp_path / "数学练习.docx"
    _build_minimal_question_doc(source)
    before = _sha256(source)

    bundle = build_preflight_bundle(source, include_docling=True)

    assert _sha256(source) == before
    assert bundle["profile"]["docling"]["status"] == "success"
    assert bundle["profile"]["docling"]["table_count"] == 1
    assert bundle["profile"]["comparison"]["native_vs_docling_tables"] == [1, 1]


def test_role_evidence_classifies_internal_heading_without_changing_plan(tmp_path: Path) -> None:
    source = tmp_path / "语文阶段练习.docx"
    _build_role_question_doc(source)
    before = _sha256(source)

    bundle = build_preflight_bundle(source, include_docling=False)

    roles = {
        item["text"]: item
        for item in bundle["profile"]["roles"]["paragraphs"]
    }
    assert _sha256(source) == before
    assert roles["阶段练习"]["role"] == "document_title"
    assert roles["第一部分 现代文阅读"]["role"] == "section_heading"
    assert roles["1．下列说法正确的一项是（　　）"]["role"] == "question_start"
    assert roles["A．第一项"]["role"] == "option"
    assert roles["（一）高考实战"]["role"] == "internal_heading"
    assert "短段落" in roles["（一）高考实战"]["evidence"]
    assert "直接粗体占比>=0.60" in roles["（一）高考实战"]["evidence"]
    assert roles["材料内容继续说明相关背景和条件。"]["role"] == "body"
    assert roles["（3）（a+b）×（c+d）=a×c+a×d+b×c+b×d"]["role"] == "body"
    assert roles["补写出下列句子中的空缺部分。(20分)"]["role"] == "body"
    assert roles["二、拓展训练"]["role"] == "section_heading"
    assert bundle["profile"]["roles"]["automatic_exclusion_enabled"] is False
    assert bundle["plan"]["execution_enabled"] is False
    assert bundle["plan"]["mode"] == "preview_only"


def test_role_evidence_uses_docling_as_support_not_override(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = tmp_path / "语文专题练习.docx"
    doc = Document()
    doc.add_paragraph("1．阅读材料并回答问题。")
    doc.add_paragraph("专题提升")
    doc.add_paragraph("2．概括材料的中心意思。")
    doc.save(source)

    def fake_docling(_path: Path, _enabled: bool) -> dict:
        return {
            "enabled": True,
            "available": True,
            "package_version": "test",
            "status": "success",
            "document_schema": "DoclingDocument",
            "document_schema_version": "test",
            "text_count": 3,
            "table_count": 0,
            "picture_count": 0,
            "group_count": 0,
            "label_counts": {"section_header": 2, "text": 1},
            "_text_items": [
                {"order": 1, "label": "section_header", "text": "1．阅读材料并回答问题。"},
                {"order": 2, "label": "section_header", "text": "专题提升"},
                {"order": 3, "label": "text", "text": "2．概括材料的中心意思。"},
            ],
        }

    monkeypatch.setattr(
        "shared_core.document_preflight._inspect_with_docling",
        fake_docling,
    )

    bundle = build_preflight_bundle(source, include_docling=True)
    roles = {
        item["text"]: item
        for item in bundle["profile"]["roles"]["paragraphs"]
    }

    assert roles["1．阅读材料并回答问题。"]["role"] == "question_start"
    assert roles["1．阅读材料并回答问题。"]["docling_label"] == "section_header"
    assert roles["专题提升"]["role"] == "internal_heading"
    assert "Docling标签=section_header" in roles["专题提升"]["evidence"]
    assert bundle["profile"]["roles"]["docling_alignment"]["aligned_count"] == 3
