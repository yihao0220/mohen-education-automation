from __future__ import annotations

import importlib.util
import re
import sys
import unittest
from pathlib import Path

from docx import Document

from shared_core import (
    build_answer_units_from_paragraph_texts,
    build_question_units_from_docx,
    detect_subject_overlay,
    should_use_native_table_input_for_context,
)


PROJECT_ROOT = Path(__file__).resolve().parent
QUESTION_DIR = Path(
    r"D:\墨痕教育题目\未来-高二-历史（王逸豪）\高二历史 选择性必修二训练案"
)
ANSWER_DIR = Path(
    r"D:\墨痕教育题目\未来-高二-历史（王逸豪）\学案·24-25选必二学案解析"
)
USER_SAMPLE = Path(r"D:\墨痕教育题目\未来-高二-历史（王逸豪）\选修二样板.md")

EXPECTED_COUNTS = {
    1: 17,
    2: 17,
    3: 17,
    4: 18,
    5: 18,
    6: 18,
    7: 18,
    8: 18,
    9: 19,
    10: 19,
    11: 18,
    12: 19,
    13: 19,
    14: 19,
    15: 19,
}


def _load_module(module_name: str, relative_path: str):
    module_path = PROJECT_ROOT / relative_path
    spec = importlib.util.spec_from_file_location(module_name, module_path)
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    spec.loader.exec_module(module)
    return module


template = _load_module(
    "template_future_history_elective_two",
    "格式处理/格式模板库/template_future_history.py",
)


def _lesson_number(path: Path) -> int:
    match = re.search(r"第\s*(\d+)\s*课", path.name)
    if not match:
        raise AssertionError(f"文件名缺少课号: {path.name}")
    return int(match.group(1))


class FutureHistoryElectiveTwoTests(unittest.TestCase):
    def test_all_question_docs_use_history_overlay_and_native_tables(self):
        files = [path for path in QUESTION_DIR.glob("*.docx") if not path.name.startswith("~$")]
        if len(files) != 15:
            self.skipTest(f"题目样本数量不是 15: {len(files)}")

        for path in files:
            text = "\n".join(paragraph.text for paragraph in Document(path).paragraphs)[:5000]
            self.assertEqual(
                detect_subject_overlay(path.name, text, base_subject="文科"),
                "history",
                path.name,
            )
            self.assertTrue(
                should_use_native_table_input_for_context(path.name, "history"),
                path.name,
            )

            lesson = _lesson_number(path)
            units = build_question_units_from_docx(path)
            self.assertEqual(len(units), EXPECTED_COUNTS[lesson], path.name)
            self.assertEqual(
                [unit.question_id for unit in units],
                [str(index) for index in range(1, EXPECTED_COUNTS[lesson] + 1)],
                path.name,
            )

    def test_all_available_answer_docs_are_complete_and_sequential(self):
        files = [
            path
            for path in ANSWER_DIR.glob("*.docx")
            if not path.name.startswith("~$") and "_已清洗" not in path.stem
        ]
        if len(files) != 14:
            self.skipTest(f"答案样本数量不是 14: {len(files)}")

        for path in files:
            lesson = _lesson_number(path)
            entries = template.parse_paragraph_texts(
                [paragraph.text for paragraph in Document(path).paragraphs],
                renumber=True,
            )
            expected_count = EXPECTED_COUNTS[lesson]
            self.assertEqual(len(entries), expected_count, path.name)
            self.assertEqual(
                [entry["qnum"] for entry in entries],
                [str(index) for index in range(1, expected_count + 1)],
                path.name,
            )
            self.assertFalse(
                [entry["qnum"] for entry in entries if not entry["answer_lines"]],
                path.name,
            )
            self.assertFalse(
                [entry["qnum"] for entry in entries if not entry["analysis_lines"]],
                path.name,
            )

    def test_inline_answer_and_analysis_are_split(self):
        entries = template.parse_paragraph_texts(
            [
                "【基础巩固】",
                "1．第一题题干（    ）",
                "【答案】C【解析】第一题解析。",
                "【能力提升】",
                "1．第二题题干（    ）",
                "【答案】A【详解】第二题详解。",
            ],
            renumber=True,
        )

        self.assertEqual([entry["qnum"] for entry in entries], ["1", "2"])
        self.assertEqual(entries[0]["answer_lines"], ["C"])
        self.assertEqual(entries[0]["analysis_lines"], ["第一题解析。"])
        self.assertEqual(entries[1]["answer_lines"], ["A"])
        self.assertEqual(entries[1]["analysis_lines"], ["第二题详解。"])

    def test_circled_answer_points_stay_in_one_answer_box(self):
        rendered = template.render_standard_lines(
            [
                {
                    "qnum": "1",
                    "answer_lines": ["①特点甲。②特点乙。"],
                    "analysis_lines": ["解析内容。"],
                }
            ]
        )
        units = build_answer_units_from_paragraph_texts(rendered)

        self.assertEqual(len(units), 1)
        self.assertEqual(units[0].answer_mode, "whole")
        self.assertIn("要点一：", units[0].answer_items[0].text)
        self.assertIn("要点二：", units[0].answer_items[0].text)

    def test_user_sample_matches_lesson_one_answer(self):
        if not USER_SAMPLE.exists():
            self.skipTest("缺少用户选修二样板")
        lesson_one = next(
            path for path in ANSWER_DIR.glob("第1课*.docx") if not path.name.startswith("~$")
        )
        sample_entries = template.parse_paragraph_texts(
            USER_SAMPLE.read_text(encoding="utf-8").splitlines(),
            renumber=True,
        )
        source_entries = template.parse_paragraph_texts(
            [paragraph.text for paragraph in Document(lesson_one).paragraphs],
            renumber=True,
        )

        self.assertEqual([entry["qnum"] for entry in source_entries], [entry["qnum"] for entry in sample_entries])
        self.assertEqual(
            [entry["answer_lines"] for entry in source_entries],
            [entry["answer_lines"] for entry in sample_entries],
        )
        self.assertFalse([entry["qnum"] for entry in source_entries if not entry["analysis_lines"]])

    def test_layer_headers_always_stop_previous_question(self):
        fast_blade_dir = PROJECT_ROOT / "墨痕快刀"
        sys.path.insert(0, str(fast_blade_dir))
        try:
            core_parser = _load_module("future_history_core_parser", "墨痕快刀/core_parser.py")
        finally:
            sys.path.remove(str(fast_blade_dir))

        unit = type(
            "Unit",
            (),
            {
                "node_type": "VIP",
                "question_type": "subjective",
                "warnings": [],
                "media_blocks": [],
                "subquestions": [],
            },
        )()
        for header in ("【基础巩固】", "【能力提升】", "【拓展延伸】"):
            self.assertFalse(core_parser.should_ignore_inline_obstacle(unit, header, r"^【.*】$"))


if __name__ == "__main__":
    unittest.main()
