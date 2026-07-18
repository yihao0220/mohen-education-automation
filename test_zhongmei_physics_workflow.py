from __future__ import annotations

from hashlib import sha256
from io import BytesIO
from pathlib import Path
import re
import sys
import tempfile
from types import SimpleNamespace
import unittest
from zipfile import ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import RGBColor

from shared_core import (
    build_answer_units_from_docx,
    build_question_units_from_docx,
    map_answers,
)
from 墨痕快刀 import config

sys.modules.setdefault("pyautogui", SimpleNamespace(press=lambda *_args, **_kwargs: None))
sys.modules.setdefault("wps_helper", SimpleNamespace(get_active_wps=lambda: None))

from 答案录入 import answer_input
from 墨痕快刀 import core_parser
import main as controller_main
from tools.clean_zhongmei_physics_answers import (
    normalize_answer_docx,
    preflight_project,
)


TINY_PNG = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01"
    b"\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde"
    b"\x00\x00\x00\x0cIDATx\x9cc\xf8\xff\xff?\x00\x05"
    b"\xfe\x02\xfeA\xe2)\xb7\x00\x00\x00\x00IEND\xaeB`\x82"
)


def _digest(path: Path) -> str:
    return sha256(path.read_bytes()).hexdigest()


def _append_formula(paragraph, text: str = "x") -> None:
    math = OxmlElement("m:oMath")
    run = OxmlElement("m:r")
    math_text = OxmlElement("m:t")
    math_text.text = text
    run.append(math_text)
    math.append(run)
    paragraph._p.append(math)


def _count_xpath(path: Path, xpath: str, namespaces: dict[str, str]) -> int:
    from lxml import etree

    with ZipFile(path) as package:
        root = etree.fromstring(package.read("word/document.xml"))
    return len(root.xpath(xpath, namespaces=namespaces))


def _write_question_doc(path: Path, count: int = 4) -> None:
    doc = Document()
    doc.add_paragraph("[1~2题，每题4分]")
    for question_id in range(1, count + 1):
        doc.add_paragraph(f"{question_id}．第{question_id}题")
    doc.save(path)


def _write_answer_source(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("第1课时 测试")
    doc.add_paragraph("1．第一题题干")
    doc.add_paragraph("A.甲 B.乙 C.丙 D.丁")
    doc.add_paragraph("答案　C")
    first_analysis = doc.add_paragraph()
    hidden_marker = first_analysis.add_run("解析　")
    hidden_marker.font.hidden = True
    hidden_marker.font.color.rgb = RGBColor(255, 255, 255)
    first_analysis.add_run("第一题解析")
    doc.add_paragraph("①第一条续行")
    doc.add_paragraph("A：甲 B：乙 C：丙 D：丁")
    doc.add_paragraph("故本题选C。")

    doc.add_paragraph("2．第二题题干")
    rich_answer = doc.add_paragraph("答案　(1)2 m/s　(2)")
    _append_formula(rich_answer, "F")
    doc.add_paragraph("解析　第二题解析")
    analysis_picture = doc.add_paragraph("解析图如下")
    analysis_picture.add_run().add_picture(BytesIO(TINY_PNG))

    score_header = doc.add_paragraph("[1~2题，每题6分]")
    score_header.add_run().add_picture(BytesIO(TINY_PNG))
    doc.add_paragraph("1．能力练习重新编号")
    doc.add_paragraph("答案　B")

    doc.add_paragraph("2．内联标志题")
    doc.add_paragraph("答案　A解析　内联解析内容")
    doc.save(path)


class ZhongmeiPhysicsQuestionRuleTests(unittest.TestCase):
    def test_science_score_cues_are_obstacles_without_fixed_numbers(self):
        score_cues = [
            "[1~8题，每题4分]",
            "[9~11题，每题6分]",
            "[1、2题，每题4分]",
            "[7题6分]",
            "[6题6分]",
            "[6分]",
            "[分值：54分]",
        ]
        for text in score_cues:
            with self.subTest(text=text):
                self.assertTrue(
                    any(re.match(pattern, text) for pattern in config.CONFIG_SCIENCE["obstacles"])
                )

        self.assertFalse(
            any(
                re.match(pattern, "6.(10分)一道正常物理题")
                for pattern in config.CONFIG_SCIENCE["obstacles"]
            )
        )

    def test_score_cue_paragraph_with_picture_is_excluded_as_one_obstacle(self):
        previous_config = config.CURRENT_CONFIG
        try:
            config.CURRENT_CONFIG = config.CONFIG_SCIENCE
            obstacle, obstacle_type = core_parser.is_obstacle(
                SimpleNamespace(
                    # WPS 会把段首图片在文本层暴露为“/”，段尾保留回车。
                    Text="/　[9~11题，每题6分]\r",
                    InlineShapes=SimpleNamespace(Count=1),
                )
            )
        finally:
            config.CURRENT_CONFIG = previous_config
        self.assertTrue(obstacle)

        image_question = SimpleNamespace(
            node_type="STD",
            question_type="choice",
            warnings=["image_related_question"],
            media_blocks=["跑道图"],
            subquestions=[],
        )
        self.assertFalse(
            core_parser.should_ignore_inline_obstacle(
                image_question,
                "/　[9~11题，每题6分]",
                obstacle_type,
            )
        )

    def test_workspace_key_ignores_blank_paper_prefix(self):
        question = controller_main._normalize_workspace_key(
            "【空白试卷】第1课时 运动的描述 课时精练.docx"
        )
        answer = controller_main._normalize_workspace_key(
            "第1课时 运动的描述 课时精练_已清洗.docx"
        )
        self.assertEqual(question, answer)


class ZhongmeiPhysicsAnswerCleanerTests(unittest.TestCase):
    def test_extracts_only_rich_answer_blocks_and_preserves_source(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir) / "众美-高三-物理" / "答案"
            root.mkdir(parents=True)
            source = root / "样例.docx"
            output = root / "样例_已清洗.docx"
            _write_answer_source(source)
            source_hash = _digest(source)

            result = normalize_answer_docx(source, output, expected_question_count=4)

            self.assertEqual(result.answer_count, 4)
            self.assertEqual(result.inserted_analysis_placeholders, 1)
            self.assertEqual(result.rich_answers_moved_to_analysis, 1)
            self.assertEqual(_digest(source), source_hash)

            cleaned = Document(output)
            texts = [paragraph.text for paragraph in cleaned.paragraphs]
            joined = "\n".join(texts)
            self.assertNotIn("第一题题干", joined)
            self.assertNotIn("能力练习重新编号", joined)
            self.assertNotIn("每题6分", joined)
            self.assertIn("①第一条续行", joined)
            self.assertIn("A：甲 B：乙 C：丙 D：丁", joined)
            self.assertIn("故本题选C。", joined)
            self.assertIn("解析： ", texts)
            self.assertIn("4．A", texts)
            self.assertIn("解析：内联解析内容", texts)
            self.assertEqual(len(cleaned.inline_shapes), 1)
            self.assertEqual(
                _count_xpath(
                    output,
                    ".//m:oMath",
                    {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"},
                ),
                1,
            )

            with ZipFile(output) as package:
                from lxml import etree

                xml_root = etree.fromstring(package.read("word/document.xml"))
            namespaces = {
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            }
            analysis_run = xml_root.xpath(
                ".//w:r[w:t[starts-with(text(), '解析：')]][1]",
                namespaces=namespaces,
            )[0]
            self.assertFalse(
                analysis_run.xpath(
                    "./w:rPr/w:vanish | ./w:rPr/w:webHidden | ./w:rPr/w:specVanish",
                    namespaces=namespaces,
                )
            )
            self.assertEqual(
                analysis_run.xpath("string(./w:rPr/w:color/@w:val)", namespaces=namespaces),
                "000000",
            )
            self.assertEqual(
                analysis_run.xpath(
                    "string(./w:rPr/w:rFonts/@w:eastAsia)", namespaces=namespaces
                ),
                "宋体",
            )

            units = build_answer_units_from_docx(output, preserve_source_positions=True)
            self.assertEqual([unit.question_id for unit in units], ["1", "2", "3", "4"])
            self.assertEqual(units[0].answer_items[0].text, "C")
            self.assertEqual(units[2].answer_items[0].text, "B")
            self.assertEqual(units[3].answer_items[0].text, "A")
            self.assertTrue(units[1].metadata["analysis_only_subanswers"])

            question_path = root / "【空白试卷】样例.docx"
            _write_question_doc(question_path, count=4)
            mapped = map_answers(build_question_units_from_docx(question_path), units)
            self.assertNotIn("answer_split_but_question_whole", mapped[1].review_flags)

    def test_project_preflight_pairs_prefix_and_checks_counts_before_writes(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir) / "众美-高三-物理"
            question_dir = root / "第一章"
            answer_dir = root / "答案" / "第一章"
            question_dir.mkdir(parents=True)
            answer_dir.mkdir(parents=True)
            _write_question_doc(question_dir / "【空白试卷】样例.docx", count=4)
            _write_answer_source(answer_dir / "样例.docx")

            pairs = preflight_project(root)

            self.assertEqual(len(pairs), 1)
            self.assertEqual(pairs[0].question_count, 4)
            self.assertEqual(pairs[0].answer_count, 4)
            self.assertFalse((answer_dir / "样例_已清洗.docx").exists())

    def test_project_preflight_stops_on_question_answer_count_mismatch(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir) / "众美-高三-物理"
            question_dir = root / "第一章"
            answer_dir = root / "答案" / "第一章"
            question_dir.mkdir(parents=True)
            answer_dir.mkdir(parents=True)
            _write_question_doc(question_dir / "【空白试卷】样例.docx", count=3)
            _write_answer_source(answer_dir / "样例.docx")

            with self.assertRaisesRegex(ValueError, "题答数量不一致"):
                preflight_project(root)


class ZhongmeiPhysicsAnswerInputTests(unittest.TestCase):
    class _Doc:
        FullName = (
            "D:\\墨痕教育题目\\众美-高三-物理\\答案\\第一章\\"
            r"第1课时 运动的描述 课时精练_已清洗.docx"
        )
        Name = "第1课时 运动的描述 课时精练_已清洗.docx"

    def test_physics_subquestions_use_one_f2_answer_box(self):
        block = {
            "answer_mode": "subquestion",
            "force_whole_answer_input": False,
        }
        self.assertFalse(answer_input.should_split_subquestion_answers(self._Doc(), block))

    def test_plain_subanswers_do_not_fail_review_when_physics_forces_one_f2_box(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir) / "众美-高三-物理"
            answer_path = root / "答案" / "样例_已清洗.docx"
            question_path = root / "【空白试卷】样例.docx"
            answer_path.parent.mkdir(parents=True)
            answer_doc = Document()
            answer_doc.add_paragraph("1．")
            answer_doc.add_paragraph("答案：(1)甲 (2)乙")
            answer_doc.add_paragraph("解析：测试")
            answer_doc.save(answer_path)
            _write_question_doc(question_path, count=1)

            answer_units = build_answer_units_from_docx(answer_path)
            self.assertEqual(answer_units[0].answer_mode, "subquestion")
            self.assertTrue(answer_units[0].metadata["force_whole_answer_input"])
            mapped = map_answers(build_question_units_from_docx(question_path), answer_units)
            self.assertNotIn("answer_split_but_question_whole", mapped[0].review_flags)


if __name__ == "__main__":
    unittest.main()
