from __future__ import annotations

from hashlib import sha256
import importlib.util
import json
import os
from pathlib import Path
import re
import shutil
import sys
import types
from zipfile import ZipFile

from docx import Document
from PIL import Image
import pytest

from shared_core.answer_core import build_answer_units_from_docx
from shared_core.review_gate import derive_review_status_path, get_review_gate_result
from tools.clean_future_biology_answers import (
    clean_answer_batch,
    clean_answer_document,
    preflight_answer_document,
    preflight_clean_batch,
)


REAL_BATCH = Path(
    os.environ.get(
        "MOHEN_FUTURE_BIOLOGY_DIR",
        Path(__file__).resolve().parents[2] / "墨痕教育" / "未来-高二-生物",
    )
)
QUESTION_PATTERN = re.compile(r"^\s*(\d+)\s*[．.、]")
ANSWER_PATTERN = re.compile(r"^\s*答案：")
ANALYSIS_PATTERN = re.compile(r"^\s*解析：")
OPTION_PATTERN = re.compile(r"^\s*[A-DＡ-Ｄ]\s*[．.、]")
STRUCTURE_PATTERN = re.compile(
    r"^\s*(?:题组[一二三四五六七八九十\d]+|[一二三四五六七八九十]+[、．.]\s*(?:选择题|非选择题))"
)


def _sha256(path: Path) -> str:
    return sha256(path.read_bytes()).hexdigest()


def _write_png(path: Path, color: tuple[int, int, int]) -> None:
    Image.new("RGB", (32, 20), color).save(path)


def _write_fixture(path: Path, tmp_path: Path) -> None:
    analysis_image = tmp_path / "analysis.png"
    answer_image = tmp_path / "answer.png"
    _write_png(analysis_image, (220, 30, 30))
    _write_png(answer_image, (30, 30, 220))

    document = Document()
    document.add_paragraph("[分值：100分]")
    document.add_paragraph("题组一　基础选择")
    document.add_paragraph("1．第一题题干")
    document.add_paragraph("A．错误选项")
    document.add_paragraph("B．正确选项")
    document.add_paragraph("答案　B")
    document.add_paragraph("解析　第一题解析。")
    document.add_paragraph().add_run().add_picture(str(analysis_image))
    document.add_paragraph("题组二　主观题")
    document.add_paragraph("2．第二题题干")
    document.add_table(rows=1, cols=1).cell(0, 0).text = "题干表格"
    document.add_paragraph("(1)第一小问")
    document.add_paragraph("(2)第二小问")
    document.add_paragraph("答案　(1)甲")
    document.add_paragraph("(2)乙")
    document.add_paragraph().add_run().add_picture(str(answer_image))
    document.save(path)


def _document_media_reference_count(path: Path) -> int:
    with ZipFile(path) as package:
        document_xml = package.read("word/document.xml")
    return document_xml.count(b"r:embed=") + document_xml.count(b"r:id=")


def _load_answer_input_module(monkeypatch):
    monkeypatch.setitem(
        sys.modules,
        "pyautogui",
        types.SimpleNamespace(press=lambda _key: None),
    )
    monkeypatch.setitem(
        sys.modules,
        "wps_helper",
        types.SimpleNamespace(get_active_wps=lambda: None),
    )
    module_path = Path(__file__).resolve().parent / "答案录入" / "answer_input.py"
    spec = importlib.util.spec_from_file_location(
        "future_biology_answer_input_under_test",
        module_path,
    )
    assert spec is not None and spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    monkeypatch.setitem(sys.modules, spec.name, module)
    spec.loader.exec_module(module)
    return module


def test_inline_biology_subanswers_trigger_one_f4_per_subquestion(monkeypatch) -> None:
    answer_input = _load_answer_input_module(monkeypatch)
    paragraphs = [
        "15．",
        "答案：(1)内环境　血浆　组织液　淋巴液　"
        "(2)消化系统　泌尿系统　(3)消化　循环　(4)A　HCO、H2CO3",
        "解析： ",
    ]
    document_text = "".join(f"{text}\r" for text in paragraphs)
    paragraph_spans = []
    cursor = 0
    for text in paragraphs:
        end = cursor + len(text) + 1
        paragraph_spans.append((cursor, end))
        cursor = end

    selected = {"text": ""}
    events: list[tuple[str, str]] = []

    class FakeRange:
        def __init__(self, start: int, end: int):
            self.Start = start
            self.End = end

        @property
        def Text(self) -> str:
            return document_text[self.Start : self.End]

        def Select(self) -> None:
            selected["text"] = self.Text

    class FakeParagraph:
        def __init__(self, index: int):
            self.Range = FakeRange(*paragraph_spans[index - 1])

    class FakeParagraphs:
        Count = len(paragraphs)

        def __call__(self, index: int):
            return FakeParagraph(index)

    class FakeDocument:
        Paragraphs = FakeParagraphs()
        FullName = r"D:\墨痕教育题目\未来-高二-生物\答案\样例_已清洗.docx"
        Name = "样例_已清洗.docx"

        @staticmethod
        def Range(start: int, end: int):
            return FakeRange(start, end)

    class FakeWindow:
        @staticmethod
        def ScrollIntoView(_range) -> None:
            return None

    class FakeApplication:
        ActiveWindow = FakeWindow()

    class FakeWps:
        Application = FakeApplication()

    monkeypatch.setattr(
        sys.modules["pyautogui"],
        "press",
        lambda key: events.append((key, selected["text"])),
    )
    monkeypatch.setattr(answer_input.time, "sleep", lambda _seconds: None)

    answer_input.execute_input(
        FakeDocument(),
        FakeWps(),
        [
            {
                "qnum": "15",
                "ans_start_p": 2,
                "ana_start_p": 3,
                "end_p": 3,
                "answer_mode": "subquestion",
                "force_whole_answer_input": False,
            }
        ],
        0,
        1,
    )

    assert [key for key, _text in events] == ["f4", "f4", "f4", "f4", "f3"]
    assert [text.strip() for _key, text in events[:4]] == [
        "内环境　血浆　组织液　淋巴液",
        "消化系统　泌尿系统",
        "消化　循环",
        "A　HCO、H2CO3",
    ]


def test_clean_document_removes_questions_and_preserves_answer_media(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    _write_fixture(source, tmp_path)
    source_hash = _sha256(source)

    preflight = preflight_answer_document(source)
    refreshed, output_hash = clean_answer_document(source, output)

    assert _sha256(source) == source_hash
    assert output_hash == _sha256(output)
    assert refreshed == preflight
    assert preflight.question_count == 2
    assert preflight.missing_analysis_count == 1
    assert preflight.excluded_structure_count == 1
    assert preflight.excluded_media_count == 0
    assert preflight.retained_media_count == 2

    document = Document(output)
    texts = [paragraph.text.strip() for paragraph in document.paragraphs]
    nonempty = [text for text in texts if text]
    assert nonempty == [
        "1．",
        "答案：B",
        "解析：第一题解析。",
        "2．",
        "答案：(1)甲 (2)乙",
        "解析：",
    ]
    assert not document.tables
    assert _document_media_reference_count(output) == 2
    assert not any(OPTION_PATTERN.match(text) for text in texts)
    assert not any(STRUCTURE_PATTERN.match(text) for text in texts)

    units = build_answer_units_from_docx(output, preserve_source_positions=True)
    assert [unit.question_id for unit in units] == ["1", "2"]
    assert units[0].answer_mode == "whole"
    assert units[1].answer_mode == "subquestion"
    assert len(units[1].answer_items) == 2


def test_batch_preflights_before_writing_and_generates_review_gate(tmp_path: Path) -> None:
    project_root = tmp_path / "未来-高二-生物"
    input_root = project_root / "答案" / "按课时截取"
    elective_one = input_root / "选必一答案"
    elective_two = input_root / "选必二答案"
    elective_one.mkdir(parents=True)
    elective_two.mkdir(parents=True)
    _write_fixture(elective_one / "第一课.docx", tmp_path)

    preflights = preflight_clean_batch(
        project_root,
        expected_source_count=1,
        expected_question_count=2,
    )
    assert len(preflights) == 1

    output_root = tmp_path / "final"
    report_root = tmp_path / "reports"
    results, manifest_path = clean_answer_batch(
        project_root,
        output_dir=output_root,
        report_dir=report_root,
        expected_source_count=1,
        expected_question_count=2,
    )

    assert len(results) == 1
    result = results[0]
    assert result.blocking_issue_count == 0
    assert result.output_path.name == "第一课_已清洗.docx"
    assert result.output_path.is_file()
    assert result.review_report_path.is_file()
    assert get_review_gate_result(result.output_path)["allowed"] is True
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    assert manifest["document_count"] == 1
    assert manifest["question_count"] == 2
    assert manifest["inserted_analysis_placeholders"] == 1
    assert manifest["approved_document_count"] == 1

    windows_root = tmp_path / "windows" / "答案" / "已清洗"
    windows_root.mkdir(parents=True)
    relocated = windows_root / result.output_path.name
    source_status = Path(derive_review_status_path(result.output_path))
    relocated_status = windows_root / source_status.name
    shutil.copyfile(result.output_path, relocated)
    shutil.copyfile(source_status, relocated_status)
    os.utime(relocated, ns=(1_800_000_000_000_000_000,) * 2)
    assert get_review_gate_result(relocated)["allowed"] is True


def test_real_batch_cleans_all_answers_without_source_changes(tmp_path: Path) -> None:
    if not REAL_BATCH.is_dir():
        pytest.skip("当前机器没有未来高二生物真实答案批次")

    preflights = preflight_clean_batch(
        REAL_BATCH,
        expected_source_count=42,
        expected_question_count=580,
    )
    source_hashes = {
        item.source_path: item.source_sha256 for item in preflights
    }
    assert sum(item.missing_analysis_count for item in preflights) == 46
    assert sum(item.excluded_structure_count for item in preflights) == 50
    assert sum(item.excluded_media_count for item in preflights) == 41
    assert sum(item.retained_media_count for item in preflights) == 5

    output_root = tmp_path / "cleaned"
    report_root = tmp_path / "reports"
    results, manifest_path = clean_answer_batch(
        REAL_BATCH,
        output_dir=output_root,
        report_dir=report_root,
        expected_source_count=42,
        expected_question_count=580,
    )

    assert len(results) == 42
    assert sum(result.question_count for result in results) == 580
    assert sum(result.inserted_analysis_placeholders for result in results) == 46
    assert sum(result.retained_media_count for result in results) == 5
    assert all(result.blocking_issue_count == 0 for result in results)
    assert all(_sha256(path) == digest for path, digest in source_hashes.items())

    output_docs = sorted(output_root.rglob("*_已清洗.docx"))
    assert len(output_docs) == 42
    total_questions = 0
    total_answers = 0
    total_analyses = 0
    total_media = 0
    for path in output_docs:
        document = Document(path)
        texts = [paragraph.text.strip() for paragraph in document.paragraphs]
        total_questions += sum(bool(QUESTION_PATTERN.match(text)) for text in texts)
        total_answers += sum(bool(ANSWER_PATTERN.match(text)) for text in texts)
        total_analyses += sum(bool(ANALYSIS_PATTERN.match(text)) for text in texts)
        total_media += _document_media_reference_count(path)
        assert not document.tables
        assert not any(OPTION_PATTERN.match(text) for text in texts)
        assert not any(STRUCTURE_PATTERN.match(text) for text in texts)
        assert get_review_gate_result(path)["allowed"] is True

    assert (total_questions, total_answers, total_analyses) == (580, 580, 580)
    assert total_media == 5
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    assert manifest["approved_document_count"] == 42
    assert manifest["retained_media_count"] == 5
