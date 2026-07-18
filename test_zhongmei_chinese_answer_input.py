# coding: utf-8
from __future__ import annotations

import importlib.util
import io
import sys
import tempfile
import unittest
from contextlib import redirect_stdout
from pathlib import Path

from docx import Document

from shared_core import (
    build_answer_units_from_docx,
    build_question_units_from_docx,
    build_review_report,
    map_answers,
)
from shared_core.answer_core import build_answer_units_from_paragraph_texts
from tools.clean_zhongmei_chinese_answers import _is_heading as is_clean_answer_heading


PROJECT_ROOT = Path(__file__).resolve().parent
ANSWER_DIR = Path(
    r"D:\墨痕教育题目\众美-高三-语文\答案\对点练案答案"
)
IGNORED_ANSWER_HEADINGS = (
    "一、专项训练",
    "(一)请比对下列选项与原文，说明选项错在何处。",
    "(二)图文解读题",
    "二、综合训练",
    "阅读下面的文字，完成文后题目。",
    "(一)分析评价信息",
    "(二)逻辑推断",
    "(三)分析论证特点",
    "(二)梳理论述思路",
    "(一)分析理据关系",
    "(一)阅读下面的文字，完成文后题目。",
    "(二)阅读下面的文字，完成文后题目。",
    "(三)阅读下面的文字，完成文后题目。",
)


def _load_answer_input_module():
    answer_dir = PROJECT_ROOT / "答案录入"
    module_path = answer_dir / "answer_input.py"
    spec = importlib.util.spec_from_file_location("zhongmei_answer_input", module_path)
    if not spec or not spec.loader:
        raise AssertionError(f"无法加载答案录入模块: {module_path}")
    module = importlib.util.module_from_spec(spec)
    sys.path.insert(0, str(answer_dir))
    try:
        spec.loader.exec_module(module)
    finally:
        sys.path.remove(str(answer_dir))
    return module


class ZhongmeiChineseAnswerInputTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.answer_input = _load_answer_input_module()

    def test_review_gate_is_read_from_current_document_each_time(self):
        class FakeDocument:
            FullName = r"D:\答案\对点练案1_已清洗.docx"

        calls = []
        original = self.answer_input.get_review_gate_result
        self.answer_input.get_review_gate_result = lambda path: calls.append(path) or {
            "allowed": True,
            "status": "approved",
            "reason": "自动检查已通过",
            "status_path": "状态.json",
        }
        try:
            result = self.answer_input.read_current_review_gate(FakeDocument())
        finally:
            self.answer_input.get_review_gate_result = original

        self.assertTrue(result["allowed"])
        self.assertEqual(calls, [FakeDocument.FullName])

    def test_zhongmei_classical_question_path_maps_to_cleaned_answer(self):
        question_path = Path(
            r"D:\墨痕教育题目\众美-高三-语文\文言文\选择性必修下册\27《种树郭橐驼传》.docx"
        )
        expected = Path(
            r"D:\墨痕教育题目\众美-高三-语文\答案\文言文答案\选择性必修下册\27《种树郭橐驼传》_已清洗.docx"
        )

        self.assertEqual(
            self.answer_input.derive_zhongmei_classical_cleaned_answer_path(
                question_path
            ),
            expected,
        )
        self.assertIsNone(
            self.answer_input.derive_zhongmei_classical_cleaned_answer_path(expected)
        )
        self.assertIsNone(
            self.answer_input.derive_zhongmei_classical_cleaned_answer_path(
                Path(r"D:\墨痕教育题目\其他项目\文言文\27《种树郭橐驼传》.docx")
            )
        )

    def test_zhongmei_classical_question_auto_opens_approved_cleaned_answer(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir) / "众美-高三-语文"
            question_path = (
                project_root / "文言文" / "选择性必修下册" / "27《种树郭橐驼传》.docx"
            )
            answer_path = (
                project_root
                / "答案"
                / "文言文答案"
                / "选择性必修下册"
                / "27《种树郭橐驼传》_已清洗.docx"
            )
            question_path.parent.mkdir(parents=True)
            answer_path.parent.mkdir(parents=True)
            Document().save(question_path)
            Document().save(answer_path)

            class FakeDocument:
                def __init__(self, path):
                    self.FullName = str(path)
                    self.Name = Path(path).name
                    self.activated = False

                def Activate(self):
                    self.activated = True

            source_doc = FakeDocument(question_path)
            target_doc = FakeDocument(answer_path)

            class FakeDocuments:
                Count = 1

                @staticmethod
                def Item(_index):
                    return source_doc

                @staticmethod
                def Open(path):
                    self.assertEqual(Path(path), answer_path)
                    return target_doc

            class FakeWps:
                Documents = FakeDocuments()

            original_gate = self.answer_input.get_review_gate_result
            self.answer_input.get_review_gate_result = lambda _path: {
                "allowed": True,
                "status": "approved",
                "reason": "自动检查已通过",
                "status_path": str(answer_path.with_name(f"{answer_path.stem}_审核状态.json")),
            }
            try:
                resolved_doc, message = self.answer_input.resolve_active_answer_document(
                    FakeWps(),
                    source_doc,
                )
            finally:
                self.answer_input.get_review_gate_result = original_gate

            self.assertIs(resolved_doc, target_doc)
            self.assertTrue(target_doc.activated)
            self.assertIn("已自动切换", message)
            self.assertIn(target_doc.Name, message)

    def test_zhongmei_classical_auto_switch_keeps_gate_blocking(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir) / "众美-高三-语文"
            question_path = project_root / "文言文" / "27《种树郭橐驼传》.docx"
            answer_path = (
                project_root
                / "答案"
                / "文言文答案"
                / "27《种树郭橐驼传》_已清洗.docx"
            )
            question_path.parent.mkdir(parents=True)
            answer_path.parent.mkdir(parents=True)
            Document().save(question_path)
            Document().save(answer_path)

            class FakeDocument:
                FullName = str(question_path)
                Name = question_path.name

            class FakeDocuments:
                Count = 1

                @staticmethod
                def Item(_index):
                    return FakeDocument()

                @staticmethod
                def Open(_path):
                    raise AssertionError("门禁未通过时不得打开答案文档")

            class FakeWps:
                Documents = FakeDocuments()

            original_gate = self.answer_input.get_review_gate_result
            self.answer_input.get_review_gate_result = lambda _path: {
                "allowed": False,
                "status": "stale",
                "reason": "答案文件已变更",
                "status_path": str(answer_path.with_name(f"{answer_path.stem}_审核状态.json")),
            }
            try:
                with self.assertRaisesRegex(ValueError, "答案文件已变更"):
                    self.answer_input.resolve_active_answer_document(
                        FakeWps(),
                        FakeDocument(),
                    )
            finally:
                self.answer_input.get_review_gate_result = original_gate

    def test_section_headings_do_not_enter_previous_answer_or_analysis(self):
        for heading in IGNORED_ANSWER_HEADINGS:
            with self.subTest(heading=heading):
                units = build_answer_units_from_paragraph_texts(
                    [
                        "1．",
                        "答案：A",
                        "解析：第一题解析。",
                        heading,
                        "2．",
                        "答案：B",
                        "解析：第二题解析。",
                    ]
                )

                self.assertEqual(len(units), 2)
                first_answer = "\n".join(item.text for item in units[0].answer_items)
                first_analysis = "\n".join(item.text for item in units[0].analysis_items)
                self.assertNotIn(heading, first_answer)
                self.assertNotIn(heading, first_analysis)
                self.assertEqual(first_analysis, "第一题解析。")

    def test_real_lesson_32_preserves_source_paragraph_positions(self):
        path = ANSWER_DIR / "对点练案32_已清洗.docx"
        if not path.exists():
            self.skipTest(f"缺少真实样本: {path}")

        paragraphs = [paragraph.text.strip() for paragraph in Document(path).paragraphs]
        units = build_answer_units_from_docx(
            path,
            preserve_source_positions=True,
        )

        self.assertEqual(len(units), 14)
        self.assertEqual(
            (
                units[0].metadata["ans_start_p"],
                units[0].metadata["ana_start_p"],
            ),
            (3, 4),
        )
        self.assertEqual(
            (
                units[6].metadata["ans_start_p"],
                units[6].metadata["ana_start_p"],
            ),
            (22, 23),
        )
        for unit in units:
            answer_text = paragraphs[unit.metadata["ans_start_p"] - 1]
            analysis_text = paragraphs[unit.metadata["ana_start_p"] - 1]
            with self.subTest(question_id=unit.question_id):
                self.assertRegex(answer_text, r"^答案[：:]")
                self.assertRegex(analysis_text, r"^解析[：:]")

    def test_real_lesson_32_execute_input_selects_real_answer_and_analysis(self):
        path = ANSWER_DIR / "对点练案32_已清洗.docx"
        if not path.exists():
            self.skipTest(f"缺少真实样本: {path}")

        paragraphs = [paragraph.text for paragraph in Document(path).paragraphs]
        units = self.answer_input.build_input_units_from_docx(path)
        blocks = self.answer_input.build_blocks_from_units(units)
        document_text = "".join(f"{text}\r" for text in paragraphs)
        paragraph_spans = []
        cursor = 0
        for text in paragraphs:
            end = cursor + len(text) + 1
            paragraph_spans.append((cursor, end))
            cursor = end

        selected = {"text": ""}
        events = []

        class FakeRange:
            def __init__(self, start, end):
                self.Start = start
                self.End = end

            @property
            def Text(self):
                return document_text[self.Start:self.End]

            def Select(self):
                selected["text"] = self.Text

        class FakeParagraph:
            def __init__(self, index):
                start, end = paragraph_spans[index - 1]
                self.Range = FakeRange(start, end)

        class FakeParagraphs:
            Count = len(paragraphs)

            def __call__(self, index):
                return FakeParagraph(index)

        class FakeDocument:
            Paragraphs = FakeParagraphs()
            FullName = str(path)
            Name = path.name

            @staticmethod
            def Range(start, end):
                return FakeRange(start, end)

        class FakeWindow:
            @staticmethod
            def ScrollIntoView(_range):
                return None

        class FakeApplication:
            ActiveWindow = FakeWindow()

        class FakeWps:
            Application = FakeApplication()

        original_press = self.answer_input.pyautogui.press
        original_sleep = self.answer_input.time.sleep
        self.answer_input.pyautogui.press = lambda key: events.append(
            (key, selected["text"])
        )
        self.answer_input.time.sleep = lambda _seconds: None
        try:
            with redirect_stdout(io.StringIO()):
                self.answer_input.execute_input(
                    FakeDocument(),
                    FakeWps(),
                    blocks,
                    0,
                    len(blocks),
                )
        finally:
            self.answer_input.pyautogui.press = original_press
            self.answer_input.time.sleep = original_sleep

        self.assertEqual(events[0][0], "f2")
        self.assertEqual(events[0][1].strip(), "C")
        self.assertEqual(events[1][0], "f3")
        self.assertIn("一愁莫展", events[1][1])
        self.assertEqual(sum(key == "f3" for key, _ in events), 14)
        for key, text in events:
            with self.subTest(key=key, text=text[:20]):
                if key in {"f2", "f4"}:
                    self.assertTrue(text.strip())
                    self.assertNotIn("答案：", text)
                if key == "f3":
                    self.assertNotIn("答案：", text)

    def test_real_lesson_37_preserves_source_paragraph_positions(self):
        path = ANSWER_DIR / "对点练案37_已清洗.docx"
        if not path.exists():
            self.skipTest(f"缺少真实样本: {path}")

        paragraphs = [paragraph.text.strip() for paragraph in Document(path).paragraphs]
        units = build_answer_units_from_docx(
            path,
            preserve_source_positions=True,
        )

        self.assertEqual(len(units), 13)
        for unit in units:
            answer_text = paragraphs[unit.metadata["ans_start_p"] - 1]
            analysis_text = paragraphs[unit.metadata["ana_start_p"] - 1]
            with self.subTest(question_id=unit.question_id):
                self.assertRegex(answer_text, r"^答案[：:]")
                self.assertRegex(analysis_text, r"^解析[：:]")

    def test_numbered_question_preserves_first_subquestion_and_strips_answer_labels(self):
        units = build_answer_units_from_paragraph_texts(
            [
                "1．",
                "(1)答案：正确。",
                "(2)答案：错误。根据原文可知。",
                "(3)答案：引用诗句的目的判断错误。",
                "(4)答案：正确。",
                "解析：这是整道大题的解析。",
            ]
        )

        self.assertEqual(len(units), 1)
        self.assertEqual(units[0].answer_mode, "subquestion")
        self.assertEqual(
            [item.item_id for item in units[0].answer_items],
            ["(1)", "(2)", "(3)", "(4)"],
        )
        self.assertEqual(
            [item.text for item in units[0].answer_items],
            [
                "正确。",
                "错误。根据原文可知。",
                "引用诗句的目的判断错误。",
                "正确。",
            ],
        )

    def test_circled_answer_points_stay_in_one_answer_box(self):
        answer_text = (
            "①语言口语化，朴实而生动。"
            "②连用动词，生动而富有情韵。"
            "③句式灵活，轻快活泼。"
        )
        units = build_answer_units_from_paragraph_texts(
            [
                "9．",
                f"答案：{answer_text}",
                "解析：这是整道题的解析。",
            ]
        )

        self.assertEqual(len(units), 1)
        self.assertEqual(units[0].answer_mode, "whole")
        self.assertEqual(len(units[0].answer_items), 1)
        self.assertEqual(units[0].answer_items[0].text, answer_text)
        self.assertEqual(self.answer_input.find_subquestion_matches(answer_text), [])

    def test_interleaved_subquestion_answers_and_analyses_are_separated(self):
        units = build_answer_units_from_paragraph_texts(
            [
                "2．",
                "(1)答案：第一问答案。",
                "(1)解析：第一问解析。",
                "(2)答案：第二问答案第一行。",
                "第二问答案续行。",
                "(2)解析：第二问解析。",
            ]
        )

        self.assertEqual(len(units), 1)
        self.assertEqual(units[0].answer_mode, "subquestion")
        self.assertEqual(
            [item.text for item in units[0].answer_items],
            ["第一问答案。", "第二问答案第一行。\n第二问答案续行。"],
        )
        self.assertEqual(
            [item.text for item in units[0].analysis_items],
            ["第一问解析。", "第二问解析。"],
        )
        self.assertTrue(units[0].metadata["interleaved_subquestion_analysis"])

    def test_real_lesson_2_groups_complete_reading_as_five_subanswers(self):
        path = ANSWER_DIR / "对点练案2_已清洗.docx"
        if not path.exists():
            self.skipTest(f"缺少真实样本: {path}")

        units = self.answer_input.build_input_units_from_docx(path)
        units_by_id = {unit.question_id: unit for unit in units}

        self.assertEqual([unit.question_id for unit in units], ["1", "2", "3", "4", "5"])
        self.assertEqual(len(units_by_id["1"].answer_items), 4)
        self.assertEqual(len(units_by_id["5"].answer_items), 5)
        self.assertEqual(len(units_by_id["5"].analysis_items), 5)
        self.assertEqual(units_by_id["5"].answer_mode, "subquestion")

    def test_materialized_reading_group_has_contiguous_f4_answers_and_one_f3_analysis(self):
        path = ANSWER_DIR / "对点练案2_已清洗.docx"
        if not path.exists():
            self.skipTest(f"缺少真实样本: {path}")

        units = self.answer_input.build_input_units_from_docx(path)
        with tempfile.TemporaryDirectory() as temp_dir:
            input_path, prepared_units = self.answer_input.prepare_material_group_input_document(
                path,
                units,
                temp_dir=temp_dir,
            )
            self.assertEqual(Path(input_path), path)
            prepared_doc = Document(input_path)
            reading = next(unit for unit in prepared_units if unit.question_id == "5")
            metadata = reading.metadata
            paragraphs = [paragraph.text for paragraph in prepared_doc.paragraphs]

            answer_paragraphs = paragraphs[
                metadata["ans_start_p"] - 1 : metadata["ana_start_p"] - 1
            ]
            analysis_paragraphs = paragraphs[
                metadata["ana_start_p"] - 1 : metadata["end_p"]
            ]

            self.assertEqual(answer_paragraphs[0], "5．")
            self.assertEqual(len(answer_paragraphs[1:]), 5)
            self.assertEqual(
                [text[:3] for text in answer_paragraphs[1:]],
                ["(1)", "(2)", "(3)", "(4)", "(5)"],
            )
            self.assertTrue(all("答案：" in text for text in answer_paragraphs[1:]))
            self.assertEqual(analysis_paragraphs[0], "解析：")
            self.assertEqual(len(analysis_paragraphs[1:]), 5)
            self.assertEqual(
                [text[:3] for text in analysis_paragraphs[1:]],
                ["(1)", "(2)", "(3)", "(4)", "(5)"],
            )

            answer_range_text = "\r".join(answer_paragraphs) + "\r"
            _, text_body = self.answer_input.strip_input_question_prefix(answer_range_text)
            matches = self.answer_input.find_subquestion_matches(text_body)
            self.assertEqual([match.group(1) for match in matches], ["(1)", "(2)", "(3)", "(4)", "(5)"])

    def test_real_interleaved_subquestion_document_is_materialized_before_input(self):
        path = ANSWER_DIR / "对点练案41_已清洗.docx"
        if not path.exists():
            self.skipTest(f"缺少真实样本: {path}")

        units = self.answer_input.build_input_units_from_docx(path)
        source_unit = next(unit for unit in units if unit.question_id == "2")
        self.assertEqual(len(source_unit.answer_items), 2)
        self.assertEqual(len(source_unit.analysis_items), 2)
        self.assertTrue(source_unit.metadata["interleaved_subquestion_analysis"])

        with tempfile.TemporaryDirectory() as temp_dir:
            input_path, prepared_units = self.answer_input.prepare_material_group_input_document(
                path,
                units,
                temp_dir=temp_dir,
            )
            self.assertNotEqual(Path(input_path), path)
            prepared_doc = Document(input_path)
            prepared = next(unit for unit in prepared_units if unit.question_id == "2")
            paragraphs = [paragraph.text for paragraph in prepared_doc.paragraphs]
            answer_paragraphs = paragraphs[
                prepared.metadata["ans_start_p"] - 1 : prepared.metadata["ana_start_p"] - 1
            ]
            analysis_paragraphs = paragraphs[
                prepared.metadata["ana_start_p"] - 1 : prepared.metadata["end_p"]
            ]

            self.assertEqual([text[:3] for text in answer_paragraphs[1:]], ["(1)", "(2)"])
            self.assertEqual([text[:3] for text in analysis_paragraphs[1:]], ["(1)", "(2)"])

    def test_real_lesson_41_question_9_uses_one_f2_answer(self):
        path = ANSWER_DIR / "对点练案41_已清洗.docx"
        if not path.exists():
            self.skipTest(f"缺少真实样本: {path}")

        units = self.answer_input.build_input_units_from_docx(path)
        target = next(unit for unit in units if unit.question_id == "9")
        self.assertEqual(target.answer_mode, "whole")
        self.assertEqual(len(target.answer_items), 1)
        answer_text = target.answer_items[0].text
        self.assertTrue(all(marker in answer_text for marker in "①②③"))
        self.assertEqual(self.answer_input.find_subquestion_matches(answer_text), [])

    def test_all_44_cleaned_answers_keep_circled_points_whole(self):
        circled = set("①②③④⑤⑥⑦⑧⑨⑩")
        incorrectly_split = []
        for lesson in range(1, 45):
            path = ANSWER_DIR / f"对点练案{lesson}_已清洗.docx"
            if not path.exists():
                self.skipTest(f"缺少第 {lesson} 份已清洗答案")
            for unit in build_answer_units_from_docx(path):
                item_ids = [item.item_id for item in unit.answer_items]
                if len(item_ids) > 1 and all(item_id in circled for item_id in item_ids):
                    incorrectly_split.append((lesson, unit.question_id, item_ids))

        self.assertEqual(incorrectly_split, [])

    def test_subquestion_scan_ignores_parenthesized_numbers_inside_answer_text(self):
        text = (
            "\r(1)第一小问答案中引用第(9)段。"
            "\r(2)第二小问答案。"
            "\r(3)第三小问答案。"
        )
        matches = self.answer_input.find_subquestion_matches(text)
        self.assertEqual([match.group(1) for match in matches], ["(1)", "(2)", "(3)"])

    def test_all_44_cleaned_answers_have_complete_reading_group_mapping(self):
        files = sorted(
            (
                path
                for path in ANSWER_DIR.glob("*_已清洗.docx")
                if not path.name.startswith("~$")
            ),
            key=lambda path: int("".join(char for char in path.stem.split("_")[0] if char.isdigit())),
        )
        if len(files) != 44:
            self.skipTest(f"真实已清洗答案样本数量不是 44: {len(files)}")

        material_group_count = 0
        for lesson, path in enumerate(files, 1):
            question_path = ANSWER_DIR.parent.parent / "对点练案" / f"对点练案{lesson}.docx"
            questions = build_question_units_from_docx(question_path)
            units = map_answers(questions, build_answer_units_from_docx(path))
            for unit in units:
                if unit.metadata.get("mapping_method") != "material_group_premerged":
                    continue
                material_group_count += 1
                grouped_ids = unit.metadata["grouped_question_ids"]
                with self.subTest(path=path.name, question_id=unit.question_id):
                    self.assertEqual(len(unit.answer_items), len(grouped_ids))
                    self.assertEqual(len(unit.analysis_items), len(grouped_ids))
                    self.assertNotIn("orphan_answer", unit.review_flags)
                    self.assertNotIn("missing_grouped_answer", unit.review_flags)

        self.assertEqual(material_group_count, 40)

    def test_real_lesson_2_review_accepts_answer_defined_subquestions_and_five_question_group(self):
        answer_path = ANSWER_DIR / "对点练案2_已清洗.docx"
        question_path = ANSWER_DIR.parent.parent / "对点练案" / "对点练案2.docx"
        if not answer_path.exists() or not question_path.exists():
            self.skipTest("缺少对点练案2题目或已清洗答案")

        questions = build_question_units_from_docx(question_path)
        answers = map_answers(questions, build_answer_units_from_docx(answer_path))
        report = build_review_report(answer_path.name, questions, answers)

        self.assertEqual(report.summary["high_risk_count"], 0)
        self.assertFalse([issue for issue in report.issues if issue.severity == "error"])

    def test_all_44_cleaned_answers_pass_question_answer_review(self):
        blocked = []
        for lesson in range(1, 45):
            answer_path = ANSWER_DIR / f"对点练案{lesson}_已清洗.docx"
            question_path = ANSWER_DIR.parent.parent / "对点练案" / f"对点练案{lesson}.docx"
            if not answer_path.exists() or not question_path.exists():
                self.skipTest(f"缺少第 {lesson} 份题目或已清洗答案")
            questions = build_question_units_from_docx(question_path)
            answers = map_answers(questions, build_answer_units_from_docx(answer_path))
            report = build_review_report(answer_path.name, questions, answers)
            errors = [issue.title for issue in report.issues if issue.severity == "error"]
            if errors:
                blocked.append((lesson, errors))

        self.assertEqual(blocked, [])

    def test_all_44_cleaned_answers_exclude_section_headings_from_payloads(self):
        contaminated = []
        for lesson in range(1, 45):
            answer_path = ANSWER_DIR / f"对点练案{lesson}_已清洗.docx"
            if not answer_path.exists():
                self.skipTest(f"缺少第 {lesson} 份已清洗答案")
            headings = [
                paragraph.text.strip()
                for paragraph in Document(answer_path).paragraphs
                if paragraph.text.strip() and is_clean_answer_heading(paragraph.text.strip())
            ]
            for unit in build_answer_units_from_docx(answer_path):
                payload = "\n".join(
                    item.text
                    for item in unit.answer_items + unit.analysis_items
                    if item.text
                )
                hits = [heading for heading in headings if heading in payload]
                if hits:
                    contaminated.append((lesson, unit.question_id, hits))

        self.assertEqual(contaminated, [])

    def test_real_lesson_2_execution_sequences_match_both_new_rules(self):
        path = ANSWER_DIR / "对点练案2_已清洗.docx"
        if not path.exists():
            self.skipTest(f"缺少真实样本: {path}")

        units = self.answer_input.build_input_units_from_docx(path)
        with tempfile.TemporaryDirectory() as temp_dir:
            input_path, prepared_units = self.answer_input.prepare_material_group_input_document(
                path,
                units,
                temp_dir=temp_dir,
            )
            paragraphs = [paragraph.text for paragraph in Document(input_path).paragraphs]
            explicit = next(unit for unit in prepared_units if unit.question_id == "1")
            reading = next(unit for unit in prepared_units if unit.question_id == "5")
            blocks = self.answer_input.build_blocks_from_units([explicit, reading])

            selected_texts = []
            document_text = "".join(f"{text}\r" for text in paragraphs)
            paragraph_spans = []
            cursor = 0
            for text in paragraphs:
                end = cursor + len(text) + 1
                paragraph_spans.append((cursor, end))
                cursor = end

            class FakeRange:
                def __init__(self, start, end):
                    self.Start = start
                    self.End = end

                @property
                def Text(self):
                    return document_text[self.Start:self.End]

                def Select(self):
                    selected_texts.append(self.Text)

            class FakeParagraph:
                def __init__(self, index):
                    start, end = paragraph_spans[index - 1]
                    self.Range = FakeRange(start, end)

            class FakeParagraphs:
                Count = len(paragraphs)

                def __call__(self, index):
                    return FakeParagraph(index)

            class FakeDocument:
                Paragraphs = FakeParagraphs()

                @staticmethod
                def Range(start, end):
                    return FakeRange(start, end)

            class FakeWindow:
                @staticmethod
                def ScrollIntoView(_range):
                    return None

            class FakeApplication:
                ActiveWindow = FakeWindow()

            class FakeWps:
                Application = FakeApplication()

            pressed_keys = []
            original_press = self.answer_input.pyautogui.press
            original_sleep = self.answer_input.time.sleep
            self.answer_input.pyautogui.press = pressed_keys.append
            self.answer_input.time.sleep = lambda _seconds: None
            try:
                with redirect_stdout(io.StringIO()):
                    self.answer_input.execute_input(FakeDocument(), FakeWps(), blocks, 0, 2)
            finally:
                self.answer_input.pyautogui.press = original_press
                self.answer_input.time.sleep = original_sleep

            self.assertEqual(
                pressed_keys,
                ["f4", "f4", "f4", "f4", "f3", "f4", "f4", "f4", "f4", "f4", "f3"],
            )
            self.assertEqual(len(selected_texts), 11)
            answer_selection_indexes = [0, 1, 2, 3, 5, 6, 7, 8, 9]
            self.assertTrue(
                all("答案：" not in selected_texts[index] for index in answer_selection_indexes)
            )
            self.assertTrue(all(marker in selected_texts[-1] for marker in ["(1)", "(2)", "(3)", "(4)", "(5)"]))
            self.assertTrue(
                all(
                    heading not in selected_text
                    for selected_text in selected_texts
                    for heading in IGNORED_ANSWER_HEADINGS
                )
            )


if __name__ == "__main__":
    unittest.main()
