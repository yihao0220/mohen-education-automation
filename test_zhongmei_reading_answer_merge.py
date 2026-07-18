# coding: utf-8
from __future__ import annotations

import hashlib
import tempfile
import unittest
from pathlib import Path

from docx import Document

from shared_core import (
    build_answer_units_from_docx,
    build_question_units_from_docx,
    build_review_report,
    map_answers,
)
from shared_core.answer_core import build_answer_units_from_paragraph_texts
from shared_core.answer_core import infer_grouped_question_ids
from tools.merge_zhongmei_reading_answer_groups import merge_document_in_place


ANSWER_DIR = Path(
    r"D:\墨痕教育题目\众美-高三-语文\答案\对点练案答案"
)
QUESTION_DIR = ANSWER_DIR.parent.parent / "对点练案"


def _write_legacy_group_document(
    source: Path,
    target: Path,
    *,
    first_question_id: int,
    grouped_question_ids: range,
) -> None:
    source_doc = Document(source)
    source_texts = [paragraph.text for paragraph in source_doc.paragraphs]
    group_start = source_texts.index(f"{first_question_id}．")
    source_unit = next(
        unit
        for unit in build_answer_units_from_docx(source)
        if unit.question_id == str(first_question_id)
    )
    legacy = Document()
    for text in source_texts[:group_start]:
        legacy.add_paragraph(text)
    for question_id, answer_item, analysis_item in zip(
        grouped_question_ids,
        source_unit.answer_items,
        source_unit.analysis_items,
    ):
        legacy.add_paragraph(f"{question_id}．")
        legacy.add_paragraph(f"答案：{answer_item.text}")
        legacy.add_paragraph(f"解析：{analysis_item.text}")
    legacy.save(target)


class ZhongmeiReadingAnswerMergeTests(unittest.TestCase):
    def test_lesson_4_uses_paired_question_document_to_find_reading_group(self):
        question_path = QUESTION_DIR / "对点练案4.docx"
        source = ANSWER_DIR / "对点练案4_已清洗.docx"
        if not question_path.exists() or not source.exists():
            self.skipTest("缺少对点练案4真实题目或已清洗答案")

        with tempfile.TemporaryDirectory() as temp_dir:
            target_dir = Path(temp_dir) / "众美-高三-语文-王逸豪" / "答案" / "对点练案答案"
            target_dir.mkdir(parents=True)
            target = target_dir / source.name
            _write_legacy_group_document(
                source,
                target,
                first_question_id=4,
                grouped_question_ids=range(4, 9),
            )

            self.assertEqual(
                merge_document_in_place(target, question_path=question_path),
                1,
            )
            paragraphs = [paragraph.text.strip() for paragraph in Document(target).paragraphs]
            group_start = paragraphs.index("4．")
            group_end = len(paragraphs)
            merged = paragraphs[group_start:group_end]
            self.assertEqual(
                [text[:3] for text in merged[1:6]],
                ["(1)", "(2)", "(3)", "(4)", "(5)"],
            )
            self.assertEqual(merged[6], "解析：")
            self.assertEqual(
                [text[:3] for text in merged[7:12]],
                ["(1)", "(2)", "(3)", "(4)", "(5)"],
            )
            self.assertFalse(any(text in {"5．", "6．", "7．", "8．"} for text in merged))

    def test_lesson_3_keeps_blank_analysis_slots_after_merge(self):
        question_path = QUESTION_DIR / "对点练案3.docx"
        source = ANSWER_DIR / "对点练案3_已清洗.docx"
        if not question_path.exists() or not source.exists():
            self.skipTest("缺少对点练案3真实题目或已清洗答案")

        with tempfile.TemporaryDirectory() as temp_dir:
            target_dir = Path(temp_dir) / "众美-高三-语文-王逸豪" / "答案" / "对点练案答案"
            target_dir.mkdir(parents=True)
            target = target_dir / source.name
            _write_legacy_group_document(
                source,
                target,
                first_question_id=8,
                grouped_question_ids=range(8, 13),
            )

            self.assertEqual(
                merge_document_in_place(target, question_path=question_path),
                1,
            )
            mapped = map_answers(
                build_question_units_from_docx(question_path),
                build_answer_units_from_docx(target),
            )
            reading = next(unit for unit in mapped if unit.question_id == "8")
            self.assertEqual(len(reading.analysis_items), 5)
            self.assertEqual([item.text for item in reading.analysis_items[-2:]], ["", ""])

    def test_all_real_reading_groups_are_permanently_merged(self):
        if not QUESTION_DIR.exists() or not ANSWER_DIR.exists():
            self.skipTest("缺少44份真实众美题目或已清洗答案")

        total_groups = 0
        legacy_groups: list[str] = []
        for number in range(1, 45):
            question_path = QUESTION_DIR / f"对点练案{number}.docx"
            answer_path = ANSWER_DIR / f"对点练案{number}_已清洗.docx"
            if not question_path.exists() or not answer_path.exists():
                self.skipTest("缺少44份真实众美题目或已清洗答案")
            questions = build_question_units_from_docx(question_path)
            mapped = map_answers(questions, build_answer_units_from_docx(answer_path))
            mapped_by_id = {unit.question_id: unit for unit in mapped}
            for question in questions:
                grouped_ids = infer_grouped_question_ids(question)
                if len(grouped_ids) <= 1:
                    continue
                total_groups += 1
                answer = mapped_by_id[question.question_id]
                if answer.metadata.get("mapping_method") != "material_group_premerged":
                    legacy_groups.append(f"对点练案{number}第{question.question_id}题")

        self.assertEqual(total_groups, 40)
        self.assertEqual(legacy_groups, [])

    def test_shared_analysis_heading_keeps_each_subquestion_analysis_separate(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            target_dir = Path(temp_dir) / "众美-高三-语文-王逸豪" / "答案" / "对点练案答案"
            target_dir.mkdir(parents=True)
            target = target_dir / "对点练案2_已清洗.docx"
            document = Document()
            for text in [
                "5．",
                "(1)答案：C",
                "(2)答案：B",
                "解析：",
                "(1)第一题解析。",
                "(2)第二题解析。",
            ]:
                document.add_paragraph(text)
            document.save(target)
            units = build_answer_units_from_docx(target)

            self.assertEqual(len(units), 1)
            self.assertEqual([item.text for item in units[0].analysis_items], ["第一题解析。", "第二题解析。"])

    def test_parenthesized_numbers_inside_subanswer_text_are_not_new_items(self):
        units = build_answer_units_from_paragraph_texts(
            [
                "1．",
                "(1)答案：第一小题引用第(9)段。",
                "(2)答案：第二小题提到第(1)句和第(2)句。",
                "(3)答案：第三小题。",
                "解析：",
                "(1)第一题解析。",
                "(2)第二题解析。",
                "(3)第三题解析。",
            ]
        )

        self.assertEqual(len(units[0].answer_items), 3)
        self.assertIn("第(9)段", units[0].answer_items[0].text)
        self.assertIn("第(1)句和第(2)句", units[0].answer_items[1].text)

    def test_lesson_2_merges_five_answers_then_five_analyses_and_is_idempotent(self):
        question_path = QUESTION_DIR / "对点练案2.docx"
        source = ANSWER_DIR / "对点练案2_已清洗.docx"
        if not question_path.exists() or not source.exists():
            self.skipTest("缺少对点练案2真实题目或已清洗答案")

        with tempfile.TemporaryDirectory() as temp_dir:
            target_dir = Path(temp_dir) / "众美-高三-语文-王逸豪" / "答案" / "对点练案答案"
            target_dir.mkdir(parents=True)
            target = target_dir / "对点练案2_已清洗.docx"
            source_doc = Document(source)
            source_texts = [paragraph.text for paragraph in source_doc.paragraphs]
            group_start = source_texts.index("5．")
            source_unit = next(
                unit
                for unit in build_answer_units_from_docx(source)
                if unit.question_id == "5"
            )
            legacy = Document()
            for text in source_texts[:group_start]:
                legacy.add_paragraph(text)
            for question_id, answer_item, analysis_item in zip(
                range(5, 10), source_unit.answer_items, source_unit.analysis_items
            ):
                legacy.add_paragraph(f"{question_id}．")
                legacy.add_paragraph(f"答案：{answer_item.text}")
                legacy.add_paragraph(f"解析：{analysis_item.text}")
            legacy.save(target)

            self.assertEqual(
                merge_document_in_place(target, question_path=question_path),
                1,
            )
            paragraphs = [paragraph.text for paragraph in Document(target).paragraphs]
            group_start = paragraphs.index("5．")
            merged = paragraphs[group_start:]

            self.assertEqual(
                [text[:3] for text in merged[1:6]],
                ["(1)", "(2)", "(3)", "(4)", "(5)"],
            )
            self.assertTrue(all("答案：" in text for text in merged[1:6]))
            self.assertEqual(merged[6], "解析：")
            self.assertEqual(
                [text[:3] for text in merged[7:12]],
                ["(1)", "(2)", "(3)", "(4)", "(5)"],
            )
            self.assertFalse(any(text in {"6．", "7．", "8．", "9．"} for text in merged))

            questions = build_question_units_from_docx(question_path)
            mapped = map_answers(questions, build_answer_units_from_docx(target))
            report = build_review_report(target.name, questions, mapped)
            reading = next(unit for unit in mapped if unit.question_id == "5")
            self.assertEqual(len(reading.answer_items), 5)
            self.assertEqual(len(reading.analysis_items), 5)
            self.assertFalse(
                [
                    issue
                    for issue in report.issues
                    if issue.severity == "error" and issue.question_id == "5"
                ]
            )

            digest_before = hashlib.sha256(target.read_bytes()).hexdigest()
            self.assertEqual(
                merge_document_in_place(target, question_path=question_path),
                0,
            )
            self.assertEqual(hashlib.sha256(target.read_bytes()).hexdigest(), digest_before)


if __name__ == "__main__":
    unittest.main()
