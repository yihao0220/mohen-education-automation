from __future__ import annotations

import importlib.util
from io import BytesIO
import tempfile
import unittest
from pathlib import Path

from docx import Document
from shared_core import (
    build_answer_units_from_docx,
    build_question_units_from_docx,
    build_question_units_from_nodes,
    build_review_report,
    map_answers,
)
from shared_core.answer_core import infer_grouped_question_ids
from shared_core.models import AnswerItem, AnswerUnit, DocNode, QuestionUnit


PROJECT_ROOT = Path(__file__).resolve().parent


def _load_module(module_name: str, relative_path: str):
    module_path = PROJECT_ROOT / relative_path
    spec = importlib.util.spec_from_file_location(module_name, module_path)
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    spec.loader.exec_module(module)
    return module


splitter = _load_module("split_future_physics_answers", "tools/split_future_physics_answers.py")
template = _load_module("template_future_physics", "格式处理/格式模板库/template_future_physics.py")
answer_input = _load_module("future_physics_answer_input", "答案录入/answer_input.py")


class FuturePhysicsSplitTests(unittest.TestCase):
    def test_split_answer_docx_uses_question_filenames(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            question_dir = root / "questions"
            question_dir.mkdir()
            (question_dir / "课时跟踪检测(一)　描述运动的基本概念.docx").touch()
            (question_dir / "课时跟踪检测(二)　匀变速直线运动的规律.docx").touch()

            source_dir = root / "答案"
            source_dir.mkdir()
            source_path = source_dir / "课时跟踪检测答案.docx"

            source_doc = Document()
            body = source_doc._element.body
            for child in list(body):
                if child.tag.endswith("sectPr"):
                    continue
                body.remove(child)

            for paragraph_text in [
                "配套检测卷参考答案",
                "课时跟踪检测(一)",
                "1.选D 第一课时解析",
                "课时跟踪检测(二)",
                "1.解析:(1)第二课时解析",
                "答案:(1)第二课时答案",
            ]:
                source_doc.add_paragraph(paragraph_text)
            picture_paragraph = source_doc.add_paragraph()
            tiny_png = (
                b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01"
                b"\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde"
                b"\x00\x00\x00\x0cIDATx\x9cc\xf8\xff\xff?\x00\x05"
                b"\xfe\x02\xfeA\xe2)\xb7\x00\x00\x00\x00IEND\xaeB`\x82"
            )
            picture_paragraph.add_run().add_picture(BytesIO(tiny_png))
            source_doc.save(source_path)

            output_paths = splitter.split_answer_docx(
                source_path,
                question_dir=question_dir,
                output_dir=source_dir / "按课时拆分",
            )

            self.assertEqual(len(output_paths), 2)
            self.assertEqual(
                [path.name for path in output_paths],
                [
                    "课时跟踪检测(一)　描述运动的基本概念-答案.docx",
                    "课时跟踪检测(二)　匀变速直线运动的规律-答案.docx",
                ],
            )

            first_doc = Document(output_paths[0])
            self.assertEqual(first_doc.paragraphs[0].text.strip(), "课时跟踪检测(一)")
            self.assertEqual(first_doc.paragraphs[1].text.strip(), "1.选D 第一课时解析")

            second_doc = Document(output_paths[1])
            self.assertTrue(
                any("<w:drawing" in paragraph._element.xml for paragraph in second_doc.paragraphs),
                "拆分后的课时答案应保留图片/公式类绘图对象",
            )


class FuturePhysicsTemplateTests(unittest.TestCase):
    def test_parse_choice_and_big_question_variants(self):
        entries = template.parse_paragraph_texts(
            [
                "配套检测卷参考答案",
                "课时跟踪检测(一)",
                "1.选D　跳水比赛中要考虑运动员的动作",
                "2.",
                "选AB　由受力分析可知",
                "3.解析:(1)根据牛顿第二定律",
                "(2)继续推导",
                "答案:(1)2 m/s　(2)4 N",
                "4．",
                "解析：先求位移",
                "答案：5 m",
            ]
        )

        rendered = template.render_standard_lines(entries)

        self.assertEqual(rendered[0], "课时跟踪检测(一)")
        self.assertEqual(rendered[1], "1．D")
        self.assertEqual(rendered[2], "解析：跳水比赛中要考虑运动员的动作")
        self.assertEqual(rendered[3], "2．AB")
        self.assertEqual(rendered[4], "解析：由受力分析可知")
        self.assertEqual(rendered[5], "3．")
        self.assertEqual(rendered[6], "答案：(1)2 m/s (2)4 N")
        self.assertEqual(rendered[7], "解析：(1)根据牛顿第二定律")
        self.assertEqual(rendered[8], "(2)继续推导")
        self.assertEqual(rendered[9], "4．")
        self.assertEqual(rendered[10], "答案：5 m")
        self.assertEqual(rendered[11], "解析：先求位移")

    def test_parse_does_not_turn_time_prefix_into_new_question(self):
        entries = template.parse_paragraph_texts(
            [
                "9.解析:(1)小轿车从开始刹车到停止所用的最短时间为t1==20 s",
                "25 s时小轿车已停止,故通过的最小距离x1==300 m。",
                "(2)小轿车在反应时间内行驶的距离为15 m",
                "答案:(1)300 m (2)265 m",
                "10.解析:(1)根据题意继续求解",
                "答案:(1)2.4 m (2)152.5 s",
            ]
        )

        rendered = template.render_standard_lines(entries)

        self.assertTrue(rendered[0].startswith("9．"))
        self.assertIn("25 s时小轿车已停止,故通过的最小距离x1==300 m。", rendered)
        self.assertTrue(any(line.startswith("10．") for line in rendered))
        self.assertFalse(any(line.startswith("25．") for line in rendered))

    def test_parse_splits_embedded_question_start_in_same_paragraph(self):
        entries = template.parse_paragraph_texts(
            [
                "8.选CD 雪圈做匀速圆周运动,故C、D正确。9.解析:(1)若圆锥体与石块均静止",
                "(2)当圆锥体与石块一起转动时",
                "答案:(1)120 N (2)121.536 N",
                "10.解析:(1)静止时弹簧伸长了L",
                "答案:(1)k (2)ω",
            ]
        )

        rendered = template.render_standard_lines(entries)

        self.assertEqual(rendered[0], "8．CD")
        self.assertTrue(any(line.startswith("9．") for line in rendered))
        self.assertTrue(any(line.startswith("10．") for line in rendered))


class FuturePhysicsAnswerInputTests(unittest.TestCase):
    class _Doc:
        def __init__(self, full_name: str):
            self.FullName = full_name
            self.Name = Path(full_name).name

    def test_future_physics_subjective_subanswers_use_whole_input(self):
        doc = self._Doc(
            r"D:\墨痕教育题目\未来-高二-物理\答案\按课时拆分\课时跟踪检测(八)　摩擦力-答案_已清洗.docx"
        )
        block = {
            "qnum": "12",
            "answer_mode": "subquestion",
            "force_whole_answer_input": False,
        }
        self.assertFalse(answer_input.should_split_subquestion_answers(doc, block))

    def test_other_docs_keep_subanswer_split(self):
        doc = self._Doc(r"D:\墨痕教育题目\其他项目\答案\样例_已清洗.docx")
        block = {
            "qnum": "12",
            "answer_mode": "subquestion",
            "force_whole_answer_input": False,
        }
        self.assertTrue(answer_input.should_split_subquestion_answers(doc, block))


class FuturePhysicsSharedCoreTests(unittest.TestCase):
    def test_infer_grouped_question_ids_ignores_subquestion_markers(self):
        question = QuestionUnit(
            question_id="10",
            subject="理科",
            subject_overlay=None,
            grade_hint=None,
            question_type="subjective",
            stem_blocks=["10.(15分)轮滑赛道模型"],
            option_blocks=[],
            subquestions=["(1)求速度", "(2)求功", "(3)求距离"],
            media_blocks=[],
            material_blocks=[],
            source_span=(1, 4),
        )

        self.assertEqual(infer_grouped_question_ids(question), ["10"])

    def test_review_allows_future_physics_whole_answer_for_subquestions(self):
        question = QuestionUnit(
            question_id="1",
            subject="理科",
            subject_overlay=None,
            grade_hint=None,
            question_type="subjective",
            stem_blocks=["1.实验题"],
            option_blocks=[],
            subquestions=["(2)填写实验结果", "(3)填写实验结论"],
            media_blocks=[],
            material_blocks=[],
            source_span=(1, 3),
        )

        with tempfile.TemporaryDirectory() as temp_dir:
            future_dir = Path(temp_dir) / "未来-高二-物理"
            future_dir.mkdir()
            answer_path = future_dir / "样例-答案_已清洗.docx"

            answer_doc = Document()
            answer_doc.add_paragraph("1．(2)CD (3)标记位置")
            answer_doc.add_paragraph("解析：(2)需要记录方向")
            answer_doc.add_paragraph("(3)需要比较大小和方向")
            answer_doc.save(answer_path)

            mapped = map_answers([question], build_answer_units_from_docx(answer_path))
            report = build_review_report(answer_path.name, [question], mapped)
            titles = [issue.title for issue in report.issues if issue.severity == "error"]

            self.assertNotIn("题目含小问但答案未拆分", titles)
            self.assertEqual(report.summary["high_risk_count"], 0)

    def test_review_allows_future_physics_analysis_only_subanswers(self):
        question = QuestionUnit(
            question_id="1",
            subject="理科",
            subject_overlay=None,
            grade_hint=None,
            question_type="subjective",
            stem_blocks=["1.计算题"],
            option_blocks=[],
            subquestions=["(1)求电场强度", "(2)求速度"],
            media_blocks=[],
            material_blocks=[],
            source_span=(1, 3),
        )

        with tempfile.TemporaryDirectory() as temp_dir:
            future_dir = Path(temp_dir) / "未来-高二-物理"
            future_dir.mkdir()
            answer_path = future_dir / "样例-答案_已清洗.docx"

            answer_doc = Document()
            answer_doc.add_paragraph("1．(1)")
            answer_doc.add_paragraph("(2)")
            answer_doc.add_paragraph("解析：(1)根据电场强度定义可得E")
            answer_doc.add_paragraph("(2)联立可得速度v")
            answer_doc.save(answer_path)

            mapped = map_answers([question], build_answer_units_from_docx(answer_path))
            report = build_review_report(answer_path.name, [question], mapped)
            error_titles = [issue.title for issue in report.issues if issue.severity == "error"]
            warning_titles = [issue.title for issue in report.issues if issue.severity == "warning"]

            self.assertEqual(error_titles, [])
            self.assertIn("小问答案转入解析承载", warning_titles)

    def test_docx_question_scan_splits_embedded_next_question(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            question_path = Path(temp_dir) / "课时跟踪检测(三十九)　四种类碰撞模型的研究.docx"

            question_doc = Document()
            question_doc.add_paragraph("9.(13分)第九题题干")
            question_doc.add_paragraph("(1)子问一")
            question_doc.add_paragraph("(2)子问二。10.(16分)第十题题干")
            question_doc.add_paragraph("(1)第十题子问一")
            question_doc.add_paragraph("(2)第十题子问二")
            question_doc.save(question_path)

            units = build_question_units_from_docx(question_path)

            self.assertEqual([unit.question_id for unit in units], ["9", "10"])

    def test_docx_question_scan_accepts_starred_question_number(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            question_path = Path(temp_dir) / "课时跟踪检测(五十三)　磁场的描述 磁场对电流的作用.docx"

            question_doc = Document()
            question_doc.add_paragraph("6.第六题题干")
            question_doc.add_paragraph("A.甲")
            question_doc.add_paragraph("B.乙")
            question_doc.add_paragraph("★7.如图所示为电磁炮的基本原理图")
            question_doc.add_paragraph("A.导轨间的磁场方向向上")
            question_doc.add_paragraph("B.导轨间的磁场磁感应强度大小为")
            question_doc.add_paragraph("8.第八题题干")
            question_doc.add_paragraph("A.丙")
            question_doc.add_paragraph("B.丁")
            question_doc.save(question_path)

            units = build_question_units_from_docx(question_path)

            self.assertEqual([unit.question_id for unit in units], ["6", "7", "8"])

    def test_wps_question_scan_does_not_split_table_decimal_cells(self):
        nodes = [
            DocNode(index=3, text="1.(6分)某一物理兴趣实验小组研究摩托车速度随时间变化的规律。"),
            DocNode(index=20, text="0.48\r\x07", metadata={"in_table": True}),
            DocNode(index=21, text="1.24\r\x07", metadata={"in_table": True}),
            DocNode(index=22, text="2.32\r\x07", metadata={"in_table": True}),
            DocNode(index=30, text="(1)由表格数据可判断摩托车近似做匀加速直线运动。"),
            DocNode(index=31, text="(2)当x=1.24 m时摩托车的速度为____m/s;"),
            DocNode(index=33, text="2.(6分)研究小组用如图甲所示的装置来研究自由落体运动。"),
        ]

        units = build_question_units_from_nodes("课时跟踪检测(六)　探究小车速度随时间变化的规律.docx", "理科", nodes)

        self.assertEqual([unit.question_id for unit in units], ["1", "2"])
        self.assertEqual(units[0].source_span, (3, 32))

    def test_review_does_not_treat_track_labels_as_embedded_choice_answers(self):
        question = QuestionUnit(
            question_id="7",
            subject="理科",
            subject_overlay=None,
            grade_hint=None,
            question_type="choice",
            stem_blocks=["7.导轨运动问题"],
            option_blocks=["A.甲", "B.乙", "C.丙", "D.丁"],
            subquestions=[],
            media_blocks=[],
            material_blocks=[],
            source_span=(1, 2),
        )
        answer = AnswerUnit(
            question_id="7",
            answer_mode="whole",
            answer_items=[AnswerItem(item_id="7", text="CD")],
            analysis_items=[
                AnswerItem(
                    item_id="7",
                    text="金属杆在AA1B1B区域和BB1C1C区域运动，故C、D正确。",
                )
            ],
        )

        report = build_review_report("样例.docx", [question], [answer])
        error_titles = [issue.title for issue in report.issues if issue.severity == "error"]

        self.assertNotIn("答案块疑似串入其他题号", error_titles)


if __name__ == "__main__":
    unittest.main()
