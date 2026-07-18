from __future__ import annotations

import re
import tempfile
import unittest
from pathlib import Path

from docx import Document

from tools.clean_zhongmei_chinese_answers import (
    UnsupportedTemplateError,
    clean_batch,
    parse_document,
    preflight_source_files,
    render_clean_lines,
    validate_clean_docx,
    write_clean_docx,
)


SAMPLE_ROOT = Path(r"D:\墨痕教育题目\众美-高三-语文")
ANSWER_DIR = SAMPLE_ROOT / "答案" / "对点练案答案"
EXPECTED_QUESTION_COUNTS = {
    1: 12,
    2: 9,
    3: 12,
    4: 8,
    5: 8,
    6: 6,
    7: 6,
    8: 7,
    9: 6,
    10: 8,
    11: 6,
    12: 8,
    13: 5,
    14: 6,
    15: 4,
    16: 7,
    17: 10,
    18: 11,
    19: 10,
    20: 9,
    21: 9,
    22: 7,
    23: 7,
    24: 8,
    25: 6,
    26: 7,
    27: 6,
    28: 6,
    29: 6,
    30: 7,
    31: 30,
    32: 14,
    33: 14,
    34: 12,
    35: 9,
    36: 10,
    37: 13,
    38: 12,
    39: 10,
    40: 12,
    41: 9,
    42: 10,
    43: 9,
    44: 9,
}


def _write_docx(path: Path, paragraphs: list[str], *, table_text: str | None = None) -> None:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    if table_text is not None:
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = table_text
    doc.save(path)


def _normalized_reference_lines(path: Path) -> list[str]:
    return [
        line.strip()
        for line in path.read_text(encoding="utf-8").splitlines()
        if line.strip() and not line.lstrip().startswith("#")
    ]


def _normalized_actual_lines(lines: list[str]) -> list[str]:
    normalized: list[str] = []
    for line in lines:
        stripped = line.strip()
        if stripped == "解析：":
            continue
        normalized.append(re.sub(r"^(\d+)．$", r"\1", stripped))
    return normalized


class ZhongmeiChineseAnswerCleanerTests(unittest.TestCase):
    def test_clean_batch_permanently_merges_complete_reading_group(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            sample_root = Path(temp_dir) / "众美-高三-语文-王逸豪"
            question_dir = sample_root / "对点练案"
            answer_dir = sample_root / "答案" / "对点练案答案"
            question_dir.mkdir(parents=True)
            answer_dir.mkdir(parents=True)

            _write_docx(
                question_dir / "对点练案1.docx",
                [
                    "对点练案1",
                    "(分值：6分)",
                    "阅读下面的文字，完成文后题目。",
                    "材料正文。",
                    "1．第一题",
                    "2．第二题",
                ],
            )
            _write_docx(
                answer_dir / "对点练案1.docx",
                [
                    "1．第一题",
                    "答案 A",
                    "解析 第一题解析。",
                    "2．第二题",
                    "答案 B",
                    "解析 第二题解析。",
                ],
            )

            [output_path] = clean_batch(answer_dir)
            paragraphs = [
                paragraph.text.strip()
                for paragraph in Document(output_path).paragraphs
                if paragraph.text.strip()
            ]

            self.assertEqual(
                paragraphs,
                [
                    "1．",
                    "(1)答案：A",
                    "(2)答案：B",
                    "解析：",
                    "(1)第一题解析。",
                    "(2)第二题解析。",
                ],
            )

    def test_normative_markdown_samples_match(self):
        missing = [
            path
            for number in (1, 2)
            for path in (
                ANSWER_DIR / f"对点练案{number}.docx",
                SAMPLE_ROOT / f"对点练{number}-清洗后文档.md",
            )
            if not path.exists()
        ]
        if missing:
            self.skipTest(f"缺少规范样例：{missing}")

        for number in (1, 2):
            clean_document = parse_document(ANSWER_DIR / f"对点练案{number}.docx")
            actual = _normalized_actual_lines(render_clean_lines(clean_document))
            expected = _normalized_reference_lines(
                SAMPLE_ROOT / f"对点练{number}-清洗后文档.md"
            )
            self.assertEqual(actual, expected, f"对点练案{number} 与规范文档不一致")

    def test_all_real_sources_preflight_and_question_counts(self):
        source_files = [
            ANSWER_DIR / f"对点练案{number}.docx"
            for number in EXPECTED_QUESTION_COUNTS
        ]
        if any(not path.exists() for path in source_files):
            self.skipTest("缺少 44 份真实答案样本")

        clean_documents = preflight_source_files(source_files)
        actual_counts = {
            int(re.search(r"\d+", clean_document.source_path.stem).group()): len(
                clean_document.questions
            )
            for clean_document in clean_documents
        }
        self.assertEqual(actual_counts, EXPECTED_QUESTION_COUNTS)
        self.assertEqual(sum(actual_counts.values()), 400)

    def test_missing_analysis_adds_blank_analysis_marker(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "对点练案1.docx"
            _write_docx(source, ["一、专项训练", "1．题干", "答案　A"])
            lines = render_clean_lines(parse_document(source))
            self.assertEqual(lines, ["一、专项训练", "1．", "答案：A", "解析： "])

    def test_multisubquestion_answers_keep_labels_and_add_one_blank_analysis(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "对点练案1.docx"
            _write_docx(
                source,
                [
                    "1．判断下列说法。",
                    "(1)第一小问",
                    "答案　正确。",
                    "(2)第二小问",
                    "答案　错误。",
                ],
            )
            lines = render_clean_lines(parse_document(source))
            self.assertEqual(
                lines,
                [
                    "1．",
                    "(1)答案：正确。",
                    "(2)答案：错误。",
                    "解析： ",
                ],
            )

    def test_numbered_analysis_continuations_do_not_start_new_questions(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "对点练案1.docx"
            _write_docx(
                source,
                [
                    "1．第一题题干",
                    "答案　A",
                    "解析　第一题解析。",
                    "1．解析要点一。",
                    "2．解析要点二。",
                    "2．第二题题干",
                    "答案　B",
                ],
            )
            clean_document = parse_document(source)
            self.assertEqual([question.number for question in clean_document.questions], [1, 2])
            self.assertEqual(
                clean_document.questions[0].occurrences[0].analysis_lines,
                ["第一题解析。", "1．解析要点一。", "2．解析要点二。"],
            )

    def test_unknown_answer_marker_stops_preflight(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "对点练案1.docx"
            _write_docx(source, ["1．题干", "【答案】A", "【解析】说明"])
            with self.assertRaisesRegex(UnsupportedTemplateError, "未发现行首“答案”标记"):
                parse_document(source)

    def test_answer_marker_inside_table_stops_preflight(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "对点练案1.docx"
            _write_docx(source, ["1．题干", "答案　A"], table_text="解析　表格解析")
            with self.assertRaisesRegex(UnsupportedTemplateError, "表格中含答案或解析标记"):
                parse_document(source)

    def test_written_docx_reopens_and_matches_expected_text(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "对点练案1.docx"
            output = Path(temp_dir) / "对点练案1_已清洗.docx"
            _write_docx(
                source,
                ["一、专项训练", "1．题干", "答案　A", "解析　完整解析。"],
            )
            clean_document = parse_document(source)
            write_clean_docx(clean_document, output)
            validate_clean_docx(output, clean_document)
            reopened = Document(output)
            self.assertEqual(
                [paragraph.text.strip() for paragraph in reopened.paragraphs if paragraph.text.strip()],
                ["一、专项训练", "1．", "答案：A", "解析：完整解析。"],
            )


if __name__ == "__main__":
    unittest.main()
