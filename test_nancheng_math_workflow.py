# coding: utf-8
from __future__ import annotations

import importlib.util
import os
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

from shared_core.answer_core import build_answer_units_from_paragraph_texts, map_answers
from shared_core.question_core import build_question_units_from_docx
from shared_core.review import build_review_report


PROJECT_ROOT = Path(__file__).resolve().parent
NANCHENG_G4_DIR = Path(r"D:\墨痕教育题目\南城二小-四年级-数学-王逸豪")
NANCHENG_G5_DIR = Path(r"D:\墨痕教育题目\南城二小-五年级-数学-王逸豪")
NANCHENG_ANSWER_DOC = NANCHENG_G4_DIR / "答案与解析数学.docx"
FORMAT_MAIN_SPEC = importlib.util.spec_from_file_location(
    "mohen_format_main_nancheng_test",
    PROJECT_ROOT / "格式处理" / "main.py",
)
format_main = importlib.util.module_from_spec(FORMAT_MAIN_SPEC)
assert FORMAT_MAIN_SPEC and FORMAT_MAIN_SPEC.loader
FORMAT_MAIN_SPEC.loader.exec_module(format_main)
SPLITTER_SPEC = importlib.util.spec_from_file_location(
    "split_nancheng_math_answers_test",
    PROJECT_ROOT / "tools" / "split_nancheng_math_answers.py",
)
splitter = importlib.util.module_from_spec(SPLITTER_SPEC)
assert SPLITTER_SPEC and SPLITTER_SPEC.loader
SPLITTER_SPEC.loader.exec_module(splitter)

sys.path.insert(0, str(PROJECT_ROOT / "格式处理"))
sys.path.insert(0, str(PROJECT_ROOT))

from 格式模板库 import template_nancheng_math  # noqa: E402


def _load_nonempty_paragraphs(path: Path) -> list[str]:
    doc = Document(path)
    return [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]


def _write_docx(path: Path, lines: list[str]) -> None:
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    doc.save(path)


def _find_g5_78_answer_doc() -> Path:
    answer_dir = NANCHENG_G5_DIR / "五年级下册Ai学导练参考答案"
    matches = sorted(answer_dir.glob("Ai学导练（六）参考答案 78单元*.docx"))
    if not matches:
        raise FileNotFoundError("未找到五年级 7-8 单元答案 docx")
    return matches[-1]


class FakeRange:
    def __init__(self, text: str):
        self.Text = text


class FakeParagraph:
    def __init__(self, text: str):
        self.Range = FakeRange(text)


class FakeParagraphCollection:
    def __init__(self, texts: list[str]):
        self._items = [FakeParagraph(text) for text in texts]
        self.Count = len(self._items)

    def __call__(self, index: int):
        return self._items[index - 1]


class FakeDoc:
    def __init__(self, name: str, texts: list[str]):
        self.Name = name
        self.Paragraphs = FakeParagraphCollection(texts)


class NanchengMathQuestionScanTests(unittest.TestCase):
    def test_scans_spaced_dot_and_section_prefixed_question_numbers(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            question_path = Path(temp_dir) / "五年级数学南城题号兼容.docx"
            _write_docx(
                question_path,
                [
                    "13.上一题",
                    "14 . 将  与 通分得到两个新的分数。",
                    "四、根据下面的统计图填空。",
                    "五、23.下面是好运公司2019年各月利润情况的折线统计图。",
                    "(1)(     )月的利润最多。",
                ],
            )

            units = build_question_units_from_docx(question_path)

            self.assertEqual([unit.question_id for unit in units], ["13", "14", "23"])

    def test_real_nancheng_fifth_grade_docs_have_continuous_question_numbers(self):
        expected = {
            "2025-2026学年度第二学期五年级数学AI学导练（4单元）.docx": (30, []),
            "2025-2026学年度第二学期五年级数学AI学导练（7-8单元）.docx": (27, []),
        }
        for filename, (expected_count, expected_missing) in expected.items():
            with self.subTest(filename=filename):
                units = build_question_units_from_docx(NANCHENG_G5_DIR / filename)
                numbers = [int(unit.question_id) for unit in units]
                missing = [n for n in range(1, max(numbers) + 1) if n not in numbers]

                self.assertEqual(len(units), expected_count)
                self.assertEqual(missing, expected_missing)


class NanchengMathAnswerTemplateTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.answer_texts = _load_nonempty_paragraphs(NANCHENG_ANSWER_DOC)

    def test_nancheng_template_wins_match_for_total_answer_doc(self):
        fake_doc = FakeDoc(NANCHENG_ANSWER_DOC.name, self.answer_texts)
        template, score = format_main.match_best_template(
            fake_doc,
            interactive=False,
            doc_name=NANCHENG_ANSWER_DOC.name,
        )

        self.assertIs(template, template_nancheng_math)
        self.assertGreater(score, 0.20)

    def test_renders_fourth_grade_unit_answers_with_global_question_ids(self):
        unit_texts = template_nancheng_math.extract_unit_texts(self.answer_texts, "第5单元")
        question_units = build_question_units_from_docx(NANCHENG_G4_DIR / "2026人教版四年级下册数学第5单元测评卷.docx")
        entries = template_nancheng_math.parse_unit_answer_texts(unit_texts, question_units)
        rendered_lines = template_nancheng_math.render_standard_lines(entries)
        answer_units = build_answer_units_from_paragraph_texts(rendered_lines)

        self.assertEqual(len(answer_units), 25)
        self.assertEqual([unit.question_id for unit in answer_units[:6]], ["1", "2", "3", "4", "5", "6"])
        self.assertEqual(answer_units[10].question_id, "11")
        self.assertEqual(answer_units[10].answer_items[0].text, "C")
        self.assertEqual(answer_units[24].question_id, "25")
        self.assertTrue(answer_units[24].analysis_items)

    def test_real_fourth_grade_answer_units_pass_review_after_mapping(self):
        unit_expectations = [
            ("第5单元", "2026人教版四年级下册数学第5单元测评卷.docx", 25),
            ("第6单元", "2026人教版四年级下册数学第6单元测评卷.docx", 23),
            ("第7单元", "2026人教版四年级下册数学第7单元测评卷.docx", 18),
            ("第8、9单元", "2026人教版四年级下册数学第89单元测评卷.docx", 22),
        ]
        for unit_title, question_name, expected_count in unit_expectations:
            with self.subTest(unit_title=unit_title):
                question_units = build_question_units_from_docx(NANCHENG_G4_DIR / question_name)
                unit_texts = template_nancheng_math.extract_unit_texts(self.answer_texts, unit_title)
                entries = template_nancheng_math.parse_unit_answer_texts(unit_texts, question_units)
                rendered_lines = template_nancheng_math.render_standard_lines(entries)
                answer_units = build_answer_units_from_paragraph_texts(rendered_lines)
                mapped_units = map_answers(question_units, answer_units)
                report = build_review_report(f"{unit_title}-答案.docx", question_units, mapped_units)

                self.assertEqual(len(answer_units), expected_count)
                self.assertEqual(report.summary["high_risk_count"], 0)

    def test_real_fifth_grade_78_answer_doc_maps_compact_choices_and_subanswers(self):
        question_path = NANCHENG_G5_DIR / "2025-2026学年度第二学期五年级数学AI学导练（7-8单元）.docx"
        answer_path = _find_g5_78_answer_doc()
        question_units = build_question_units_from_docx(question_path)
        entries = template_nancheng_math.parse_unit_answer_texts(_load_nonempty_paragraphs(answer_path), question_units)
        rendered_lines = template_nancheng_math.render_standard_lines(entries)
        answer_units = build_answer_units_from_paragraph_texts(rendered_lines)
        mapped_units = map_answers(question_units, answer_units)
        report = build_review_report(answer_path.name, question_units, mapped_units)

        self.assertEqual(len(answer_units), len(question_units))
        self.assertEqual(answer_units[15].question_id, "16")
        self.assertEqual(answer_units[15].answer_items[0].text, "B")
        unit_23 = next(unit for unit in answer_units if unit.question_id == "23")
        self.assertEqual(unit_23.answer_mode, "subquestion")
        self.assertEqual([item.text for item in unit_23.answer_items], [
            "10  50",
            "4  20",
            "4  10  1  4",
            "44-40=4(万元)",
            "(32+25+21)÷3=26(万元)",
            "(36+40+48+50+40+44)÷6=43(万元)",
        ])
        self.assertEqual(report.summary["high_risk_count"], 0)


class NanchengMathSplitToolTests(unittest.TestCase):
    def test_split_total_answer_docx_writes_reviewable_unit_answer_docs(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output_dir = Path(temp_dir) / "南城拆分输出"
            output_paths = splitter.split_answer_docx(
                NANCHENG_ANSWER_DOC,
                question_dir=NANCHENG_G4_DIR,
                output_dir=output_dir,
            )

            self.assertEqual(
                [path.name for path in output_paths],
                [
                    "2026人教版四年级下册数学第5单元测评卷-答案_已清洗.docx",
                    "2026人教版四年级下册数学第6单元测评卷-答案_已清洗.docx",
                    "2026人教版四年级下册数学第7单元测评卷-答案_已清洗.docx",
                    "2026人教版四年级下册数学第89单元测评卷-答案_已清洗.docx",
                ],
            )

            for output_path in output_paths:
                question_name = output_path.name.replace("-答案_已清洗.docx", ".docx")
                question_units = build_question_units_from_docx(NANCHENG_G4_DIR / question_name)
                answer_units = build_answer_units_from_paragraph_texts(_load_nonempty_paragraphs(output_path))
                report = build_review_report(output_path.name, question_units, map_answers(question_units, answer_units))

                self.assertEqual(len(answer_units), len(question_units))
                self.assertEqual(report.summary["high_risk_count"], 0)

    def test_standardize_single_answer_docx_writes_reviewable_answer_doc(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output_dir = Path(temp_dir) / "南城单卷输出"
            question_path = NANCHENG_G5_DIR / "2025-2026学年度第二学期五年级数学AI学导练（7-8单元）.docx"
            answer_path = _find_g5_78_answer_doc()

            output_path = splitter.standardize_answer_docx(
                answer_path,
                question_doc=question_path,
                output_dir=output_dir,
            )

            answer_units = build_answer_units_from_paragraph_texts(_load_nonempty_paragraphs(output_path))
            report = build_review_report(
                output_path.name,
                build_question_units_from_docx(question_path),
                map_answers(build_question_units_from_docx(question_path), answer_units),
            )

            self.assertEqual(output_path.name, "2025-2026学年度第二学期五年级数学AI学导练（7-8单元）-答案_已清洗.docx")
            self.assertTrue(output_path.with_name(f"{output_path.stem}_审核清单.md").exists())
            self.assertTrue(output_path.with_name(f"{output_path.stem}_审核状态.json").exists())
            self.assertEqual(len(answer_units), 27)
            self.assertEqual(report.summary["high_risk_count"], 0)


if __name__ == "__main__":
    unittest.main()
