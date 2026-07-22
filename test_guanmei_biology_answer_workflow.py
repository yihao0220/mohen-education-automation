from __future__ import annotations

import importlib.util
from pathlib import Path
import sys
import tempfile
from types import SimpleNamespace
import unittest
from zipfile import ZipFile

from docx import Document

from shared_core.answer_core import (
    build_answer_units_from_paragraph_texts,
    build_answer_units_from_wps,
)
from shared_core.review_gate import get_review_gate_result
from tools.package_guanmei_biology_answers import package_guanmei_biology_answers
from tools.refresh_guanmei_biology_review_status import (
    refresh_guanmei_biology_review_statuses,
)


def _load_answer_input_module():
    sys.modules.setdefault(
        "pyautogui",
        SimpleNamespace(press=lambda *_args, **_kwargs: None),
    )
    sys.modules.setdefault(
        "wps_helper",
        SimpleNamespace(get_active_wps=lambda: None),
    )
    module_path = Path(__file__).resolve().parent / "答案录入" / "answer_input.py"
    spec = importlib.util.spec_from_file_location(
        "guanmei_biology_answer_input",
        module_path,
    )
    if not spec or not spec.loader:
        raise AssertionError(f"无法加载答案录入模块: {module_path}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


answer_input = _load_answer_input_module()


class _MockRange:
    def __init__(self, text: str) -> None:
        self.Text = text

    def Information(self, _code: int) -> bool:
        return False


class _MockParagraph:
    def __init__(self, text: str) -> None:
        self.Range = _MockRange(text)


class _MockParagraphs:
    def __init__(self, texts: list[str]) -> None:
        self._items = [_MockParagraph(text) for text in texts]
        self.Count = len(self._items)

    def __call__(self, index: int) -> _MockParagraph:
        return self._items[index - 1]


class _MockDoc:
    def __init__(self, texts: list[str]) -> None:
        self.FullName = (
            "D:/墨痕教育题目/莞美-高二-生物/答案/第1章/"
            "第1节　细胞的生活环境-答案_已清洗.docx"
        )
        self.Name = "第1节　细胞的生活环境-答案_已清洗.docx"
        self.Paragraphs = _MockParagraphs(texts)


class GuanmeiBiologyAnswerWorkflowTest(unittest.TestCase):
    def test_standalone_compact_choices_keep_source_paragraph_positions(self) -> None:
        units = build_answer_units_from_paragraph_texts(
            [
                "1B",
                "解析：第一题解析",
                "2D",
                "解析：第二题解析",
            ],
            preserve_source_positions=True,
        )

        self.assertEqual([unit.question_id for unit in units], ["1", "2"])
        self.assertEqual([unit.answer_items[0].text for unit in units], ["B", "D"])
        self.assertEqual([unit.source_span for unit in units], [(1, 2), (3, 4)])
        self.assertEqual([unit.answer_span for unit in units], [(1, 1), (3, 3)])
        self.assertEqual([unit.analysis_span for unit in units], [(2, 2), (4, 4)])

    def test_compact_choice_requires_the_whole_paragraph(self) -> None:
        units = build_answer_units_from_paragraph_texts(
            [
                "1B细胞生活在内环境中",
                "解析：这不是标准答案行",
            ],
            preserve_source_positions=True,
        )

        self.assertEqual(units, [])

    def test_dotted_choice_format_remains_unchanged(self) -> None:
        units = build_answer_units_from_paragraph_texts(
            ["1．B", "解析：原格式解析"],
            preserve_source_positions=True,
        )

        self.assertEqual(len(units), 1)
        self.assertEqual(units[0].question_id, "1")
        self.assertEqual(units[0].answer_items[0].text, "B")

    def test_wps_guanmei_subjective_answer_restores_inline_subquestions(self) -> None:
        units = build_answer_units_from_wps(
            _MockDoc(
                [
                    "13．（1）甲、乙、丙　（2）上升　（3）b→c→d",
                    "解析：三问共用一段解析",
                ]
            )
        )

        self.assertEqual(len(units), 1)
        self.assertEqual(units[0].question_id, "13")
        self.assertEqual(units[0].answer_mode, "subquestion")
        self.assertEqual(
            [item.text for item in units[0].answer_items],
            ["甲、乙、丙", "上升", "b→c→d"],
        )
        self.assertEqual(units[0].answer_span, (1, 1))
        self.assertEqual(units[0].analysis_span, (2, 2))

    def test_cross_paragraph_answer_collects_inline_third_subquestion(self) -> None:
        units = build_answer_units_from_wps(
            _MockDoc(
                [
                    "12．（1）体液免疫",
                    "（2）遗传倾向和个体差异　（3）避免接触过敏原",
                    "解析：三问共用一段解析",
                ]
            )
        )

        self.assertEqual(len(units), 1)
        self.assertEqual(units[0].question_id, "12")
        self.assertEqual(
            [item.item_id for item in units[0].answer_items],
            ["（1）", "（2）", "（3）"],
        )
        self.assertEqual(
            [item.text for item in units[0].answer_items],
            ["体液免疫", "遗传倾向和个体差异", "避免接触过敏原"],
        )
        self.assertEqual(units[0].answer_span, (1, 2))
        self.assertEqual(units[0].analysis_span, (3, 3))

    def test_execution_strips_compact_choice_question_number(self) -> None:
        offset, text_body = answer_input.strip_input_question_prefix("1B\r")

        self.assertEqual(offset, 1)
        self.assertEqual(text_body, "B\r")
        self.assertEqual(answer_input.find_subquestion_matches(text_body), [])

    def test_execution_preserves_inline_subquestion_one(self) -> None:
        answer_text = "13．（1）甲（2）乙（3）丙\r"
        offset, text_body = answer_input.strip_input_question_prefix(answer_text)
        matches = answer_input.find_subquestion_matches(text_body)

        self.assertEqual(offset, len("13．"))
        self.assertTrue(text_body.startswith("（1）"))
        self.assertEqual(
            [match.group(1) for match in matches],
            ["（1）", "（2）", "（3）"],
        )

    def test_execution_collects_inline_third_subquestion_after_line_two(self) -> None:
        answer_text = "12．（1）甲\r（2）乙　（3）丙\r"
        _, text_body = answer_input.strip_input_question_prefix(answer_text)
        matches = answer_input.find_subquestion_matches(text_body)

        self.assertEqual(
            [match.group(1) for match in matches],
            ["（1）", "（2）", "（3）"],
        )

    def test_execution_ignores_parenthesized_reference_inside_subanswer(self) -> None:
        answer_text = (
            "12．\r(1)第一小问答案中引用第(9)段。"
            "\r(2)第二小问答案。"
            "\r(3)第三小问答案。"
        )
        _, text_body = answer_input.strip_input_question_prefix(answer_text)
        matches = answer_input.find_subquestion_matches(text_body)

        self.assertEqual(
            [match.group(1) for match in matches],
            ["(1)", "(2)", "(3)"],
        )

    def test_execution_keeps_single_subquestion_and_physics_policies(self) -> None:
        _, text_body = answer_input.strip_input_question_prefix(
            "20．\r(1)单个翻译小问\r"
        )
        matches = answer_input.find_subquestion_matches(text_body)
        block = {
            "answer_mode": "subquestion",
            "force_whole_answer_input": False,
        }

        self.assertEqual([match.group(1) for match in matches], ["(1)"])
        self.assertTrue(answer_input.should_split_subquestion_answers(_MockDoc([]), block))

        physics_doc = SimpleNamespace(
            FullName=(
                "D:/墨痕教育题目/未来-高二-物理/答案/"
                "样例_已清洗.docx"
            ),
            Name="样例_已清洗.docx",
        )
        self.assertFalse(
            answer_input.should_split_subquestion_answers(physics_doc, block)
        )

    def test_execution_keeps_circled_points_in_one_answer(self) -> None:
        self.assertEqual(
            answer_input.find_subquestion_matches("①要点一②要点二③要点三"),
            [],
        )

    def test_review_status_and_windows_package_remain_portable(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_dir:
            project_root = Path(temporary_dir) / "莞美-高二-生物"
            question_dir = project_root / "第1章"
            answer_dir = project_root / "答案" / "按章节拆分" / "第1章"
            question_dir.mkdir(parents=True)
            answer_dir.mkdir(parents=True)

            question_path = question_dir / "第1节　细胞的生活环境-排版终稿.docx"
            question_doc = Document()
            question_doc.add_paragraph("1．下列说法正确的是（　　）")
            question_doc.add_paragraph("A．选项A")
            question_doc.add_paragraph("B．选项B")
            question_doc.add_paragraph("C．选项C")
            question_doc.add_paragraph("D．选项D")
            question_doc.save(question_path)

            answer_path = answer_dir / "第1节　细胞的生活环境-答案_已清洗.docx"
            answer_doc = Document()
            answer_doc.add_paragraph("1B")
            answer_doc.add_paragraph("解析：选项B正确。")
            answer_doc.save(answer_path)

            results = refresh_guanmei_biology_review_statuses(
                project_root,
                expected_count=1,
            )
            self.assertEqual(len(results), 1)
            self.assertTrue(results[0]["allowed"])
            self.assertTrue(get_review_gate_result(answer_path)["allowed"])

            archive_path, document_count, file_count = package_guanmei_biology_answers(
                project_root,
                expected_count=1,
            )
            self.assertEqual(document_count, 1)
            self.assertEqual(file_count, 2)
            with ZipFile(archive_path) as package:
                names = package.namelist()
                self.assertEqual(len(names), 2)
                self.assertTrue(all(name.startswith("答案/第1章/") for name in names))
                self.assertIsNone(package.testzip())


if __name__ == "__main__":
    unittest.main()
