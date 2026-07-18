from pathlib import Path

from docx import Document

from shared_core import get_review_gate_result
from tools.clean_zhongmei_chinese_answers import clean_batch
from tools.zhongmei_review_status import refresh_review_statuses


def _write_docx(path: Path, paragraphs: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(path)


def test_refresh_review_status_approves_valid_pair_and_detects_later_change(tmp_path):
    project_dir = tmp_path / "众美-高三-语文-测试"
    question_dir = project_dir / "对点练案"
    answer_dir = project_dir / "答案" / "对点练案答案"
    question_path = question_dir / "对点练案1.docx"
    answer_path = answer_dir / "对点练案1_已清洗.docx"

    _write_docx(
        question_path,
        ["1．下列说法正确的一项是（　　）", "A．甲", "B．乙", "C．丙", "D．丁"],
    )
    _write_docx(answer_path, ["1．", "答案：A", "解析：根据原文可知。"])

    result = refresh_review_statuses([answer_path])

    assert result[0]["status"] == "approved"
    assert get_review_gate_result(answer_path)["allowed"] is True

    document = Document(answer_path)
    document.add_paragraph("人工改动")
    document.save(answer_path)

    stale = get_review_gate_result(answer_path)
    assert stale["allowed"] is False
    assert stale["status"] == "stale"


def test_refresh_review_status_refuses_missing_question_without_writing_state(tmp_path):
    answer_dir = tmp_path / "项目" / "答案" / "对点练案答案"
    answer_path = answer_dir / "对点练案1_已清洗.docx"
    _write_docx(answer_path, ["1．", "答案：A", "解析：说明。"])

    try:
        refresh_review_statuses([answer_path])
    except FileNotFoundError as exc:
        assert "题目文档" in str(exc)
    else:
        raise AssertionError("缺少题目文档时应停止生成审核状态")

    assert not answer_path.with_name("对点练案1_已清洗_审核状态.json").exists()


def test_clean_batch_refreshes_review_status_when_paired_question_exists(tmp_path):
    project_dir = tmp_path / "众美-高三-语文-测试"
    question_path = project_dir / "对点练案" / "对点练案1.docx"
    answer_dir = project_dir / "答案" / "对点练案答案"
    source_path = answer_dir / "对点练案1.docx"

    _write_docx(
        question_path,
        ["1．下列说法正确的一项是（　　）", "A．甲", "B．乙", "C．丙", "D．丁"],
    )
    _write_docx(source_path, ["1．", "答案：A", "解析：根据原文可知。"])

    output_paths = clean_batch(answer_dir)

    assert len(output_paths) == 1
    assert get_review_gate_result(output_paths[0])["allowed"] is True
