from __future__ import annotations

import hashlib
import json
import re
from collections import Counter
from importlib.metadata import PackageNotFoundError, version
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn

from .answer_core import infer_grouped_question_ids
from .document_roles import build_document_role_profile
from .question_core import build_question_units_from_nodes, scan_docx_nodes
from .strategies import choose_strategy
from .subject_overlay import detect_subject_overlay, get_subject_overlay


PROFILE_SCHEMA_VERSION = "1.1"
ACTION_PLAN_SCHEMA_VERSION = "1.0"
GENERATOR_NAME = "mohen-document-preflight-poc"


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as source:
        for chunk in iter(lambda: source.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _package_version(package_name: str) -> str | None:
    try:
        return version(package_name)
    except PackageNotFoundError:
        return None


def _inspect_native_docx(path: Path) -> dict[str, Any]:
    document = Document(path)
    paragraphs = list(document.paragraphs)
    tables = list(document.tables)
    paragraph_count = len(paragraphs)

    table_anchors: list[dict[str, Any]] = []
    paragraph_index = 0
    table_index = 0
    for body_index, child in enumerate(document.element.body.iterchildren(), 1):
        if child.tag == qn("w:p"):
            paragraph_index += 1
            continue
        if child.tag != qn("w:tbl"):
            continue
        table_index += 1
        table = tables[table_index - 1]
        table_anchors.append(
            {
                "table_index": table_index,
                "body_index": body_index,
                "paragraph_before": paragraph_index or None,
                "paragraph_after": paragraph_index + 1 if paragraph_index < paragraph_count else None,
                "row_count": len(table.rows),
                "column_count": max((len(row.cells) for row in table.rows), default=0),
            }
        )

    heading_candidates: list[dict[str, Any]] = []
    media_paragraphs: list[int] = []
    formula_paragraphs: list[int] = []
    for index, paragraph in enumerate(paragraphs, 1):
        text = (paragraph.text or "").strip()
        style_name = str(getattr(paragraph.style, "name", "") or "")
        if text and (
            style_name.lower().startswith("heading")
            or style_name.startswith("标题")
        ):
            heading_candidates.append(
                {"paragraph_index": index, "style": style_name, "text": text[:120]}
            )

        xml = paragraph._p.xml
        paragraph_drawings = xml.count("<w:drawing")
        paragraph_vml = xml.count("<w:pict")
        paragraph_formulas = len(re.findall(r"<m:oMath(?:\s|>)", xml))
        if paragraph_drawings or paragraph_vml:
            media_paragraphs.append(index)
        if paragraph_formulas:
            formula_paragraphs.append(index)

    document_xml = document.element.xml
    drawing_count = document_xml.count("<w:drawing")
    vml_count = document_xml.count("<w:pict")
    formula_count = len(re.findall(r"<m:oMath(?:\s|>)", document_xml))

    return {
        "paragraph_count": paragraph_count,
        "nonempty_paragraph_count": sum(bool((paragraph.text or "").strip()) for paragraph in paragraphs),
        "table_count": len(tables),
        "inline_shape_count": len(document.inline_shapes),
        "drawing_count": drawing_count,
        "vml_count": vml_count,
        "formula_count": formula_count,
        "heading_count": len(heading_candidates),
        "heading_candidates": heading_candidates,
        "table_anchors": table_anchors,
        "media_paragraphs": media_paragraphs,
        "formula_paragraphs": formula_paragraphs,
    }


def _inspect_with_docling(path: Path, enabled: bool) -> dict[str, Any]:
    package_version = _package_version("docling")
    if not enabled:
        return {
            "enabled": False,
            "available": package_version is not None,
            "package_version": package_version,
            "status": "skipped",
        }
    if package_version is None:
        return {
            "enabled": True,
            "available": False,
            "package_version": None,
            "status": "unavailable",
            "error": "未安装 docling；已保留原生扫描结果。",
        }

    try:
        from docling.document_converter import DocumentConverter

        result = DocumentConverter().convert(path)
        exported = result.document.export_to_dict()
        texts = exported.get("texts") or []
        tables = exported.get("tables") or []
        pictures = exported.get("pictures") or []
        groups = exported.get("groups") or []
        label_counts = Counter(
            str(item.get("label") or "unknown")
            for item in [*texts, *tables, *pictures]
            if isinstance(item, dict)
        )
        result_status = getattr(result, "status", "success")
        status_value = getattr(result_status, "value", str(result_status))
        return {
            "enabled": True,
            "available": True,
            "package_version": package_version,
            "status": str(status_value).lower(),
            "document_schema": exported.get("schema_name"),
            "document_schema_version": exported.get("version"),
            "text_count": len(texts),
            "table_count": len(tables),
            "picture_count": len(pictures),
            "group_count": len(groups),
            "label_counts": dict(sorted(label_counts.items())),
            "_text_items": [
                {
                    "order": order,
                    "label": str(item.get("label") or "unknown"),
                    "text": str(item.get("text") or ""),
                }
                for order, item in enumerate(texts, 1)
                if isinstance(item, dict) and item.get("text")
            ],
        }
    except Exception as exc:
        return {
            "enabled": True,
            "available": True,
            "package_version": package_version,
            "status": "failed",
            "error": f"{type(exc).__name__}: {exc}",
        }


def _build_profile_issues(native: dict[str, Any], docling: dict[str, Any]) -> list[dict[str, str]]:
    issues: list[dict[str, str]] = []
    if docling.get("enabled") and docling.get("status") in {"unavailable", "failed"}:
        issues.append(
            {
                "code": "DOCLING_UNAVAILABLE",
                "severity": "warning",
                "message": docling.get("error", "Docling 结构对照不可用。"),
            }
        )
        return issues

    if not docling.get("enabled") or docling.get("status") == "skipped":
        return issues

    if native["table_count"] != docling.get("table_count"):
        issues.append(
            {
                "code": "TABLE_COUNT_MISMATCH",
                "severity": "warning",
                "message": (
                    f"原生扫描发现 {native['table_count']} 个表格，"
                    f"Docling 发现 {docling.get('table_count')} 个。"
                ),
            }
        )

    native_media_count = native["drawing_count"] + native["vml_count"]
    if native_media_count and native_media_count != docling.get("picture_count"):
        issues.append(
            {
                "code": "MEDIA_COUNT_MISMATCH",
                "severity": "warning",
                "message": (
                    f"原生 OOXML 发现 {native_media_count} 个媒体节点，"
                    f"Docling 发现 {docling.get('picture_count')} 个图片元素。"
                ),
            }
        )
    return issues


def _resolve_subject(doc_name: str, nodes: list) -> tuple[str, str | None]:
    sample_text = " ".join(node.text for node in nodes[:20])
    overlay_name = detect_subject_overlay(doc_name, sample_text, base_subject="文科")
    if overlay_name:
        overlay = get_subject_overlay(overlay_name)
        return (overlay.base_subject if overlay else "文科"), overlay_name
    return choose_strategy(doc_name, sample_text).name, None


def _tables_for_paragraph_span(
    table_anchors: list[dict[str, Any]],
    paragraph_start: int,
    paragraph_end: int,
) -> list[int]:
    selected: list[int] = []
    for anchor in table_anchors:
        paragraph_after = anchor.get("paragraph_after")
        paragraph_before = anchor.get("paragraph_before")
        if paragraph_after is not None and paragraph_start <= paragraph_after <= paragraph_end:
            selected.append(anchor["table_index"])
        elif paragraph_after is None and paragraph_before is not None:
            if paragraph_start <= paragraph_before <= paragraph_end:
                selected.append(anchor["table_index"])
    return selected


def _compile_actions(path: Path, nodes: list, native: dict[str, Any]) -> tuple[list[dict[str, Any]], list[dict[str, str]]]:
    subject_name, overlay_name = _resolve_subject(path.name, nodes)
    units = build_question_units_from_nodes(
        path.name,
        subject_name,
        nodes,
        overlay_name=overlay_name,
    )
    nodes_by_index = {node.index: node for node in nodes}
    actions: list[dict[str, Any]] = []
    for sequence, unit in enumerate(units, 1):
        selected_nodes = [
            nodes_by_index[index]
            for index in range(unit.source_span[0], unit.source_span[1] + 1)
            if index in nodes_by_index
        ]
        paragraph_indexes = [
            int(node.metadata.get("source_paragraph_index", node.index))
            for node in selected_nodes
        ]
        if not paragraph_indexes:
            paragraph_indexes = [unit.source_span[0], unit.source_span[1]]
        paragraph_start = min(paragraph_indexes)
        paragraph_end = max(paragraph_indexes)
        media_paragraphs = [
            index
            for index in native["media_paragraphs"]
            if paragraph_start <= index <= paragraph_end
        ]
        formula_paragraphs = [
            index
            for index in native["formula_paragraphs"]
            if paragraph_start <= index <= paragraph_end
        ]
        actions.append(
            {
                "sequence": sequence,
                "action": "question_input_preview",
                "key": "F1",
                "question_ids": infer_grouped_question_ids(unit),
                "question_type": unit.question_type,
                "node_type": unit.node_type,
                "subject": unit.subject,
                "subject_overlay": unit.subject_overlay,
                "preview": unit.preview,
                "confidence": unit.confidence,
                "warnings": list(unit.warnings),
                "source_ref": {
                    "virtual_node_start": unit.source_span[0],
                    "virtual_node_end": unit.source_span[1],
                    "paragraph_start": paragraph_start,
                    "paragraph_end": paragraph_end,
                    "table_indexes": _tables_for_paragraph_span(
                        native["table_anchors"], paragraph_start, paragraph_end
                    ),
                    "media_paragraphs": media_paragraphs,
                    "formula_paragraphs": formula_paragraphs,
                },
            }
        )

    blockers = [
        {
            "code": "POC_WPS_EXECUTION_DISABLED",
            "severity": "blocker",
            "message": "PoC 尚未绑定 WPS Range，仅生成预检动作，不允许执行按键。",
        }
    ]
    if not actions:
        blockers.append(
            {
                "code": "NO_QUESTION_ACTIONS",
                "severity": "blocker",
                "message": "没有编译出可预检的题目录入动作。",
            }
        )
    return actions, blockers


def build_preflight_bundle(
    docx_path: str | Path,
    *,
    include_docling: bool = True,
) -> dict[str, Any]:
    path = Path(docx_path).resolve()
    if not path.is_file():
        raise FileNotFoundError(path)
    if path.suffix.lower() != ".docx":
        raise ValueError(f"只支持 DOCX 文件: {path}")

    hash_before = _sha256(path)
    native = _inspect_native_docx(path)
    nodes = scan_docx_nodes(path)
    docling = _inspect_with_docling(path, include_docling)
    docling_text_items = docling.pop("_text_items", [])
    roles = build_document_role_profile(path, docling_text_items)
    actions, blockers = _compile_actions(path, nodes, native)
    hash_after = _sha256(path)
    if hash_after != hash_before:
        raise RuntimeError(f"原始 DOCX 在只读画像期间发生变化，已拒绝生成计划: {path}")

    profile_issues = _build_profile_issues(native, docling)
    profile = {
        "schema_version": PROFILE_SCHEMA_VERSION,
        "generator": GENERATOR_NAME,
        "source": {
            "name": path.name,
            "path": path.as_posix(),
            "size_bytes": path.stat().st_size,
            "sha256": hash_before,
            "readonly_input": True,
            "hash_verified_unchanged": True,
        },
        "native": native,
        "docling": docling,
        "roles": roles,
        "comparison": {
            "paragraphs_vs_docling_texts": [
                native["paragraph_count"],
                docling.get("text_count"),
            ],
            "native_vs_docling_tables": [
                native["table_count"],
                docling.get("table_count"),
            ],
            "native_media_vs_docling_pictures": [
                native["drawing_count"] + native["vml_count"],
                docling.get("picture_count"),
            ],
        },
        "fingerprint": {
            "paragraph_count": native["paragraph_count"],
            "nonempty_paragraph_count": native["nonempty_paragraph_count"],
            "table_count": native["table_count"],
            "media_node_count": native["drawing_count"] + native["vml_count"],
            "formula_count": native["formula_count"],
            "heading_count": native["heading_count"],
            "question_action_count": len(actions),
            "role_counts": roles["role_counts"],
        },
        "issues": profile_issues,
    }
    plan = {
        "schema_version": ACTION_PLAN_SCHEMA_VERSION,
        "generator": GENERATOR_NAME,
        "source_name": path.name,
        "source_sha256": hash_before,
        "profile_schema_version": PROFILE_SCHEMA_VERSION,
        "execution_enabled": False,
        "mode": "preview_only",
        "actions": actions,
        "blocking_issues": blockers,
    }
    return {"profile": profile, "plan": plan}


def _profile_markdown(profile: dict[str, Any]) -> str:
    native = profile["native"]
    docling = profile["docling"]
    roles = profile["roles"]
    issue_lines = (
        [f"- [{issue['severity']}] {issue['code']}：{issue['message']}" for issue in profile["issues"]]
        or ["- 无"]
    )
    role_summary = "、".join(
        f"{role}={count}"
        for role, count in roles["role_counts"].items()
    ) or "无"
    candidate_lines = [
        (
            f"- 第 {item['paragraph_index']} 段｜{item['role']}｜"
            f"置信度 {item['confidence']:.2f}｜{item['text']}｜"
            f"证据：{'；'.join(item['evidence'])}"
        )
        for item in roles["heading_candidates"]
    ] or ["- 无"]
    lines = [
            f"# 文档画像：{profile['source']['name']}",
            "",
            "## 原始文件",
            "",
            f"- SHA256：`{profile['source']['sha256']}`",
            f"- 大小：{profile['source']['size_bytes']} 字节",
            "- 只读校验：通过，生成前后哈希一致",
            "",
            "## 原生 DOCX 结构",
            "",
            f"- 段落：{native['paragraph_count']}（非空 {native['nonempty_paragraph_count']}）",
            f"- 原生表格：{native['table_count']}",
            f"- Drawing/VML：{native['drawing_count']}/{native['vml_count']}",
            f"- 公式：{native['formula_count']}",
            f"- 样式标题：{native['heading_count']}",
            "",
            "## Docling 对照",
            "",
            f"- 状态：{docling['status']}",
            f"- 版本：{docling.get('package_version') or '未安装'}",
            f"- 文本/表格/图片：{docling.get('text_count', '-')}/{docling.get('table_count', '-')}/{docling.get('picture_count', '-')}",
            "",
            "## 段落角色证据",
            "",
            f"- 角色计数：{role_summary}",
            (
                "- Docling 文本对齐："
                f"{roles['docling_alignment']['aligned_count']}/"
                f"{roles['docling_alignment']['source_count']}"
            ),
            "- 自动排除题内标题：否（仅供人工审核）",
            "",
            "### 标题角色候选",
            "",
        ]
    lines.extend(candidate_lines)
    lines.extend(
        [
            "",
            "## 结构指纹",
            "",
            "```json",
            json.dumps(profile["fingerprint"], ensure_ascii=False, indent=2),
            "```",
            "",
            "## 风险",
            "",
            *issue_lines,
            "",
        ]
    )
    return "\n".join(lines)


def _plan_markdown(plan: dict[str, Any]) -> str:
    lines = [
        f"# 动作计划：{plan['source_name']}",
        "",
        "> 仅供预检，不允许执行 WPS。JSON 是权威机器事实，本文件是派生人工视图。",
        "",
        f"- 源文件 SHA256：`{plan['source_sha256']}`",
        f"- 模式：{plan['mode']}",
        f"- 计划动作数：{len(plan['actions'])}",
        f"- 是否允许执行：{'是' if plan['execution_enabled'] else '否'}",
        "",
        "## 阻断问题",
        "",
    ]
    lines.extend(
        f"- [{issue['severity']}] {issue['code']}：{issue['message']}"
        for issue in plan["blocking_issues"]
    )
    lines.extend(["", "## 动作", ""])
    if not plan["actions"]:
        lines.append("- 无")
    for action in plan["actions"]:
        source_ref = action["source_ref"]
        lines.extend(
            [
                f"### {action['sequence']}. F1 题号 {', '.join(action['question_ids'])}",
                "",
                f"- 类型：{action['question_type']} / {action['node_type']}",
                f"- 原始段落：{source_ref['paragraph_start']}—{source_ref['paragraph_end']}",
                f"- 原生表格：{source_ref['table_indexes'] or '无'}",
                f"- 媒体段落：{source_ref['media_paragraphs'] or '无'}",
                f"- 公式段落：{source_ref['formula_paragraphs'] or '无'}",
                f"- 置信度：{action['confidence']:.2f}",
                f"- 预览：{action['preview']}",
                "",
            ]
        )
    return "\n".join(lines) + "\n"


def _safe_stem(path: Path) -> str:
    return re.sub(r"[<>:\"/\\|?*]+", "_", path.stem).strip(" .") or "document"


def write_preflight_artifacts(
    docx_path: str | Path,
    output_dir: str | Path,
    *,
    include_docling: bool = True,
) -> dict[str, Path]:
    path = Path(docx_path)
    target_dir = Path(output_dir)
    target_dir.mkdir(parents=True, exist_ok=True)
    bundle = build_preflight_bundle(path, include_docling=include_docling)
    stem = _safe_stem(path)
    paths = {
        "profile_json": target_dir / f"{stem}_DocumentProfile.json",
        "profile_md": target_dir / f"{stem}_DocumentProfile.md",
        "plan_json": target_dir / f"{stem}_ActionPlan.json",
        "plan_md": target_dir / f"{stem}_ActionPlan.md",
    }
    paths["profile_json"].write_text(
        json.dumps(bundle["profile"], ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    paths["plan_json"].write_text(
        json.dumps(bundle["plan"], ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    paths["profile_md"].write_text(_profile_markdown(bundle["profile"]), encoding="utf-8")
    paths["plan_md"].write_text(_plan_markdown(bundle["plan"]), encoding="utf-8")
    return paths
