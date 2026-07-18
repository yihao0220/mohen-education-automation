from __future__ import annotations

import re
from dataclasses import replace
from pathlib import Path

from docx import Document

from .models import AnswerItem, AnswerUnit
from .strategies import extract_normalized_question_id


Q_START_PATTERN = re.compile(r"^\s*(\d+)(?:[．.]|【答案】|[💡★☆])\s*")
ANSWER_BLOCK_START_PATTERN = re.compile(
    r"^\s*(?:(\d+\s*[．.]\s*[（(]\s*\d+\s*[）)])\s*[．.]?|(\d+)(?:[．.]|【答案】|[💡★☆]))\s*"
)
ANSWER_PREFIX_PATTERN = re.compile(r"^\s*答案[：:]\s*")
ANSWER_QUESTION_PREFIX_PATTERN = re.compile(
    r"^[ \t]*\d+[ \t]*[．.][ \t]*(?:[（(][ \t]*\d+[ \t]*[）)][ \t]*)?"
)
BIG_QUESTION_PATTERN = re.compile(r"^\s*[（(][一二三四五六七八九十]+[）)]\s*$")
SUB_QUESTION_PATTERN = re.compile(r"^\s*[（(](\d+)[）)]\s*")
ANALYSIS_START_PATTERN = re.compile(r"^\s*(?:【(?:解析|详解|分析|解答)】|(?:解析|详解|分析|解答)[：:]?)")
SUB_ANALYSIS_START_PATTERN = re.compile(
    r"^\s*\d+[．.]?\s*(?:【(?:解析|详解|分析|解答)】|(?:解析|详解|分析|解答)[：:]?)"
)
INLINE_SUB_PATTERN = re.compile(r"([\(（]\d+[\)）])")
INTERLEAVED_SUB_ENTRY_PATTERN = re.compile(
    r"^\s*([\(（]\d+[\)）])\s*(答案|解析)[：:]\s*(.*)$"
)
COMPACT_CHOICE_PATTERN = re.compile(r"(\d+)[．.]?\s*([A-D])(?=(?:\s*\d+[．.]?\s*[A-D])|$)")
ANALYSIS_ENTRY_PATTERN = re.compile(
    r"^\s*(?:【(?:解析|详解|分析)】\s*|(?:解析|详解|分析)[：:]?\s*)?(\d+)[．.]\s*(.*)$"
)
HEADING_PATTERN = re.compile(r"^\s*[一二三四五六七八九十]+、")
ANSWER_TITLE_PATTERN = re.compile(r"^\s*.*参考答案\s*$")
SECTION_HEADING_PATTERN = re.compile(
    r"^\s*[一二三四五六七八九十]+、.*(?:选择题|单选题|填空题|综合题|非选择题|材料题|材料分析题|解答题|简答题|判断题|实验题|探究题).*$"
)
MATH_SECTION_HEADING_PATTERN = re.compile(
    r"^\s*[一二三四五六七八九十]+、.*(?:基础知识讲解|基础训练|思维冲浪|能力拓展|能力提升|巩固练习|拓展训练).*[：:]?\s*$"
)
PLAIN_SECTION_HEADING_PATTERN = re.compile(
    r"^\s*(?:基础训练|能力训练|能力拓展|能力提升|思维冲浪|巩固练习|拓展训练|提高训练)[：:]?\s*$"
)
PAPER_SECTION_PATTERN = re.compile(r"^\s*第[ⅠⅡⅢⅣⅤIVX一二三四五六七八九十]+卷\s*$")
QUESTION_RANGE_PATTERN = re.compile(r"(\d+)\s*[~～\-—、至]\s*(\d+)\s*(?:小)?题")
TOP_LEVEL_QUESTION_LINE_PATTERN = re.compile(
    r"^\s*(?:[一二三四五六七八九十]+[、．.]\s*)?(?:[★☆/]\s*)?(?:【\s*[GYC]\s*】\s*|[GYC]\s*)?\d+\s*[．.、]"
)
PHYSICS_WHOLE_ANSWER_MARKERS = (
    "未来-高二-物理",
    "众美-高三-物理",
)
ZHONGMEI_CHINESE_ANSWER_DEFINED_SUBQUESTION_MARKERS = (
    "众美-高三-语文",
    "对点练案",
)
ZHONGMEI_CHINESE_PROJECT_MARKER = "众美-高三-语文"
ZHONGMEI_CLASSICAL_CHINESE_DIR_NAMES = frozenset({"文言文", "文言文答案"})
ANSWER_INPUT_BOUNDARY_HEADINGS = frozenset(
    {
        "一、专项训练",
        "二、综合训练",
        "(一)请比对下列选项与原文，说明选项错在何处",
        "(二)图文解读题",
        "(一)分析评价信息",
        "(二)逻辑推断",
        "(三)分析论证特点",
        "(二)梳理论述思路",
        "(一)分析理据关系",
        "阅读下面的文字，完成文后题目",
        "(一)阅读下面的文字，完成文后题目",
        "(二)阅读下面的文字，完成文后题目",
        "(三)阅读下面的文字，完成文后题目",
    }
)
ZHONGMEI_ANSWER_BOUNDARY_PATTERNS = (
    re.compile(r"^\s*[一二三四五六七八九十]+、\S+"),
    re.compile(r"^\s*[（(][一二三四五六七八九十]+[）)]\S+"),
    re.compile(r"^\s*阅读下面.*完成(?:文后|后面)?题目[。.]*\s*$"),
)
ZHONGMEI_CLASSICAL_TITLE_BOUNDARY_PATTERN = re.compile(r"^\s*《[^《》]+》\s*$")


def _clean_text(raw_text: str) -> str:
    return raw_text.replace("\r", "").replace("\n", "").replace("\x07", "").replace("\u00a0", " ").strip()


def _is_answer_input_boundary_heading(
    text: str,
    *,
    use_zhongmei_heading_boundaries: bool = False,
    use_zhongmei_classical_title_boundaries: bool = False,
) -> bool:
    normalized = re.sub(r"\s+", "", text or "")
    normalized = normalized.replace("（", "(").replace("）", ")").rstrip("。.")
    if normalized in ANSWER_INPUT_BOUNDARY_HEADINGS:
        return True
    return bool(
        (
            use_zhongmei_heading_boundaries
            and any(pattern.match(text or "") for pattern in ZHONGMEI_ANSWER_BOUNDARY_PATTERNS)
        )
        or (
            use_zhongmei_classical_title_boundaries
            and ZHONGMEI_CLASSICAL_TITLE_BOUNDARY_PATTERN.match(text or "")
        )
    )


def _strip_trailing_answer_boundary_headings(
    text: str,
    boundary_headings: tuple[str, ...],
) -> str:
    stripped = (text or "").rstrip()
    changed = True
    while changed and stripped:
        changed = False
        for heading in boundary_headings:
            if stripped.endswith(heading):
                stripped = stripped[: -len(heading)].rstrip()
                changed = True
                break
    return stripped


def _extract_compact_choice_pairs(text: str) -> list[tuple[str, str]]:
    normalized = (text or "").replace(" ", "").replace("\t", "")
    return [(qid, answer) for qid, answer in COMPACT_CHOICE_PATTERN.findall(normalized)]


def _extract_heading_question_number(text: str) -> int | None:
    match = Q_START_PATTERN.match(text or "")
    if not match:
        return None
    return int(match.group(1))


def _normalize_answer_question_id(raw_qid: str) -> str:
    qid = re.sub(r"\s+", "", str(raw_qid or ""))
    qid = qid.replace(".", "．").replace("(", "（").replace(")", "）")
    return qid


def _match_answer_block_start(text: str):
    return ANSWER_BLOCK_START_PATTERN.match(text or "")


def _answer_start_question_id(match) -> str:
    return _normalize_answer_question_id(match.group(1) or match.group(2) or "")


def _looks_like_answer_block_start(text: str) -> bool:
    return bool(
        _match_answer_block_start(text)
        or BIG_QUESTION_PATTERN.match(text)
        or HEADING_PATTERN.match(text)
    )


def _parse_compact_analysis_entry(text: str, combo_ids: set[str], allow_marker_only: bool) -> tuple[str, str] | None:
    match = ANALYSIS_ENTRY_PATTERN.match(text)
    if match and match.group(1) in combo_ids:
        return match.group(1), match.group(2).strip()

    if allow_marker_only and ANALYSIS_START_PATTERN.match(text):
        return next(iter(combo_ids)), ANALYSIS_START_PATTERN.sub("", text, count=1).strip()

    return None


def _expand_compact_choice_block(lines: list[str], start_index: int) -> tuple[list[str], int]:
    combos = _extract_compact_choice_pairs(lines[start_index])
    combo_ids = [qid for qid, _ in combos]
    combo_id_set = set(combo_ids)
    analysis_by_qid: dict[str, str] = {}
    current_qid: str | None = None
    current_chunks: list[str] = []
    cursor = start_index + 1

    while cursor < len(lines):
        text = lines[cursor]
        if not text:
            cursor += 1
            continue

        parsed = _parse_compact_analysis_entry(
            text,
            combo_id_set,
            allow_marker_only=len(combo_ids) == 1,
        )
        if parsed:
            if current_qid is not None:
                analysis_by_qid[current_qid] = " ".join(chunk for chunk in current_chunks if chunk).strip()
            current_qid, first_chunk = parsed
            current_chunks = [first_chunk] if first_chunk else []
            cursor += 1
            continue

        if current_qid is not None and not _looks_like_answer_block_start(text):
            current_chunks.append(text)
            cursor += 1
            continue

        break

    if current_qid is not None:
        analysis_by_qid[current_qid] = " ".join(chunk for chunk in current_chunks if chunk).strip()

    normalized_lines: list[str] = []
    for qid, answer in combos:
        normalized_lines.append(f"{qid}．{answer}")
        analysis_text = analysis_by_qid.get(qid)
        if analysis_text:
            normalized_lines.append(f"解析：{analysis_text}")

    return normalized_lines, cursor


def _normalize_paragraph_texts(
    paragraph_texts: list[str],
    *,
    preserve_source_positions: bool = False,
) -> list[str]:
    cleaned_lines = [_clean_text(raw_text) for raw_text in paragraph_texts]
    normalized_lines: list[str] = []
    index = 0

    while index < len(cleaned_lines):
        text = cleaned_lines[index]
        if not text:
            if preserve_source_positions:
                normalized_lines.append("")
            index += 1
            continue

        compact_pairs = _extract_compact_choice_pairs(text)
        if len(compact_pairs) > 1:
            if preserve_source_positions:
                normalized_lines.append(text)
                index += 1
                continue
            current_heading_number = _extract_heading_question_number(text)
            compact_numbers = [int(qid) for qid, _ in compact_pairs]
            if (
                current_heading_number is not None
                and compact_numbers
                and compact_numbers[0] == 1
                and current_heading_number > compact_numbers[-1]
            ):
                normalized_lines.append(text)
                index += 1
                continue

            previous_nonempty = ""
            probe = index - 1
            while probe >= 0:
                if cleaned_lines[probe]:
                    previous_nonempty = cleaned_lines[probe]
                    break
                probe -= 1

            previous_heading_number = _extract_heading_question_number(previous_nonempty)
            if (
                previous_heading_number is not None
                and compact_numbers
                and compact_numbers[0] == 1
                and previous_heading_number > compact_numbers[-1]
            ):
                normalized_lines.append(text)
                index += 1
                continue

            expanded_lines, next_index = _expand_compact_choice_block(cleaned_lines, index)
            normalized_lines.extend(expanded_lines)
            index = next_index
            continue

        if _is_ignorable_answer_heading(text):
            if preserve_source_positions:
                normalized_lines.append(text)
            index += 1
            continue

        normalized_lines.append(text)
        index += 1

    return normalized_lines


def _is_ignorable_answer_heading(text: str) -> bool:
    return bool(
        ANSWER_TITLE_PATTERN.match(text)
        or SECTION_HEADING_PATTERN.match(text)
        or MATH_SECTION_HEADING_PATTERN.match(text)
        or PLAIN_SECTION_HEADING_PATTERN.match(text)
        or PAPER_SECTION_PATTERN.match(text)
    )


def _extract_table_answer_lines(doc: Document) -> list[str]:
    def cell_text(cell) -> str:
        parts = [_clean_text(paragraph.text) for paragraph in cell.paragraphs]
        return " ".join(part for part in parts if part).strip()

    def strip_label(cells: list[str], labels: set[str]) -> list[str]:
        if cells and cells[0] in labels:
            return cells[1:]
        return cells

    extracted_lines: list[str] = []
    for table in doc.tables:
        rows = [[cell_text(cell) for cell in row.cells] for row in table.rows]
        row_index = 0
        while row_index + 1 < len(rows):
            question_row = strip_label(rows[row_index], {"题号", "题目"})
            answer_row = strip_label(rows[row_index + 1], {"答案", "参考答案"})

            question_ids = [cell for cell in question_row if re.fullmatch(r"\d+", cell)]
            answers = [cell for cell in answer_row if cell]

            if question_ids and len(question_ids) == len(answers):
                extracted_lines.extend(
                    f"{question_id}．{answer}"
                    for question_id, answer in zip(question_ids, answers)
                )
                row_index += 2
                continue

            row_index += 1

    return extracted_lines


def _docx_uses_whole_answer_input_for_subquestions(docx_path: str | Path) -> bool:
    path_text = str(docx_path)
    return any(marker in path_text for marker in PHYSICS_WHOLE_ANSWER_MARKERS)


def _docx_uses_zhongmei_classical_chinese_context(docx_path: str | Path) -> bool:
    path_text = str(docx_path)
    path_parts = {
        part.strip()
        for part in re.split(r"[\\/]", path_text)
        if part.strip()
    }
    return bool(
        ZHONGMEI_CHINESE_PROJECT_MARKER in path_text
        and path_parts.intersection(ZHONGMEI_CLASSICAL_CHINESE_DIR_NAMES)
    )


def _docx_allows_answer_defined_subquestions(docx_path: str | Path) -> bool:
    path_text = str(docx_path)
    return bool(
        all(
            marker in path_text
            for marker in ZHONGMEI_CHINESE_ANSWER_DEFINED_SUBQUESTION_MARKERS
        )
        or _docx_uses_zhongmei_classical_chinese_context(docx_path)
    )


def _answer_items_all_empty(unit: AnswerUnit) -> bool:
    return bool(unit.answer_items) and not any(item.text and item.text.strip() for item in unit.answer_items)


def _answer_items_are_subquestion_placeholders(unit: AnswerUnit) -> bool:
    texts = [item.text.strip() for item in unit.answer_items if item.text and item.text.strip()]
    return bool(texts) and all(
        re.fullmatch(r"(?:[（(]\s*\d+\s*[）)]\s*)+", text)
        for text in texts
    )


def _analysis_contains_subquestion_markers(unit: AnswerUnit) -> bool:
    analysis_text = " ".join(
        item.text.strip() for item in unit.analysis_items if item.text and item.text.strip()
    ).strip()
    if not analysis_text:
        return False
    return bool(re.search(r"[（(]1[）)]", analysis_text))


def _is_single_classical_translation_subquestion(unit: AnswerUnit) -> bool:
    if unit.question_id != "2" or len(unit.answer_items) != 1 or len(unit.analysis_items) != 1:
        return False
    answer_marker = unit.answer_items[0].item_id or ""
    analysis_text = unit.analysis_items[0].text or ""
    return bool(
        re.fullmatch(r"[（(]1[）)]", answer_marker)
        and re.match(r"^\s*[（(]1[）)]\s*得分点\s*[：:]", analysis_text)
    )


def _apply_docx_context_to_units(units: list[AnswerUnit], docx_path: str | Path) -> list[AnswerUnit]:
    force_whole_answer_input = _docx_uses_whole_answer_input_for_subquestions(docx_path)
    allow_answer_defined_subquestions = _docx_allows_answer_defined_subquestions(docx_path)
    ordered_occurrence_mapping = _docx_uses_zhongmei_classical_chinese_context(docx_path)
    source_path = str(Path(docx_path))

    for unit in units:
        if ordered_occurrence_mapping and _is_single_classical_translation_subquestion(unit):
            unit.answer_mode = "subquestion"
        unit_force_whole_answer_input = force_whole_answer_input or (
            ordered_occurrence_mapping and unit.answer_mode == "whole"
        )
        unit.metadata["source_docx_path"] = source_path
        unit.metadata["force_whole_answer_input"] = unit_force_whole_answer_input
        unit.metadata["allow_answer_defined_subquestions"] = allow_answer_defined_subquestions
        unit.metadata["ordered_occurrence_mapping"] = ordered_occurrence_mapping

        if (
            allow_answer_defined_subquestions
            and not ordered_occurrence_mapping
            and len(unit.analysis_items) == 1
            and re.match(r"^\s*[（(]1[）)]", unit.analysis_items[0].text or "")
        ):
            sub_analysis_items = _extract_sub_answer_items(unit.analysis_items[0].text)
            if len(sub_analysis_items) > 1:
                unit.analysis_items = sub_analysis_items

        if not unit_force_whole_answer_input:
            continue

        if _answer_items_all_empty(unit) or _answer_items_are_subquestion_placeholders(unit):
            if "empty_sub_answers" in unit.review_flags:
                unit.review_flags = [flag for flag in unit.review_flags if flag != "empty_sub_answers"]
                unit.confidence = min(1.0, unit.confidence + 0.3)
            if _analysis_contains_subquestion_markers(unit):
                unit.metadata["analysis_only_subanswers"] = True

    return units


def _extract_sub_answer_items(text: str) -> list[AnswerItem]:
    line_marker_pattern = re.compile(
        r"(?m)^[ \t]*([\(（]\d+[\)）])"
    )
    line_matches = list(line_marker_pattern.finditer(text or ""))
    line_numbers = [
        _sub_marker_number(match.group(1))
        for match in line_matches
    ]
    if (
        len(line_matches) >= 2
        and line_numbers == list(range(1, len(line_matches) + 1))
    ):
        items: list[AnswerItem] = []
        for index, match in enumerate(line_matches):
            start = match.end()
            end = line_matches[index + 1].start() if index + 1 < len(line_matches) else len(text)
            item_text = text[start:end].strip()
            item_text = ANSWER_PREFIX_PATTERN.sub("", item_text, count=1).strip()
            items.append(AnswerItem(item_id=match.group(1), text=item_text))
        return items

    stripped_text = text.lstrip()
    first = INLINE_SUB_PATTERN.search(stripped_text)
    if first and first.start() == 0:
        marker = first.group(1)
        if re.match(r"^[\(（]1[\)）]$", marker):
            pattern = re.compile(r"([\(（]\d+[\)）])")
        else:
            pattern = None

        if pattern is not None:
            matches = list(pattern.finditer(text))
            items: list[AnswerItem] = []
            for idx, match in enumerate(matches):
                start = match.end()
                end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
                item_text = text[start:end].strip()
                item_text = ANSWER_PREFIX_PATTERN.sub("", item_text, count=1).strip()
                items.append(AnswerItem(item_id=match.group(1), text=item_text))
            return items

        compact_pairs = _extract_compact_choice_pairs(text)
        if len(compact_pairs) >= 2:
            return [
                AnswerItem(item_id=f"({qid})", text=answer_text.strip())
                for qid, answer_text in compact_pairs
            ]

    compact_pairs = _extract_compact_choice_pairs(text)
    if len(compact_pairs) >= 2:
        return [
            AnswerItem(item_id=f"({qid})", text=answer_text.strip())
            for qid, answer_text in compact_pairs
        ]

    return _extract_sequential_line_sub_answer_items(text)


def _sub_marker_number(marker: str) -> int | None:
    bracket_match = re.fullmatch(r"[（(](\d+)[）)]", marker or "")
    if bracket_match:
        return int(bracket_match.group(1))
    return None


def _extract_interleaved_sub_answer_analysis_items(
    text: str,
) -> tuple[list[AnswerItem], list[AnswerItem]]:
    entries: list[tuple[str, str, str]] = []
    current_marker: str | None = None
    current_kind: str | None = None
    current_chunks: list[str] = []

    def flush_current() -> None:
        if current_marker is None or current_kind is None:
            return
        entries.append(
            (
                current_marker,
                current_kind,
                "\n".join(chunk for chunk in current_chunks if chunk).strip(),
            )
        )

    for raw_line in (text or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        match = INTERLEAVED_SUB_ENTRY_PATTERN.match(line)
        if match:
            flush_current()
            current_marker = match.group(1)
            current_kind = match.group(2)
            current_chunks = [match.group(3).strip()] if match.group(3).strip() else []
            continue
        if current_marker is None:
            return [], []
        current_chunks.append(line)

    flush_current()
    answer_entries = [entry for entry in entries if entry[1] == "答案"]
    analysis_entries = [entry for entry in entries if entry[1] == "解析"]
    if not answer_entries or not analysis_entries:
        return [], []

    answer_numbers = [_sub_marker_number(marker) for marker, _, _ in answer_entries]
    if answer_numbers != list(range(1, len(answer_entries) + 1)):
        return [], []

    answer_items = [
        AnswerItem(item_id=marker, text=entry_text)
        for marker, _, entry_text in answer_entries
    ]
    analysis_items = [
        AnswerItem(item_id=marker, text=entry_text)
        for marker, _, entry_text in analysis_entries
    ]
    return answer_items, analysis_items


def _extract_sequential_line_sub_answer_items(text: str) -> list[AnswerItem]:
    lines = [line.strip() for line in (text or "").splitlines() if line.strip()]
    markers: list[int] = []
    parts: list[tuple[int, str]] = []
    for line in lines:
        match = re.match(r"^\s*(\d+)[．.、]?\s*(.*)$", line)
        if not match:
            return []
        number = int(match.group(1))
        markers.append(number)
        parts.append((number, match.group(2).strip()))

    if len(parts) < 2:
        return []

    expected = list(range(1, len(parts) + 1))
    if markers != expected:
        return []

    return [
        AnswerItem(item_id=f"({number})", text=content)
        for number, content in parts
    ]


def _looks_like_calculation_process_answer(text: str) -> bool:
    lines = [line.strip() for line in (text or "").splitlines() if line.strip()]
    if len(lines) < 2:
        return False
    math_line_count = sum(1 for line in lines if re.search(r"[=＝×÷+]|^\s*原式", line))
    return math_line_count >= 2


def build_answer_units_from_paragraph_texts(
    paragraph_texts: list[str],
    *,
    use_zhongmei_heading_boundaries: bool = False,
    use_zhongmei_classical_title_boundaries: bool = False,
    preserve_source_positions: bool = False,
) -> list[AnswerUnit]:
    boundary_headings: tuple[str, ...] = ()
    if use_zhongmei_heading_boundaries:
        boundary_headings = tuple(
            sorted(
                {
                    _clean_text(raw_text)
                    for raw_text in paragraph_texts
                    if any(
                        pattern.match(_clean_text(raw_text))
                        for pattern in ZHONGMEI_ANSWER_BOUNDARY_PATTERNS
                    )
                    or (
                        use_zhongmei_classical_title_boundaries
                        and ZHONGMEI_CLASSICAL_TITLE_BOUNDARY_PATTERN.match(
                            _clean_text(raw_text)
                        )
                    )
                },
                key=len,
                reverse=True,
            )
        )
    paragraph_texts = _normalize_paragraph_texts(
        paragraph_texts,
        preserve_source_positions=preserve_source_positions,
    )
    units: list[AnswerUnit] = []
    current_q = None
    current_big_blocks: list[dict] = []
    in_big_question = False
    big_question_qnum = None
    sub_q_index = 0

    for i, raw_text in enumerate(paragraph_texts, 1):
        text = _clean_text(raw_text)
        if not text:
            continue

        if _is_ignorable_answer_heading(text) or _is_answer_input_boundary_heading(
            text,
            use_zhongmei_heading_boundaries=use_zhongmei_heading_boundaries,
            use_zhongmei_classical_title_boundaries=use_zhongmei_classical_title_boundaries,
        ):
            if current_q is not None:
                current_q["end_p"] = i - 1
                units.append(_block_to_answer_unit(current_q, paragraph_texts))
                current_q = None
            if current_big_blocks:
                for block in current_big_blocks:
                    block["end_p"] = i - 1
                    units.append(_block_to_answer_unit(block, paragraph_texts))
                current_big_blocks = []
            in_big_question = False
            big_question_qnum = None
            sub_q_index = 0
            continue

        if BIG_QUESTION_PATTERN.match(text):
            if current_q is not None:
                current_q["end_p"] = i - 1
                units.append(_block_to_answer_unit(current_q, paragraph_texts))
                current_q = None
            if current_big_blocks:
                for block in current_big_blocks:
                    units.append(_block_to_answer_unit(block, paragraph_texts))
                current_big_blocks = []
            in_big_question = True
            big_question_qnum = text.strip()[1:-1]
            sub_q_index = 0
            continue

        match_sub = SUB_QUESTION_PATTERN.match(text)
        if match_sub and in_big_question:
            sub_q_index += 1
            current_big_blocks.append(
                {
                    "qnum": f"{big_question_qnum}.{sub_q_index}",
                    "ans_start_p": i,
                    "ana_start_p": None,
                    "end_p": None,
                    "is_sub_question": True,
                    "boundary_headings": boundary_headings,
                }
            )
            continue

        match_q = _match_answer_block_start(text)
        if match_q:
            current_q_number = _extract_primary_question_number(current_q["qnum"]) if current_q else None
            matched_qid = _answer_start_question_id(match_q)
            matched_q_number = _extract_primary_question_number(matched_qid) or 0
            if (
                current_q
                and current_q["ana_start_p"] is not None
                and current_q_number is not None
                and matched_q_number <= current_q_number
                and not ("（" in matched_qid and matched_qid != current_q["qnum"])
            ):
                continue
            if current_q_number is not None and current_q_number >= 10 and matched_q_number < current_q_number:
                if current_q and current_q["ana_start_p"] is None and SUB_ANALYSIS_START_PATTERN.match(text):
                    current_q["ana_start_p"] = i
                continue
            if current_big_blocks:
                for block in current_big_blocks:
                    units.append(_block_to_answer_unit(block, paragraph_texts))
                current_big_blocks = []
            if current_q is not None:
                current_q["end_p"] = i - 1
                units.append(_block_to_answer_unit(current_q, paragraph_texts))
            in_big_question = False
            big_question_qnum = None
            ans_start_p = i
            q_content = text[len(match_q.group(0)):]
            if not q_content.strip() and i < len(paragraph_texts):
                next_text = paragraph_texts[i].strip()
                if ANSWER_PREFIX_PATTERN.match(next_text):
                    ans_start_p = i + 1
            current_q = {
                "qnum": matched_qid,
                "ans_start_p": ans_start_p,
                "ana_start_p": None,
                "end_p": None,
                "boundary_headings": boundary_headings,
            }
            continue

        if ANALYSIS_START_PATTERN.match(text) or SUB_ANALYSIS_START_PATTERN.match(text):
            if current_big_blocks:
                for block in current_big_blocks:
                    if block["ana_start_p"] is None:
                        block["ana_start_p"] = i
            elif current_q and current_q["ana_start_p"] is None:
                current_q["ana_start_p"] = i

    if current_big_blocks:
        for block in current_big_blocks:
            block["end_p"] = len(paragraph_texts)
            units.append(_block_to_answer_unit(block, paragraph_texts))
    if current_q is not None:
        current_q["end_p"] = len(paragraph_texts)
        units.append(_block_to_answer_unit(current_q, paragraph_texts))

    units.sort(key=lambda unit: unit.source_span[0])
    return units


def _block_to_answer_unit(block: dict, paragraph_texts: list[str]) -> AnswerUnit:
    qnum = block["qnum"]
    ans_start_p = block["ans_start_p"]
    ana_start_p = block["ana_start_p"]
    end_p = block["end_p"] or len(paragraph_texts)
    ans_end_p = (ana_start_p - 1) if ana_start_p else end_p

    answer_text = "\n".join(paragraph_texts[max(ans_start_p - 1, 0):ans_end_p]).strip()
    answer_text_wo_prefix = ANSWER_PREFIX_PATTERN.sub("", answer_text).strip()
    answer_text_wo_prefix = _strip_trailing_answer_boundary_headings(
        answer_text_wo_prefix,
        block.get("boundary_headings", ()),
    )
    # 只允许题号与小问号位于同一行时一起剥离。若文档是“1．\n(1)答案：…”，
    # 必须保留下一段的 (1)，否则整组会被误判为普通答案。
    answer_text_wo_prefix = ANSWER_QUESTION_PREFIX_PATTERN.sub(
        "", answer_text_wo_prefix, count=1
    ).strip()
    analysis_items: list[AnswerItem] = []
    interleaved_analysis = False
    answer_items, interleaved_analysis_items = _extract_interleaved_sub_answer_analysis_items(
        answer_text_wo_prefix
    )
    if interleaved_analysis_items:
        analysis_items = interleaved_analysis_items
        interleaved_analysis = True
    else:
        answer_items = _extract_sub_answer_items(answer_text_wo_prefix)
        if not answer_items:
            answer_items = [AnswerItem(item_id=qnum, text=answer_text_wo_prefix)]

    if ana_start_p and not interleaved_analysis:
        analysis_text = "\n".join(paragraph_texts[max(ana_start_p - 1, 0):end_p]).strip()
        analysis_text = ANALYSIS_START_PATTERN.sub("", analysis_text, count=1).strip()
        analysis_text = _strip_trailing_answer_boundary_headings(
            analysis_text,
            block.get("boundary_headings", ()),
        )
        analysis_items = [AnswerItem(item_id=qnum, text=analysis_text)]

    review_flags: list[str] = []
    confidence = 1.0
    if len(answer_items) > 1 and not any(item.text for item in answer_items):
        review_flags.append("empty_sub_answers")
        confidence -= 0.3
    if not analysis_items and not _looks_like_calculation_process_answer(answer_text_wo_prefix):
        review_flags.append("missing_analysis")
        confidence -= 0.05

    return AnswerUnit(
        question_id=qnum,
        answer_mode="subquestion" if len(answer_items) > 1 else "whole",
        answer_items=answer_items,
        analysis_items=analysis_items,
        confidence=max(0.1, confidence),
        review_flags=review_flags,
        source_span=(ans_start_p, end_p),
        answer_span=(ans_start_p, ans_end_p),
        analysis_span=(ana_start_p or 0, end_p if ana_start_p else 0),
        metadata={
            "ans_start_p": ans_start_p,
            "ans_end_p": ans_end_p,
            "ana_start_p": ana_start_p,
            "end_p": end_p,
            "is_sub_question": block.get("is_sub_question", False),
            "interleaved_subquestion_analysis": interleaved_analysis,
        },
    )


def build_answer_units_from_wps(doc) -> list[AnswerUnit]:
    texts = []
    paras = doc.Paragraphs
    for i in range(1, paras.Count + 1):
        try:
            if paras(i).Range.Information(12):
                texts.append("")
            else:
                texts.append(paras(i).Range.Text)
        except Exception:
            texts.append("")
    context = " ".join(
        str(getattr(doc, attr, "") or "") for attr in ("FullName", "Name")
    )
    units = build_answer_units_from_paragraph_texts(
        texts,
        use_zhongmei_heading_boundaries=_docx_allows_answer_defined_subquestions(context),
        use_zhongmei_classical_title_boundaries=_docx_uses_zhongmei_classical_chinese_context(
            context
        ),
        preserve_source_positions=True,
    )
    if _docx_uses_zhongmei_classical_chinese_context(context):
        source_path = str(getattr(doc, "FullName", "") or getattr(doc, "Name", "") or context)
        return _apply_docx_context_to_units(units, source_path)
    return units


def build_answer_units_from_docx(
    docx_path: str | Path,
    *,
    preserve_source_positions: bool = False,
) -> list[AnswerUnit]:
    try:
        doc = Document(docx_path)
    except Exception as exc:
        raise ValueError(f"无效 docx 文件: {Path(docx_path).name}") from exc

    table_units = build_answer_units_from_paragraph_texts(_extract_table_answer_lines(doc))
    use_classical_chinese_context = _docx_uses_zhongmei_classical_chinese_context(docx_path)
    paragraph_units = build_answer_units_from_paragraph_texts(
        [para.text for para in doc.paragraphs],
        use_zhongmei_heading_boundaries=_docx_allows_answer_defined_subquestions(docx_path),
        use_zhongmei_classical_title_boundaries=use_classical_chinese_context,
        preserve_source_positions=preserve_source_positions,
    )
    merged_units = _merge_docx_answer_units(
        table_units,
        paragraph_units,
        preserve_paragraph_occurrences=use_classical_chinese_context,
    )
    return _apply_docx_context_to_units(merged_units, docx_path)


def _merge_docx_answer_units(
    table_units: list[AnswerUnit],
    paragraph_units: list[AnswerUnit],
    *,
    preserve_paragraph_occurrences: bool = False,
) -> list[AnswerUnit]:
    if preserve_paragraph_occurrences:
        paragraph_question_ids = {unit.question_id for unit in paragraph_units}
        table_fallbacks = [
            unit
            for unit in table_units
            if unit.question_id not in paragraph_question_ids
        ]
        return [*paragraph_units, *table_fallbacks]

    merged_by_qid: dict[str, AnswerUnit] = {}
    for unit in table_units:
        merged_by_qid[unit.question_id] = unit
    for unit in _filter_non_monotonic_paragraph_units(paragraph_units):
        merged_by_qid[unit.question_id] = unit

    return sorted(
        merged_by_qid.values(),
        key=lambda unit: _question_id_sort_key(unit.question_id),
    )


def _extract_primary_question_number(question_id: str) -> int | None:
    match = re.match(r"\s*(\d+)", str(question_id or ""))
    if not match:
        return None
    return int(match.group(1))


def _filter_non_monotonic_paragraph_units(paragraph_units: list[AnswerUnit]) -> list[AnswerUnit]:
    filtered: list[AnswerUnit] = []
    max_seen = -1
    reset_mode = False

    for unit in paragraph_units:
        current_number = _extract_primary_question_number(unit.question_id)
        if current_number is None:
            filtered.append(unit)
            continue

        if reset_mode:
            if current_number > max_seen:
                reset_mode = False
            else:
                continue

        if current_number < max_seen:
            reset_mode = True
            continue

        filtered.append(unit)
        max_seen = max(max_seen, current_number)

    return filtered


def _question_id_sort_key(question_id: str) -> tuple[int, ...]:
    parts = [int(part) for part in re.findall(r"\d+", question_id)]
    return tuple(parts or [10**9])


def infer_grouped_question_ids(question) -> list[str]:
    for material_block in getattr(question, "material_blocks", []) or []:
        range_match = QUESTION_RANGE_PATTERN.search(material_block or "")
        if not range_match:
            continue
        start = int(range_match.group(1))
        end = int(range_match.group(2))
        if start <= end:
            return [str(number) for number in range(start, end + 1)]

    grouped_ids: list[str] = []
    seen: set[str] = set()

    def add_qid(candidate: str | None) -> None:
        normalized = str(candidate or "").strip()
        if not normalized or normalized in seen:
            return
        seen.add(normalized)
        grouped_ids.append(normalized)

    def add_top_level_numeric_qid(candidate: str | None) -> None:
        normalized = str(candidate or "").strip()
        if re.fullmatch(r"\d+", normalized):
            add_qid(normalized)

    add_qid(getattr(question, "question_id", None))

    for block in (getattr(question, "stem_blocks", []) or []) + (getattr(question, "subquestions", []) or []):
        block_text = (block or "").strip()
        if not TOP_LEVEL_QUESTION_LINE_PATTERN.match(block_text):
            continue
        add_top_level_numeric_qid(extract_normalized_question_id(block_text))

    return grouped_ids


def question_expects_subquestion_answers(question) -> bool:
    return bool(getattr(question, "subquestions", []) or len(infer_grouped_question_ids(question)) > 1)


def _copy_answer_unit_with_question_id(
    answer_unit: AnswerUnit,
    question_id: str,
    *,
    confidence_delta: float = 0.0,
    extra_flags: list[str] | None = None,
    extra_metadata: dict | None = None,
) -> AnswerUnit:
    merged_flags = list(answer_unit.review_flags)
    for flag in extra_flags or []:
        if flag not in merged_flags:
            merged_flags.append(flag)

    merged_metadata = dict(answer_unit.metadata)
    if extra_metadata:
        merged_metadata.update(extra_metadata)

    return replace(
        answer_unit,
        question_id=question_id,
        confidence=max(0.1, answer_unit.confidence + confidence_delta),
        review_flags=merged_flags,
        metadata=merged_metadata,
    )


def _normalize_answer_source(answer_source) -> list[AnswerUnit]:
    if isinstance(answer_source, (str, Path)):
        return build_answer_units_from_docx(answer_source)
    return list(answer_source or [])


def _answer_mapping_question_id(unit: AnswerUnit) -> str:
    if not unit.metadata.get("force_whole_answer_input"):
        return unit.question_id
    compound_match = re.fullmatch(
        r"(\d+)\s*[．.]\s*[（(]\s*\d+\s*[）)]",
        unit.question_id or "",
    )
    return compound_match.group(1) if compound_match else unit.question_id


def _pop_next_unconsumed_answer(sequential_pool: list[AnswerUnit], consumed_ids: set[str]) -> AnswerUnit | None:
    while sequential_pool:
        candidate = sequential_pool.pop(0)
        if candidate.question_id not in consumed_ids:
            return candidate
    return None


def _join_answer_item_text(unit: AnswerUnit) -> str:
    return " ".join(item.text.strip() for item in unit.answer_items if item.text and item.text.strip()).strip()


def _join_analysis_item_text(unit: AnswerUnit) -> str:
    return " ".join(item.text.strip() for item in unit.analysis_items if item.text and item.text.strip()).strip()


def _analysis_can_stand_in_for_subanswers(unit: AnswerUnit) -> bool:
    if unit.metadata.get("analysis_only_subanswers"):
        return True
    if _join_answer_item_text(unit):
        return False
    analysis_text = _join_analysis_item_text(unit)
    if not analysis_text:
        return False
    return bool(re.search(r"[（(]1[）)]", analysis_text))


def _build_grouped_material_answer_unit(
    question,
    grouped_ids: list[str],
    answer_map: dict[str, AnswerUnit],
    sequential_pool: list[AnswerUnit],
    consumed_ids: set[str],
    question_index: int,
) -> AnswerUnit | None:
    pre_grouped = answer_map.get(question.question_id)
    if (
        pre_grouped
        and pre_grouped.question_id not in consumed_ids
        and pre_grouped.answer_mode == "subquestion"
        and len(pre_grouped.answer_items) == len(grouped_ids)
        and len(pre_grouped.analysis_items) == len(grouped_ids)
        and [
            _sub_marker_number(item.item_id)
            for item in pre_grouped.answer_items
        ]
        == list(range(1, len(grouped_ids) + 1))
        and [
            _sub_marker_number(item.item_id)
            for item in pre_grouped.analysis_items
        ]
        == list(range(1, len(grouped_ids) + 1))
    ):
        consumed_ids.add(pre_grouped.question_id)
        return _copy_answer_unit_with_question_id(
            pre_grouped,
            question.question_id,
            extra_metadata={
                "mapping_method": "material_group_premerged",
                "grouped_question_ids": grouped_ids,
                "original_question_ids": grouped_ids,
                "question_index": question_index,
                "is_material_group": True,
            },
        )

    slots: list[tuple[str, AnswerUnit | None, str]] = []

    for grouped_id in grouped_ids:
        exact_match = answer_map.get(grouped_id)
        if exact_match and exact_match.question_id not in consumed_ids:
            consumed_ids.add(exact_match.question_id)
            slots.append((grouped_id, exact_match, "exact"))
            continue

        fallback = _pop_next_unconsumed_answer(sequential_pool, consumed_ids)
        if fallback is not None:
            consumed_ids.add(fallback.question_id)
            slots.append((grouped_id, fallback, "sequential"))
            continue

        slots.append((grouped_id, None, "missing"))

    if not any(unit is not None for _, unit, _ in slots):
        return None

    answer_items: list[AnswerItem] = []
    analysis_items: list[AnswerItem] = []
    review_flags: list[str] = []
    confidence = 1.0
    source_starts: list[int] = []
    source_ends: list[int] = []
    answer_starts: list[int] = []
    answer_ends: list[int] = []
    analysis_starts: list[int] = []
    analysis_ends: list[int] = []
    original_question_ids: list[str] = []

    for sub_index, (grouped_id, unit, mapping_method) in enumerate(slots, 1):
        marker = f"({sub_index})"
        if unit is None:
            answer_items.append(AnswerItem(item_id=marker, text=""))
            if "missing_grouped_answer" not in review_flags:
                review_flags.append("missing_grouped_answer")
            confidence -= 0.35
            continue

        original_question_ids.append(unit.question_id)
        source_starts.append(unit.source_span[0])
        source_ends.append(unit.source_span[1])
        if unit.answer_span != (0, 0):
            answer_starts.append(unit.answer_span[0])
            answer_ends.append(unit.answer_span[1])
        if unit.analysis_span != (0, 0):
            analysis_starts.append(unit.analysis_span[0])
            analysis_ends.append(unit.analysis_span[1])

        answer_text = _join_answer_item_text(unit)
        answer_items.append(AnswerItem(item_id=marker, text=answer_text))

        analysis_text = _join_analysis_item_text(unit)
        if analysis_text:
            analysis_items.append(AnswerItem(item_id=marker, text=analysis_text))

        confidence = min(confidence, unit.confidence)
        for flag in unit.review_flags:
            if flag not in review_flags:
                review_flags.append(flag)

        if mapping_method == "sequential":
            if "sequential_mapping" not in review_flags:
                review_flags.append("sequential_mapping")
            confidence -= 0.2

    source_span = (
        min(source_starts) if source_starts else 0,
        max(source_ends) if source_ends else 0,
    )
    answer_span = (
        min(answer_starts) if answer_starts else 0,
        max(answer_ends) if answer_ends else 0,
    )
    analysis_span = (
        min(analysis_starts) if analysis_starts else 0,
        max(analysis_ends) if analysis_ends else 0,
    )

    metadata = {
        "mapping_method": "material_group",
        "grouped_question_ids": grouped_ids,
        "original_question_ids": original_question_ids,
        "question_index": question_index,
        "is_material_group": True,
    }
    if any(method == "sequential" and unit is not None for _, unit, method in slots):
        metadata["original_question_id"] = next(
            unit.question_id
            for _, unit, method in slots
            if method == "sequential" and unit is not None
        )

    return AnswerUnit(
        question_id=question.question_id,
        answer_mode="subquestion",
        answer_items=answer_items,
        analysis_items=analysis_items,
        confidence=max(0.1, confidence),
        review_flags=review_flags,
        source_span=source_span,
        answer_span=answer_span,
        analysis_span=analysis_span,
        metadata=metadata,
    )


def map_answers(
    question_units,
    answer_source,
) -> list[AnswerUnit]:
    """
    以题目结构为主，对答案块做题号对齐。

    策略：
    1. 众美文言文答案由路径 metadata 启用按出现顺序映射，允许原题号重置
    2. 其他文档优先按 question_id 精确匹配
    3. 剩余未匹配项按顺序兜底映射，并强制降置信、加 review flag
    4. 题目结构与答案结构冲突时继续降置信并标记
    5. 无法消费的答案块保留为孤立块，供审核清单拦截
    """
    question_units = list(question_units)
    raw_answers = _normalize_answer_source(answer_source)
    ordered_occurrence_mapping = bool(raw_answers) and all(
        unit.metadata.get("ordered_occurrence_mapping")
        for unit in raw_answers
    )
    answer_map: dict[str, AnswerUnit] = {}
    if not ordered_occurrence_mapping:
        for unit in raw_answers:
            answer_map[unit.question_id] = unit
        for unit in raw_answers:
            answer_map.setdefault(_answer_mapping_question_id(unit), unit)
    consumed_ids: set[str] = set()
    mapped_units: list[AnswerUnit] = []

    unmatched_answers = (
        []
        if ordered_occurrence_mapping
        else [
            unit
            for unit in raw_answers
            if _answer_mapping_question_id(unit)
            not in {q.question_id for q in question_units}
        ]
    )
    sequential_pool = list(unmatched_answers)

    for index, question in enumerate(question_units, 1):
        grouped_ids = infer_grouped_question_ids(question)
        mapped = None

        if ordered_occurrence_mapping:
            if index <= len(raw_answers):
                occurrence_answer = raw_answers[index - 1]
                mapped = _copy_answer_unit_with_question_id(
                    occurrence_answer,
                    question.question_id,
                    extra_metadata={
                        "mapping_method": "ordered_occurrence",
                        "original_question_id": occurrence_answer.question_id,
                        "question_index": index,
                    },
                )
        elif len(grouped_ids) > 1:
            mapped = _build_grouped_material_answer_unit(
                question,
                grouped_ids,
                answer_map,
                sequential_pool,
                consumed_ids,
                index,
            )
        else:
            exact_match = answer_map.get(question.question_id)
            if exact_match and question.question_id not in consumed_ids:
                original_question_id = exact_match.question_id
                mapped = _copy_answer_unit_with_question_id(
                    exact_match,
                    question.question_id,
                    extra_metadata={
                        "mapping_method": "exact",
                        "original_question_id": original_question_id,
                        "question_index": index,
                    },
                )
                consumed_ids.add(question.question_id)
                consumed_ids.add(original_question_id)
            elif sequential_pool:
                fallback = _pop_next_unconsumed_answer(sequential_pool, consumed_ids)
                if fallback is not None:
                    consumed_ids.add(fallback.question_id)
                    mapped = _copy_answer_unit_with_question_id(
                        fallback,
                        question.question_id,
                        confidence_delta=-0.35,
                        extra_flags=["sequential_mapping"],
                        extra_metadata={
                            "mapping_method": "sequential",
                            "original_question_id": fallback.question_id,
                            "question_index": index,
                        },
                    )

        if mapped is None:
            continue

        if question_expects_subquestion_answers(question) and mapped.answer_mode != "subquestion":
            if mapped.metadata.get("force_whole_answer_input"):
                extra_metadata = {"whole_answer_subquestion_mode": True}
                if _analysis_can_stand_in_for_subanswers(mapped):
                    extra_metadata["analysis_only_subanswers"] = True
                mapped = _copy_answer_unit_with_question_id(
                    mapped,
                    question.question_id,
                    extra_metadata=extra_metadata,
                )
            elif _analysis_can_stand_in_for_subanswers(mapped):
                mapped = _copy_answer_unit_with_question_id(
                    mapped,
                    question.question_id,
                    extra_metadata={"analysis_only_subanswers": True},
                )
            else:
                mapped = _copy_answer_unit_with_question_id(
                    mapped,
                    question.question_id,
                    confidence_delta=-0.2,
                    extra_flags=["question_has_subquestions_but_answer_whole"],
                )
        elif (
            question_expects_subquestion_answers(question)
            and _answer_items_all_empty(mapped)
            and _analysis_can_stand_in_for_subanswers(mapped)
        ):
            mapped = _copy_answer_unit_with_question_id(
                mapped,
                question.question_id,
                extra_metadata={"analysis_only_subanswers": True},
            )
        elif not question_expects_subquestion_answers(question) and mapped.answer_mode == "subquestion":
            if mapped.metadata.get("force_whole_answer_input"):
                mapped = _copy_answer_unit_with_question_id(
                    mapped,
                    question.question_id,
                    extra_metadata={"whole_answer_subquestion_mode": True},
                )
            elif mapped.metadata.get("allow_answer_defined_subquestions"):
                mapped = _copy_answer_unit_with_question_id(
                    mapped,
                    question.question_id,
                    extra_metadata={"answer_defined_subquestions": True},
                )
            else:
                mapped = _copy_answer_unit_with_question_id(
                    mapped,
                    question.question_id,
                    confidence_delta=-0.2,
                    extra_flags=["answer_split_but_question_whole"],
                )

        mapped_units.append(mapped)

    if ordered_occurrence_mapping:
        for raw_answer in raw_answers[len(question_units):]:
            mapped_units.append(
                _copy_answer_unit_with_question_id(
                    raw_answer,
                    raw_answer.question_id,
                    confidence_delta=-0.4,
                    extra_flags=["orphan_answer"],
                    extra_metadata={"mapping_method": "orphan"},
                )
            )
        return mapped_units

    mapped_question_ids = {unit.question_id for unit in mapped_units}
    consumed_original_ids: set[str] = set()
    for unit in mapped_units:
        original_qid = unit.metadata.get("original_question_id", unit.question_id)
        consumed_original_ids.add(original_qid)
        for grouped_original_id in unit.metadata.get("original_question_ids", []) or []:
            consumed_original_ids.add(grouped_original_id)

    for raw_answer in raw_answers:
        if raw_answer.question_id in consumed_original_ids or raw_answer.question_id in mapped_question_ids:
            continue
        mapped_units.append(
            _copy_answer_unit_with_question_id(
                raw_answer,
                raw_answer.question_id,
                confidence_delta=-0.4,
                extra_flags=["orphan_answer"],
                extra_metadata={"mapping_method": "orphan"},
            )
        )

    return mapped_units
