from __future__ import annotations

import sys
from pathlib import Path

import pytest
from docx import Document
from PIL import Image

from shared_core.document_render import sha256_file
from shared_core.macos_quicklook_render import export_docx_to_pdf_with_macos_quicklook


@pytest.mark.skipif(sys.platform != "darwin", reason="macOS Quick Look 集成测试")
def test_macos_quicklook_exports_visible_images_without_mouse(tmp_path: Path) -> None:
    source = tmp_path / "source" / "sample.docx"
    source.parent.mkdir()
    image_path = tmp_path / "banner.png"
    Image.new("RGB", (300, 80), color=(180, 180, 180)).save(image_path)
    document = Document()
    document.add_picture(str(image_path))
    document.add_heading("第1章 章节横幅", level=1)
    document.add_paragraph("对点训练")
    document.add_paragraph("1. 题目正文")
    document.add_picture(str(image_path))
    document.save(source)
    before = sha256_file(source)
    output_pdf = tmp_path / "render" / "preview.pdf"

    metadata = export_docx_to_pdf_with_macos_quicklook(source, output_pdf)

    assert output_pdf.is_file()
    assert output_pdf.stat().st_size > 0
    assert sha256_file(source) == before
    assert metadata["provider"] == "macos_quicklook_webkit"
    assert metadata["renderer_role"] == "development_preview"
    assert metadata["page_truth_authority"] is False
    assert metadata["layout_mode"] == "continuous_preview"
    assert metadata["visible_attachment_count"] >= 2
    assert metadata["source_hash_verified_unchanged"] is True
