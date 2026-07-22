import json
import importlib.util
import tempfile
from pathlib import Path, PureWindowsPath

from docx import Document

from shared_core import (
    AnswerItem,
    AnswerUnit,
    QuestionUnit,
    build_answer_units_from_docx,
    build_answer_units_from_paragraph_texts,
    build_review_report,
    get_review_gate_result,
    initialize_review_status,
    map_answers,
    update_review_status,
)


def _make_question(
    qid: str,
    *,
    question_type: str = "choice",
    subquestions: list[str] | None = None,
    stem_blocks: list[str] | None = None,
    material_blocks: list[str] | None = None,
    confidence: float = 1.0,
    warnings: list[str] | None = None,
) -> QuestionUnit:
    return QuestionUnit(
        question_id=qid,
        subject="理科",
        subject_overlay=None,
        grade_hint=None,
        question_type=question_type,
        stem_blocks=stem_blocks or [f"{qid}题题干"],
        option_blocks=["A", "B", "C", "D"],
        subquestions=subquestions or [],
        media_blocks=[],
        material_blocks=material_blocks or [],
        source_span=(1, 2),
        confidence=confidence,
        warnings=warnings or [],
    )


def _make_answer(
    qid: str,
    *,
    mode: str = "whole",
    answer_text: str = "A",
    confidence: float = 1.0,
    review_flags: list[str] | None = None,
) -> AnswerUnit:
    return AnswerUnit(
        question_id=qid,
        answer_mode=mode,
        answer_items=[AnswerItem(item_id=qid, text=answer_text)],
        analysis_items=[AnswerItem(item_id=qid, text=f"{qid}题解析")],
        confidence=confidence,
        review_flags=review_flags or [],
        source_span=(1, 2),
        answer_span=(1, 1),
        analysis_span=(2, 2),
    )


def test_map_answers_exact_match_keeps_ids():
    questions = [_make_question("1"), _make_question("2")]
    answers = [_make_answer("1", answer_text="B"), _make_answer("2", answer_text="C")]

    mapped = map_answers(questions, answers)

    assert [unit.question_id for unit in mapped] == ["1", "2"]
    assert not mapped[0].review_flags
    assert not mapped[1].review_flags


def test_map_answers_sequential_fallback_marks_review():
    questions = [_make_question("13"), _make_question("14")]
    answers = [_make_answer("1", answer_text="甲"), _make_answer("2", answer_text="乙")]

    mapped = map_answers(questions, answers)

    assert [unit.question_id for unit in mapped[:2]] == ["13", "14"]
    assert "sequential_mapping" in mapped[0].review_flags
    assert mapped[0].confidence < 0.75
    assert mapped[0].metadata["original_question_id"] == "1"


def test_review_report_blocks_mapping_risks():
    questions = [
        _make_question("15", subquestions=["(1)", "(2)"]),
        _make_question("16", confidence=0.6, warnings=["cross_page_media"]),
    ]
    answers = [
        _make_answer("1", answer_text="甲"),
        _make_answer("16", mode="whole", answer_text="乙", review_flags=["manual_check"]),
        _make_answer("99", answer_text="孤立答案"),
    ]

    mapped = map_answers(questions, answers)
    report = build_review_report("零错流程测试", questions, mapped)

    assert report.summary["high_risk_count"] >= 2
    issue_titles = {issue.title for issue in report.issues}
    assert "题答顺序映射，需人工确认" in issue_titles
    assert "答案块置信度较低" in issue_titles
    assert "题块置信度较低" in issue_titles


def test_map_answers_groups_material_range_into_subquestions():
    questions = [
        _make_question(
            "3",
            question_type="material_choice",
            stem_blocks=[
                "3.在西藏发现鱼龙化石，说明（   ）",
                "4.能反映西藏这种地表形态变化的词语是（   ）",
                "5.造成西藏这种地表形态变化的根本原因是（   ）",
            ],
            material_blocks=["阅读材料，回答3-5题。"],
        ),
        _make_question("6"),
    ]
    answers = [
        _make_answer("3", answer_text="D"),
        _make_answer("4", answer_text="C"),
        _make_answer("5", answer_text="C"),
        _make_answer("6", answer_text="A"),
    ]

    mapped = map_answers(questions, answers)

    assert [unit.question_id for unit in mapped] == ["3", "6"]
    assert mapped[0].answer_mode == "subquestion"
    assert [item.text for item in mapped[0].answer_items] == ["D", "C", "C"]
    assert [item.item_id for item in mapped[0].answer_items] == ["(1)", "(2)", "(3)"]
    assert mapped[0].metadata["grouped_question_ids"] == ["3", "4", "5"]
    assert not any(flag == "answer_split_but_question_whole" for flag in mapped[0].review_flags)


def test_map_answers_groups_material_questions_with_prefixed_question_ids():
    questions = [
        _make_question(
            "2",
            question_type="material_choice",
            stem_blocks=[
                "【C】2．日本和新西兰（   ）",
                "A．甲 B．乙",
                "【G】3．与日本相比，新西兰出口商品以乳制品和肉类为主，主要自然原因是（   ）",
                "A．甲 B．乙",
            ],
            material_blocks=["图为日本和新西兰示意图。完成下面小题。"],
        ),
    ]
    answers = [
        _make_answer("2", answer_text="B"),
        _make_answer("3", answer_text="C"),
    ]

    mapped = map_answers(questions, answers)

    assert [unit.question_id for unit in mapped] == ["2"]
    assert mapped[0].answer_mode == "subquestion"
    assert [item.text for item in mapped[0].answer_items] == ["B", "C"]
    assert mapped[0].metadata["grouped_question_ids"] == ["2", "3"]


def test_review_report_flags_contaminated_answer_block():
    questions = [_make_question("1")]
    contaminated = AnswerUnit(
        question_id="1",
        answer_mode="whole",
        answer_items=[AnswerItem(item_id="1", text="D")],
        analysis_items=[AnswerItem(item_id="1", text="正常解析后混入 19【答案】 1.B 2.D 3.C 4.B 5.D")],
        confidence=1.0,
        review_flags=[],
        source_span=(1, 3),
        answer_span=(1, 1),
        analysis_span=(2, 3),
    )

    report = build_review_report("污染答案测试", questions, [contaminated])

    assert any(issue.title == "答案块疑似串入其他题号" for issue in report.issues)
    assert report.summary["high_risk_count"] >= 1


def test_docx_answer_merge_keeps_table_answers_when_paragraph_numbering_resets():
    with tempfile.TemporaryDirectory() as temp_dir:
        answer_path = Path(temp_dir) / "答案.docx"

        doc = Document()
        table = doc.add_table(rows=2, cols=4)
        headers = ["题号", "1", "2", "3"]
        answers = ["答案", "A", "D", "B"]
        for idx, value in enumerate(headers):
            table.cell(0, idx).text = value
        for idx, value in enumerate(answers):
            table.cell(1, idx).text = value

        doc.add_paragraph("16.")
        doc.add_paragraph("综合题答案")
        doc.add_paragraph("17.")
        doc.add_paragraph("另一题答案")
        doc.add_paragraph("1.小问一")
        doc.add_paragraph("2.小问二")
        doc.save(answer_path)

        units = build_answer_units_from_docx(str(answer_path))

        assert [unit.question_id for unit in units] == ["1", "2", "3", "16", "17"]
        assert [unit.answer_items[0].text for unit in units[:3]] == ["A", "D", "B"]


def test_build_answer_units_supports_numeric_subanswers_and_answer_marker_titles():
    units = build_answer_units_from_paragraph_texts(
        [
            "16.",
            "1第一空",
            "2第二空",
            "3第三空",
            "18💡",
            "1.D 2.C 3.D",
            "19【答案】 1.B 2.D 3.C 4.B 5.D",
        ]
    )

    assert [unit.question_id for unit in units] == ["16", "18", "19"]
    assert units[0].answer_mode == "subquestion"
    assert [item.text for item in units[0].answer_items] == ["第一空", "第二空", "第三空"]
    assert units[1].answer_mode == "subquestion"
    assert [item.text for item in units[1].answer_items] == ["D", "C", "D"]
    assert units[2].answer_mode == "subquestion"
    assert [item.text for item in units[2].answer_items] == ["B", "D", "C", "B", "D"]


def test_format_alignment_rewrites_docx_using_question_structure():
    project_root = Path(__file__).resolve().parent
    format_main_path = project_root / "格式处理" / "main.py"
    spec = importlib.util.spec_from_file_location("mohen_format_main_test", format_main_path)
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    spec.loader.exec_module(module)

    with tempfile.TemporaryDirectory() as temp_dir:
        question_path = Path(temp_dir) / "地理题目.docx"
        answer_path = Path(temp_dir) / "地理答案_已清洗.docx"

        question_doc = Document()
        question_doc.add_paragraph("阅读材料，回答3-5题。")
        question_doc.add_paragraph("3.在西藏发现鱼龙化石，说明（   ）")
        question_doc.add_paragraph("A．甲 B．乙")
        question_doc.add_paragraph("4.能反映西藏这种地表形态变化的词语是（   ）")
        question_doc.add_paragraph("A．甲 B．乙")
        question_doc.add_paragraph("5.造成西藏这种地表形态变化的根本原因是（   ）")
        question_doc.add_paragraph("A．甲 B．乙")
        question_doc.add_paragraph("6.下一题")
        question_doc.save(question_path)

        answer_doc = Document()
        answer_doc.add_paragraph("3．D")
        answer_doc.add_paragraph("解析：第3题解析")
        answer_doc.add_paragraph("4．C")
        answer_doc.add_paragraph("解析：第4题解析")
        answer_doc.add_paragraph("5．C")
        answer_doc.add_paragraph("解析：第5题解析")
        answer_doc.add_paragraph("6．A")
        answer_doc.add_paragraph("解析：第6题解析")
        answer_doc.save(answer_path)

        aligned = module._align_cleaned_doc_with_question_doc(str(question_path), str(answer_path))

        assert aligned is not None
        rewritten = Document(answer_path)
        texts = [para.text for para in rewritten.paragraphs if para.text.strip()]
        assert texts[:4] == ["3．(1)D", "(2)C", "(3)C", "解析：(1)第3题解析"]
        assert texts[4:6] == ["(2)第4题解析", "(3)第5题解析"]
        assert texts[6:8] == ["6．A", "解析：第6题解析"]


def _load_controller_main():
    project_root = Path(__file__).resolve().parent
    controller_main_path = project_root / "main.py"
    spec = importlib.util.spec_from_file_location("mohen_controller_main_test", controller_main_path)
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    spec.loader.exec_module(module)
    return module


def test_controller_helpers():
    controller_main = _load_controller_main()

    report = build_review_report(
        "helper-test",
        [_make_question("1")],
        map_answers([_make_question("1")], [_make_answer("9", answer_text="A")]),
    )

    assert controller_main.derive_cleaned_output_path(r"E:\samples\答案.docx").endswith("答案_已清洗.docx")
    assert controller_main.has_blocking_review_issues(report) is True


def test_build_manifest_entry():
    controller_main = _load_controller_main()

    entry = controller_main.build_manifest_entry(
        file_path=PureWindowsPath(r"E:\samples\高一地理限训4.docx"),
        sample_kind="question_doc",
        subject="地理",
        question_count=12,
        tags=["material", "image"],
        requires_subquestion_split=False,
    )

    assert entry["file_name"] == "高一地理限训4.docx"
    assert entry["sample_kind"] == "question_doc"
    assert entry["tags"] == ["material", "image"]


def test_workspace_config_auto_create_and_collect_overview():
    controller_main = _load_controller_main()

    with tempfile.TemporaryDirectory() as temp_dir:
        project_root = Path(temp_dir)
        (project_root / "题目库").mkdir()
        (project_root / "答案库").mkdir()
        (project_root / "清洗库").mkdir()
        (project_root / "审核库").mkdir()

        (project_root / "题目库" / "第1课题目.docx").write_text("q", encoding="utf-8")
        (project_root / "答案库" / "第1课答案.docx").write_text("a", encoding="utf-8")
        (project_root / "答案库" / "第1课答案_已清洗.docx").write_text("skip", encoding="utf-8")
        (project_root / "清洗库" / "第1课答案_已清洗.docx").write_text("c", encoding="utf-8")
        (project_root / "审核库" / "第1课答案_已清洗_审核清单.md").write_text("r", encoding="utf-8")

        config_path = project_root / "工作台路径配置.json"
        config_data = {
            "question_dirs": ["题目库"],
            "raw_answer_dirs": ["答案库"],
            "clean_answer_dirs": ["清洗库"],
            "review_dirs": ["审核库"],
        }
        config_path.write_text(json.dumps(config_data, ensure_ascii=False, indent=2), encoding="utf-8")

        loaded = controller_main.load_workspace_config(config_path, project_root=project_root)
        overview = controller_main.collect_workspace_overview(loaded)

        assert Path(loaded["resolved"]["question_dirs"][0]) == (project_root / "题目库").resolve()
        assert overview["buckets"]["question_dirs"][0]["count"] == 1
        assert overview["buckets"]["raw_answer_dirs"][0]["count"] == 1
        assert overview["buckets"]["clean_answer_dirs"][0]["count"] == 1
        assert overview["buckets"]["review_dirs"][0]["count"] == 1


def test_workspace_config_is_created_when_missing():
    controller_main = _load_controller_main()

    with tempfile.TemporaryDirectory() as temp_dir:
        config_path = Path(temp_dir) / "工作台路径配置.json"
        created_path = controller_main.ensure_workspace_config(config_path)
        loaded = controller_main.load_workspace_config(config_path, project_root=Path(temp_dir))

        assert Path(created_path) == config_path
        assert config_path.exists()
        assert loaded["raw"]["question_dirs"] == controller_main.DEFAULT_WORKSPACE_CONFIG["question_dirs"]


def test_workspace_records_pair_files_and_summarize_status():
    controller_main = _load_controller_main()

    with tempfile.TemporaryDirectory() as temp_dir:
        project_root = Path(temp_dir)
        (project_root / "题目库").mkdir()
        (project_root / "答案库").mkdir()
        (project_root / "清洗库").mkdir()
        (project_root / "审核库").mkdir()

        question_path = project_root / "题目库" / "限训11.docx"
        raw_answer_path = project_root / "答案库" / "限训11答案.docx"
        clean_answer_path = project_root / "清洗库" / "限训11答案_已清洗.docx"
        review_report_path = project_root / "审核库" / "限训11答案_已清洗_审核清单.md"

        question_path.write_text("q", encoding="utf-8")
        raw_answer_path.write_text("a", encoding="utf-8")
        review_report_path.write_text("# report", encoding="utf-8")

        doc = Document()
        doc.add_paragraph("1．A")
        doc.add_paragraph("解析：")
        doc.save(clean_answer_path)
        initialize_review_status(clean_answer_path, report_path=str(review_report_path), report=None)
        update_review_status(clean_answer_path, status="approved", reviewer="system", note="ok")

        other_question_path = project_root / "题目库" / "限训12题目.docx"
        other_raw_answer_path = project_root / "答案库" / "限训12答案.docx"
        other_clean_answer_path = project_root / "清洗库" / "限训12答案_已清洗.docx"
        other_question_path.write_text("q2", encoding="utf-8")
        other_raw_answer_path.write_text("a2", encoding="utf-8")
        doc2 = Document()
        doc2.add_paragraph("1．B")
        doc2.add_paragraph("解析：")
        doc2.save(other_clean_answer_path)
        initialize_review_status(other_clean_answer_path, report=None)
        update_review_status(other_clean_answer_path, status="rejected", reviewer="system", note="bad")

        config_path = project_root / "工作台路径配置.json"
        config_data = {
            "question_dirs": ["题目库"],
            "raw_answer_dirs": ["答案库"],
            "clean_answer_dirs": ["清洗库"],
            "review_dirs": ["审核库"],
        }
        config_path.write_text(json.dumps(config_data, ensure_ascii=False, indent=2), encoding="utf-8")

        loaded = controller_main.load_workspace_config(config_path, project_root=project_root)
        records = controller_main.collect_workspace_records(loaded)
        summary = controller_main.summarize_workspace_records(records)
        matched = next(record for record in records if record["display_name"] == "限训11")

        assert matched["question_doc_path"] == str(question_path.resolve())
        assert matched["raw_answer_path"] == str(raw_answer_path.resolve())
        assert matched["clean_answer_path"] == str(clean_answer_path.resolve())
        assert matched["review_report_path"] == str(review_report_path.resolve())
        assert matched["gate_status"] == "approved"
        assert matched["stage"] == "可录答案"
        assert summary["可录答案"] == 1
        assert summary["自动检查未通过"] == 1


def test_find_workspace_record_by_question_path():
    controller_main = _load_controller_main()

    with tempfile.TemporaryDirectory() as temp_dir:
        project_root = Path(temp_dir)
        (project_root / "题目库").mkdir()
        (project_root / "答案库").mkdir()

        question_path = project_root / "题目库" / "第1课题目.docx"
        raw_answer_path = project_root / "答案库" / "第1课答案.docx"
        question_path.write_text("q", encoding="utf-8")
        raw_answer_path.write_text("a", encoding="utf-8")

        config_path = project_root / "工作台路径配置.json"
        config_data = {
            "question_dirs": ["题目库"],
            "raw_answer_dirs": ["答案库"],
            "clean_answer_dirs": ["清洗库"],
            "review_dirs": ["审核库"],
        }
        config_path.write_text(json.dumps(config_data, ensure_ascii=False, indent=2), encoding="utf-8")

        loaded = controller_main.load_workspace_config(config_path, project_root=project_root)
        record = controller_main.find_workspace_record_by_question_path(str(question_path.resolve()), loaded)

        assert record is not None
        assert record["raw_answer_path"] == str(raw_answer_path.resolve())
        assert record["stage"] == "待清洗"


def test_review_gate_requires_approval_and_invalidates_on_change():
    with tempfile.TemporaryDirectory() as temp_dir:
        answer_path = Path(temp_dir) / "答案_已清洗.docx"
        report_path = Path(temp_dir) / "答案_已清洗_审核清单.md"

        doc = Document()
        doc.add_paragraph("1．A")
        doc.add_paragraph("解析：")
        doc.save(answer_path)
        report_path.write_text("# report", encoding="utf-8")

        initialize_review_status(answer_path, report_path=str(report_path), report=None)
        pending_result = get_review_gate_result(answer_path)
        assert pending_result["allowed"] is False
        assert pending_result["status"] == "pending"

        update_review_status(answer_path, status="approved", reviewer="tester", note="ok")
        approved_result = get_review_gate_result(answer_path)
        assert approved_result["allowed"] is True

        doc = Document(str(answer_path))
        doc.add_paragraph("2．B")
        doc.save(answer_path)

        stale_result = get_review_gate_result(answer_path)
        assert stale_result["allowed"] is False
        assert stale_result["status"] == "stale"


def test_auto_review_decision_marks_approved_or_rejected():
    controller_main = _load_controller_main()

    with tempfile.TemporaryDirectory() as temp_dir:
        answer_path = Path(temp_dir) / "答案_已清洗.docx"
        doc = Document()
        doc.add_paragraph("1．A")
        doc.add_paragraph("解析：")
        doc.save(answer_path)

        report_ok = build_review_report(
            "auto-ok",
            [_make_question("1")],
            map_answers([_make_question("1")], [_make_answer("1", answer_text="A")]),
        )
        controller_main.initialize_review_status(answer_path, report=None)
        controller_main.apply_auto_review_decision(str(answer_path), report_ok)
        gate_ok = get_review_gate_result(answer_path)
        assert gate_ok["allowed"] is True
        assert gate_ok["status"] == "approved"

        report_bad = build_review_report(
            "auto-bad",
            [_make_question("2")],
            map_answers([_make_question("2")], [_make_answer("9", answer_text="B")]),
        )
        controller_main.apply_auto_review_decision(str(answer_path), report_bad)
        gate_bad = get_review_gate_result(answer_path)
        assert gate_bad["allowed"] is False
        assert gate_bad["status"] == "rejected"


if __name__ == "__main__":
    test_map_answers_exact_match_keeps_ids()
    test_map_answers_sequential_fallback_marks_review()
    test_review_report_blocks_mapping_risks()
    test_map_answers_groups_material_range_into_subquestions()
    test_map_answers_groups_material_questions_with_prefixed_question_ids()
    test_review_report_flags_contaminated_answer_block()
    test_docx_answer_merge_keeps_table_answers_when_paragraph_numbering_resets()
    test_build_answer_units_supports_numeric_subanswers_and_answer_marker_titles()
    test_format_alignment_rewrites_docx_using_question_structure()
    test_controller_helpers()
    test_build_manifest_entry()
    test_workspace_config_auto_create_and_collect_overview()
    test_workspace_config_is_created_when_missing()
    test_workspace_records_pair_files_and_summarize_status()
    test_find_workspace_record_by_question_path()
    test_review_gate_requires_approval_and_invalidates_on_change()
    test_auto_review_decision_marks_approved_or_rejected()
    print("ALL PASSED")
