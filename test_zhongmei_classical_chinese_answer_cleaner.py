from __future__ import annotations

import base64
import re
import tempfile
import unittest
from io import BytesIO
from pathlib import Path

from docx import Document

from shared_core import build_answer_units_from_docx
from tools.clean_zhongmei_classical_chinese_answers import (
    CleanDocument,
    DEFAULT_ANSWER_DIR,
    DEFAULT_QUESTION_DIR,
    UnsupportedTemplateError,
    clean_batch,
    count_score_points,
    discover_source_pairs,
    parse_document,
    preflight_source_pairs,
    render_clean_lines,
    validate_clean_docx,
)


ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk"
    "+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def _write_docx(path: Path, paragraphs: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)


def _real_samples_available() -> bool:
    return DEFAULT_ANSWER_DIR.is_dir() and DEFAULT_QUESTION_DIR.is_dir()


@unittest.skipUnless(_real_samples_available(), "缺少众美文言文真实样本")
class ZhongmeiClassicalChineseRealSampleTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.pairs = discover_source_pairs()
        cls.clean_documents = preflight_source_pairs(cls.pairs)
        cls.by_number = {
            int(clean_document.source_path.name[:2]): clean_document
            for clean_document in cls.clean_documents
        }

    def test_discovers_all_recursive_pairs_and_groups_196_markers_into_154_blocks(self):
        self.assertEqual(len(self.pairs), 28)
        self.assertEqual(
            [pair.number for pair in self.pairs],
            list(range(1, 29)),
        )
        self.assertTrue(
            all(pair.answer_path.relative_to(DEFAULT_ANSWER_DIR) == pair.relative_path for pair in self.pairs)
        )
        self.assertTrue(
            all(pair.question_path.relative_to(DEFAULT_QUESTION_DIR) == pair.relative_path for pair in self.pairs)
        )
        self.assertEqual(
            sum(clean_document.source_answer_marker_count for clean_document in self.clean_documents),
            196,
        )
        self.assertEqual(
            sum(len(clean_document.blocks) for clean_document in self.clean_documents),
            154,
        )
        translation_blocks = [
            block
            for clean_document in self.clean_documents
            for block in clean_document.blocks
            if block.is_translation
        ]
        self.assertEqual(len(translation_blocks), 33)
        self.assertEqual(sum(len(block.occurrences) for block in translation_blocks), 74)
        self.assertEqual(
            sum(
                max(0, len(occurrence.answer_lines) - 1)
                for clean_document in self.clean_documents
                for block in clean_document.blocks
                for occurrence in block.occurrences
            ),
            219,
        )

    def test_score_point_oracle_counts_330_and_document_19_plain_lists_as_4_3_4(self):
        analysis_lines = [
            analysis_line
            for clean_document in self.clean_documents
            for block in clean_document.blocks
            if block.is_translation
            for occurrence in block.occurrences
            for analysis_line in occurrence.analysis_lines
        ]
        self.assertEqual(len(analysis_lines), 74)
        self.assertEqual(sum(count_score_points(line) for line in analysis_lines), 330)

        document_19_plain_lists = [
            occurrence.analysis_lines[0]
            for block in self.by_number[19].blocks
            if block.is_translation
            for occurrence in block.occurrences
            if occurrence.analysis_lines
            and not re.search(r"[；;]", occurrence.analysis_lines[0])
            and count_score_points(occurrence.analysis_lines[0]) > 1
        ]
        self.assertEqual(
            [count_score_points(line) for line in document_19_plain_lists],
            [4, 3, 4],
        )

    def test_preserves_duplicate_original_numbers_and_only_repairs_suwu_whitelist(self):
        expected_ids = {
            7: ["1", "2", "3", "4", "5", "1", "2", "3", "4"],
            10: ["1", "2", "3", "4", "1", "2", "3"],
            16: ["1", "2", "3", "4", "5", "1", "2", "3", "4"],
            17: ["1", "2", "3", "4", "1", "2", "3", "4"],
            19: ["1", "2", "3", "4", "5", "1", "2", "3", "4"],
            20: ["1", "2", "3", "4"],
        }
        for number, expected in expected_ids.items():
            actual = [block.question_id for block in self.by_number[number].blocks]
            self.assertEqual(actual, expected, self.by_number[number].source_path.name)

    def test_translation_score_points_move_to_analysis_and_circled_lists_stay_in_answer(self):
        clean_document = self.by_number[1]
        translation_block = clean_document.blocks[1]
        self.assertEqual(translation_block.question_id, "2")
        self.assertEqual(len(translation_block.occurrences), 3)
        for occurrence in translation_block.occurrences:
            self.assertNotIn("得分点", "\n".join(occurrence.answer_lines))
            self.assertTrue(occurrence.analysis_lines)
            self.assertTrue(occurrence.analysis_lines[0].startswith("得分点："))

        circled_list_block = clean_document.blocks[2]
        self.assertIn("①", "\n".join(circled_list_block.occurrences[0].answer_lines))
        self.assertEqual(circled_list_block.occurrences[0].analysis_lines, [])

    def test_translation_blocks_render_subanswers_before_one_merged_analysis(self):
        translation_block_count = 0
        translation_answer_count = 0

        for clean_document in self.clean_documents:
            for block in clean_document.blocks:
                if not block.is_translation:
                    continue
                translation_block_count += 1
                translation_answer_count += len(block.occurrences)
                isolated_document = CleanDocument(
                    source_path=clean_document.source_path,
                    question_path=clean_document.question_path,
                    relative_path=clean_document.relative_path,
                    titles=[],
                    blocks=[block],
                    source_answer_marker_count=len(block.occurrences),
                )
                lines = render_clean_lines(isolated_document)
                answer_entries = [
                    (index, int(match.group(1)))
                    for index, line in enumerate(lines)
                    if (match := re.match(r"^\((\d+)\)答案：", line))
                ]
                analysis_headers = [
                    index for index, line in enumerate(lines) if line == "解析："
                ]

                self.assertEqual(
                    [marker for _, marker in answer_entries],
                    list(range(1, len(block.occurrences) + 1)),
                    clean_document.source_path.name,
                )
                self.assertEqual(len(analysis_headers), 1, clean_document.source_path.name)
                analysis_index = analysis_headers[0]
                self.assertTrue(
                    all(index < analysis_index for index, _ in answer_entries)
                )
                self.assertFalse(
                    [line for line in lines if re.match(r"^\(\d+\)解析：", line)]
                )
                merged_analysis_markers = [
                    int(match.group(1))
                    for line in lines[analysis_index + 1 :]
                    if (match := re.match(r"^\((\d+)\)", line))
                ]
                self.assertEqual(
                    merged_analysis_markers,
                    list(range(1, len(block.occurrences) + 1)),
                    clean_document.source_path.name,
                )

        self.assertEqual(translation_block_count, 33)
        self.assertEqual(translation_answer_count, 74)

    def test_keeps_work_titles_drops_task_headings_and_restores_known_image_glyphs(self):
        self.assertIn("《劝学》", [title.text for title in self.by_number[1].titles])
        self.assertEqual(
            [title.text for title in self.by_number[7].titles],
            ["《庖丁解牛》《烛之武退秦师》", "《庖丁解牛》", "《烛之武退秦师》"],
        )

        all_lines = {
            number: render_clean_lines(clean_document)
            for number, clean_document in self.by_number.items()
        }
        forbidden = re.compile(r"^(?:[一二三四五六七八九十]+、|[（(][一二三四五六七八九十]+[）)])")
        self.assertFalse(
            [line for lines in all_lines.values() for line in lines if forbidden.match(line)]
        )
        self.assertEqual("\n".join(all_lines[1]).count("輮"), 4)
        self.assertEqual("\n".join(all_lines[17]).count("絖"), 1)
        self.assertEqual("\n".join(all_lines[19]).count("餔"), 1)
        self.assertNotIn(
            "INCLUDEPICTURE",
            "\n".join(line for lines in all_lines.values() for line in lines),
        )

    def test_clean_batch_writes_recursive_outputs_to_temp_root_and_reopens(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            output_root = (
                Path(temp_dir)
                / "众美-高三-语文"
                / "答案"
                / "文言文答案"
            )
            output_paths = clean_batch(output_dir=output_root)
            self.assertEqual(len(output_paths), 28)
            self.assertTrue(all(path.is_file() for path in output_paths))
            self.assertTrue(
                (output_root / "必修上册" / "01《劝学》_已清洗.docx").is_file()
            )
            self.assertFalse(
                [path for path in output_root.iterdir() if path.name.startswith(".文言文答案清洗临时_")]
            )
            for output_path, expected_document in zip(output_paths, self.clean_documents):
                validate_clean_docx(output_path, expected_document)
                units = build_answer_units_from_docx(output_path)
                self.assertEqual(len(units), len(expected_document.blocks))
                for block, unit in zip(expected_document.blocks, units):
                    if not block.is_translation:
                        continue
                    self.assertEqual(unit.answer_mode, "subquestion")
                    self.assertEqual(len(unit.answer_items), len(block.occurrences))
                    self.assertEqual(len(unit.analysis_items), 1)
                    self.assertFalse(unit.metadata["interleaved_subquestion_analysis"])
                    self.assertFalse(unit.metadata["force_whole_answer_input"])


class ZhongmeiClassicalChineseSafetyTests(unittest.TestCase):
    def test_unknown_inline_image_inside_extracted_answer_stops_parsing(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            answer_path = root / "答案" / "01《测试》.docx"
            question_path = root / "题目" / "01《测试》.docx"
            _write_docx(question_path, ["1．测试题干"])

            answer_path.parent.mkdir(parents=True)
            answer_doc = Document()
            paragraph = answer_doc.add_paragraph("答案：测试答案")
            paragraph.add_run().add_picture(BytesIO(ONE_PIXEL_PNG))
            answer_doc.save(answer_path)

            with self.assertRaisesRegex(UnsupportedTemplateError, "未知内联图片"):
                parse_document(
                    answer_path,
                    question_path,
                    relative_path=Path("01《测试》.docx"),
                )

    def test_single_synthetic_block_gets_blank_analysis_and_title_from_filename(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            answer_path = root / "答案" / "01《测试》.docx"
            question_path = root / "题目" / "01《测试》.docx"
            _write_docx(question_path, ["1．测试题干"])
            _write_docx(answer_path, ["一、微点夯实", "答案：①甲 ②乙"])

            clean_document = parse_document(
                answer_path,
                question_path,
                relative_path=Path("01《测试》.docx"),
            )
            self.assertEqual(
                render_clean_lines(clean_document),
                ["《测试》", "1．", "答案：①甲 ②乙", "解析： "],
            )


if __name__ == "__main__":
    unittest.main()
