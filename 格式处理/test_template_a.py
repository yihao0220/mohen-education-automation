# coding: utf-8
"""
模板A离线测试脚本
测试虎鹰-四年级-语文格式清洗逻辑
"""

import re
import sys
import os

# 添加路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# 模拟模板A的正则模式
QUESTION_PATTERN = re.compile(r"^(\d+)\.\s*(.+)")
CHOICE_ANSWER_PATTERN = re.compile(r"^(\d+)\.\s*([A-D])\b\s*(.*)")
QUESTION_WITH_SUB_PATTERN = re.compile(r"^(\d+)\.\s*[(（](\d+)[)）]\s*([A-D]?)\s*(.*)")
SUB_QUESTION_PATTERN = re.compile(r"^[(（](\d+)[)）]\s*(.*)")
MULTI_SUB_PATTERN = re.compile(r"[(（](\d+)[)）]([^（(]*?)(?=[(（]\d+[)）]|$)")

# 测试数据（从四下语文1-2答案2.docx提取的样本）
TEST_CASES = [
    # 章节标题
    ("1 古诗词三首", "section"),
    ("2 乡下人家", "section"),
    ("阅读素养练习(一)", "section"),
    ("语文园地", "section"),
    
    # 大题答案
    ("1. 蜻蜓 蝶 稀 shū 茅檐 翁 bō", "big_question"),
    ("5. 杏子 菜花 篱笆 小路 小溪 青草 夏季", "big_question"),
    
    # 选择题
    ("2. D", "choice"),
    ("3. A", "choice"),
    
    # 带小问的选择题（答案必须是字母A-D）
    ("7. (1)B", "choice_with_sub"),
    ("(2)D", "sub_only"),
    ("(3)C A B", "sub_only"),
    
    # 大题带小问（答案不是字母，是文字）
    ("6. (1)锄豆 天真活泼 手巧 顽皮可爱", "big_with_sub"),
    ("(2)急切、欢快 闲适、自在 天真可爱 喜爱之情", "sub_only"),
    ("(3)示例:一片金黄的菜花地里...", "sub_only"),
    
    # 多小问同行（选择题格式）
    ("3. (1)B (2)D (3)C A B", "multi_sub_choice"),
    
    # 大题带多小问同行
    ("4.(1)舒缓、轻柔", "big_with_sub"),
]


def test_patterns():
    """测试正则模式匹配"""
    print("=" * 60)
    print("模板A正则模式测试")
    print("=" * 60)
    
    for text, expected_type in TEST_CASES:
        print(f"\n测试: '{text[:50]}...' (期望: {expected_type})")
        
        # 测试 QUESTION_WITH_SUB_PATTERN
        match = QUESTION_WITH_SUB_PATTERN.match(text)
        if match:
            print(f"  ✓ QUESTION_WITH_SUB_PATTERN 匹配: 题号={match.group(1)}, 小问={match.group(2)}, 答案={match.group(3)}, 剩余={match.group(4)[:30] if match.group(4) else ''}")
            continue
            
        # 测试 CHOICE_ANSWER_PATTERN
        match = CHOICE_ANSWER_PATTERN.match(text)
        if match:
            print(f"  ✓ CHOICE_ANSWER_PATTERN 匹配: 题号={match.group(1)}, 答案={match.group(2)}, 剩余={match.group(3)[:30] if match.group(3) else ''}")
            continue
            
        # 测试 QUESTION_PATTERN
        match = QUESTION_PATTERN.match(text)
        if match:
            print(f"  ✓ QUESTION_PATTERN 匹配: 题号={match.group(1)}, 内容={match.group(2)[:40]}")
            
            # 检查内容是否包含小问
            content = match.group(2)
            sub_match = SUB_QUESTION_PATTERN.match(content)
            if sub_match:
                print(f"    └─ 包含小问: ({sub_match.group(1)}){sub_match.group(2)[:30]}")
                
                # 检查多小问
                subs = MULTI_SUB_PATTERN.findall(content)
                if len(subs) > 1:
                    print(f"    └─ 多小问: {subs}")
            continue
            
        # 测试 SUB_QUESTION_PATTERN
        match = SUB_QUESTION_PATTERN.match(text)
        if match:
            print(f"  ✓ SUB_QUESTION_PATTERN 匹配: ({match.group(1)}){match.group(2)[:40]}")
            continue
            
        # 测试章节标题
        if re.match(r"^\d+\s+[^0-9]", text) or re.match(r"^阅读素养练习", text) or re.match(r"^单元素养练习", text) or re.match(r"^语文园地", text):
            print(f"  ✓ 章节标题识别")
            continue
            
        print(f"  ✗ 未匹配任何模式")


def test_multi_sub_extraction():
    """测试多小问提取"""
    print("\n" + "=" * 60)
    print("多小问提取测试")
    print("=" * 60)
    
    test_cases = [
        "(1)B (2)D (3)C A B",
        "(1)振动 空气 音色 (2)510 340",
        "(1)B (2)C",
    ]
    
    for text in test_cases:
        print(f"\n测试: '{text}'")
        subs = MULTI_SUB_PATTERN.findall(text)
        print(f"  提取结果: {subs}")


def test_docx_reading():
    """测试读取实际docx文件"""
    print("\n" + "=" * 60)
    print("实际文档读取测试")
    print("=" * 60)
    
    try:
        from docx import Document
        
        doc_path = os.path.join(os.path.dirname(__file__), "待清洗文件", "四下语文1-2答案2.docx")
        if not os.path.exists(doc_path):
            print(f"  ! 文件不存在: {doc_path}")
            return
            
        doc = Document(doc_path)
        print(f"  ✓ 成功读取文档，共 {len(doc.paragraphs)} 个段落")
        
        # 显示前30行
        print("\n  前30行内容:")
        for i, para in enumerate(doc.paragraphs[:30]):
            text = para.text.strip()
            if text:
                print(f"    {i+1:2d}: {text[:70]}")
                
    except ImportError:
        print("  ! 未安装 python-docx，跳过文档读取测试")
    except Exception as e:
        print(f"  ! 读取文档出错: {e}")


def test_section_scan():
    """测试章节扫描功能"""
    print("\n" + "=" * 60)
    print("章节扫描功能测试")
    print("=" * 60)
    
    try:
        from docx import Document
        
        doc_path = os.path.join(os.path.dirname(__file__), "待清洗文件", "四下语文1-2答案2.docx")
        if not os.path.exists(doc_path):
            print(f"  ! 文件不存在: {doc_path}")
            return
            
        doc = Document(doc_path)
        
        # 模拟扫描章节
        sections = []
        SECTION_PATTERNS = [
            r"^\d+\s+[^0-9（(].*",
            r"^阅读素养练习[（(].*[)）]",
            r"^单元素养练习[（(].*[)）]",
            r"^语文园地",
        ]
        
        def is_section_title(text):
            for pattern in SECTION_PATTERNS:
                if re.match(pattern, text):
                    return True
            return False
        
        for i, para in enumerate(doc.paragraphs, 1):
            text = para.text.strip()
            if text and is_section_title(text):
                sections.append((i, text))
        
        print(f"  ✓ 扫描到 {len(sections)} 个章节：")
        for idx, (line_num, title) in enumerate(sections, 1):
            print(f"    [{idx}] 行{line_num}: {title[:50]}")
                
    except ImportError:
        print("  ! 未安装 python-docx，跳过章节扫描测试")
    except Exception as e:
        print(f"  ! 测试出错: {e}")


if __name__ == "__main__":
    test_patterns()
    test_multi_sub_extraction()
    test_docx_reading()
    test_section_scan()
    
    print("\n" + "=" * 60)
    print("测试完成")
    print("=" * 60)
