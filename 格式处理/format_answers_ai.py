# coding: utf-8
"""
答案格式清洗 - AI 智能版 v2.0

功能：
  - 使用 DeepSeek-V3 AI 智能识别任意格式的答案文档
  - 自动提取选择题答案、大题答案及解析
  - 输出标准化格式 docx，可直接供答案录入模块使用

使用方法：
  1. 将待清洗文件放入 待清洗文件/ 目录（或直接运行后输入路径）
  2. 运行: python format_answers_ai.py
  3. 清洗结果保存在 已清洗文件/ 目录

输出格式（Standard output format）：
  选择题: N．A　  （答案字母后跟全角空格）
          解析：...
  大题:   N．
          (1)答案内容
          (2)答案内容
          解析：(1)...（2）...
"""

import os
import sys
import json
import re
import time
import urllib.request

# 确保 UTF-8 输出
sys.stdout = __import__("io").TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

# ── 尝试导入 python-docx ───────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ 缺少依赖库，请先运行: pip install python-docx")
    sys.exit(1)


# ═══════════════════════════════════════════════════════════════
#  配置区（按需修改）
# ═══════════════════════════════════════════════════════════════

API_URL  = "https://api.siliconflow.cn/v1/chat/completions"
MODEL    = "deepseek-ai/DeepSeek-V3.2"

# 每批最多发送的非空行数（防止超 context）
MAX_LINES_PER_BATCH = 250

# 目录配置
BASE_DIR   = os.path.dirname(os.path.abspath(__file__))
INPUT_DIR  = os.path.join(BASE_DIR, "待清洗文件")
OUTPUT_DIR = os.path.join(BASE_DIR, "已清洗文件")


# ═══════════════════════════════════════════════════════════════
#  文档读取（支持 .docx 和 .doc）
# ═══════════════════════════════════════════════════════════════

def read_docx_lines(path):
    """
    按顺序读取 docx 中所有段落和表格单元格的文本。
    返回 [(原始索引, 文本), ...] 的列表（只含非空行）。
    """
    doc = Document(path)
    lines = []
    idx = 0

    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            # 普通段落：拼接所有 <w:t> 文本
            text = "".join(
                (node.text or "")
                for node in child.iter()
                if node.tag.endswith("}t")
            ).strip()
            idx += 1
            if text:
                lines.append((idx, text))

        elif tag == "tbl":
            # 表格：按行拼接各单元格
            for tr in child.findall(".//{%s}tr" % "http://schemas.openxmlformats.org/wordprocessingml/2006/main"):
                row_parts = []
                for tc in tr.findall(".//{%s}tc" % "http://schemas.openxmlformats.org/wordprocessingml/2006/main"):
                    cell_text = "".join(
                        (node.text or "")
                        for node in tc.iter()
                        if node.tag.endswith("}t")
                    ).strip()
                    if cell_text:
                        row_parts.append(cell_text)
                idx += 1
                if row_parts:
                    lines.append((idx, "　".join(row_parts)))

    return lines


def read_doc_via_wps(doc_path):
    """
    使用 WPS COM 接口读取 .doc 旧格式文件。
    返回 [(原始索引, 文本), ...] 的列表（只含非空行）。
    """
    try:
        import win32com.client
    except ImportError:
        print("   ❌ 缺少 pywin32，无法读取 .doc 文件。请运行: pip install pywin32")
        return None

    wps = None
    doc = None
    lines = []

    try:
        wps = win32com.client.Dispatch("Kwps.Application")
        wps.Visible = False

        # 打开文档
        abs_path = os.path.abspath(doc_path)
        doc = wps.Documents.Open(abs_path)

        # 读取所有段落
        for i, para in enumerate(doc.Paragraphs, 1):
            text = para.Range.Text.strip()
            if text:
                lines.append((i, text))

        # 读取表格内容
        for table in doc.Tables:
            for row in table.Rows:
                row_parts = []
                for cell in row.Cells:
                    text = cell.Range.Text.strip().replace('\r', '').replace('\x07', '')
                    if text:
                        row_parts.append(text)
                if row_parts:
                    lines.append((len(lines) + 1, "　".join(row_parts)))

        return lines

    except Exception as e:
        print(f"   ❌ WPS 读取失败: {e}")
        return None

    finally:
        try:
            if doc:
                doc.Close(SaveChanges=False)
        except:
            pass
        try:
            if wps:
                wps.Quit()
        except:
            pass


def read_document_lines(path):
    """
    智能读取文档，自动判断 .docx 或 .doc 格式。
    返回 [(原始索引, 文本), ...] 的列表。
    """
    ext = os.path.splitext(path)[1].lower()

    if ext == ".docx":
        try:
            return read_docx_lines(path)
        except Exception as e:
            print(f"   ⚠️ python-docx 读取失败，尝试 WPS: {e}")
            return read_doc_via_wps(path)

    elif ext == ".doc":
        return read_doc_via_wps(path)

    else:
        print(f"   ❌ 不支持的文件格式: {ext}")
        return None


# ═══════════════════════════════════════════════════════════════
#  DeepSeek API 调用
# ═══════════════════════════════════════════════════════════════

_EXTRACT_PROMPT_PREFIX = """\
你是专业试卷答案提取工具。请从以下中国高中/初中语文试题答案文档中，提取所有题目的答案和解析。

【提取规则】
1. 选择题：格式为 "1." → "答案：X" → "解析：..."，答案为单个字母A/B/C/D
2. 大题（非选择题）：格式为 "（一）" 章节标题，下面有多个小问
3. 小问识别：每个段落以 "1."、"2." 或 "（1）"、"（2）" 开头的视为独立小问，必须分别提取
   - 重要：保留原始编号（如"1."、"2."）作为答案的一部分
   - 每个小问作为一个独立的数组元素
4. 章节小标题：如 "一、基础巩固层"、"二、能力提升层" 等，保留作为章节标记
5. 解析识别：以【解析】、解析：、【详解】、【分析】等开头的内容视为解析
6. 忽略内容：题干正文、A./B./C./D. 选项行、页码、学生姓名班级等
7. 答案续行：不以新题号开头、不是选项行、不是解析行的普通文字，追加到前一题答案

【输出格式】
直接输出纯 JSON，不要加代码块标记（```json...```），不要有任何说明文字：
{"questions": [
  {"num": 1, "type": "choice", "answer": "B", "analysis": "解析内容，没有则填空串"},
  {"num": "一、基础巩固层", "type": "section", "answer": "", "analysis": ""},
  {"num": "（一）", "type": "big", "answer": ["（1）答案内容", "（2）答案内容"], "analysis": "（1）解析内容\n（2）解析内容"}
]}

【字段说明】
- num: 题号（整数如1/2/3，或章节标题如"一、基础巩固层"、"（一）"）
- type: "choice"（选择题）、"big"（大题/非选择题）、"section"（章节小标题）
- answer: 选择题填字符串如"A"，大题填字符串数组，章节标题填空串
- analysis: 解析内容（字符串，无解析填空串""）

文档内容（"""


def _build_prompt(text_lines):
    """构建提示词（使用拼接而非 .format()，避免文档中 { } 干扰）。"""
    text_content = "\n".join(f"[{i}] {t}" for i, t in text_lines)
    n = len(text_lines)
    return _EXTRACT_PROMPT_PREFIX + str(n) + " 行）：\n" + text_content


def call_api(text_lines, retry=3):
    """
    调用 DeepSeek API 提取答案结构。
    text_lines: [(idx, text), ...]
    返回解析后的 dict，失败返回 None。
    """
    api_key = os.getenv("SILICONFLOW_API_KEY") or os.getenv("DEEPSEEK_API_KEY")
    if not api_key:
        print("   ❌ 未检测到 API Key，请先设置环境变量 SILICONFLOW_API_KEY")
        return None

    prompt = _build_prompt(text_lines)

    data = {
        "model": MODEL,
        "messages": [
            {
                "role": "system",
                "content": "你是专业的JSON数据提取程序，只输出纯JSON，不加任何说明文字或代码块标记。",
            },
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.0,
        "max_tokens": 12000,
    }

    req = urllib.request.Request(
        API_URL,
        data=json.dumps(data, ensure_ascii=False).encode("utf-8"),
        headers={
            "Content-Type": "application/json; charset=utf-8",
            "Authorization": f"Bearer {api_key}",
        },
    )

    for attempt in range(retry):
        try:
            print(f"   🤖 正在调用 AI（第 {attempt + 1}/{retry} 次）...")
            with urllib.request.urlopen(req, timeout=180) as resp:
                result = json.loads(resp.read().decode("utf-8"))
                content = result["choices"][0]["message"]["content"]
                print(f"   ✅ AI 返回 {len(content)} 字符")
                return _parse_json_response(content)
        except Exception as e:
            print(f"   ❌ 请求失败 (尝试 {attempt + 1}/{retry}): {e}")
            if attempt < retry - 1:
                time.sleep(3)

    return None


def _fix_json_errors(text):
    """修复常见的 AI JSON 生成错误。"""
    # 修复 1: "num": 6,etype" -> "num": 6, "type"
    text = re.sub(r'",\s*([a-zA-Z_]+)"', r'", "\1"', text)
    # 修复 2: 数字后缺少逗号 "num": 6,etype -> "num": 6,"etype"
    text = re.sub(r'(\d),([a-zA-Z_])', r'\1, "\2', text)
    # 修复 3: 缺少开头的引号 "num":6, -> "num": 6,
    text = re.sub(r'"(\w+)":(\d)', r'"\1": \2', text)
    return text


def _parse_json_response(text):
    """从 AI 返回文本中提取并解析 JSON。"""
    original = text

    # 1. 去掉代码块标记
    m = re.search(r"```(?:json)?\s*(.*?)\s*```", text, re.DOTALL | re.IGNORECASE)
    if m:
        text = m.group(1)

    # 2. 直接解析
    try:
        return json.loads(text)
    except json.JSONDecodeError as e1:
        pass

    # 3. 去掉开头的非 JSON 说明文字，找 { 开始的位置
    start = text.find("{")
    if start > 0:
        text = text[start:]
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            pass

    # 4. 贪婪匹配最外层 JSON 对象
    m = re.search(r"\{.*\}", text, re.DOTALL)
    if m:
        try:
            return json.loads(m.group(0))
        except json.JSONDecodeError:
            pass

    # 5. 尝试修复常见的 JSON 错误
    fixed_text = _fix_json_errors(text)
    try:
        return json.loads(fixed_text)
    except json.JSONDecodeError:
        pass

    # 6. 所有策略失败，打印诊断信息
    print(f"   ❌ JSON 解析失败。原始返回内容（前 500 字）:")
    print("   " + original[:500].replace("\n", "\n   "))
    return None


# ═══════════════════════════════════════════════════════════════
#  输出文档写入（python-docx）
# ═══════════════════════════════════════════════════════════════

def _make_run(para, text):
    """
    向段落添加 run，并统一设置字体：
    宋体（中文）/ Times New Roman（英文），12pt，黑色，不加粗不斜体。
    """
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = False
    run.font.italic = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = "Times New Roman"

    # 设置中文字体
    try:
        rpr = run._r.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rpr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), "宋体")
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
    except Exception:
        pass

    return run


def _add_para(doc, text):
    """添加一个新段落，设置行距和字体，返回段落对象。"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    _make_run(para, text)
    return para


def write_output_docx(questions, output_path):
    """
    将提取的题目列表写入标准格式 docx。

    标准格式：
      选择题:
        N．A　
        解析：...（无解析则写 "解析： " 含一个空格）

      大题:
        N．
        (1)答案内容
        (2)答案内容
        解析：...（无解析则写 "解析： "）

      英语填空题:
        Ⅰ．
        （1）in
        （2）of
        解析：

      英语大题（Ⅳ部分）:
        Ⅳ．
        （1）D
        （2）D
        （3）B
        解析：
        （1）推理判断题...
        （2）细节理解题...

      语文格式:
        一、基础巩固层（章节小标题）
        1．B　
        解析：...
        （一）
        （1）答案内容
        （2）答案内容
        解析：...
    """
    doc = Document()

    # 移除默认空段落
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    for q in questions:
        num      = q.get("num", 0)
        q_type   = q.get("type", "choice")
        answer   = q.get("answer", "")
        analysis = (q.get("analysis") or "").strip()

        if q_type == "section":
            # ── 章节小标题（如"一、基础巩固层"）────────────────
            _add_para(doc, f"{num}")

        elif q_type == "choice":
            # ── 选择题 ──────────────────────────────────────
            # 第1行: N．A　（全角空格结尾，供答案录入定位）
            _add_para(doc, f"{num}.")
            _add_para(doc, f"答案：{answer}")
            _add_para(doc, f"解析：{analysis}" if analysis else "解析：")

        elif q_type == "english_fill":
            # ── 英语填空题（Ⅰ、Ⅱ、Ⅲ）────────────────────────
            # 第1行: Ⅰ．（罗马数字章节）
            _add_para(doc, f"{num}．")

            # 答案部分（小问编号如（1）（2）（3））
            if isinstance(answer, list):
                for line in answer:
                    line = line.strip()
                    if line:
                        _add_para(doc, line)
            elif answer:
                _add_para(doc, str(answer).strip())

            # 解析行（英语填空题通常无解析，但保留格式）
            _add_para(doc, f"解析：{analysis}" if analysis else "解析： ")

        elif q_type == "english_big":
            # ── 英语大题（Ⅳ部分）────────────────────────────
            # 第1行: Ⅳ．
            _add_para(doc, f"{num}．")

            # 答案部分（如（1）D、（2）D、（3）B）
            if isinstance(answer, list):
                for line in answer:
                    line = line.strip()
                    if line:
                        _add_para(doc, line)
            elif answer:
                _add_para(doc, str(answer).strip())

            # 解析行
            _add_para(doc, f"解析：{analysis}" if analysis else "解析： ")

        elif q_type == "big":
            # ── 大题 ────────────────────────────────────────
            # 判断是否为语文大题格式（章节标题如"（一）"）
            is_chinese_big = isinstance(num, str) and num.startswith("（")
            if is_chinese_big:
                # 语文大题格式：直接输出章节标题
                _add_para(doc, f"{num}")
            else:
                # 普通大题格式：N．
                _add_para(doc, f"{num}．")

            # 答案部分
            if isinstance(answer, list):
                # 如果是数组，先合并所有内容（因为AI可能把多个小问合并到一个数组元素）
                full_text = "".join(str(item) for item in answer if item)
                # 语文大题格式：先输出 "答案：" 标题行
                if is_chinese_big:
                    _add_para(doc, "答案：")
                    # 检测并拆分小问（1.、2.、（1）、（2）等格式）
                    # 使用正则匹配小问编号：1.、2.、（1）、（2）等（支持任意位置）
                    parts = re.split(r'(?=\d+\.|（\d+）|\(\d+\))', full_text)
                    filtered_parts = [p.strip() for p in parts if p.strip() and not re.match(r'^\d+\.$', p.strip())]
                    
                    if len(filtered_parts) > 1:
                        # 有多個小问，分别输出
                        for part in filtered_parts:
                            # 将 "1." 格式转换为 "（1）" 格式
                            part = re.sub(r'^(\d+)\.', r'（\1）', part)
                            _add_para(doc, part)
                    else:
                        # 只有一个小问或无法拆分，直接输出
                        _add_para(doc, full_text)
                else:
                    # 普通大题：按行分割，检测小问格式
                    lines = full_text.split('\n')
                    for line in lines:
                        line = line.strip()
                        if line:
                            # 检测是否以数字编号开头（1.、2.、（1）、（2）等）
                            if re.match(r'^(\d+)\.', line):
                                # 将 "1." 转换为 "（1）"
                                line = re.sub(r'^(\d+)\.', r'（\1）', line)
                                _add_para(doc, line)
                            elif re.match(r'^（\d+）', line) or re.match(r'^\(\d+\)', line):
                                _add_para(doc, line)
                            else:
                                # 非编号行，直接输出
                                _add_para(doc, line)
            elif answer:
                ans_str = str(answer).strip()
                # 语文大题格式：先输出 "答案：" 标题行
                if is_chinese_big:
                    _add_para(doc, "答案：")
                    # 直接输出答案内容
                    _add_para(doc, ans_str)
                else:
                    # 普通大题：智能拆分小问标记
                    if "（1）" in ans_str or "(1)" in ans_str or re.search(r'(?:^|\s)1\.', ans_str):
                        # 使用正则拆分小问：匹配（1）、(1)、1.、2. 等格式
                        parts = re.split(r'(?=（\d+）|\(\d+\)|(?:^|\s)\d+\.)', ans_str)
                        for part in parts:
                            part = part.strip()
                            if part:
                                # 将 "1." 格式转换为 "（1）" 格式
                                part = re.sub(r'^(\d+)\.', r'（\1）', part)
                                _add_para(doc, part)
                    else:
                        _add_para(doc, ans_str)

            # 解析行
            _add_para(doc, f"解析：{analysis}" if analysis else "解析： ")

        else:
            # ── 其他类型（默认按大题处理）────────────────────
            # 第1行: N．
            _add_para(doc, f"{num}．")

            # 答案部分
            if isinstance(answer, list):
                for line in answer:
                    line = line.strip()
                    if line:
                        _add_para(doc, line)
            elif answer:
                _add_para(doc, str(answer).strip())

            # 解析行
            _add_para(doc, f"解析：{analysis}" if analysis else "解析： ")

    doc.save(output_path)
    print(f"   ✅ 文档已保存: {os.path.basename(output_path)}")


# ═══════════════════════════════════════════════════════════════
#  单文件处理主流程
# ═══════════════════════════════════════════════════════════════

def process_file(input_path, output_dir):
    """
    处理单个 docx 文件：读取 → AI提取 → 写入标准格式。
    返回 True 表示成功。
    """
    filename = os.path.basename(input_path)
    name, ext = os.path.splitext(filename)
    # 避免重复添加 _已清洗
    if "_已清洗" in name:
        output_filename = filename
    else:
        output_filename = f"{name}_已清洗{ext}"
    output_path = os.path.join(output_dir, output_filename)

    print(f"\n{'=' * 60}")
    print(f"📄 处理文件: {filename}")
    print("=" * 60)

    # ── 1. 读取文档 ───────────────────────────────────────────
    lines = read_document_lines(input_path)
    if lines is None:
        return False

    if not lines:
        print("   ❌ 文档为空或无可读文本")
        return False

    print(f"   📊 共读取 {len(lines)} 行非空内容")

    # 打印前 10 行预览
    print("   📋 内容预览（前10行）:")
    for i, t in lines[:10]:
        print(f"      [{i:3d}] {t[:80]}")
    if len(lines) > 10:
        print(f"      ... 共 {len(lines)} 行")

    # ── 2. 分批调用 AI ────────────────────────────────────────
    all_questions = []
    batches = [
        lines[s : s + MAX_LINES_PER_BATCH]
        for s in range(0, len(lines), MAX_LINES_PER_BATCH)
    ]
    print(f"\n   🔢 共分 {len(batches)} 批发送给 AI")

    for batch_idx, batch in enumerate(batches, 1):
        if len(batches) > 1:
            print(f"\n   📦 第 {batch_idx}/{len(batches)} 批（{len(batch)} 行）")

        result = call_api(batch)
        if not result:
            print(f"   ❌ 批次 {batch_idx} 失败，跳过")
            continue

        questions = result.get("questions", [])
        print(f"   📝 批次 {batch_idx} 提取到 {len(questions)} 道题")
        all_questions.extend(questions)

    if not all_questions:
        print("   ❌ 未提取到任何题目，请检查文档内容或 API Key")
        return False

    # ── 3. 去重（保留原文档顺序，不按题号排序）─────────────────────────────
    seen = set()
    unique_qs = []
    for q in all_questions:
        num = q.get("num", 0)
        # 用题号+类型作为唯一标识
        key = f"{num}_{q.get('type', 'unknown')}"
        if num and key not in seen:
            seen.add(key)
            unique_qs.append(q)

    sorted_qs = unique_qs  # 保持原文档顺序

    choice_count = sum(1 for q in sorted_qs if q.get("type") == "choice")
    big_count    = sum(1 for q in sorted_qs if q.get("type") == "big")
    english_fill_count = sum(1 for q in sorted_qs if q.get("type") == "english_fill")
    english_big_count = sum(1 for q in sorted_qs if q.get("type") == "english_big")
    section_count = sum(1 for q in sorted_qs if q.get("type") == "section")
    print(f"\n   📊 提取统计: 共 {len(sorted_qs)} 道题")
    print(f"      选择题 {choice_count} 道，大题 {big_count} 道，章节 {section_count} 个")
    print(f"      英语填空 {english_fill_count} 道，英语大题 {english_big_count} 道")

    # 打印提取预览
    print("\n   📋 提取结果预览（前 15 题）:")
    for q in sorted_qs[:15]:
        num      = q.get("num", "?")
        q_type   = q.get("type", "?")
        answer   = q.get("answer", "")
        has_ana  = bool((q.get("analysis") or "").strip())
        icon     = "○" if q_type == "choice" else "◎"
        ans_str  = str(answer)[:35] if isinstance(answer, str) else str(answer)[:35]
        ana_mark = " [有解析]" if has_ana else ""
        print(f"      {icon} 第{num:3}题  {ans_str}{ana_mark}")
    if len(sorted_qs) > 15:
        print(f"      ... 共 {len(sorted_qs)} 题")

    # ── 4. 写入输出文档 ───────────────────────────────────────
    os.makedirs(output_dir, exist_ok=True)
    try:
        write_output_docx(sorted_qs, output_path)
    except Exception as e:
        import traceback
        print(f"   ❌ 写入文档失败: {e}")
        traceback.print_exc()
        return False

    return True


# ═══════════════════════════════════════════════════════════════
#  主入口
# ═══════════════════════════════════════════════════════════════

def main():
    print("=" * 60)
    print("🤖 墨痕快刀 · AI 智能格式清洗 v2.0")
    print("   引擎: MiniMax-M2.5 | 无需 WPS | 支持任意答案格式")
    print("=" * 60)

    os.makedirs(INPUT_DIR, exist_ok=True)
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # ── 扫描待清洗文件夹 ──────────────────────────────────────
    raw_files = [
        f for f in os.listdir(INPUT_DIR)
        if f.endswith((".doc", ".docx")) and not f.startswith("~$")
    ]

    if not raw_files:
        print(f"\n📂 待清洗文件夹为空: {INPUT_DIR}")
        print("   请输入文件路径，或将 .docx 文件放入该目录后重新运行")
        manual = input("\n👉 直接输入文件路径（留空退出）: ").strip().strip('"')
        if not manual or not os.path.exists(manual):
            print("❌ 文件不存在，退出")
            return
        target_files = [(manual, os.path.dirname(manual))]
    else:
        print(f"\n📁 发现 {len(raw_files)} 个待处理文件:")
        for i, f in enumerate(raw_files, 1):
            print(f"   [{i}] {f}")
        print(f"\n📤 输出目录: {OUTPUT_DIR}")
        print()
        print("请选择操作:")
        print("  [0] 处理全部文件")
        for i, f in enumerate(raw_files, 1):
            print(f"  [{i}] 仅处理: {f}")
        print()

        choice = input("👉 输入编号（默认 0）: ").strip()

        if choice == "" or choice == "0":
            selected = raw_files
        else:
            try:
                idx = int(choice) - 1
                if 0 <= idx < len(raw_files):
                    selected = [raw_files[idx]]
                else:
                    print("❌ 无效编号")
                    return
            except ValueError:
                print("❌ 无效输入")
                return

        target_files = [
            (os.path.join(INPUT_DIR, f), OUTPUT_DIR) for f in selected
        ]

    # ── 逐文件处理 ────────────────────────────────────────────
    success = 0
    failed  = 0
    for input_path, out_dir in target_files:
        if process_file(input_path, out_dir):
            success += 1
        else:
            failed += 1

    print(f"\n{'=' * 60}")
    print(f"🎉 全部处理完成！成功 {success} 个，失败 {failed} 个")
    print(f"📁 输出文件位于: {OUTPUT_DIR}")
    print("=" * 60)


if __name__ == "__main__":
    main()
