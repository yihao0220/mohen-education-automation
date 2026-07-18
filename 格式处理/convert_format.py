#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
语文答案格式转换工具
将原格式转换为标准格式，并生成Word文档
支持单个文件或整个文件夹批量处理
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn


def set_chinese_font(run, font_name='宋体', font_size=12):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)


def set_paragraph_format(paragraph, line_spacing=1.5, space_before=0, space_after=0):
    """设置段落格式：1.5倍行距"""
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)


def split_mixed_line(line):
    """
    拆分混合行（多题在同一行的情况）
    例如：1.答案：xxx。解析：2.答案：yyy。解析：3.答案：zzz。
    """
    questions = []
    
    # 模式：数字.答案：...解析：（可选）...（后面可能跟下一个题号）
    # 使用前瞻来分割多个题目
    pattern = r'(\d+)\.\s*答案：\s*(.*?)(?=\d+\.\s*答案：|$)'
    matches = list(re.finditer(pattern, line, re.DOTALL))
    
    for i, match in enumerate(matches):
        q_num = match.group(1)
        content = match.group(2).strip()
        
        # 在这个内容里查找解析
        analysis_pattern = r'(.*?)\s*解析：\s*(.*)'
        analysis_match = re.match(analysis_pattern, content, re.DOTALL)
        
        if analysis_match:
            answer = analysis_match.group(1).strip()
            analysis = analysis_match.group(2).strip()
        else:
            answer = content
            analysis = ""
        
        questions.append({
            'num': q_num,
            'answer': answer,
            'analysis': analysis
        })
    
    return questions


def parse_original_format(content):
    """
    解析原格式内容
    支持多种格式：
    1. 紧凑格式：1.答案：D解析：xxx
    2. 分行格式：1.答案：xxx。 和 解析：xxx（分开的行）
    3. 混合格式：多题在同一行
    """
    questions = []
    lines = content.strip().split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        
        # 检查是否是混合行（包含多个题号）
        # 统计行中的题号数量
        q_num_count = len(re.findall(r'\d+\.\s*答案：', line))
        
        if q_num_count > 1:
            # 混合行，需要拆分
            split_questions = split_mixed_line(line)
            questions.extend(split_questions)
            i += 1
        else:
            # 尝试匹配紧凑格式：数字.答案：xxx解析：xxx
            compact_pattern = r'^(\d+)\.\s*答案：\s*(.*?)\s*解析：\s*(.*)$'
            compact_match = re.match(compact_pattern, line, re.DOTALL)
            
            if compact_match:
                # 紧凑格式
                q_num = compact_match.group(1)
                answer = compact_match.group(2).strip()
                analysis = compact_match.group(3).strip()
                
                questions.append({
                    'num': q_num,
                    'answer': answer,
                    'analysis': analysis
                })
                i += 1
            else:
                # 尝试匹配分行格式：数字.答案：xxx
                answer_pattern = r'^(\d+)\.\s*答案：\s*(.+)$'
                answer_match = re.match(answer_pattern, line, re.DOTALL)
                
                if answer_match:
                    q_num = answer_match.group(1)
                    answer = answer_match.group(2).strip()
                    analysis = ""
                    
                    # 查找下一行的解析
                    if i + 1 < len(lines):
                        next_line = lines[i + 1].strip()
                        if next_line.startswith('解析：'):
                            analysis = next_line[3:].strip()  # 去掉"解析："前缀
                            i += 2
                        else:
                            i += 1
                    else:
                        i += 1
                    
                    questions.append({
                        'num': q_num,
                        'answer': answer,
                        'analysis': analysis
                    })
                else:
                    i += 1
    
    return questions


def convert_to_standard_format(questions, section_title="一、积累与运用"):
    """
    转换为标准格式
    标准格式：
    1.
    答案：xxx
    解析：xxx
    """
    result = []
    result.append("")
    result.append(section_title)
    
    for q in questions:
        # 题号行
        result.append(f"{q['num']}.")
        
        # 答案行
        result.append(f"答案：{q['answer']}")
        
        # 解析行
        result.append(f"解析：{q['analysis']}")
    
    return '\n'.join(result)


def create_word_doc(questions, output_path, section_title="一、积累与运用"):
    """
    创建Word文档
    字体：宋体，小四（12pt），1.5倍行距
    """
    doc = Document()
    
    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    
    # 添加空行
    p = doc.add_paragraph()
    set_paragraph_format(p)
    
    # 添加章节标题
    p = doc.add_paragraph()
    run = p.add_run(section_title)
    set_chinese_font(run, '宋体', 12)
    set_paragraph_format(p)
    
    # 添加每道题
    for q in questions:
        # 题号行
        p = doc.add_paragraph()
        run = p.add_run(f"{q['num']}.")
        set_chinese_font(run, '宋体', 12)
        set_paragraph_format(p)
        
        # 答案行
        p = doc.add_paragraph()
        run = p.add_run(f"答案：{q['answer']}")
        set_chinese_font(run, '宋体', 12)
        set_paragraph_format(p)
        
        # 解析行
        p = doc.add_paragraph()
        run = p.add_run(f"解析：{q['analysis']}")
        set_chinese_font(run, '宋体', 12)
        set_paragraph_format(p)
    
    # 保存文档
    doc.save(output_path)
    print(f"Word文档已保存至: {output_path}")


def read_input_file(input_path):
    """读取输入文件内容，支持 .md/.txt 和 .docx"""
    if input_path.endswith('.docx'):
        # 读取 Word 文档
        try:
            doc = Document(input_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return '\n'.join(paragraphs)
        except Exception as e:
            print(f"  错误：无法读取Word文件 - {e}")
            return ""
    else:
        # 读取文本文件
        with open(input_path, 'r', encoding='utf-8') as f:
            return f.read()


def process_single_file(input_path, output_dir, section_title):
    """处理单个文件"""
    # 读取原文件
    content = read_input_file(input_path)
    
    if not content:
        print(f"  警告：文件内容为空或读取失败，跳过")
        return False
    
    # 获取文件名（不含扩展名）
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    
    # 构建输出路径
    output_docx = os.path.join(output_dir, f"{base_name}_转换后.docx")
    
    print(f"\n处理文件: {input_path}")
    print(f"  输出DOCX: {output_docx}")
    
    # 解析原格式
    questions = parse_original_format(content)
    print(f"  共解析到 {len(questions)} 道题目")
    
    if not questions:
        print(f"  警告：未解析到任何题目，跳过")
        return False
    
    # 创建Word文档
    create_word_doc(questions, output_docx, section_title)
    
    return True


def main():
    # ============================================
    # 用户配置区域 - 请修改以下路径
    # ============================================
    
    # 输入路径（可以是单个文件或文件夹）
    # 示例1：单个文件
    # input_path = r'D:\墨痕教育题目\郴州金海-初一语文-王逸豪\原格式.md'
    # 示例2：文件夹（会处理文件夹内所有.md和.txt文件）
    input_path = r'D:\墨痕教育题目\郴州金海-初一语文-王逸豪\原格式'
    
    # 输出文件夹路径（留空则使用输入文件所在文件夹）
    output_dir = r'D:\墨痕教育题目\郴州金海-初一语文-王逸豪\输出文件'  # 例如：r'D:\输出文件夹'
    
    # 章节标题（可选修改）
    section_title = "一、积累与运用"
    
    # ============================================
    # 配置结束
    # ============================================
    
    # 检查输入路径是否存在
    if not os.path.exists(input_path):
        print(f"错误：输入路径不存在 - {input_path}")
        print("请修改代码中的 input_path 路径")
        return
    
    print("=" * 50)
    print("开始格式转换...")
    print("=" * 50)
    
    # 判断是文件还是文件夹
    if os.path.isfile(input_path):
        # 单个文件处理
        if not output_dir:
            output_dir = os.path.dirname(input_path) or '.'
        os.makedirs(output_dir, exist_ok=True)
        process_single_file(input_path, output_dir, section_title)
        
    elif os.path.isdir(input_path):
        # 文件夹批量处理
        if not output_dir:
            output_dir = input_path
        os.makedirs(output_dir, exist_ok=True)
        
        # 获取所有 .md, .txt, .docx 文件（排除临时文件）
        files = [f for f in os.listdir(input_path) 
                 if f.endswith(('.md', '.txt', '.docx')) 
                 and not f.startswith('~$')  # 排除Word临时文件
                 and os.path.isfile(os.path.join(input_path, f))]
        
        if not files:
            print(f"错误：文件夹中没有找到 .md, .txt 或 .docx 文件 - {input_path}")
            return
        
        print(f"\n找到 {len(files)} 个文件待处理")
        
        success_count = 0
        for filename in files:
            file_path = os.path.join(input_path, filename)
            if process_single_file(file_path, output_dir, section_title):
                success_count += 1
        
        print(f"\n{'=' * 50}")
        print(f"批量处理完成！成功: {success_count}/{len(files)}")
        print(f"输出文件夹: {output_dir}")
    
    print("=" * 50)


if __name__ == '__main__':
    main()
