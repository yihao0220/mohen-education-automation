# coding: utf-8
"""
格式处理入口 - 渐进式模板化清洗

工作流程:
1. 扫描 待清洗文件/ 目录，列出所有 .doc/.docx 文件
2. 识别文档格式，匹配最佳模板
3. 调用模板进行清洗
4. 输出到 已清洗文件/ 目录

支持的模板:
- 模板A: 方括号解析格式 (template_a.py)
- 模板B: 选项解析混合格式 (template_b.py)
- 模板C: 道法【答案】【解析】格式 (template_c.py)
"""

import io
import os
import sys
import time
import re
import zipfile

from docx import Document

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)

# 处理导入路径，支持直接运行和模块运行两种方式
try:
    # 尝试相对导入（作为模块运行时）
    from .common import get_active_wps, debug_log
    from .格式模板库 import template_a, template_b, template_c, template_d, template_e, template_chinese, template_math, template_future_physics, template_future_history, template_nancheng_math
except ImportError:
    # 相对导入失败时，使用绝对导入（直接运行时）
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    from common import get_active_wps, debug_log
    from 格式模板库 import template_a, template_b, template_c, template_d, template_e, template_chinese, template_math, template_future_physics, template_future_history, template_nancheng_math

from shared_core import (
    build_answer_units_from_docx,
    build_question_units_from_docx,
    build_review_report,
    export_review_report,
    map_answers,
)


# 注册所有可用模板
TEMPLATES = [
    template_a,
    template_b,
    template_c,
    template_d,
    template_e,
    template_future_physics,
    template_future_history,
    template_math,
    template_nancheng_math,
    template_chinese,
]


def _should_align_cleaned_output(template) -> bool:
    rich_text_templates = {
        "安乡金海 - 初二 - 数学",
        "未来高二 - 物理总答案",
    }
    return template.TEMPLATE_FEATURES.get("name") not in rich_text_templates


def _ensure_utf8_stdout():
    current_stdout = getattr(sys, "stdout", None)
    buffer = getattr(current_stdout, "buffer", None)
    encoding = getattr(current_stdout, "encoding", "") or ""
    if buffer is not None and encoding.lower() != "utf-8":
        sys.stdout = io.TextIOWrapper(buffer, encoding="utf-8")


def _is_valid_docx_file(path):
    return not str(path).lower().endswith(".docx") or zipfile.is_zipfile(path)


def _extract_numeric_tokens(name):
    return re.findall(r"\d+", name or "")


def _find_matching_question_doc(answer_filename, question_dir):
    if not question_dir or not os.path.isdir(question_dir):
        return None

    answer_tokens = _extract_numeric_tokens(os.path.splitext(answer_filename)[0])
    candidates = []
    for filename in os.listdir(question_dir):
        if not filename.endswith((".doc", ".docx")):
            continue
        if "答案" in filename:
            continue
        score = 0
        question_tokens = _extract_numeric_tokens(os.path.splitext(filename)[0])
        if answer_tokens and question_tokens:
            score = len(set(answer_tokens) & set(question_tokens))
        candidates.append((score, filename))

    if not candidates:
        return None

    candidates.sort(key=lambda item: item[0], reverse=True)
    if candidates[0][0] == 0:
        return None
    return os.path.join(question_dir, candidates[0][1])


def _generate_review_report_if_possible(question_doc_path, answer_doc_path):
    if not question_doc_path or not os.path.exists(question_doc_path):
        return None, None
    if not answer_doc_path or not os.path.exists(answer_doc_path):
        return None, None

    try:
        question_units = build_question_units_from_docx(question_doc_path)
        raw_answer_units = build_answer_units_from_docx(answer_doc_path)
        answer_units = map_answers(question_units, raw_answer_units)
    except ValueError as exc:
        print(f"⚠️ 审核清单跳过: {exc}")
        return None, None

    report = build_review_report(os.path.basename(answer_doc_path), question_units, answer_units)
    base_name = os.path.splitext(answer_doc_path)[0]
    report_path = f"{base_name}_审核清单.md"
    export_review_report(report, report_path)
    return report_path, report


def _has_blocking_review_issues(report) -> bool:
    return bool(report and any(issue.severity == "error" for issue in report.issues))


def _print_review_gate_status(report_path, report):
    if not report_path:
        return
    print(f"📝 已生成审核清单: {os.path.basename(report_path)}")
    if _has_blocking_review_issues(report):
        print("⚠️ 审核未通过：检测到高风险问题，请先处理审核清单，不建议直接录入。")
    else:
        print("✅ 审核通过：未发现高风险问题，可继续进入答案录入。")


def _build_question_answer_alignment(question_doc_path, answer_doc_path):
    if not question_doc_path or not os.path.exists(question_doc_path):
        return None
    if not answer_doc_path or not os.path.exists(answer_doc_path):
        return None
    if not str(question_doc_path).lower().endswith(".docx"):
        print("⚠️ 题目对齐跳过: 题目文档不是 .docx，暂不支持结构对齐")
        return None
    if not str(answer_doc_path).lower().endswith(".docx"):
        print("⚠️ 题目对齐跳过: 清洗输出不是 .docx，暂不支持结构重写")
        return None

    try:
        question_units = build_question_units_from_docx(question_doc_path)
        raw_answer_units = build_answer_units_from_docx(answer_doc_path)
    except ValueError as exc:
        print(f"⚠️ 题目对齐跳过: {exc}")
        return None

    return question_units, map_answers(question_units, raw_answer_units)


def _clear_docx_body(doc_obj):
    body = doc_obj._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def _render_unit_item_text(unit, item, fallback_index):
    marker = str(getattr(item, "item_id", "") or "").strip()
    text = str(getattr(item, "text", "") or "").strip()
    if not marker or marker == unit.question_id:
        return text
    return f"{marker}{text}"


def _rewrite_standardized_docx(answer_doc_path, answer_units):
    doc_obj = Document(answer_doc_path)
    _clear_docx_body(doc_obj)

    for unit in answer_units:
        if unit.answer_mode == "subquestion":
            if unit.answer_items:
                for idx, item in enumerate(unit.answer_items, 1):
                    line = _render_unit_item_text(unit, item, idx)
                    if idx == 1:
                        doc_obj.add_paragraph(f"{unit.question_id}．{line}")
                    else:
                        doc_obj.add_paragraph(line)
            else:
                doc_obj.add_paragraph(f"{unit.question_id}．")

            if unit.analysis_items:
                first_analysis = _render_unit_item_text(unit, unit.analysis_items[0], 1)
                doc_obj.add_paragraph(f"解析：{first_analysis}")
                for idx, item in enumerate(unit.analysis_items[1:], 2):
                    doc_obj.add_paragraph(_render_unit_item_text(unit, item, idx))
            else:
                doc_obj.add_paragraph("解析：")
            continue

        answer_text = " ".join(
            _render_unit_item_text(unit, item, idx)
            for idx, item in enumerate(unit.answer_items, 1)
            if _render_unit_item_text(unit, item, idx)
        ).strip()
        doc_obj.add_paragraph(f"{unit.question_id}．{answer_text}")

        analysis_text = " ".join(
            _render_unit_item_text(unit, item, idx)
            for idx, item in enumerate(unit.analysis_items, 1)
            if _render_unit_item_text(unit, item, idx)
        ).strip()
        doc_obj.add_paragraph(f"解析：{analysis_text}" if analysis_text else "解析：")

    doc_obj.save(answer_doc_path)


def _align_cleaned_doc_with_question_doc(question_doc_path, answer_doc_path):
    aligned_payload = _build_question_answer_alignment(question_doc_path, answer_doc_path)
    if not aligned_payload:
        return None

    _, aligned_answer_units = aligned_payload
    _rewrite_standardized_docx(answer_doc_path, aligned_answer_units)
    print("已按题目结构重写标准答案")
    return aligned_payload


def _collect_nonempty_paragraph_texts(doc):
    """
    一次性读取文档非空段落文本，减少模板匹配阶段的重复 COM 扫描。

    Returns:
        list[str]: 非空段落文本列表
    """
    texts = []
    paras = doc.Paragraphs
    for i in range(1, paras.Count + 1):
        try:
            text = paras(i).Range.Text.strip()
        except Exception:
            continue
        if text:
            texts.append(text)
    return texts


def get_input_source():
    """
    获取输入源
    优先级: 1. WPS活动文档 2. 待清洗文件目录
    
    Returns:
        tuple: (输入类型, 文件路径或文档对象列表, 输入目录)
               返回 (None, None, None) 表示未找到输入源
    """
    # 首先检查 WPS 活动文档
    wps = get_active_wps()
    if wps:
        try:
            doc = wps.ActiveDocument
            doc_name = doc.Name
            if doc_name.endswith(('.doc', '.docx')):
                print(f"✅ 检测到 WPS 活动文档: {doc_name}")
                return 'wps', [doc], None
        except Exception:
            pass
    
    # 其次扫描待清洗文件目录
    input_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), '待清洗文件')
    if os.path.exists(input_dir):
        files = [f for f in os.listdir(input_dir) if f.endswith(('.doc', '.docx'))]
        if files:
            print(f"📁 发现 {len(files)} 个待清洗文件")
            return 'folder', files, input_dir
    
    return None, None, None


def match_best_template(doc, interactive=True, doc_name=""):
    """
    为文档匹配最佳模板，低置信度时询问用户
    
    改进的匹配策略：
    1. 计算基础匹配得分
    2. 检测排他性特征（某个模板独有的标记）
    3. 根据排他性特征调整得分，避免模板B/C混淆
    4. 文件名辅助判断（当文档特征不明确时）
    
    Args:
        doc: WPS 文档对象
        interactive: 是否启用低置信度询问
        doc_name: 文档文件名（用于辅助判断）
    
    Returns:
        (module, score): 最佳匹配的模板模块和得分，无匹配则返回 (None, 0)
    """
    import re
    
    # 收集文档文本用于排他性特征检测
    all_texts = _collect_nonempty_paragraph_texts(doc)
    doc_text = '\n'.join(all_texts[:100])
    
    # 排他性特征定义
    # 模板B的独有特征（如果文档有这些，就不应该是模板C）
    TEMPLATE_B_EXCLUSIVE = [
        r'【详解】',           # 模板B用【详解】，模板C用【解析】 - 最强区分特征
        r'故选[A-D]。',        # 模板B结尾：故选A。/故选B。
    ]
    # 模板C的独有特征（如果文档有这些，就不应该是模板B）
    TEMPLATE_C_EXCLUSIVE = [
        r'【解析】',           # 模板C用【解析】，模板B用【详解】 - 最强区分特征
        r'故本题选[A-D]',      # 模板C结尾：故本题选A（无句号）
        r'考点考查',           # 道法文档特有
        r'能力考查',           # 道法文档特有
        r'核心素养',           # 道法文档特有
    ]
    
    # 文件名辅助关键词（政治/道法相关）
    POLITICAL_KEYWORDS = [
        '青春', '男生女生', '思想', '人文', '美德', 
        '法律', '民法', '犯罪', '刑罚', '律己', '违法'
    ]
    
    # 检测排他性特征
    has_b_exclusive = any(re.search(p, doc_text) for p in TEMPLATE_B_EXCLUSIVE)
    has_c_exclusive = any(re.search(p, doc_text) for p in TEMPLATE_C_EXCLUSIVE)
    
    # 文件名辅助判断
    doc_name_lower = doc_name.lower()
    has_political_keyword = any(kw in doc_name_lower for kw in POLITICAL_KEYWORDS)
    
    scores = []
    
    for template in TEMPLATES:
        base_score = template.match_score(doc, cached_texts=all_texts)
        threshold = template.TEMPLATE_FEATURES.get("match_threshold", 0.03)
        name = template.TEMPLATE_FEATURES['name']
        
        # 根据排他性特征调整得分
        adjusted_score = base_score
        template_module_name = template.__name__ if hasattr(template, '__name__') else str(template)
        
        if 'template_b' in template_module_name:
            # 模板B：如果有C的独有特征，大幅降分
            if has_c_exclusive:
                adjusted_score = base_score * 0.1
                print(f"   📊 {name}: 基础匹配度 {base_score:.2%} → 调整后 {adjusted_score:.2%} (检测到模板C特征)")
            # 如果文件名有政治关键词但没有B的排他特征，降分
            elif has_political_keyword and not has_b_exclusive:
                adjusted_score = base_score * 0.3
                print(f"   📊 {name}: 基础匹配度 {base_score:.2%} → 调整后 {adjusted_score:.2%} (文件名含政治关键词)")
            else:
                print(f"   📊 {name}: 匹配度 {adjusted_score:.2%} (阈值 {threshold:.2%})")
        elif 'template_c' in template_module_name:
            # 模板C：如果有B的独有特征，大幅降分
            if has_b_exclusive:
                adjusted_score = base_score * 0.1
                print(f"   📊 {name}: 基础匹配度 {base_score:.2%} → 调整后 {adjusted_score:.2%} (检测到模板B特征)")
            # 如果文件名有政治关键词，加分
            elif has_political_keyword:
                adjusted_score = min(base_score * 1.5, 1.0)  # 最高不超过100%
                print(f"   📊 {name}: 基础匹配度 {base_score:.2%} → 调整后 {adjusted_score:.2%} (文件名含政治关键词)")
            # 兜底：没有排他性特征但文件名有政治关键词，给基础分
            elif not has_b_exclusive and not has_c_exclusive and has_political_keyword:
                adjusted_score = base_score
                print(f"   📊 {name}: 匹配度 {adjusted_score:.2%} (无排他特征，文件名辅助)")
            else:
                print(f"   📊 {name}: 匹配度 {adjusted_score:.2%} (阈值 {threshold:.2%})")
        else:
            print(f"   📊 {name}: 匹配度 {adjusted_score:.2%} (阈值 {threshold:.2%})")
        
        scores.append((template, adjusted_score, name))
    
    # 按得分排序
    scores.sort(key=lambda x: x[1], reverse=True)
    
    if not scores or scores[0][1] == 0:
        return None, 0
    
    # 平局决胜：如果前两名得分相同（差距小于0.001），且有政治关键词，优先选模板C
    if len(scores) >= 2:
        first_score = scores[0][1]
        second_score = scores[1][1]
        if abs(first_score - second_score) < 0.001 and has_political_keyword:
            # 检查第二名是否是模板C
            second_module_name = scores[1][0].__name__ if hasattr(scores[1][0], '__name__') else str(scores[1][0])
            if 'template_c' in second_module_name:
                # 交换第一名和第二名
                scores[0], scores[1] = scores[1], scores[0]
                print(f"   🏆 平局决胜：根据文件名政治关键词，优先选择模板C")
    
    best_template, best_score, best_name = scores[0]
    
    # 输出排他性特征检测结果
    if has_b_exclusive or has_c_exclusive:
        print(f"\n   🔍 排他性特征检测:")
        if has_b_exclusive:
            print(f"      - 检测到模板B独有特征: 【详解】 或 故选X。")
        if has_c_exclusive:
            print(f"      - 检测到模板C独有特征: 故本题选X 或 考点/能力/素养")
    
    # 低置信度检测：与第二名差距小于10%，或最高得分低于15%
    if interactive and len(scores) >= 2:
        second_score = scores[1][1]
        gap = best_score - second_score
        
        if gap < 0.10 or best_score < 0.15:
            print(f"\n⚠️  自动匹配置信度较低（差距 {gap:.2%}，最高 {best_score:.2%}）")
            print("   请手动选择模板：")
            for idx, (t, s, n) in enumerate(scores[:3], 1):
                marker = " ← 推荐" if idx == 1 else ""
                print(f"   [{idx}] {n} ({s:.2%}){marker}")
            print(f"   [0] 取消清洗")
            
            while True:
                choice = input("\n👉 请选择模板编号: ").strip()
                if choice == '0':
                    return None, 0
                try:
                    idx = int(choice) - 1
                    if 0 <= idx < len(scores[:3]):
                        return scores[idx][0], scores[idx][1]
                except ValueError:
                    pass
                print("   ❌ 无效输入，请重新选择")
    
    return best_template, best_score


def process_wps_document(doc, wps, question_doc_path=None):
    """处理 WPS 活动文档，清洗后另存为新文档（保存在原文档相同目录）"""
    original_name = doc.Name
    original_path = doc.Path
    print(f"\n📄 处理文档: {original_name}")
    print(f"📁 原文档路径: {original_path}")
    print("-" * 60)
    
    # 匹配最佳模板
    print("🔍 识别文档格式...")
    template, score = match_best_template(doc, doc_name=original_name)
    
    if not template:
        print("❌ 未能识别文档格式，无法匹配任何模板")
        return False
    
    print(f"✅ 匹配模板: {template.TEMPLATE_FEATURES['name']} (置信度: {score:.2%})")
    
    # 检查模板是否支持章节选择（模板A）
    if hasattr(template, 'scan_sections'):
        # 先获取章节信息
        sections_info = template.clean_document(doc, return_sections_info=True)
        if isinstance(sections_info, dict) and sections_info.get('selected_sections'):
            selected_sections = sections_info['selected_sections']
            sections = sections_info['sections']
            ranges = sections_info['ranges']
            
            # 如果选择了多个章节，询问是否分别保存
            if len(selected_sections) > 1:
                print(f"\n💡 检测到选择了 {len(selected_sections)} 个章节")
                save_mode = input("   是否将每个章节保存为单独文件？(y/n，默认n): ").strip().lower()
                
                if save_mode == 'y':
                    # 逐个章节清洗并保存
                    for i, (idx, (start, end)) in enumerate(zip(selected_sections, ranges), 1):
                        section_title = sections[idx][1][:20]  # 章节标题前20字
                        print(f"\n{'='*60}")
                        print(f"📄 清洗章节 {i}/{len(selected_sections)}: {section_title}")
                        print(f"{'='*60}")
                        
                        # 重新加载文档（因为每次清洗会修改文档）
                        doc = wps.ActiveDocument
                        
                        # 清洗单个章节
                        template.clean_document(doc, selected_sections=[idx])
                        
                        # 保存，带章节序号后缀
                        output_dir = original_path
                        name, ext = os.path.splitext(original_name)
                        if '_已清洗' not in name:
                            name = f"{name}_已清洗"
                        output_filename = f"{name}{i}{ext}"
                        output_path = os.path.join(output_dir, output_filename)
                        
                        try:
                            if os.path.exists(output_path):
                                os.remove(output_path)
                            doc.SaveAs(output_path)
                            print(f"✅ 已保存: {output_filename}")
                        except Exception as e:
                            print(f"⚠️ 保存出错: {e}")
                            continue
                    
                    print(f"\n🎉 所有章节清洗完成！共保存 {len(selected_sections)} 个文件")
                    return True
    
    # 执行普通清洗（单文件模式）
    print("\n🧹 开始清洗...")
    result = template.clean_document(doc)
    
    if result:
        # 准备输出路径（与原文档相同目录）
        output_dir = original_path
        name, ext = os.path.splitext(original_name)
        output_ext = ".docx" if ext.lower() == ".doc" else ext
        # 如果原文档名已经包含_已清洗，则不再添加
        if '_已清洗' in name:
            output_filename = f"{name}{output_ext}"
        else:
            output_filename = f"{name}_已清洗{output_ext}"
        output_path = os.path.join(output_dir, output_filename)
        
        print(f"\n💾 保存清洗后的文档...")
        try:
            # 如果文件已存在，先删除
            if os.path.exists(output_path):
                os.remove(output_path)
            
            # 另存为新文档
            if output_path.lower().endswith(".docx"):
                try:
                    doc.SaveAs2(output_path, FileFormat=16)
                except Exception:
                    doc.SaveAs(output_path)
            else:
                doc.SaveAs(output_path)
            print(f"✅ 已保存到: {output_filename}")
            print(f"📁 输出目录: {output_dir}")
            if _should_align_cleaned_output(template):
                aligned_payload = _align_cleaned_doc_with_question_doc(question_doc_path, output_path)
                if aligned_payload:
                    try:
                        doc.Close(SaveChanges=False)
                        reopened_doc = wps.Documents.Open(output_path)
                        reopened_doc.Activate()
                        print("已重新打开按题目结构重写后的标准答案文档")
                    except Exception as reopen_exc:
                        print(f"⚠️ 重载标准答案文档失败: {reopen_exc}")
            else:
                print("⚠️ 当前模板检测到公式/图片承载需求，已跳过题答重写，避免清洗后富文本丢失。")
            report_path, report = _generate_review_report_if_possible(question_doc_path, output_path)
            _print_review_gate_status(report_path, report)
            
            # 询问是否关闭原文档
            print("\n📝 原文档仍保持打开状态")
            
        except Exception as e:
            print(f"⚠️ 保存文档时出错: {e}")
            return False
    
    return result


def process_folder_files(files, input_dir, question_dir=None):
    """处理文件夹中的文件：先全部打开，再逐一清洗保存"""
    # 获取 WPS 应用
    wps = get_active_wps()
    if not wps:
        print("❌ 未找到运行中的 WPS，无法处理文件")
        return False

    # 过滤掉 ~$ 临时锁定文件
    valid_files = [f for f in files if not f.startswith('~$')]
    skipped = len(files) - len(valid_files)
    if skipped:
        print(f"⚠️  跳过 {skipped} 个临时锁定文件（~$开头）")

    if not valid_files:
        print("❌ 没有有效文件可处理")
        return False

    # ── 收集 WPS 中已打开的文档（文件名 → doc 对象）─────────────────
    already_open = {}
    try:
        for d in wps.Documents:
            already_open[d.Name] = d
    except Exception:
        pass

    # ── 第一阶段：让 Windows 直接打开文件（比 COM 稳定）─────────────
    print(f"\n📂 第一阶段：打开全部 {len(valid_files)} 个文件...\n")
    need_open = []
    processable_files = []
    for i, filename in enumerate(valid_files, 1):
        if filename in already_open:
            print(f"   ♻️  ({i}/{len(valid_files)}) 已在WPS中: {filename}")
            processable_files.append(filename)
        else:
            input_path = os.path.join(input_dir, filename)
            if not _is_valid_docx_file(input_path):
                print(f"   ⚠️  ({i}/{len(valid_files)}) 无效 docx，占位文件已跳过: {filename}")
                continue
            os.startfile(input_path)          # 让 Windows 用默认程序（WPS）打开
            need_open.append(filename)
            processable_files.append(filename)
            print(f"   📂 ({i}/{len(valid_files)}) 已发送打开指令: {filename}")
            time.sleep(0.3)                   # 稍作间隔，避免 WPS 被淹没

    # 等待所有文件出现在 WPS 文档列表中（最多 120 秒）
    if need_open:
        print(f"\n⏳ 等待 {len(need_open)} 个文档在 WPS 中加载完成...")
        target_names = set(processable_files)
        timeout = 120
        elapsed = 0
        while elapsed < timeout:
            try:
                loaded = {d.Name for d in wps.Documents}
            except Exception:
                loaded = set()
            remaining = target_names - loaded
            if not remaining:
                break
            print(f"   还未就绪: {len(remaining)} 个，继续等待...")
            time.sleep(3)
            elapsed += 3
        else:
            print("⚠️  等待超时，部分文档可能未加载成功，将处理已就绪的文件")

    # 收集实际已打开的文档
    opened_docs = []
    try:
        doc_map = {d.Name: d for d in wps.Documents}
    except Exception:
        doc_map = {}
    for filename in processable_files:
        if filename in doc_map:
            opened_docs.append((doc_map[filename], filename))
        else:
            print(f"   ⚠️  未能加载: {filename}，跳过")

    print(f"\n✅ 就绪 {len(opened_docs)} 个文件，开始清洗...\n")

    # ── 第二阶段：逐一清洗 ─────────────────────────────────────────
    processed = 0
    failed = 0

    for idx, (doc, filename) in enumerate(opened_docs, 1):
        print(f"\n{'=' * 60}")
        print(f"🧹 清洗文件 ({idx}/{len(opened_docs)}): {filename}")
        print(f"📁 输出目录: {input_dir}")
        print("=" * 60)

        try:
            # 匹配模板
            print("🔍 识别文档格式...")
            template, score = match_best_template(doc, interactive=False, doc_name=filename)

            if not template:
                print(f"❌ 未能识别格式，跳过: {filename}")
                doc.Close(SaveChanges=False)
                failed += 1
                continue

            print(f"✅ 匹配模板: {template.TEMPLATE_FEATURES['name']} (置信度: {score:.2%})")

            # 执行清洗
            print("\n🧹 开始清洗...")
            if template.clean_document(doc):
                name, ext = os.path.splitext(filename)
                output_ext = ".docx" if ext.lower() == ".doc" else ext
                output_filename = f"{name}_已清洗{output_ext}"
                output_path = os.path.join(input_dir, output_filename)

                if os.path.exists(output_path):
                    os.remove(output_path)

                if output_path.lower().endswith(".docx"):
                    try:
                        doc.SaveAs2(output_path, FileFormat=16)
                    except Exception:
                        doc.SaveAs(output_path)
                else:
                    doc.SaveAs(output_path)
                doc.Close(SaveChanges=False)
                print(f"\n✅ 清洗完成: {output_filename}")
                question_doc_path = _find_matching_question_doc(filename, question_dir)
                if _should_align_cleaned_output(template):
                    _align_cleaned_doc_with_question_doc(question_doc_path, output_path)
                else:
                    print("⚠️ 当前模板检测到公式/图片承载需求，已跳过题答重写，避免清洗后富文本丢失。")
                report_path, report = _generate_review_report_if_possible(question_doc_path, output_path)
                _print_review_gate_status(report_path, report)
                processed += 1
            else:
                print(f"\n❌ 清洗失败: {filename}")
                doc.Close(SaveChanges=False)
                failed += 1

        except Exception as e:
            print(f"\n❌ 处理出错: {filename} - {e}")
            try:
                doc.Close(SaveChanges=False)
            except Exception:
                pass
            failed += 1

    print(f"\n{'=' * 60}")
    print(f"📊 批量处理完成: 成功 {processed} 个, 失败 {failed} 个")
    print(f"📁 文件保存在: {input_dir}")
    print("=" * 60)

    return failed == 0


def show_menu(active_doc_name=None):
    """显示主菜单"""
    print("\n" + "=" * 60)
    print("🚀 墨痕快刀 - 答案格式清洗工具 (模板化版 v2.0)")
    print("=" * 60)
    doc_info = f" ({active_doc_name})" if active_doc_name else ""
    print(f"   [1] 📄 清洗 WPS 当前活动文档{doc_info}")
    print("   [2] 📂 批量清洗待清洗文件目录")
    print("   [r] 🔄 刷新/切换活动文档")
    print("   [q] ❌ 退出")
    print("-" * 60)


def main():
    """主入口函数"""
    _ensure_utf8_stdout()
    print("\n" + "=" * 60)
    print("🔌 正在初始化...")
    
    # 检查输入源并缓存当前活动文档名
    source_type, source_data, source_dir = get_input_source()
    active_doc_name = None
    if source_type == 'wps' and source_data:
        try:
            active_doc_name = source_data[0].Name
        except Exception:
            pass
    
    if not source_type:
        print("❌ 未找到输入源:")
        print("   1. 请在 WPS 中打开要清洗的文档")
        print("   2. 或将文件放入 '待清洗文件' 目录")
        input("\n按回车键退出...")
        return
    
    while True:
        show_menu(active_doc_name)
        user_input = input("👉 请选择: ").strip().lower()
        
        if user_input == 'q':
            print("👋 再见！")
            break
        
        if user_input == 'r':
            # 刷新：重新检测 WPS 活动文档
            print("\n🔄 正在刷新活动文档...")
            new_source_type, new_source_data, new_source_dir = get_input_source()
            if new_source_type == 'wps' and new_source_data:
                try:
                    active_doc_name = new_source_data[0].Name
                    print(f"✅ 已切换到活动文档: {active_doc_name}")
                except Exception as e:
                    print(f"⚠️ 获取活动文档失败: {e}")
                    active_doc_name = None
            else:
                print("⚠️ 未检测到 WPS 活动文档")
                active_doc_name = None
            continue
        
        if user_input == '1':
            # 清洗当前活动文档
            wps = get_active_wps()
            if wps:
                try:
                    doc = wps.ActiveDocument
                    question_doc_path = input("📘 如需生成审核清单，请输入对应题目文档路径（回车跳过）: ").strip()
                    process_wps_document(doc, wps, question_doc_path=question_doc_path or None)
                    # 更新缓存的文档名（可能已保存为新文件名）
                    try:
                        active_doc_name = wps.ActiveDocument.Name
                    except Exception:
                        pass
                except Exception as e:
                    print(f"❌ 获取活动文档失败: {e}")
            else:
                print("❌ 未找到运行中的 WPS")
        
        elif user_input == '2':
            # 批量清洗文件夹
            print("\n" + "-" * 60)
            folder_path = input("📁 请输入待清洗文件夹路径: ").strip()
            print("-" * 60)
            
            if not folder_path:
                print("❌ 路径不能为空")
                continue
            
            if not os.path.exists(folder_path):
                print(f"❌ 路径不存在: {folder_path}")
                continue
            
            if not os.path.isdir(folder_path):
                print(f"❌ 不是有效的文件夹: {folder_path}")
                continue
            
            # 扫描文件夹中的doc/docx文件
            files = [f for f in os.listdir(folder_path) if f.endswith(('.doc', '.docx'))]
            if not files:
                print(f"❌ 文件夹中没有找到 .doc 或 .docx 文件")
                continue
            
            print(f"✅ 找到 {len(files)} 个待清洗文件")
            question_dir = input("📘 如需批量生成审核清单，请输入题目文档文件夹路径（回车跳过）: ").strip()
            process_folder_files(files, folder_path, question_dir=question_dir or None)
        
        else:
            print("❌ 无效的输入，请重新选择")


if __name__ == "__main__":
    main()
