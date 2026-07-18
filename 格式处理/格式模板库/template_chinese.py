# coding: utf-8
"""
格式模板 - 语文试卷答案格式

适用文档特征：
- 章节标题: `一、基础巩固层`、`二、能力提升层`、`三、拓展延伸层`
- 子章节: `（一）非连续性文本阅读题`、`（二）阅读理解训练题`
- 选择题格式1: `1.B。A 项 "鲜为人知"...` (题号.答案。解析)
- 选择题格式2: `D 解析：...` (无题号，答案+解析)
- 大题格式: `1.邓稼先具有无私奉献...` (题号.答案内容)

转换规则：
- 保留章节标题和子章节标题
- 选择题: 提取答案和解析，分行输出
- 大题: 提取答案内容，添加空解析行
- 统一重新编号
"""

import re
import os

# ══════════════════════════════════════════════════════════════════════════════
# 模板特征定义
# ══════════════════════════════════════════════════════════════════════════════

TEMPLATE_FEATURES = {
    "name": "语文试卷-答案清洗",
    "patterns": [
        r"^\d+\.[A-D]。",                  # 选择题格式1: 1.B。
        r"^[A-D]\s*解析[：:]",              # 选择题格式2: D 解析：
        r"^\d+\.[A-D]\s*解析[：:]",         # 选择题格式3: 1.D 解析：
        r"^[一二三四五六七八九十]+、",       # 中文章节标题: 一、
        r"^（[一二三四五六七八九十\d]+）",   # 子章节: （一）、（1）
        r"^\d+\.\s*[\u4e00-\u9fa5]",       # 大题: 1.中文内容
    ],
    "match_threshold": 0.05,
}

# ══════════════════════════════════════════════════════════════════════════════
# 垃圾行过滤模式
# ══════════════════════════════════════════════════════════════════════════════

GARBAGE_PATTERNS = [
    r"^\s*$",                              # 空行
    r"^参考答案\s*$",                       # 参考答案标题
    r"^答案[:：]?\s*$",                     # 纯"答案"行
    r"^解析[:：]?\s*$",                     # 纯"解析"行
]

# ══════════════════════════════════════════════════════════════════════════════
# 章节识别模式
# ══════════════════════════════════════════════════════════════════════════════

# 一级章节: 一、基础巩固层
SECTION_LEVEL1_PATTERN = re.compile(r"^([一二三四五六七八九十]+)、\s*(.*)$")

# 二级章节: （一）非连续性文本阅读题
# 注意: 必须有标题文字，不能是纯答案内容
SECTION_LEVEL2_PATTERN = re.compile(r"^[（(]([\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+)[)）]\s*([\u4e00-\u9fa5]{2,}.*)$")

# ══════════════════════════════════════════════════════════════════════════════
# 答案识别模式
# ══════════════════════════════════════════════════════════════════════════════

# 选择题格式1: 1.B。解析内容 或 1.B。A 项...
CHOICE_FORMAT1 = re.compile(r"^(\d+)\.([A-D])。(.*)$")

# 选择题格式2: D 解析：内容 (无题号)
CHOICE_FORMAT2 = re.compile(r"^([A-D])\s*解析[：:]\s*(.*)$")

# 选择题格式3: 1.D 解析：内容
CHOICE_FORMAT3 = re.compile(r"^(\d+)\.([A-D])\s*解析[：:]\s*(.*)$")

# 选择题格式4: 1.D（无解析，纯答案）
CHOICE_FORMAT4 = re.compile(r"^(\d+)\.([A-D])\s*$")

# 选择题格式5: 纯字母答案 A / B / C / D
CHOICE_FORMAT5 = re.compile(r"^([A-D])\s*$")

# 选择题格式6: 字母+括号解析 C（"惟妙惟肖"...） 或 B（解析内容）
CHOICE_FORMAT6 = re.compile(r"^([A-D])[（(](.+)[）)]$")

# 大题格式: 1.答案内容
BIG_QUESTION_PATTERN = re.compile(r"^(\d+)\.\s*(.+)$")

# 小问格式: (1)xxx 或 （1）xxx
SUB_QUESTION_PATTERN = re.compile(r"^[（(](\d+)[)）]\s*(.*)$")

# ══════════════════════════════════════════════════════════════════════════════
# 辅助函数
# ══════════════════════════════════════════════════════════════════════════════

def is_garbage_line(text):
    """判断是否是垃圾行"""
    for pattern in GARBAGE_PATTERNS:
        if re.match(pattern, text):
            return True
    return False


def is_section_title(text):
    """判断是否是章节标题（一级或二级）"""
    if SECTION_LEVEL1_PATTERN.match(text):
        return True
    if SECTION_LEVEL2_PATTERN.match(text):
        return True
    return False


def match_score(doc, cached_texts=None):
    """
    计算文档与此模板的匹配度
    
    Args:
        doc: WPS 文档对象
    
    Returns:
        float: 匹配分数 (0-1)
    """
    total_lines = 0
    matched_lines = 0
    
    texts = cached_texts if cached_texts is not None else [
        p.Range.Text.strip() for p in doc.Paragraphs
    ]

    for text in texts:
        if not text:
            continue
        total_lines += 1
        
        for pattern in TEMPLATE_FEATURES["patterns"]:
            if re.search(pattern, text):
                matched_lines += 1
                break
    
    if total_lines == 0:
        return 0
    return matched_lines / total_lines


def set_font_format(doc):
    """
    设置文档字体格式
    - 小四（12号）
    - 黑色
    - 宋体（中文）/ Times New Roman（英文/数字）
    - 不加粗
    """
    try:
        font = doc.Content.Font
        font.Size = 12
        font.Color = 0
        font.Name = "Times New Roman"
        font.NameFarEast = "宋体"
        font.Bold = False
        font.Italic = False
        print("   ✓ 字体格式设置完成：小四、黑色、宋体/Times New Roman、不加粗")
    except Exception as e:
        print(f"   ! 字体设置失败: {e}")


# ══════════════════════════════════════════════════════════════════════════════
# 核心清洗函数
# ══════════════════════════════════════════════════════════════════════════════

def clean_document(doc):
    """
    清洗文档（WPS COM 对象）
    
    Args:
        doc: WPS 文档对象
    
    Returns:
        bool: 是否成功
    """
    print(f"   ▶ 使用语文模板清洗: {doc.Name}")
    print(f"   📊 文档共 {doc.Paragraphs.Count} 个段落")
    
    # 1. 删除所有图片
    try:
        shapes_count = doc.Shapes.Count
        if shapes_count > 0:
            for i in range(shapes_count, 0, -1):
                doc.Shapes(i).Delete()
            print(f"   ✓ 删除 {shapes_count} 个悬浮图片")
        inline_count = doc.InlineShapes.Count
        if inline_count > 0:
            for i in range(inline_count, 0, -1):
                doc.InlineShapes(i).Delete()
            print(f"   ✓ 删除 {inline_count} 个嵌入图片")
    except Exception as e:
        print(f"   ! 删除图片失败: {e}")
    
    # 2. 遍历段落，提取内容
    paras = doc.Paragraphs
    results = []  # 存储解析结果
    current_section = None  # 当前一级章节
    current_subsection = None  # 当前二级章节
    question_counter = 0  # 章节内题号计数器
    
    print(f"\n   🔍 开始解析文档内容...\n")
    
    i = 1
    while i <= paras.Count:
        p = paras(i)
        text = p.Range.Text.replace("\r", "").replace("\n", "").strip()
        
        if not text:
            i += 1
            continue
        
        debug_text = text[:60] + "..." if len(text) > 60 else text
        
        # 垃圾行过滤
        if is_garbage_line(text):
            print(f"   [行{i:3d}] 🚫 跳过垃圾行: {debug_text}")
            i += 1
            continue
        
        # ── 一级章节标题 ─────────────────────────────────────────
        m1 = SECTION_LEVEL1_PATTERN.match(text)
        if m1:
            section_num = m1.group(1)
            section_title = m1.group(2).strip()
            current_section = f"{section_num}、{section_title}"
            current_subsection = None
            question_counter = 0  # 重置题号计数器
            results.append({
                'type': 'section',
                'level': 1,
                'content': current_section
            })
            print(f"   [行{i:3d}] 📑 一级章节: {current_section}")
            i += 1
            continue
        
        # ── 二级章节标题 ─────────────────────────────────────────
        m2 = SECTION_LEVEL2_PATTERN.match(text)
        if m2:
            subsection_num = m2.group(1)
            subsection_title = m2.group(2).strip()
            current_subsection = f"（{subsection_num}）{subsection_title}"
            question_counter = 0  # 重置题号计数器
            results.append({
                'type': 'section',
                'level': 2,
                'content': current_subsection
            })
            print(f"   [行{i:3d}] 📑 二级章节: {current_subsection}")
            i += 1
            continue
        
        # ── 选择题格式1: 1.B。解析内容 ─────────────────────────────
        c1 = CHOICE_FORMAT1.match(text)
        if c1:
            q_num = int(c1.group(1))
            answer = c1.group(2)
            analysis = c1.group(3).strip()
            question_counter += 1
            results.append({
                'type': 'choice',
                'original_num': q_num,
                'output_num': question_counter,
                'answer': answer,
                'analysis': analysis
            })
            print(f"   [行{i:3d}] ✅ 选择题(格式1): 第{q_num}题 = {answer}")
            i += 1
            continue
        
        # ── 选择题格式3: 1.D 解析：内容 ───────────────────────────
        c3 = CHOICE_FORMAT3.match(text)
        if c3:
            q_num = int(c3.group(1))
            answer = c3.group(2)
            analysis = c3.group(3).strip()
            question_counter += 1
            results.append({
                'type': 'choice',
                'original_num': q_num,
                'output_num': question_counter,
                'answer': answer,
                'analysis': analysis
            })
            print(f"   [行{i:3d}] ✅ 选择题(格式3): 第{q_num}题 = {answer}")
            i += 1
            continue
        
        # ── 选择题格式4: 1.D（纯答案无解析）─────────────────────────
        c4 = CHOICE_FORMAT4.match(text)
        if c4:
            q_num = int(c4.group(1))
            answer = c4.group(2)
            question_counter += 1
            results.append({
                'type': 'choice',
                'original_num': q_num,
                'output_num': question_counter,
                'answer': answer,
                'analysis': ''
            })
            print(f"   [行{i:3d}] ✅ 选择题(格式4): 第{q_num}题 = {answer}")
            i += 1
            continue
        
        # ── 选择题格式2: D 解析：内容（无题号）─────────────────────
        c2 = CHOICE_FORMAT2.match(text)
        if c2:
            answer = c2.group(1)
            analysis = c2.group(2).strip()
            question_counter += 1
            results.append({
                'type': 'choice',
                'original_num': question_counter,
                'output_num': question_counter,
                'answer': answer,
                'analysis': analysis
            })
            print(f"   [行{i:3d}] ✅ 选择题(格式2/无题号): = {answer}")
            i += 1
            continue
        
        # ── 大题格式: 1.答案内容 ────────────────────────────────────
        bq = BIG_QUESTION_PATTERN.match(text)
        if bq:
            q_num = int(bq.group(1))
            content = bq.group(2).strip()
            question_counter += 1
            
            # 收集续行
            answer_lines = [content]
            i += 1
            while i <= paras.Count:
                next_text = paras(i).Range.Text.replace("\r", "").replace("\n", "").strip()
                if not next_text:
                    i += 1
                    continue
                
                # 遇到新题号、章节标题等则停止
                if BIG_QUESTION_PATTERN.match(next_text):
                    break
                if CHOICE_FORMAT1.match(next_text):
                    break
                if CHOICE_FORMAT2.match(next_text):
                    break
                if CHOICE_FORMAT3.match(next_text):
                    break
                if CHOICE_FORMAT4.match(next_text):
                    break
                if SECTION_LEVEL1_PATTERN.match(next_text):
                    break
                if SECTION_LEVEL2_PATTERN.match(next_text):
                    break
                if is_garbage_line(next_text):
                    break
                
                # 检查是否是小问格式 (1)xxx
                sub_m = SUB_QUESTION_PATTERN.match(next_text)
                if sub_m:
                    sub_num = sub_m.group(1)
                    sub_content = sub_m.group(2).strip()
                    answer_lines.append(f"（{sub_num}）{sub_content}")
                    print(f"   [行{i:3d}]    └─ 小问({sub_num}): {sub_content[:40]}")
                else:
                    # 续行内容
                    answer_lines.append(next_text)
                    print(f"   [行{i:3d}]    └─ 续行: {next_text[:40]}")
                i += 1
            
            results.append({
                'type': 'big',
                'original_num': q_num,
                'output_num': question_counter,
                'answer_lines': answer_lines,
                'analysis': ''
            })
            print(f"   [行{i-1:3d}] 📌 大题: 第{q_num}题 ({len(answer_lines)} 行答案)")
            continue
        
        # ── 未识别行 ─────────────────────────────────────────────
        print(f"   [行{i:3d}] ⚪ 未识别: {debug_text}")
        i += 1
    
    # 3. 统计
    print(f"\n   📊 解析完成统计:")
    section_count = sum(1 for r in results if r['type'] == 'section')
    choice_count = sum(1 for r in results if r['type'] == 'choice')
    big_count = sum(1 for r in results if r['type'] == 'big')
    print(f"      - 章节标题: {section_count} 个")
    print(f"      - 选择题: {choice_count} 道")
    print(f"      - 大题: {big_count} 道")
    
    # 4. 清空文档并重新写入
    doc.Content.Text = ""
    
    # 全局题号计数器（跨章节连续编号）
    global_counter = 0
    
    for item in results:
        if item['type'] == 'section':
            # 章节标题直接输出
            doc.Content.InsertAfter(f"{item['content']}\r")
        
        elif item['type'] == 'choice':
            global_counter += 1
            answer = item['answer']
            analysis = item['analysis']
            
            # 输出格式:
            # 题号.
            # 答案：X
            # 解析：内容
            doc.Content.InsertAfter(f"{global_counter}.\r")
            doc.Content.InsertAfter(f"答案：{answer}\r")
            if analysis:
                doc.Content.InsertAfter(f"解析：{analysis}\r")
            else:
                doc.Content.InsertAfter(f"解析：\r")
        
        elif item['type'] == 'big':
            global_counter += 1
            answer_lines = item['answer_lines']
            
            # 输出格式:
            # 题号.
            # 答案：
            # (答案内容，可能多行)
            # 解析：
            doc.Content.InsertAfter(f"{global_counter}.\r")
            doc.Content.InsertAfter(f"答案：\r")
            for line in answer_lines:
                doc.Content.InsertAfter(f"{line}\r")
            doc.Content.InsertAfter(f"解析：\r")
    
    print(f"\n   ✓ 共输出 {global_counter} 道题目，{section_count} 个章节标题")
    
    # 5. 设置字体格式
    set_font_format(doc)
    
    return True


# ══════════════════════════════════════════════════════════════════════════════
# 独立测试函数（使用 python-docx，不依赖 WPS COM）
# ══════════════════════════════════════════════════════════════════════════════

def clean_document_standalone(input_path, output_path=None):
    """
    独立清洗函数（使用 python-docx）
    
    Args:
        input_path: 输入文件路径
        output_path: 输出文件路径（可选，默认在同目录生成_已清洗.docx）
    
    Returns:
        bool: 是否成功
    """
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    
    print(f"\n{'='*60}")
    print(f"🔍 开始清洗: {os.path.basename(input_path)}")
    print(f"{'='*60}")
    
    doc = Document(input_path)
    
    results = []
    question_counter = 0
    in_subsection = False  # 是否在二级章节（阅读大题）中
    subsection_sub_counter = 0  # 二级章节内的小问计数器
    
    for i, para in enumerate(doc.paragraphs, 1):
        text = para.text.strip()
        
        if not text:
            continue
        
        debug_text = text[:60] + "..." if len(text) > 60 else text
        
        # 垃圾行过滤
        if is_garbage_line(text):
            print(f"   [行{i:3d}] 🚫 跳过: {debug_text}")
            continue
        
        # 一级章节：退出二级章节模式
        m1 = SECTION_LEVEL1_PATTERN.match(text)
        if m1:
            section_num = m1.group(1)
            section_title = m1.group(2).strip()
            results.append({
                'type': 'section',
                'level': 1,
                'content': f"{section_num}、{section_title}"
            })
            question_counter = 0
            in_subsection = False  # 退出二级章节模式
            print(f"   [行{i:3d}] 📑 章节: {section_num}、{section_title}")
            continue
        
        # 二级章节：进入阅读大题模式
        m2 = SECTION_LEVEL2_PATTERN.match(text)
        if m2:
            subsection_num = m2.group(1)
            subsection_title = m2.group(2).strip()
            results.append({
                'type': 'subsection',  # 特殊类型：二级章节大题
                'content': f"（{subsection_num}）{subsection_title}",
                'sub_answers': [],  # 小问答案列表
                'analysis': ''
            })
            in_subsection = True  # 进入二级章节模式
            subsection_sub_counter = 0  # 重置小问计数器
            print(f"   [行{i:3d}] 📑 阅读大题: （{subsection_num}）{subsection_title}")
            continue
        
        # 在二级章节中，小题转换为小问
        if in_subsection:
            # 选择题格式1: 1.B。解析
            c1 = CHOICE_FORMAT1.match(text)
            if c1:
                subsection_sub_counter += 1
                answer = c1.group(2)
                analysis = c1.group(3).strip()
                results[-1]['sub_answers'].append({
                    'num': subsection_sub_counter,
                    'answer': answer,
                    'analysis': analysis
                })
                print(f"   [行{i:3d}] ✅ 小问({subsection_sub_counter}): 选择 {answer}")
                continue
            
            # 选择题格式3: 1.D 解析：
            c3 = CHOICE_FORMAT3.match(text)
            if c3:
                subsection_sub_counter += 1
                answer = c3.group(2)
                analysis = c3.group(3).strip()
                results[-1]['sub_answers'].append({
                    'num': subsection_sub_counter,
                    'answer': answer,
                    'analysis': analysis
                })
                print(f"   [行{i:3d}] ✅ 小问({subsection_sub_counter}): 选择 {answer}")
                continue
            
            # 选择题格式4: 1.D
            c4 = CHOICE_FORMAT4.match(text)
            if c4:
                subsection_sub_counter += 1
                answer = c4.group(2)
                results[-1]['sub_answers'].append({
                    'num': subsection_sub_counter,
                    'answer': answer,
                    'analysis': ''
                })
                print(f"   [行{i:3d}] ✅ 小问({subsection_sub_counter}): 选择 {answer}")
                continue
            
            # 大题格式: 1.内容（转换为小问）
            bq = BIG_QUESTION_PATTERN.match(text)
            if bq:
                subsection_sub_counter += 1
                content = bq.group(2).strip()
                results[-1]['sub_answers'].append({
                    'num': subsection_sub_counter,
                    'answer': content,
                    'analysis': ''
                })
                print(f"   [行{i:3d}] 📌 小问({subsection_sub_counter}): {content[:40]}...")
                continue
            
            # 续行处理：追加到上一个小问
            if results[-1]['sub_answers']:
                results[-1]['sub_answers'][-1]['answer'] += '\n' + text
                print(f"   [行{i:3d}]    └─ 续行")
                continue
        
        # 不在二级章节中的正常处理
        # 选择题格式1: 1.B。解析
        c1 = CHOICE_FORMAT1.match(text)
        if c1:
            question_counter += 1
            results.append({
                'type': 'choice',
                'answer': c1.group(2),
                'analysis': c1.group(3).strip()
            })
            print(f"   [行{i:3d}] ✅ 选择题: {c1.group(1)} → {c1.group(2)}")
            continue
        
        # 选择题格式3: 1.D 解析：
        c3 = CHOICE_FORMAT3.match(text)
        if c3:
            question_counter += 1
            results.append({
                'type': 'choice',
                'answer': c3.group(2),
                'analysis': c3.group(3).strip()
            })
            print(f"   [行{i:3d}] ✅ 选择题: {c3.group(1)} → {c3.group(2)}")
            continue
        
        # 选择题格式4: 1.D
        c4 = CHOICE_FORMAT4.match(text)
        if c4:
            question_counter += 1
            results.append({
                'type': 'choice',
                'answer': c4.group(2),
                'analysis': ''
            })
            print(f"   [行{i:3d}] ✅ 选择题: {c4.group(1)} → {c4.group(2)}")
            continue
        
        # 选择题格式2: D 解析：
        c2 = CHOICE_FORMAT2.match(text)
        if c2:
            question_counter += 1
            results.append({
                'type': 'choice',
                'answer': c2.group(1),
                'analysis': c2.group(2).strip()
            })
            print(f"   [行{i:3d}] ✅ 选择题(无题号): → {c2.group(1)}")
            continue
        
        # 选择题格式5: 纯字母答案 A / B / C / D
        c5 = CHOICE_FORMAT5.match(text)
        if c5:
            question_counter += 1
            results.append({
                'type': 'choice',
                'answer': c5.group(1),
                'analysis': ''
            })
            print(f"   [行{i:3d}] ✅ 选择题(纯字母): → {c5.group(1)}")
            continue
        
        # 选择题格式6: 字母+括号解析 C（解析内容）
        c6 = CHOICE_FORMAT6.match(text)
        if c6:
            question_counter += 1
            results.append({
                'type': 'choice',
                'answer': c6.group(1),
                'analysis': c6.group(2).strip()
            })
            print(f"   [行{i:3d}] ✅ 选择题(括号解析): → {c6.group(1)}")
            continue
        
        # 大题格式: 1.内容
        bq = BIG_QUESTION_PATTERN.match(text)
        if bq:
            question_counter += 1
            results.append({
                'type': 'big',
                'answer': bq.group(2).strip(),
                'analysis': '',
                'has_num': True  # 标记有明确题号
            })
            print(f"   [行{i:3d}] 📌 大题: {bq.group(1)}")
            continue
        
        # 续行处理：只对有明确题号的大题收集续行
        if results and results[-1]['type'] == 'big' and results[-1].get('has_num'):
            results[-1]['answer'] += '\n' + text
            print(f"   [行{i:3d}]    └─ 续行")
            continue
        
        # 未识别的行作为大题答案处理（无题号，按顺序编号）
        question_counter += 1
        results.append({
            'type': 'big',
            'answer': text,
            'analysis': '',
            'has_num': False  # 标记无明确题号
        })
        print(f"   [行{i:3d}] 📝 大题(无题号): 第{question_counter}题")
    
    # 生成输出文档
    output_doc = Document()
    
    global_counter = 0
    for item in results:
        if item['type'] == 'section':
            p = output_doc.add_paragraph(item['content'])
        elif item['type'] == 'subsection':
            # 二级章节（阅读大题）：输出章节标题 + 答案：+ 小问列表 + 解析：
            output_doc.add_paragraph(item['content'])
            output_doc.add_paragraph("答案：")
            for sub in item['sub_answers']:
                sub_num = sub['num']
                sub_answer = sub['answer']
                # 输出小问答案
                for idx, line in enumerate(sub_answer.split('\n')):
                    if idx == 0:
                        output_doc.add_paragraph(f"（{sub_num}）{line}")
                    else:
                        output_doc.add_paragraph(line)
            output_doc.add_paragraph("解析：")
        elif item['type'] == 'choice':
            global_counter += 1
            output_doc.add_paragraph(f"{global_counter}.")
            output_doc.add_paragraph(f"答案：{item['answer']}")
            analysis = item['analysis'] if item['analysis'] else ''
            output_doc.add_paragraph(f"解析：{analysis}")
        elif item['type'] == 'big':
            global_counter += 1
            output_doc.add_paragraph(f"{global_counter}.")
            output_doc.add_paragraph(f"答案：")
            for line in item['answer'].split('\n'):
                output_doc.add_paragraph(line)
            output_doc.add_paragraph(f"解析：")
    
    # 设置字体
    for para in output_doc.paragraphs:
        for run in para.runs:
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 保存
    if output_path is None:
        name, ext = os.path.splitext(input_path)
        output_path = f"{name}_已清洗.docx"
    
    output_doc.save(output_path)
    
    print(f"\n{'='*60}")
    print(f"✅ 清洗完成！")
    print(f"   输入: {os.path.basename(input_path)}")
    print(f"   输出: {os.path.basename(output_path)}")
    print(f"   题目数: {global_counter}")
    print(f"{'='*60}")
    
    return True


# 模块导出
TEMPLATE_INFO = TEMPLATE_FEATURES


if __name__ == "__main__":
    # 独立测试
    import sys
    
    if len(sys.argv) > 1:
        input_file = sys.argv[1]
    else:
        # 默认测试文件
        script_dir = os.path.dirname(os.path.abspath(__file__))
        parent_dir = os.path.dirname(script_dir)
        input_file = os.path.join(parent_dir, "待清洗文件", "1《邓稼先》分层作业（含答案）.docx")
    
    if os.path.exists(input_file):
        clean_document_standalone(input_file)
    else:
        print(f"文件不存在: {input_file}")
