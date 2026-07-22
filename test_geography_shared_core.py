from docx import Document
import unittest

from shared_core import (
    build_answer_units_from_docx,
    build_answer_units_from_paragraph_texts,
    build_question_units_from_nodes,
)
from shared_core.models import DocNode


class GeographySharedCoreTests(unittest.TestCase):
    def test_geography_prefixed_question_ids_are_normalized(self):
        nodes = [
            DocNode(index=1, text="【G】1.呈“山河相间、纵列分布”特点的地形区是（　 　）"),
            DocNode(index=2, text="A．中南半岛\tB．阿拉伯半岛\tC．印度半岛\tD．马来半岛"),
            DocNode(index=3, text="Y2.东南亚唯一的内陆国是（     ）"),
            DocNode(index=4, text="A．泰国\tB．缅甸\tC．越南\tD．老挝"),
        ]

        units = build_question_units_from_nodes("限训11：5.7.docx", "文科", nodes)

        self.assertEqual([unit.question_id for unit in units], ["1", "2"])
        self.assertTrue(all(unit.subject_overlay == "geography" for unit in units))

    def test_material_block_uses_first_real_question_id_not_year(self):
        nodes = [
            DocNode(
                index=1,
                text="2025年1月23日13时15分，长征六号改运载火箭在太原卫星发射中心点火升空。下图为太原卫星发射基地位置及地球公转示意图。据此完成下面小题。",
            ),
            DocNode(index=2, text="", has_inline_media=True),
            DocNode(index=3, text="1．长征六号改运载火箭发射地位于（   ）"),
            DocNode(index=4, text="A．琼\tB．晋\tC．甘\tD．蜀"),
            DocNode(index=5, text="2．卫星被准确送入预定轨道时，我国大部分地区（   ）"),
            DocNode(index=6, text="A．春风拂过\tB．气温攀升\tC．金叶飘落\tD．银装素裹"),
        ]

        units = build_question_units_from_nodes("地理限训14.docx", "文科", nodes)

        self.assertEqual(len(units), 1)
        self.assertEqual(units[0].question_id, "1")
        self.assertEqual(units[0].question_type, "material_choice")

    def test_duplicate_material_paragraphs_do_not_create_fake_year_question(self):
        nodes = [
            DocNode(
                index=1,
                text="2025年1月23日13时15分，长征六号改运载火箭在太原卫星发射中心点火升空。下图为太原卫星发射基地位置及地球公转示意图。据此完成下面小题。",
            ),
            DocNode(
                index=2,
                text="2025年1月23日13时15分，长征六号改运载火箭在太原卫星发射中心点火升空。下图为太原卫星发射基地位置及地球公转示意图。据此完成下面小题。",
            ),
            DocNode(index=3, text="1．长征六号改运载火箭发射地位于（   ）"),
            DocNode(index=4, text="A．琼\tB．晋\tC．甘\tD．蜀"),
            DocNode(index=5, text="2．卫星被准确送入预定轨道时，我国大部分地区（   ）"),
            DocNode(index=6, text="A．春风拂过\tB．气温攀升\tC．金叶飘落\tD．银装素裹"),
        ]

        units = build_question_units_from_nodes("地理限训14.docx", "文科", nodes)

        self.assertEqual([unit.question_id for unit in units], ["1"])

    def test_compressed_choice_answers_split_into_multiple_units(self):
        paragraphs = [
            "一、选择题",
            "1.A2.D3.B4.C",
            "【详解】1．第一题解析。",
            "2．第二题解析。",
            "3．第三题解析。",
            "4．第四题解析。",
        ]

        units = build_answer_units_from_paragraph_texts(paragraphs)

        self.assertEqual([unit.question_id for unit in units], ["1", "2", "3", "4"])
        self.assertEqual([unit.answer_items[0].text for unit in units], ["A", "D", "B", "C"])
        self.assertTrue(all(unit.analysis_items for unit in units))

    def test_answer_units_can_be_built_from_table_docx(self):
        import tempfile

        doc = Document()
        table = doc.add_table(rows=2, cols=4)
        table.cell(0, 0).text = "题号"
        table.cell(0, 1).text = "1"
        table.cell(0, 2).text = "2"
        table.cell(0, 3).text = "3"
        table.cell(1, 0).text = "答案"
        table.cell(1, 1).text = "B"
        table.cell(1, 2).text = "D"
        table.cell(1, 3).text = "A"

        with tempfile.TemporaryDirectory() as temp_dir:
            path = f"{temp_dir}\\限训14答案.docx"
            doc.save(path)
            units = build_answer_units_from_docx(path)

        self.assertEqual([unit.question_id for unit in units], ["1", "2", "3"])
        self.assertEqual([unit.answer_items[0].text for unit in units], ["B", "D", "A"])

    def test_invalid_docx_raises_clear_value_error(self):
        import tempfile

        with tempfile.TemporaryDirectory() as temp_dir:
            path = f"{temp_dir}\\损坏答案.docx"
            with open(path, "w", encoding="utf-8") as handle:
                handle.write('{"code":4404,"error":"fail to get resource"}')

            with self.assertRaisesRegex(ValueError, "无效 docx 文件"):
                build_answer_units_from_docx(path)

    def test_section_headings_are_not_merged_into_previous_answer(self):
        paragraphs = [
            "1．A",
            "2．B",
            "12．D",
            "八年级地理学科限时训练（4）参考答案",
            "一、单项选择题（每小题4分，共60分）",
            "二、填空题（每小空2分，共40分）",
            "13. (1) A；冬冷夏热，降水较少，气温年较差大",
            "(2) 沙漠",
        ]

        units = build_answer_units_from_paragraph_texts(paragraphs)

        self.assertEqual(
            [unit.question_id for unit in units],
            ["1", "2", "12", "13．（1）"],
        )
        self.assertEqual(units[2].answer_items[0].text, "D")
        self.assertNotIn("参考答案", units[2].answer_items[0].text)

    def test_docx_prefers_paragraph_answers_over_duplicate_table_answers(self):
        import tempfile

        doc = Document()
        table = doc.add_table(rows=2, cols=3)
        table.cell(0, 0).text = "题号"
        table.cell(0, 1).text = "1"
        table.cell(0, 2).text = "2"
        table.cell(1, 0).text = "答案"
        table.cell(1, 1).text = "B"
        table.cell(1, 2).text = "D"
        doc.add_paragraph("1．A")
        doc.add_paragraph("【解析】1．第一题解析")
        doc.add_paragraph("2．C")
        doc.add_paragraph("【解析】2．第二题解析")
        doc.add_paragraph("3．D")

        with tempfile.TemporaryDirectory() as temp_dir:
            path = f"{temp_dir}\\重复答案.docx"
            doc.save(path)
            units = build_answer_units_from_docx(path)

        self.assertEqual([unit.question_id for unit in units], ["1", "2", "3"])
        self.assertEqual([unit.answer_items[0].text for unit in units], ["A", "C", "D"])
        self.assertTrue(units[0].analysis_items)

    def test_subjective_question_with_subparts_stays_as_one_question_unit(self):
        nodes = [
            DocNode(index=1, text="13.亚洲气候类型多样，各地的气候差异很大。读图，回答下列问题。"),
            DocNode(index=2, text="(1)亚洲气候复杂多样，其中分布最广的气候类型是____。"),
            DocNode(index=3, text="(2)同纬度的E和F分布区，气候差异较大。"),
            DocNode(index=4, text="(3)由于受到地形的影响，形成了高原山地气候。"),
            DocNode(index=5, text="14.非洲气候分布特点如下。"),
        ]

        units = build_question_units_from_nodes("限训13.docx", "文科", nodes)

        self.assertEqual([unit.question_id for unit in units], ["13", "14"])
        self.assertEqual(units[0].source_span, (1, 4))
        self.assertEqual(units[0].subquestions, ["(1)亚洲气候复杂多样，其中分布最广的气候类型是____。", "(2)同纬度的E和F分布区，气候差异较大。", "(3)由于受到地形的影响，形成了高原山地气候。"])

    def test_material_block_recognizes_range_with_xiaoti(self):
        nodes = [
            DocNode(index=1, text="2023年5月，我国科研人员在西藏发现了一批三叠纪喜马拉雅鱼龙化石，这是古生物学上的又一次重大发现。完成3-5小题。"),
            DocNode(index=2, text="3.在西藏发现鱼龙化石，说明（   ）"),
            DocNode(index=3, text="A．鱼龙曾经生活在陆地\tB．此地曾经是一片汪洋"),
            DocNode(index=4, text="4.能反映西藏这种地表形态变化的词语是（   ）"),
            DocNode(index=5, text="A．愚公移山\tB．刻舟求剑"),
            DocNode(index=6, text="5.造成西藏这种地表形态变化的根本原因是（   ）"),
            DocNode(index=7, text="A．地壳运动\tB．季节变化"),
            DocNode(index=8, text="下图为世界部分区域海陆分布示意图，据图完成6-8小题。"),
            DocNode(index=9, text="6.亚马孙河是世界上流量最大的河流，刚果河是世界第二大河，两河都注入（   ）"),
            DocNode(index=10, text="A．太平洋\tB．大西洋"),
            DocNode(index=11, text="7.甲地理事物为（   ）"),
            DocNode(index=12, text="A．大陆\tB．半岛"),
            DocNode(index=13, text="8.①山脉的说法错误的是（   ）"),
            DocNode(index=14, text="A．世界上最长的山脉\tB．是板块碰撞挤压形成"),
        ]

        units = build_question_units_from_nodes("限训8.docx", "文科", nodes)

        self.assertEqual([unit.question_id for unit in units], ["3", "6"])
        self.assertEqual([unit.source_span for unit in units], [(1, 7), (8, 14)])
        self.assertTrue(all(unit.question_type == "material_choice" for unit in units))
        self.assertTrue(all(unit.subject_overlay == "geography" for unit in units))

    def test_material_block_range_stops_before_next_numbered_question_group(self):
        nodes = [
            DocNode(index=1, text="下图为世界部分区域海陆分布示意图，据图完成6-8小题。"),
            DocNode(index=2, text="Y6.亚马孙河是世界上流量最大的河流，刚果河是世界第二大河，两河都注入（   ）"),
            DocNode(index=3, text="A．太平洋\tB．大西洋\tC．印度洋\tD．北冰洋"),
            DocNode(index=4, text="G7.甲地理事物为（   ）"),
            DocNode(index=5, text="A．大陆\tB．半岛\tC．海峡\tD．岛屿"),
            DocNode(index=6, text="G8.①山脉的说法错误的是（   ）"),
            DocNode(index=7, text="A．世界上最长的山脉\tB．是板块碰撞挤压形成"),
            DocNode(index=8, text="C9.我国最高点是位于喜马拉雅山脉的珠穆朗玛峰。"),
            DocNode(index=9, text="A．山峰甲\tB．山峰乙"),
            DocNode(index=10, text="Y10.下列地理现象或事实，能够说明地球表面海陆处在不断运动和变化之中的是（   ）"),
        ]

        units = build_question_units_from_nodes("限训4.docx", "文科", nodes)

        self.assertEqual([unit.question_id for unit in units], ["6", "9", "10"])
        self.assertEqual(units[0].source_span, (1, 7))
        self.assertEqual(units[1].source_span, (8, 9))


if __name__ == "__main__":
    unittest.main()
