import os
import sys
import re
import time
import tempfile
import pyautogui
from dataclasses import replace
from pathlib import Path

from docx import Document

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)

# 把墨痕快刀目录加到 sys.path 中，以便导入 wps_helper
sys.path.append(
    os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "墨痕快刀",
    )
)
from wps_helper import get_active_wps
from shared_core import (
    build_answer_units_from_docx,
    build_answer_units_from_wps,
    build_question_units_from_docx,
    format_review_gate_message,
    get_review_gate_result,
    map_answers,
)
from shared_core.answer_core import infer_grouped_question_ids

# ----------------- 配置区 -----------------
DEBUG_MODE = True
KEY_ANSWER = "f2"  # 录入普通题答案的快捷键
KEY_SUB_ANSWER = "f4"  # 录入大题小问答案的快捷键
KEY_ANALYSIS = "f3"  # 录入解析的快捷键
WAIT_TIME = 0.5  # 每次按键后的等待时间 (秒)
ZHONGMEI_CHINESE_PROJECT_PREFIX = "众美-高三-语文"
ZHONGMEI_CLASSICAL_QUESTION_DIR = "文言文"
ZHONGMEI_CLASSICAL_ANSWER_DIR = ("答案", "文言文答案")
INPUT_QUESTION_PREFIX_PATTERN = re.compile(
    r"^[ \t]*\d+[．.][ \t]*(?:[（(][ \t]*\d+[ \t]*[）)][ \t]*)?"
)

# ------------------------------------------
def debug_log(msg):
    if DEBUG_MODE:
        print(f"   ℹ️ {msg}")


def read_current_review_gate(doc) -> dict:
    """每次执行录入动作前重新读取状态，避免继续使用启动时的旧结果。"""
    try:
        return get_review_gate_result(doc.FullName)
    except Exception as exc:
        return {
            "allowed": False,
            "status": "unreadable",
            "reason": f"无法读取自动检查状态：{exc}",
            "status_path": "",
        }


def derive_zhongmei_classical_cleaned_answer_path(
    active_doc_path: str | Path,
) -> Path | None:
    """把众美文言文原题路径映射到同层级的已清洗答案路径。"""
    source_path = Path(active_doc_path)
    if source_path.suffix.lower() != ".docx" or source_path.stem.endswith("_已清洗"):
        return None

    parts = source_path.parts
    project_index = next(
        (
            index
            for index, part in enumerate(parts)
            if part.startswith(ZHONGMEI_CHINESE_PROJECT_PREFIX)
        ),
        None,
    )
    if project_index is None:
        return None

    question_index = project_index + 1
    if (
        question_index >= len(parts)
        or parts[question_index] != ZHONGMEI_CLASSICAL_QUESTION_DIR
    ):
        return None

    relative_path = Path(*parts[question_index + 1 :])
    if not relative_path.name:
        return None
    cleaned_relative_path = relative_path.with_name(
        f"{source_path.stem}_已清洗.docx"
    )
    project_root = Path(*parts[: project_index + 1])
    return project_root.joinpath(
        *ZHONGMEI_CLASSICAL_ANSWER_DIR,
        cleaned_relative_path,
    )


def _same_document_path(left: str | Path, right: str | Path) -> bool:
    return os.path.normcase(os.path.abspath(str(left))) == os.path.normcase(
        os.path.abspath(str(right))
    )


def _find_open_document(wps, target_path: str | Path):
    try:
        count = wps.Documents.Count
    except Exception:
        return None

    for index in range(1, count + 1):
        try:
            candidate = wps.Documents.Item(index)
            if _same_document_path(candidate.FullName, target_path):
                return candidate
        except Exception:
            continue
    return None


def resolve_active_answer_document(wps, doc):
    """众美文言文原题误成活动文档时，安全切换到已放行的清洗答案。"""
    try:
        source_path = Path(doc.FullName)
    except Exception:
        return doc, None

    cleaned_path = derive_zhongmei_classical_cleaned_answer_path(source_path)
    if cleaned_path is None:
        return doc, None
    if not cleaned_path.is_file():
        raise ValueError(
            "检测到当前活动文档是众美文言文原题，但未找到对应的已清洗答案："
            f"{cleaned_path}"
        )

    gate_result = get_review_gate_result(cleaned_path)
    if not gate_result["allowed"]:
        status_path = gate_result.get("status_path") or ""
        status_suffix = f" 状态文件：{status_path}" if status_path else ""
        raise ValueError(
            "检测到当前活动文档是众美文言文原题，但对应清洗答案的"
            f"自动检查门禁未通过：{gate_result['reason']}{status_suffix}"
        )

    answer_doc = _find_open_document(wps, cleaned_path)
    if answer_doc is None:
        answer_doc = wps.Documents.Open(str(cleaned_path))
    answer_doc.Activate()
    try:
        answer_doc.ActiveWindow.Activate()
    except Exception:
        pass

    return (
        answer_doc,
        f"🔄 已自动切换到清洗答案：{answer_doc.Name}",
    )


def _doc_uses_whole_answer_input_for_subquestions(doc) -> bool:
    """物理解答题虽有(1)(2)，但插件只留一个答案框。"""
    doc_markers = []
    for attr in ("FullName", "Name"):
        try:
            value = getattr(doc, attr, "")
        except Exception:
            value = ""
        if value:
            doc_markers.append(str(value))

    joined = " ".join(doc_markers)
    return any(marker in joined for marker in ("未来-高二-物理", "众美-高三-物理"))


def should_split_subquestion_answers(doc, block) -> bool:
    if block.get("force_whole_answer_input"):
        return False
    if block.get("answer_mode") != "subquestion":
        return False
    if _doc_uses_whole_answer_input_for_subquestions(doc):
        return False
    return True


def _group_material_answer_units(question_units, raw_answer_units):
    """只把完整阅读组改成一个录入单元，其他答案块保持原样。"""
    grouped_questions = []
    member_to_group_start = {}
    for question in question_units:
        grouped_ids = infer_grouped_question_ids(question)
        if len(grouped_ids) <= 1:
            continue
        grouped_questions.append((question, grouped_ids))
        for grouped_id in grouped_ids:
            member_to_group_start[grouped_id] = question.question_id

    if not grouped_questions:
        return list(raw_answer_units)

    mapped_units = map_answers(question_units, raw_answer_units)
    mapped_groups = {
        unit.question_id: unit
        for unit in mapped_units
        if unit.metadata.get("is_material_group")
    }

    grouped_units = []
    emitted_groups = set()
    for unit in raw_answer_units:
        group_start = member_to_group_start.get(unit.question_id)
        if not group_start:
            grouped_units.append(unit)
            continue
        if group_start in emitted_groups:
            continue
        mapped_group = mapped_groups.get(group_start)
        if mapped_group is not None:
            grouped_units.append(mapped_group)
            emitted_groups.add(group_start)
        else:
            grouped_units.append(unit)

    return grouped_units


def build_input_units_from_docx(docx_path):
    """从已清洗答案中识别普通题、显式小问和完整阅读组。"""
    raw_units = build_answer_units_from_docx(
        docx_path,
        preserve_source_positions=True,
    )
    question_units = build_question_units_from_docx(docx_path)
    return _group_material_answer_units(question_units, raw_units)


def parse_answer_units(doc):
    raw_units = build_answer_units_from_wps(doc)
    try:
        question_units = build_question_units_from_docx(doc.FullName)
    except Exception:
        return raw_units
    return _group_material_answer_units(question_units, raw_units)


def parse_answer_blocks(doc):
    """
    扫描当前文档，提取所有标准化的答案块位置。
    底层使用 shared_core AnswerUnit，当前函数保留旧字典格式以兼容执行层。
    """
    return build_blocks_from_units(parse_answer_units(doc))


def build_blocks_from_units(units):
    blocks = []
    for unit in units:
        blocks.append(
            {
                "qnum": unit.question_id,
                "ans_start_p": unit.metadata.get("ans_start_p", unit.source_span[0]),
                "ana_start_p": unit.metadata.get("ana_start_p"),
                "end_p": unit.metadata.get("end_p", unit.source_span[1]),
                "is_sub_question": unit.metadata.get("is_sub_question", False),
                "answer_mode": unit.answer_mode,
                "review_flags": unit.review_flags,
                "confidence": unit.confidence,
                "force_whole_answer_input": unit.metadata.get("force_whole_answer_input", False),
            }
        )
    return blocks


def _strip_answer_label(text):
    return re.sub(r"^\s*答案[：:]\s*", "", text or "", count=1).strip()


def _joined_answer_text(unit, include_item_ids=False):
    parts = []
    for index, item in enumerate(unit.answer_items, 1):
        text = _strip_answer_label(item.text)
        if include_item_ids:
            marker = item.item_id or f"({index})"
            parts.append(f"{marker}{text}")
        elif text:
            parts.append(text)
    return "\n".join(parts).strip()


def _requires_contiguous_input_document(unit):
    return bool(
        unit.metadata.get("is_material_group")
        or unit.metadata.get("interleaved_subquestion_analysis")
    )


def prepare_material_group_input_document(source_path, units, *, temp_dir):
    """
    将含完整阅读组的语文答案重组为临时连续录入文档。

    原清洗文档不修改；每个阅读组变为连续 F4 答案段和一个连续 F3 解析段。
    """
    source = Path(source_path)
    if not any(_requires_contiguous_input_document(unit) for unit in units):
        return source, list(units)

    target_dir = Path(temp_dir)
    target_dir.mkdir(parents=True, exist_ok=True)
    target_path = target_dir / f"{source.stem}_录入临时.docx"
    input_doc = Document()
    prepared_units = []

    for unit in units:
        metadata = dict(unit.metadata)
        ans_start_p = len(input_doc.paragraphs) + 1
        input_doc.add_paragraph(f"{unit.question_id}．")

        force_whole = bool(metadata.get("force_whole_answer_input"))
        split_answers = unit.answer_mode == "subquestion" and not force_whole
        if split_answers:
            for index, item in enumerate(unit.answer_items, 1):
                marker = item.item_id or f"({index})"
                input_doc.add_paragraph(f"{marker}{_strip_answer_label(item.text)}")
        else:
            answer_text = _joined_answer_text(
                unit,
                include_item_ids=len(unit.answer_items) > 1,
            )
            input_doc.add_paragraph(f"答案：{answer_text}")

        ana_start_p = len(input_doc.paragraphs) + 1
        input_doc.add_paragraph("解析：")
        if unit.analysis_items:
            if unit.metadata.get("is_material_group") or len(unit.analysis_items) > 1:
                for index, item in enumerate(unit.analysis_items, 1):
                    marker = item.item_id or f"({index})"
                    input_doc.add_paragraph(f"{marker}{item.text.strip()}")
            else:
                analysis_text = "\n".join(
                    item.text.strip()
                    for item in unit.analysis_items
                    if item.text and item.text.strip()
                )
                if analysis_text:
                    input_doc.add_paragraph(analysis_text)

        end_p = len(input_doc.paragraphs)
        metadata.update(
            {
                "ans_start_p": ans_start_p,
                "ans_end_p": ana_start_p - 1,
                "ana_start_p": ana_start_p,
                "end_p": end_p,
                "prepared_input_document": True,
            }
        )
        prepared_units.append(
            replace(
                unit,
                source_span=(ans_start_p, end_p),
                answer_span=(ans_start_p, ana_start_p - 1),
                analysis_span=(ana_start_p, end_p),
                metadata=metadata,
            )
        )

    input_doc.save(target_path)
    return target_path, prepared_units


def find_review_required_units(units):
    return [unit for unit in units if unit.requires_review]


def ensure_standardized_answer_units(answer_source):
    if isinstance(answer_source, (str, Path)):
        gate_result = get_review_gate_result(answer_source)
        if not gate_result["allowed"]:
            raise ValueError(format_review_gate_message(gate_result))
        units = build_answer_units_from_docx(
            answer_source,
            preserve_source_positions=True,
        )
    else:
        units = list(answer_source or [])

    blocking_units = find_review_required_units(units)
    return units, blocking_units


def execute_input_from_units(doc, wps, units, start_idx, end_idx, strict=True):
    blocking_units = find_review_required_units(units)
    if strict and blocking_units:
        preview = "、".join(unit.question_id for unit in blocking_units[:5])
        raise ValueError(
            f"检测到 {len(blocking_units)} 个高风险答案块，默认拒绝直接录入。"
            f" 题号示例: {preview}。请先走总控入口的清洗+自动检查流程。"
        )

    if not any(_requires_contiguous_input_document(unit) for unit in units):
        execute_input(doc, wps, build_blocks_from_units(units), start_idx, end_idx)
        return

    source_path = Path(doc.FullName)
    original_doc = doc
    temp_doc = None
    with tempfile.TemporaryDirectory(
        prefix="_answer_input_",
        dir=PROJECT_ROOT,
    ) as temp_dir:
        input_path, prepared_units = prepare_material_group_input_document(
            source_path,
            units,
            temp_dir=temp_dir,
        )
        try:
            temp_doc = wps.Documents.Open(str(input_path))
            execute_input(
                temp_doc,
                wps,
                build_blocks_from_units(prepared_units),
                start_idx,
                end_idx,
            )
        finally:
            if temp_doc is not None:
                try:
                    temp_doc.Close(False)
                except Exception:
                    pass
            try:
                original_doc.Activate()
            except Exception:
                pass


def _subquestion_marker_number(marker):
    bracket_match = re.fullmatch(r"[（(](\d+)[）)]", marker or "")
    if bracket_match:
        return int(bracket_match.group(1))
    return None


def _sequential_subquestion_matches(matches):
    if not matches:
        return []
    numbers = [_subquestion_marker_number(match.group(1)) for match in matches]
    if numbers != list(range(1, len(numbers) + 1)):
        return []
    return matches if all(re.fullmatch(r"[（(]\d+[）)]", match.group(1)) for match in matches) else []


def find_subquestion_matches(text_body):
    """只识别括号数字小问；①②等圆圈序号属于整题答案内部列点。"""
    marker_pattern = r"([\(（]\d+[\)）])"
    line_start_pattern = re.compile(
        rf"(?:(?<=\r)|(?<=\n)|\A)[ \t]*{marker_pattern}[ \t]*"
    )
    line_matches = _sequential_subquestion_matches(
        list(line_start_pattern.finditer(text_body))
    )
    if line_matches:
        return line_matches

    inline_matches = list(re.finditer(marker_pattern, text_body))
    if not inline_matches:
        return []
    first_content_pos = len(text_body) - len(text_body.lstrip())
    if inline_matches[0].start() != first_content_pos:
        return []
    return _sequential_subquestion_matches(inline_matches)


def strip_input_question_prefix(answer_text):
    """剥离顶层题号；跨段出现的 (1) 必须保留给 F4 拆分。"""
    question_prefix = INPUT_QUESTION_PREFIX_PATTERN.match(answer_text)
    offset_start = question_prefix.end() if question_prefix else 0
    text_body = answer_text[offset_start:]

    answer_prefix = re.match(r"^\s*答案[：:]\s*", text_body)
    if answer_prefix:
        offset_start += answer_prefix.end()
        text_body = text_body[answer_prefix.end():]
    return offset_start, text_body


def execute_input(doc, wps, blocks, start_idx, end_idx):
    """
    逐题录入：每题先录答案(F2)，再录解析(F3)，完成后再进入下一题
    """
    target_blocks = blocks[start_idx:end_idx]
    total = len(target_blocks)

    print(f"\n🚀 开始自动化录入，共 {total} 题")
    print(f"   答案快捷键: {KEY_ANSWER}，解析快捷键: {KEY_ANALYSIS}")
    print(f"   录入过程中请勿触碰鼠标和键盘！")
    print(f"   2秒后开始，请切换到 WPS 窗口...\n")

    time.sleep(2)

    for idx, b in enumerate(target_blocks, 1):
        qnum = b["qnum"]
        ans_start_p = b["ans_start_p"]
        ana_start_p = b["ana_start_p"]
        end_p = b["end_p"]

        ans_end_p = (ana_start_p - 1) if ana_start_p else end_p

        print(f"   ▶️ 第 {qnum} 题 ({idx}/{total})")

        # ------ 录入答案 (F2 / F4) ------
        try:
            ans_start_pos = doc.Paragraphs(ans_start_p).Range.Start
            ans_end_pos = doc.Paragraphs(ans_end_p).Range.End
            rng_ans = doc.Range(ans_start_pos, ans_end_pos)
            ans_text = rng_ans.Text

            # 跳过顶层题号或普通答案前缀；跨段的 (1) 不得随题号一起吞掉。
            offset_start, text_body = strip_input_question_prefix(ans_text)

            matches = find_subquestion_matches(text_body)

            if matches and should_split_subquestion_answers(doc, b):
                debug_log(
                    f"第 {qnum} 题包含 {len(matches)} 个小题，将自动拆分并触发 {KEY_SUB_ANSWER}"
                )
                for i in range(len(matches)):
                    m_curr = matches[i]
                    # 本小问内容的起点：当前题号之后
                    content_start = m_curr.end()
                    # 本小问内容的终点：下一个题号之前，或是整个答案的结尾
                    content_end = (
                        matches[i + 1].start()
                        if i + 1 < len(matches)
                        else len(text_body)
                    )

                    sub_text = text_body[content_start:content_end]
                    answer_label = re.match(r"^\s*答案[：:]\s*", sub_text)
                    if answer_label:
                        content_start += answer_label.end()
                        sub_text = text_body[content_start:content_end]
                    # 计算去尾后的长度差，以便在 WPS 中不选中多余的换行和空格
                    stripped_sub = sub_text.rstrip()
                    trailing_len = len(sub_text) - len(stripped_sub)

                    # 映射回 WPS 文档的绝对位置
                    abs_start = ans_start_pos + offset_start + content_start
                    abs_end = ans_start_pos + offset_start + content_end - trailing_len

                    if abs_start >= abs_end:
                        debug_log(
                            f"第 {qnum} 题 小题 {m_curr.group(1)} 内容为空，跳过录入"
                        )
                        continue

                    sub_rng = doc.Range(abs_start, abs_end)
                    sub_rng.Select()
                    wps.Application.ActiveWindow.ScrollIntoView(sub_rng)

                    print(f"   ▶️ 第 {qnum} 题 答案小问 {m_curr.group(1)} (F4)")
                    pyautogui.press(KEY_SUB_ANSWER)
                    time.sleep(WAIT_TIME)
            else:
                if matches and not should_split_subquestion_answers(doc, b):
                    debug_log(
                        f"第 {qnum} 题虽含小题标记，但当前文档场景使用整题答案框，改为整段触发 {KEY_ANSWER}"
                    )
                # 正常单题，只选中剩余所有内容
                stripped_body = text_body.rstrip()
                trailing_len = len(text_body) - len(stripped_body)
                abs_start = ans_start_pos + offset_start
                abs_end = ans_start_pos + offset_start + len(text_body) - trailing_len

                if abs_start < abs_end:
                    rng_normal = doc.Range(abs_start, abs_end)
                    rng_normal.Select()
                    wps.Application.ActiveWindow.ScrollIntoView(rng_normal)

                    debug_log(
                        f"选中第 {qnum} 题 [答案]（已跳过题号），触发 {KEY_ANSWER}"
                    )
                    pyautogui.press(KEY_ANSWER)
                    time.sleep(WAIT_TIME)
                else:
                    print(f"      ⚠️ 第 {qnum} 题 答案内容为空")

        except Exception as e:
            if hasattr(e, "excepinfo") and e.excepinfo[5] == -2146827864:
                print(f"      ❌ 致命错误：检测到当前文档已被关闭、替换或大幅删减！")
                print(
                    f"      💡 提示：您刚刚可能清理了文档或换了新文档。请在菜单输入 'r' 刷新连接！"
                )
                return  # 终止当前所有录入，直接返回主菜单
            else:
                print(f"      ❌ 答案录入失败: {e}")

        # ------ 录入解析 (F3) ------
        # 无论解析是否有内容，都要执行F3，因为WPS插件需要F3才能跳转到下一题
        try:
            # ✨ 修复：确保 end_p 不为 None
            safe_end_p = end_p if end_p is not None else ans_start_p

            if ana_start_p and ana_start_p <= safe_end_p:
                # 有"解析："段落，选择"解析："之后的内容（可能为空）
                ana_para = doc.Paragraphs(ana_start_p)
                ana_text = ana_para.Range.Text.replace("\r", "").replace("\n", "").replace("\x07", "")

                # 找到"解析："或"解析:"之后的位置
                match = re.match(r"^\s*解析[：:]", ana_text)
                if match:
                    # 跳过"解析："前缀，从后面开始选择
                    prefix_len = len(match.group(0))
                    ana_start_pos = ana_para.Range.Start + prefix_len
                else:
                    # 没匹配到前缀，从段落开头选择
                    ana_start_pos = ana_para.Range.Start

                # 从 safe_end_p 倒退，找到最后一个非空段落
                real_end_p = safe_end_p
                while real_end_p >= ana_start_p:
                    tail_text = (
                        doc.Paragraphs(real_end_p)
                        .Range.Text.replace("\r", "")
                        .replace("\n", "")
                        .replace("\x07", "")
                        .strip()
                    )
                    if tail_text:
                        break
                    real_end_p -= 1

                if real_end_p < ana_start_p:
                    # 解析区全为空，在"解析："后面创建空选择
                    ana_end_pos = ana_start_pos
                    debug_log(f"第 {qnum} 题解析为空，选择'解析：'后空白位置")
                else:
                    ana_end_pos = doc.Paragraphs(real_end_p).Range.End
                    debug_log(
                        f"选中第 {qnum} 题 [解析] 段落{ana_start_p}~{real_end_p}（已跳过尾部{safe_end_p - real_end_p}个空段落）"
                    )

                rng = doc.Range(ana_start_pos, ana_end_pos)
                rng.Select()
                wps.Application.ActiveWindow.ScrollIntoView(rng)
            else:
                # 没有"解析："段落，在答案末尾创建空选择
                # ✨ 修复：确保 ans_end_p 不为 None
                safe_ans_end_p = ans_end_p if ans_end_p is not None else ans_start_p
                ans_end_pos = doc.Paragraphs(safe_ans_end_p).Range.End
                rng = doc.Range(ans_end_pos, ans_end_pos)
                rng.Select()
                wps.Application.ActiveWindow.ScrollIntoView(rng)
                debug_log(f"第 {qnum} 题没有解析段落，在答案后执行F3跳转")

            # 无论是否有解析内容，都执行F3按键以跳转到下一题
            pyautogui.press(KEY_ANALYSIS)
            time.sleep(WAIT_TIME)
        except Exception as e:
            if hasattr(e, "excepinfo") and e.excepinfo[5] == -2146827864:
                print(
                    f"      ❌ 致命错误：检测到当前文档已被关闭、替换或大幅删减！"
                )
                print(f"      💡 提示：请在菜单输入 'r' 刷新连接！")
                return
            else:
                print(f"      ❌ 解析录入失败: {e}")

    print(f"\n🎉 全部完成！共处理 {total} 题。")


def main():
    import io

    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

    print("========================================")
    print("      🔪 墨痕快刀 - 答案录入工具 (初版)")
    print("========================================")

    while True:
        wps = get_active_wps()
        if not wps:
            print("\n❌ 未找到运行中的 WPS 或 Word (5秒后重试)...")
            time.sleep(5)
            continue

        try:
            doc = wps.ActiveDocument
        except:
            time.sleep(2)
            continue

        try:
            doc, switch_message = resolve_active_answer_document(wps, doc)
        except ValueError as exc:
            print(f"\n❌ {exc}")
            retry_choice = input(
                "\n👉 按回车键重新扫描（或输入 'q' 退出）："
            ).strip().lower()
            if retry_choice == "q":
                sys.exit(0)
            continue
        except Exception as exc:
            print(f"\n❌ 自动切换清洗答案失败：{exc}")
            retry_choice = input(
                "\n👉 按回车键重新扫描（或输入 'q' 退出）："
            ).strip().lower()
            if retry_choice == "q":
                sys.exit(0)
            continue

        if switch_message:
            print(f"\n{switch_message}")
        doc_name = doc.Name

        print(f"\n📄 当前活动文档: {doc_name}")
        print("🔍 正在扫描答案块，请稍候...")

        answer_units = parse_answer_units(doc)
        blocks = build_blocks_from_units(answer_units)
        total_q = len(blocks)

        if total_q == 0:
            print("\n⚠️ 未在该文档中发现标准格式的答案 (如 '1．A')。")
            print("   (提示: 请先使用【答案格式清洗】工具将文档标准化！)")
            input("\n👉 按回车键重新扫描 (或输入 'q' 退出): ")
            continue

        print(f"\n✅ 扫描完毕！共识别到 {total_q} 题。")
        print(f"   起止题号: 第 {blocks[0]['qnum']} 题 -> 第 {blocks[-1]['qnum']} 题")
        print("-" * 40)

        gate_result = read_current_review_gate(doc)
        if not gate_result["allowed"]:
            print("\n🛂 自动检查门禁未通过。")
            print(f"   {gate_result['reason']}")
            print(f"   状态文件: {gate_result['status_path']}")

        risky_blocks = [b for b in blocks if b["confidence"] < 0.75 or b["review_flags"]]
        if risky_blocks:
            print("\n⚠️ 当前文档仍包含高风险答案块，默认不建议直接录入。")
            print("   请优先使用项目根目录 main.py 总控入口执行“只做答案清洗+自动检查”或“全流程录入”。")
            print(f"   风险题号示例: {', '.join(b['qnum'] for b in risky_blocks[:5])}")

        while True:
            print("\n请选择要录入的范围:")
            print("  [1] 录入所有题目")
            print("  [2] 自定义录入范围 (如: 1-10)")
            print("  [r] 刷新/切换文档")
            print("  [q] 退出程序")

            choice = input("\n👉 请输入选项: ").strip().lower()

            if choice == "q":
                sys.exit(0)
            elif choice == "r":
                break  # 跳出内层循环，重新获取 WPS 文档
            elif choice == "1":
                gate_result = read_current_review_gate(doc)
                if not gate_result["allowed"]:
                    print(f"❌ 已拦截：{gate_result['reason']}")
                    if gate_result.get("status_path"):
                        print(f"   状态文件: {gate_result['status_path']}")
                    continue
                if risky_blocks:
                    print("❌ 已拦截：当前答案文档存在高风险结构，拒绝直接录入。")
                    continue
                execute_input_from_units(doc, wps, answer_units, 0, total_q)
            elif choice == "2":
                gate_result = read_current_review_gate(doc)
                if not gate_result["allowed"]:
                    print(f"❌ 已拦截：{gate_result['reason']}")
                    if gate_result.get("status_path"):
                        print(f"   状态文件: {gate_result['status_path']}")
                    continue
                if risky_blocks:
                    print("❌ 已拦截：当前答案文档存在高风险结构，拒绝直接录入。")
                    continue
                range_str = input(
                    "👉 请输入起始和结束索引(如 1-10 代表前10题): "
                ).strip()
                try:
                    s, e = map(int, range_str.split("-"))
                    if 1 <= s <= e <= total_q:
                        execute_input_from_units(doc, wps, answer_units, s - 1, e)
                    else:
                        print("❌ 范围无效，超出总题数！")
                except:
                    print("❌ 格式错误，请输入如 1-10 的格式。")
            else:
                print("❌ 无效的输入。")


if __name__ == "__main__":
    main()
