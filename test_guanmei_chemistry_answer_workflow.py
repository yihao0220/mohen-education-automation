from __future__ import annotations

from hashlib import sha256
from io import BytesIO
from pathlib import Path
import tempfile
import unittest
from zipfile import ZipFile

from docx import Document
from lxml import etree

from shared_core import build_answer_units_from_docx, get_review_gate_result
from shared_core.models import QuestionUnit
from tools.clean_guanmei_chemistry_answers import (
    _prepare_question_units_for_review,
    clean_answer_batch,
    normalize_answer_docx,
)


TINY_PNG = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01"
    b"\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde"
    b"\x00\x00\x00\x0cIDATx\x9cc\xf8\xff\xff?\x00\x05"
    b"\xfe\x02\xfeA\xe2)\xb7\x00\x00\x00\x00IEND\xaeB`\x82"
)


def _digest(path: Path) -> str:
    return sha256(path.read_bytes()).hexdigest()


def _unresolved_ignorable_prefixes(path: Path) -> set[str]:
    with ZipFile(path) as package:
        root = etree.fromstring(package.read("word/document.xml"))
    ignorable = root.get(
        "{http://schemas.openxmlformats.org/markup-compatibility/2006}Ignorable",
        "",
    )
    return {prefix for prefix in ignorable.split() if prefix not in root.nsmap}


def _write_answer_source(path: Path, *, numbered: bool = True) -> None:
    doc = Document()
    doc.add_paragraph("第一章 第一节 第1课时　测试")
    if numbered:
        doc.add_paragraph("1.答案　C")
        doc.add_paragraph("解析　第一题解析")
        doc.add_paragraph("2.答案　B")
        doc.add_paragraph("3答案　D")
        doc.add_paragraph("解析　第三题解析")
    else:
        doc.add_paragraph("答案　A")
        doc.add_paragraph("解析　第一题解析")
        doc.add_paragraph("答案　B")
        doc.add_paragraph("解析　第二题解析")
    picture_paragraph = doc.add_paragraph()
    picture_paragraph.add_run().add_picture(BytesIO(TINY_PNG))
    doc.save(path)


class GuanmeiChemistryNormalizationTests(unittest.TestCase):
    def test_normalize_labels_infer_missing_analysis_and_preserve_picture(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = root / "01 第一章 第一节 第1课时 测试-答案.docx"
            target = root / "01 第一章 第一节 第1课时 测试-答案_已清洗.docx"
            _write_answer_source(source)
            source_hash = _digest(source)

            result = normalize_answer_docx(source, target)

            self.assertEqual(result.answer_count, 3)
            self.assertEqual(result.inserted_analysis_placeholders, 1)
            self.assertEqual(_digest(source), source_hash)
            self.assertEqual(_unresolved_ignorable_prefixes(target), set())

            cleaned = Document(target)
            texts = [paragraph.text for paragraph in cleaned.paragraphs]
            self.assertIn("1．C", texts)
            self.assertIn("解析：第一题解析", texts)
            self.assertIn("2．B", texts)
            self.assertIn("解析： ", texts)
            self.assertIn("3．D", texts)
            self.assertEqual(len(cleaned.inline_shapes), 1)

            units = build_answer_units_from_docx(
                target,
                preserve_source_positions=True,
            )
            self.assertEqual([unit.question_id for unit in units], ["1", "2", "3"])
            self.assertEqual(
                [unit.answer_items[0].text for unit in units],
                ["C", "B", "D"],
            )
            self.assertFalse(any(unit.review_flags for unit in units))

    def test_unnumbered_answer_markers_are_numbered_sequentially(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = root / "04 第一章 第二节 第2课时 测试-答案.docx"
            target = root / "04 第一章 第二节 第2课时 测试-答案_已清洗.docx"
            _write_answer_source(source, numbered=False)

            result = normalize_answer_docx(source, target)

            self.assertEqual(result.answer_count, 2)
            units = build_answer_units_from_docx(target)
            self.assertEqual([unit.question_id for unit in units], ["1", "2"])
            self.assertEqual(
                [unit.answer_items[0].text for unit in units],
                ["A", "B"],
            )


class GuanmeiChemistryBatchTests(unittest.TestCase):
    def test_batch_cleaning_builds_approved_review_state(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            input_dir = root / "按章节拆分"
            output_dir = input_dir / "已清洗"
            input_dir.mkdir()

            answer_path = input_dir / "01 第一章 第一节 第1课时 测试-答案.docx"
            _write_answer_source(answer_path)

            question_path = root / "选择性必修1课时作业.docx"
            question_doc = Document()
            question_doc.add_paragraph("第一章 第一节 第1课时　测试")
            question_doc.add_paragraph("1．第一题")
            question_doc.add_paragraph("2．第二题")
            question_doc.add_paragraph("3．第三题")
            question_doc.save(question_path)

            results = clean_answer_batch(
                input_dir,
                question_docx=question_path,
                output_dir=output_dir,
            )

            self.assertEqual(len(results), 1)
            result = results[0]
            self.assertEqual(result.answer_count, 3)
            self.assertEqual(result.question_count, 3)
            self.assertEqual(result.blocking_issue_count, 0)
            self.assertTrue(result.output_path.exists())
            self.assertTrue(result.review_report_path.exists())
            gate = get_review_gate_result(result.output_path)
            self.assertTrue(gate["allowed"])
            self.assertEqual(gate["status"], "approved")

    def test_batch_supplements_confirmed_missing_first_section_answer(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            input_dir = root / "按章节拆分"
            output_dir = input_dir / "已清洗"
            input_dir.mkdir()

            answer_path = input_dir / "01 第一章 第一节 第1课时 反应热 焓变-答案.docx"
            answer_doc = Document()
            answer_doc.add_paragraph("第一章 第一节 第1课时　反应热　焓变")
            for question_id in range(1, 10):
                answer_doc.add_paragraph(f"{question_id}.答案　A")
                answer_doc.add_paragraph("解析　原答案解析")
            answer_doc.save(answer_path)

            question_path = root / "选择性必修1课时作业.docx"
            question_doc = Document()
            question_doc.add_paragraph("第一章 第一节 第1课时　反应热　焓变")
            for question_id in range(1, 11):
                question_doc.add_paragraph(f"{question_id}．第{question_id}题")
            question_doc.save(question_path)

            results = clean_answer_batch(
                input_dir,
                question_docx=question_path,
                output_dir=output_dir,
            )

            self.assertEqual(len(results), 1)
            self.assertEqual(results[0].supplemented_answer_count, 1)
            self.assertEqual(results[0].answer_count, 10)
            units = build_answer_units_from_docx(results[0].output_path)
            self.assertEqual(units[-1].question_id, "10")
            self.assertEqual(units[-1].answer_items[0].text, "B")
            self.assertIn("1.50 mol", units[-1].analysis_items[0].text)
            self.assertTrue(get_review_gate_result(results[0].output_path)["allowed"])

    def test_choice_question_statements_do_not_require_f4_answers(self):
        question = QuestionUnit(
            question_id="4",
            subject="理科",
            subject_overlay="chemistry",
            grade_hint="高二",
            question_type="choice",
            stem_blocks=["按步骤(1)(2)完成实验后，选择正确结论"],
            option_blocks=["A.甲", "B.乙", "C.丙", "D.丁"],
            subquestions=["(1)步骤一", "(2)步骤二"],
            media_blocks=[],
            material_blocks=[],
            source_span=(1, 4),
        )
        prepared = _prepare_question_units_for_review([question])

        self.assertEqual(len(prepared), 1)
        self.assertEqual(prepared[0].subquestions, [])
        self.assertEqual(prepared[0].question_type, "choice")


if __name__ == "__main__":
    unittest.main()
