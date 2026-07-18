from __future__ import annotations

import argparse
from copy import deepcopy
from dataclasses import dataclass, replace
from pathlib import Path
import re
import sys
import tempfile
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from lxml import etree as ET

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from shared_core import (
    build_answer_units_from_docx,
    build_question_units_from_docx,
    build_review_report,
    export_review_report,
    initialize_review_status,
    map_answers,
    update_review_status,
)


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
DOCUMENT_XML = "word/document.xml"

ET.register_namespace("w", W_NS)

ANSWER_MARKER_PATTERN = re.compile(
    r"^\s*(?:(?P<qid>\d+)\s*[．.]?\s*)?答案(?:(?:\s*[：:]\s*)|\s+|(?=$))"
)
ANALYSIS_MARKER_PATTERN = re.compile(
    r"^\s*(?:解析|【解析】)(?:(?:\s*[：:]\s*)|\s+|(?=$))"
)
NUMBERED_ANALYSIS_MARKER_PATTERN = re.compile(
    r"^\s*(?P<qid>\d+)\s*[．.]?\s*(?:解析|【解析】)"
    r"(?:(?:\s*[：:]\s*)|\s+|(?=$))"
)
SUBQUESTION_START_PATTERN = re.compile(r"^\s*[（(]1[）)]")
SECTION_HEADING_PATTERN = re.compile(
    r"^第[一二三四五六七八九十]+章\s+第[一二三四五六七八九十]+节"
    r"(?:\s+第\s*\d+\s*课时)?"
)
SPLIT_ANSWER_FILENAME_PATTERN = re.compile(r"^(?P<section>\d{2})\s+.+-答案\.docx$")
SUPPLEMENTAL_ANSWERS_BY_SECTION = {
    "第一章第一节第1课时反应热焓变": {
        "10": (
            "B",
            "由图可知中和反应放热，说明化学能可以转化为热能；最高温度对应"
            "V1=30 mL、V2=20 mL，NaOH溶液浓度约为1.50 mol·L-1，故选B。",
        ),
    },
}


@dataclass(frozen=True)
class NormalizationResult:
    source_path: Path
    output_path: Path
    answer_count: int
    inserted_analysis_placeholders: int


@dataclass(frozen=True)
class BatchResult:
    source_path: Path
    output_path: Path
    review_report_path: Path
    answer_count: int
    question_count: int
    inserted_analysis_placeholders: int
    supplemented_answer_count: int
    blocking_issue_count: int


def _paragraph_text(paragraph) -> str:
    return "".join(node.text or "" for node in paragraph.iter(f"{W}t"))


def _set_text(node, value: str) -> None:
    node.text = value
    if value.startswith(" ") or value.endswith(" "):
        node.set(XML_SPACE, "preserve")
    else:
        node.attrib.pop(XML_SPACE, None)


def _replace_text_prefix(paragraph, prefix_length: int, replacement: str) -> None:
    text_nodes = list(paragraph.iter(f"{W}t"))
    if not text_nodes:
        raise ValueError("答案或解析标签所在段落没有可编辑文本节点")

    cursor = 0
    replacement_written = False
    for node in text_nodes:
        value = node.text or ""
        start = cursor
        end = start + len(value)
        cursor = end

        if not replacement_written:
            suffix = value[max(0, prefix_length - start) :] if prefix_length < end else ""
            _set_text(node, replacement + suffix)
            replacement_written = True
            continue

        if end <= prefix_length:
            _set_text(node, "")
        elif start < prefix_length:
            _set_text(node, value[prefix_length - start :])


def _new_text_paragraph(reference_paragraph, value: str):
    paragraph = ET.Element(f"{W}p")
    paragraph_properties = reference_paragraph.find(f"{W}pPr")
    if paragraph_properties is not None:
        paragraph.append(deepcopy(paragraph_properties))

    run = ET.SubElement(paragraph, f"{W}r")
    reference_run_properties = reference_paragraph.find(f".//{W}rPr")
    if reference_run_properties is not None:
        run.append(deepcopy(reference_run_properties))
    text = ET.SubElement(run, f"{W}t")
    _set_text(text, value)
    return paragraph


def _new_analysis_placeholder(reference_paragraph):
    return _new_text_paragraph(reference_paragraph, "解析： ")


def _write_docx_with_document_xml(
    source_path: Path,
    output_path: Path,
    document_xml: bytes,
    *,
    allow_in_place: bool = False,
) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    if source_path.resolve() == output_path.resolve() and not allow_in_place:
        raise ValueError("清洗输出不能覆盖源答案文档")

    with tempfile.NamedTemporaryFile(
        prefix=f"{output_path.stem}_",
        suffix=".docx",
        dir=output_path.parent,
        delete=False,
    ) as temporary_file:
        temporary_path = Path(temporary_file.name)

    try:
        with ZipFile(source_path, "r") as source_zip, ZipFile(
            temporary_path,
            "w",
            compression=ZIP_DEFLATED,
        ) as target_zip:
            for info in source_zip.infolist():
                payload = document_xml if info.filename == DOCUMENT_XML else source_zip.read(info.filename)
                target_zip.writestr(info, payload)

        with ZipFile(temporary_path, "r") as check_zip:
            bad_member = check_zip.testzip()
            if bad_member:
                raise ValueError(f"清洗后的 docx 压缩包损坏: {bad_member}")
        Document(temporary_path)
        temporary_path.replace(output_path)
    finally:
        if temporary_path.exists():
            temporary_path.unlink()


def normalize_answer_docx(source_path: str | Path, output_path: str | Path) -> NormalizationResult:
    source = Path(source_path)
    output = Path(output_path)
    if not source.is_file():
        raise FileNotFoundError(f"未找到答案文档: {source}")

    with ZipFile(source, "r") as package:
        root = ET.fromstring(package.read(DOCUMENT_XML))

    body = root.find(f"{W}body")
    if body is None:
        raise ValueError(f"文档缺少正文节点: {source.name}")

    answer_paragraphs = []
    analysis_paragraphs = set()
    expected_qid = 1

    for child in list(body):
        if child.tag != f"{W}p":
            continue
        text = _paragraph_text(child)
        answer_match = ANSWER_MARKER_PATTERN.match(text)
        if answer_match:
            explicit_qid = answer_match.group("qid")
            qid = int(explicit_qid) if explicit_qid else expected_qid
            if qid != expected_qid:
                raise ValueError(
                    f"{source.name} 答案题号不连续: 期望 {expected_qid}，实际 {qid}"
                )
            payload = text[answer_match.end() :]
            if SUBQUESTION_START_PATTERN.match(payload):
                marker_paragraph = _new_text_paragraph(child, f"{qid}．")
                body.insert(list(body).index(child), marker_paragraph)
                _replace_text_prefix(child, answer_match.end(), "")
                answer_paragraphs.append(marker_paragraph)
            else:
                _replace_text_prefix(child, answer_match.end(), f"{qid}．")
                answer_paragraphs.append(child)
            expected_qid += 1
            continue

        numbered_analysis_match = NUMBERED_ANALYSIS_MARKER_PATTERN.match(text)
        if numbered_analysis_match:
            numbered_qid = int(numbered_analysis_match.group("qid"))
            if numbered_qid != expected_qid:
                raise ValueError(
                    f"{source.name} 编号解析位置异常: 期望标签 {expected_qid}，实际 {numbered_qid}"
                )
            _replace_text_prefix(child, numbered_analysis_match.end(), "解析：")
            analysis_paragraphs.add(child)
            continue

        analysis_match = ANALYSIS_MARKER_PATTERN.match(text)
        if analysis_match:
            _replace_text_prefix(child, analysis_match.end(), "解析：")
            analysis_paragraphs.add(child)

    if not answer_paragraphs:
        raise ValueError(f"{source.name} 未识别到任何答案标记")

    inserted_placeholders = 0
    for answer_index in range(len(answer_paragraphs) - 1, -1, -1):
        current_children = list(body)
        answer_paragraph = answer_paragraphs[answer_index]
        answer_position = current_children.index(answer_paragraph)
        if answer_index + 1 < len(answer_paragraphs):
            next_answer_position = current_children.index(answer_paragraphs[answer_index + 1])
        else:
            section_properties = body.find(f"{W}sectPr")
            next_answer_position = (
                current_children.index(section_properties)
                if section_properties is not None
                else len(current_children)
            )

        has_analysis = any(
            candidate in analysis_paragraphs
            for candidate in current_children[answer_position + 1 : next_answer_position]
        )
        if has_analysis:
            continue

        body.insert(next_answer_position, _new_analysis_placeholder(answer_paragraph))
        inserted_placeholders += 1

    document_xml = ET.tostring(
        root,
        encoding="utf-8",
        xml_declaration=True,
        standalone=True,
    )
    _write_docx_with_document_xml(source, output, document_xml)
    return NormalizationResult(
        source_path=source,
        output_path=output,
        answer_count=len(answer_paragraphs),
        inserted_analysis_placeholders=inserted_placeholders,
    )


def _normalize_section_title(title: str) -> str:
    return re.sub(r"\s+", "", title or "")


def _append_supplemental_answers(
    answer_docx: Path,
    supplements: dict[str, tuple[str, str]],
) -> int:
    if not supplements:
        return 0

    with ZipFile(answer_docx, "r") as package:
        root = ET.fromstring(package.read(DOCUMENT_XML))
    body = root.find(f"{W}body")
    if body is None:
        raise ValueError(f"文档缺少正文节点: {answer_docx.name}")

    body_children = list(body)
    section_properties = body.find(f"{W}sectPr")
    insert_position = (
        body_children.index(section_properties)
        if section_properties is not None
        else len(body_children)
    )
    reference_paragraph = next(
        (child for child in reversed(body_children[:insert_position]) if child.tag == f"{W}p"),
        None,
    )
    if reference_paragraph is None:
        raise ValueError(f"无法为补齐答案继承段落格式: {answer_docx.name}")

    inserted = 0
    for question_id, (answer_text, analysis_text) in sorted(
        supplements.items(),
        key=lambda item: int(item[0]),
    ):
        body.insert(
            insert_position,
            _new_text_paragraph(reference_paragraph, f"{question_id}．{answer_text}"),
        )
        insert_position += 1
        body.insert(
            insert_position,
            _new_text_paragraph(reference_paragraph, f"解析：{analysis_text}"),
        )
        insert_position += 1
        inserted += 1

    document_xml = ET.tostring(
        root,
        encoding="utf-8",
        xml_declaration=True,
        standalone=True,
    )
    _write_docx_with_document_xml(
        answer_docx,
        answer_docx,
        document_xml,
        allow_in_place=True,
    )
    return inserted


def _build_question_sections(question_docx: Path):
    document = Document(question_docx)
    headings = [
        (index, paragraph.text.strip())
        for index, paragraph in enumerate(document.paragraphs, 1)
        if SECTION_HEADING_PATTERN.match(paragraph.text.strip())
    ]
    if not headings:
        raise ValueError(f"题目文档未识别到章节标题: {question_docx.name}")

    all_units = build_question_units_from_docx(question_docx)
    sections = {}
    for section_index, (start_paragraph, title) in enumerate(headings, 1):
        end_paragraph = (
            headings[section_index][0] - 1
            if section_index < len(headings)
            else len(document.paragraphs)
        )
        section_units = [
            unit
            for unit in all_units
            if start_paragraph < unit.source_span[0] <= end_paragraph
        ]
        sections[section_index] = (title, section_units)
    return sections


def _prepare_question_units_for_review(question_units):
    return [
        replace(unit, subquestions=[])
        if unit.question_type == "choice" and unit.subquestions
        else unit
        for unit in question_units
    ]


def clean_answer_batch(
    input_dir: str | Path,
    *,
    question_docx: str | Path,
    output_dir: str | Path,
) -> list[BatchResult]:
    input_root = Path(input_dir)
    question_path = Path(question_docx)
    output_root = Path(output_dir)
    sections = _build_question_sections(question_path)

    source_files = []
    for candidate in sorted(input_root.glob("*.docx")):
        match = SPLIT_ANSWER_FILENAME_PATTERN.match(candidate.name)
        if match:
            source_files.append((int(match.group("section")), candidate))

    if not source_files:
        raise ValueError(f"未在目录中找到严格命名的章节答案: {input_root}")
    if [index for index, _ in source_files] != list(range(1, len(source_files) + 1)):
        raise ValueError("章节答案文件序号不连续，已停止整批清洗")
    if len(source_files) != len(sections):
        raise ValueError(
            f"章节数量不一致: 答案 {len(source_files)} 份，题目 {len(sections)} 章"
        )

    results = []
    for section_index, source_path in source_files:
        output_path = output_root / f"{source_path.stem}_已清洗.docx"
        normalization = normalize_answer_docx(source_path, output_path)

        answer_units = build_answer_units_from_docx(
            output_path,
            preserve_source_positions=True,
        )
        section_title, question_units = sections[section_index]
        question_ids = {unit.question_id for unit in question_units}
        answer_ids = {unit.question_id for unit in answer_units}
        section_supplements = SUPPLEMENTAL_ANSWERS_BY_SECTION.get(
            _normalize_section_title(section_title),
            {},
        )
        missing_supplements = {
            question_id: payload
            for question_id, payload in section_supplements.items()
            if question_id in question_ids and question_id not in answer_ids
        }
        supplemented_answer_count = _append_supplemental_answers(
            output_path,
            missing_supplements,
        )
        if supplemented_answer_count:
            answer_units = build_answer_units_from_docx(
                output_path,
                preserve_source_positions=True,
            )
        review_question_units = _prepare_question_units_for_review(question_units)
        mapped_units = map_answers(review_question_units, answer_units)
        report = build_review_report(
            output_path.name,
            review_question_units,
            mapped_units,
        )
        report_path = output_root / f"{output_path.stem}_审核清单.md"
        export_review_report(report, report_path)
        initialize_review_status(
            output_path,
            report_path=str(report_path),
            report=report,
        )
        blocking_issue_count = sum(
            1 for issue in report.issues if issue.severity == "error"
        )
        update_review_status(
            output_path,
            status="approved" if blocking_issue_count == 0 else "rejected",
            reviewer="system",
            note=(
                (
                    "莞美高二化学章节答案自动检查通过；"
                    f"补齐源答案缺失题 {supplemented_answer_count} 道"
                )
                if blocking_issue_count == 0
                else f"自动检查发现 {blocking_issue_count} 个高风险问题"
            ),
        )
        results.append(
            BatchResult(
                source_path=source_path,
                output_path=output_path,
                review_report_path=report_path,
                answer_count=len(answer_units),
                question_count=len(question_units),
                inserted_analysis_placeholders=normalization.inserted_analysis_placeholders,
                supplemented_answer_count=supplemented_answer_count,
                blocking_issue_count=blocking_issue_count,
            )
        )

    return results


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="莞美高二化学章节答案标准化、审核并生成录入门禁状态",
    )
    parser.add_argument("input_dir", help="按章节拆分后的原答案目录")
    parser.add_argument("--question-docx", required=True, help="对应的总题目 DOCX")
    parser.add_argument("--output-dir", required=True, help="最终已清洗答案输出目录")
    return parser


def main() -> int:
    args = _build_parser().parse_args()
    results = clean_answer_batch(
        args.input_dir,
        question_docx=args.question_docx,
        output_dir=args.output_dir,
    )
    approved = sum(1 for result in results if result.blocking_issue_count == 0)
    answer_total = sum(result.answer_count for result in results)
    placeholder_total = sum(result.inserted_analysis_placeholders for result in results)
    supplemented_total = sum(result.supplemented_answer_count for result in results)
    print(f"章节文件: {len(results)}")
    print(f"答案总数: {answer_total}")
    print(f"补空解析占位: {placeholder_total}")
    print(f"补齐源答案缺失题: {supplemented_total}")
    print(f"自动检查通过: {approved}/{len(results)}")
    print(f"输出目录: {Path(args.output_dir).resolve()}")
    return 0 if approved == len(results) else 1


if __name__ == "__main__":
    raise SystemExit(main())
