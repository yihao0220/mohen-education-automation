from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from shared_core.question_core import build_question_units_from_docx
from shared_core.answer_core import build_answer_units_from_paragraph_texts, map_answers
from shared_core.review import build_review_report, export_review_report
from shared_core.review_gate import initialize_review_status, update_review_status
from 格式处理.格式模板库 import template_nancheng_math


UNIT_TITLE_PATTERN = re.compile(r"^第\d+(?:[、,]\d+)?单元$")


def _clean_text(text: str) -> str:
    return (text or "").replace("\r", "").replace("\n", "").replace("\x07", "").strip()


def _load_nonempty_paragraphs(path: str | Path) -> list[str]:
    doc = Document(path)
    return [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]


def _unit_digits(unit_title: str) -> list[str]:
    return re.findall(r"\d+", unit_title)


def _question_path_score(unit_title: str, question_path: Path) -> int:
    digits = _unit_digits(unit_title)
    stem = question_path.stem
    if not digits:
        return 0

    if len(digits) == 1:
        return 10 if f"第{digits[0]}单元" in stem else 0

    combined = "".join(digits)
    score = 0
    if f"第{combined}单元" in stem:
        score += 20
    for digit in digits:
        if digit in stem:
            score += 3
    return score


def _build_unit_question_map(question_dir: str | Path, unit_titles: list[str]) -> dict[str, Path]:
    question_paths = [
        path
        for path in sorted(Path(question_dir).glob("*.docx"))
        if not path.name.startswith("~$") and "答案" not in path.stem and "_已清洗" not in path.stem
    ]
    mapping: dict[str, Path] = {}

    for unit_title in unit_titles:
        scored = sorted(
            ((_question_path_score(unit_title, path), path) for path in question_paths),
            key=lambda item: item[0],
            reverse=True,
        )
        if not scored or scored[0][0] <= 0:
            raise ValueError(f"未找到 {unit_title} 对应的题目文档")
        mapping[unit_title] = scored[0][1]

    return mapping


def extract_unit_titles(paragraph_texts: list[str]) -> list[str]:
    titles: list[str] = []
    for raw_text in paragraph_texts:
        text = _clean_text(raw_text).replace(",", "、")
        if UNIT_TITLE_PATTERN.match(text) and text not in titles:
            titles.append(text)
    return titles


def _clear_doc_body(doc_obj: Document) -> None:
    body = doc_obj._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def write_standard_answer_docx(lines: list[str], output_path: str | Path) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _clear_doc_body(doc)
    for line in lines:
        doc.add_paragraph(line)
    doc.save(output_path)
    return output_path


def _write_review_artifacts(question_path: Path, answer_path: Path, lines: list[str]) -> None:
    question_units = build_question_units_from_docx(question_path)
    answer_units = build_answer_units_from_paragraph_texts(lines)
    mapped_units = map_answers(question_units, answer_units)
    report = build_review_report(answer_path.name, question_units, mapped_units)
    report_path = answer_path.with_name(f"{answer_path.stem}_审核清单.md")
    export_review_report(report, report_path)
    initialize_review_status(answer_path, report_path=str(report_path), report=report)
    update_review_status(
        answer_path,
        status="rejected" if report.summary["high_risk_count"] else "approved",
        reviewer="system",
        note="自动检查未通过" if report.summary["high_risk_count"] else "自动检查通过",
    )


def standardize_answer_docx(
    source_path: str | Path,
    *,
    question_doc: str | Path,
    output_dir: str | Path | None = None,
    write_review: bool = True,
) -> Path:
    source_path = Path(source_path)
    question_path = Path(question_doc)
    output_dir = Path(output_dir) if output_dir else source_path.parent / "按单元拆分"

    question_units = build_question_units_from_docx(question_path)
    answer_texts = _load_nonempty_paragraphs(source_path)
    entries = template_nancheng_math.parse_unit_answer_texts(answer_texts, question_units)
    lines = template_nancheng_math.render_standard_lines(entries)
    output_path = output_dir / f"{question_path.stem}-答案_已清洗.docx"
    write_standard_answer_docx(lines, output_path)
    if write_review:
        _write_review_artifacts(question_path, output_path, lines)
    return output_path


def split_answer_docx(
    source_path: str | Path,
    *,
    question_dir: str | Path,
    output_dir: str | Path | None = None,
    write_review: bool = True,
) -> list[Path]:
    source_path = Path(source_path)
    question_dir = Path(question_dir)
    output_dir = Path(output_dir) if output_dir else source_path.parent / "按单元拆分"

    answer_texts = _load_nonempty_paragraphs(source_path)
    unit_titles = extract_unit_titles(answer_texts)
    if not unit_titles:
        raise ValueError("未识别到单元标题；单份答案请使用 --question-doc 指定题目文档。")
    unit_question_map = _build_unit_question_map(question_dir, unit_titles)
    output_paths: list[Path] = []

    for unit_title in unit_titles:
        question_path = unit_question_map[unit_title]
        question_units = build_question_units_from_docx(question_path)
        unit_texts = template_nancheng_math.extract_unit_texts(answer_texts, unit_title)
        entries = template_nancheng_math.parse_unit_answer_texts(unit_texts, question_units)
        lines = template_nancheng_math.render_standard_lines(entries)
        output_path = output_dir / f"{question_path.stem}-答案_已清洗.docx"
        write_standard_answer_docx(lines, output_path)
        if write_review:
            _write_review_artifacts(question_path, output_path, lines)
        output_paths.append(output_path)

    return output_paths


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="清洗南城二小小学数学答案，可按单元拆分或指定单份题目文档。")
    parser.add_argument("source", help="答案 docx 路径")
    parser.add_argument("--question-dir", help="题目文档目录；总答案按单元拆分时使用")
    parser.add_argument("--question-doc", help="题目 docx 路径；单份答案无单元标题时使用")
    parser.add_argument("--output-dir", help="拆分输出目录，默认在总答案同级下创建 按单元拆分")
    parser.add_argument("--no-review", action="store_true", help="只生成清洗 docx，不生成审核清单和审核状态")
    return parser


def main() -> None:
    parser = build_arg_parser()
    args = parser.parse_args()

    if args.question_doc:
        output_paths = [
            standardize_answer_docx(
                args.source,
                question_doc=args.question_doc,
                output_dir=args.output_dir,
                write_review=not args.no_review,
            )
        ]
    else:
        if not args.question_dir:
            parser.error("总答案拆分模式必须提供 --question-dir；单份答案模式请提供 --question-doc。")
        output_paths = split_answer_docx(
            args.source,
            question_dir=args.question_dir,
            output_dir=args.output_dir,
            write_review=not args.no_review,
        )

    print(f"已拆分 {len(output_paths)} 份答案：")
    for output_path in output_paths:
        print(f"- {output_path}")


if __name__ == "__main__":
    main()
