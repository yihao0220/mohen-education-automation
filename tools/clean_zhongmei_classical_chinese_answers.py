from __future__ import annotations

import argparse
import hashlib
import re
import shutil
import sys
import tempfile
from dataclasses import dataclass, field
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from shared_core.question_core import build_question_units_from_docx
from tools.zhongmei_review_status import refresh_review_statuses


DEFAULT_PROJECT_DIR = Path(r"D:\墨痕教育题目\众美-高三-语文")
DEFAULT_ANSWER_DIR = DEFAULT_PROJECT_DIR / "答案" / "文言文答案"
DEFAULT_QUESTION_DIR = DEFAULT_PROJECT_DIR / "文言文"

EXPECTED_SOURCE_NUMBERS = tuple(range(1, 29))
EXPECTED_DOCUMENT_COUNT = 28
EXPECTED_BLOCK_COUNT = 154
EXPECTED_ANSWER_MARKER_COUNT = 196
EXPECTED_TRANSLATION_BLOCK_COUNT = 33
EXPECTED_TRANSLATION_OCCURRENCE_COUNT = 74
EXPECTED_NONEMPTY_CONTINUATION_COUNT = 219
EXPECTED_SCORE_POINT_COUNT = 330

SOURCE_FILE_PATTERN = re.compile(
    r"^(?P<number>\d{2})(?P<title>(?:《[^《》]+》)+)\.docx$"
)
PURE_TITLE_PATTERN = re.compile(r"^(?:《[^《》]+》)+$")
ANSWER_MARKER_PATTERN = re.compile(
    r"^\s*答案(?:[\s　]*[：:]\s*|[\s　]+|$)"
)
UNKNOWN_ANSWER_MARKER_PATTERN = re.compile(r"^\s*(?:【|\[)?答案")
SOURCE_ANALYSIS_MARKER_PATTERN = re.compile(
    r"^\s*(?:【|\[)?(?:解析|详解|分析)(?:】|\])?[：:]?"
)
SECTION_HEADING_PATTERN = re.compile(
    r"^\s*[一二三四五六七八九十百]+、"
)
CHINESE_SUBHEADING_PATTERN = re.compile(
    r"^\s*[（(][一二三四五六七八九十百]+[）)]"
)
NUMBERED_QUESTION_PATTERN = re.compile(r"^\s*\d{1,3}\s*[．.]\s*")
TRANSLATION_PLACEHOLDER_PATTERN = re.compile(r"^\s*译文\s*[：:]")
BRACKET_HEADING_PATTERN = re.compile(r"^\s*\[[^\]]+\]\s*$")
PLAIN_TASK_HEADING_PATTERN = re.compile(
    r"^\s*(?:补写出下列句子中的空缺部分。?|多义实词|重点虚词|特殊句式)\s*(?:[：:].*)?$"
)
PLAIN_WORK_TITLE_BOUNDARIES = frozenset({"礼运(节选)"})
PURE_SCORE_KEYWORD_LIST_PATTERN = re.compile(
    r"^[\u3400-\u9fff]+(?:[、，,][\u3400-\u9fff]+)+$"
)
OUTPUT_QUESTION_LINE_PATTERN = re.compile(r"^(\d+)．$")
OUTPUT_ANSWER_LINE_PATTERN = re.compile(r"^(?:\(\d+\))?答案：")
OUTPUT_ANALYSIS_LINE_PATTERN = re.compile(r"^(?:\(\d+\))?解析：")

KNOWN_GLYPH_BY_IMAGE_SHA256 = {
    "5f05c3b7b2574594908b05abc1f920e30ab3683f9129ef089efde9aa4286fb02": "輮",
    "501a651bfc10a2044bd503ae5016dfb4de3156bff49c7d532198eaf2ac781d97": "絖",
    "a1bafd3ff484f89afc07525938230404b4c1b8a003c584f8c9e2082573d97c60": "餔",
}
EXPECTED_KNOWN_GLYPH_COUNTS = {
    "01《劝学》.docx": {"輮": 4},
    "17《〈老子〉四章》《五石之瓠》.docx": {"絖": 1},
    "19《屈原列传》《报任安书(节选)》.docx": {"餔": 1},
}

DOUBLE_MICRO_ANSWER_FILENAME = "18《兼爱》《礼运(节选)》.docx"
SUWU_FILENAME = "20《苏武传》.docx"
SUWU_SOURCE_IDS = ["1", "2", "5", "6"]
SUWU_OUTPUT_IDS = ["1", "2", "3", "4"]


class UnsupportedTemplateError(ValueError):
    """源文档不符合已确认的众美文言文答案模板。"""


@dataclass(frozen=True)
class SourcePair:
    number: int
    relative_path: Path
    answer_path: Path
    question_path: Path


@dataclass
class CleanTitle:
    text: str
    source_index: int


@dataclass
class AnswerOccurrence:
    source_index: int
    answer_lines: list[str] = field(default_factory=list)
    analysis_lines: list[str] = field(default_factory=list)


@dataclass
class CleanBlock:
    question_id: str
    source_index: int
    occurrences: list[AnswerOccurrence]
    is_translation: bool = False


@dataclass
class CleanDocument:
    source_path: Path
    question_path: Path
    relative_path: Path
    titles: list[CleanTitle]
    blocks: list[CleanBlock]
    source_answer_marker_count: int


def _clean_text(text: str) -> str:
    return (
        (text or "")
        .replace("\r", "")
        .replace("\n", " ")
        .replace("\x07", "")
        .replace("\u00a0", " ")
        .strip()
    )


def _source_match(path: Path) -> re.Match[str] | None:
    return SOURCE_FILE_PATTERN.fullmatch(path.name)


def _source_number(path: Path) -> int:
    match = _source_match(path)
    if not match:
        raise UnsupportedTemplateError(f"{path.name}：文件名不符合“01《篇名》.docx”格式")
    return int(match.group("number"))


def _filename_title(path: Path) -> str:
    match = _source_match(path)
    if not match:
        raise UnsupportedTemplateError(f"{path.name}：无法从文件名确定篇名")
    return match.group("title")


def discover_source_pairs(
    answer_dir: str | Path = DEFAULT_ANSWER_DIR,
    question_dir: str | Path = DEFAULT_QUESTION_DIR,
    *,
    expected_numbers: tuple[int, ...] | None = EXPECTED_SOURCE_NUMBERS,
) -> list[SourcePair]:
    answer_dir = Path(answer_dir)
    question_dir = Path(question_dir)
    if not answer_dir.is_dir():
        raise FileNotFoundError(f"答案目录不存在：{answer_dir}")
    if not question_dir.is_dir():
        raise FileNotFoundError(f"题目目录不存在：{question_dir}")

    source_files: list[Path] = []
    unknown_files: list[Path] = []
    for path in answer_dir.rglob("*.docx"):
        if path.name.startswith("~$") or path.stem.endswith("_已清洗"):
            continue
        if _source_match(path):
            source_files.append(path)
        else:
            unknown_files.append(path)
    if unknown_files:
        details = "、".join(str(path.relative_to(answer_dir)) for path in unknown_files)
        raise UnsupportedTemplateError(f"发现未确认命名的 DOCX，已停止：{details}")
    if not source_files:
        raise FileNotFoundError(f"未在 {answer_dir} 递归找到文言文答案 DOCX")

    pairs: list[SourcePair] = []
    seen_numbers: dict[int, Path] = {}
    for answer_path in source_files:
        number = _source_number(answer_path)
        if number in seen_numbers:
            raise UnsupportedTemplateError(
                f"编号 {number:02d} 出现多份答案：{seen_numbers[number]}、{answer_path}"
            )
        seen_numbers[number] = answer_path
        relative_path = answer_path.relative_to(answer_dir)
        question_path = question_dir / relative_path
        if not question_path.is_file():
            raise UnsupportedTemplateError(
                f"{relative_path}：缺少同相对路径题目文档 {question_path}"
            )
        pairs.append(
            SourcePair(
                number=number,
                relative_path=relative_path,
                answer_path=answer_path,
                question_path=question_path,
            )
        )

    pairs.sort(key=lambda pair: pair.number)
    actual_numbers = tuple(pair.number for pair in pairs)
    if expected_numbers is not None and actual_numbers != expected_numbers:
        raise UnsupportedTemplateError(
            f"文档编号不完整，实际 {actual_numbers}，预期 {expected_numbers}"
        )
    return pairs


def _is_graphical_paragraph(paragraph) -> bool:
    return bool(
        paragraph._p.xpath(
            './/*[local-name()="drawing" or local-name()="object" or local-name()="pict" or local-name()="imagedata"]'
        )
    )


def _relationship_blob(document: Document, relationship_id: str) -> bytes | None:
    relationship = document.part.rels.get(relationship_id)
    if relationship is None or relationship.is_external:
        return None
    target_part = relationship.target_part
    return getattr(target_part, "blob", None)


def _image_text_from_container(
    container,
    document: Document,
    source_path: Path,
    paragraph_index: int,
) -> str:
    relationship_ids: list[str] = []
    for blip in container.xpath('.//*[local-name()="blip"]'):
        relationship_id = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
        if relationship_id:
            relationship_ids.append(relationship_id)
    for image_data in container.xpath('.//*[local-name()="imagedata"]'):
        relationship_id = image_data.get(qn("r:id")) or image_data.get(qn("r:href"))
        if relationship_id:
            relationship_ids.append(relationship_id)

    relationship_ids = list(dict.fromkeys(relationship_ids))
    if not relationship_ids:
        raise UnsupportedTemplateError(
            f"{source_path.name}：第 {paragraph_index + 1} 段答案抽取区存在无法读取的内联对象"
        )

    restored: list[str] = []
    for relationship_id in relationship_ids:
        blob = _relationship_blob(document, relationship_id)
        digest = hashlib.sha256(blob).hexdigest() if blob else ""
        glyph = KNOWN_GLYPH_BY_IMAGE_SHA256.get(digest)
        if not glyph:
            raise UnsupportedTemplateError(
                f"{source_path.name}：第 {paragraph_index + 1} 段答案抽取区存在未知内联图片"
            )
        restored.append(glyph)
    return "".join(restored)


def _paragraph_text_with_images(
    paragraph,
    document: Document,
    source_path: Path,
    paragraph_index: int,
) -> str:
    chunks: list[str] = []
    for run in paragraph._p.xpath(".//w:r"):
        for child in run:
            local_name = child.tag.rsplit("}", 1)[-1]
            if local_name in {"t", "delText"}:
                chunks.append(child.text or "")
            elif local_name == "tab":
                chunks.append("\t")
            elif local_name in {"br", "cr"}:
                chunks.append(" ")
            elif local_name == "sym":
                char_code = child.get(qn("w:char"))
                if char_code:
                    try:
                        chunks.append(chr(int(char_code, 16)))
                    except ValueError:
                        pass
            elif local_name in {"drawing", "object", "pict"}:
                chunks.append(
                    _image_text_from_container(
                        child,
                        document,
                        source_path,
                        paragraph_index,
                    )
                )
    return _clean_text("".join(chunks))


def _is_answer_boundary(text: str) -> bool:
    return bool(
        ANSWER_MARKER_PATTERN.match(text)
        or PURE_TITLE_PATTERN.fullmatch(text)
        or SECTION_HEADING_PATTERN.match(text)
        or CHINESE_SUBHEADING_PATTERN.match(text)
        or NUMBERED_QUESTION_PATTERN.match(text)
        or TRANSLATION_PLACEHOLDER_PATTERN.match(text)
        or BRACKET_HEADING_PATTERN.match(text)
        or PLAIN_TASK_HEADING_PATTERN.match(text)
        or text in PLAIN_WORK_TITLE_BOUNDARIES
    )


def _split_score_points(
    answer_lines: list[str],
    source_path: Path,
    paragraph_index: int,
) -> tuple[list[str], list[str]]:
    cleaned_answer_lines = list(answer_lines)
    analysis_lines: list[str] = []
    score_start_pattern = re.compile(r"[\(\[]\s*得分点\s*[：:]")

    for line_index, line in enumerate(list(cleaned_answer_lines)):
        if "得分点" not in line:
            continue
        if analysis_lines:
            raise UnsupportedTemplateError(
                f"{source_path.name}：第 {paragraph_index + 1} 段同一答案出现多个得分点块"
            )
        match = score_start_pattern.search(line)
        if not match:
            raise UnsupportedTemplateError(
                f"{source_path.name}：第 {paragraph_index + 1} 段得分点边界无法确认"
            )
        opening = line[match.start()]
        closing = ")" if opening == "(" else "]"
        stripped = line.rstrip()
        if not stripped.endswith(closing):
            raise UnsupportedTemplateError(
                f"{source_path.name}：第 {paragraph_index + 1} 段得分点缺少闭合括号"
            )
        score_text = stripped[match.end() : -1].strip()
        if not score_text:
            raise UnsupportedTemplateError(
                f"{source_path.name}：第 {paragraph_index + 1} 段得分点为空"
            )
        answer_text = line[: match.start()].rstrip()
        if answer_text:
            cleaned_answer_lines[line_index] = answer_text
        else:
            cleaned_answer_lines.pop(line_index)
        analysis_lines = [f"得分点：{score_text}"]

    cleaned_answer_lines = [line for line in cleaned_answer_lines if _clean_text(line)]
    if not cleaned_answer_lines:
        raise UnsupportedTemplateError(
            f"{source_path.name}：第 {paragraph_index + 1} 段答案正文为空"
        )
    return cleaned_answer_lines, analysis_lines


def _parse_answer_occurrence(
    document: Document,
    source_path: Path,
    start_index: int,
) -> AnswerOccurrence:
    marker_paragraph = document.paragraphs[start_index]
    marker_text = _paragraph_text_with_images(
        marker_paragraph,
        document,
        source_path,
        start_index,
    )
    marker_match = ANSWER_MARKER_PATTERN.match(marker_text)
    if not marker_match:
        raise UnsupportedTemplateError(
            f"{source_path.name}：第 {start_index + 1} 段答案标记无法解析"
        )
    answer_lines = [marker_text[marker_match.end() :].strip()]

    for paragraph_index in range(start_index + 1, len(document.paragraphs)):
        paragraph = document.paragraphs[paragraph_index]
        plain_text = _clean_text(paragraph.text)
        if plain_text and _is_answer_boundary(plain_text):
            break
        if not plain_text and not _is_graphical_paragraph(paragraph):
            continue
        answer_text = _paragraph_text_with_images(
            paragraph,
            document,
            source_path,
            paragraph_index,
        )
        if answer_text:
            answer_lines.append(answer_text)

    answer_lines, analysis_lines = _split_score_points(
        answer_lines,
        source_path,
        start_index,
    )
    return AnswerOccurrence(
        source_index=start_index,
        answer_lines=answer_lines,
        analysis_lines=analysis_lines,
    )


def _extract_answer_occurrences(
    document: Document,
    source_path: Path,
) -> list[AnswerOccurrence]:
    for table_index, table in enumerate(document.tables):
        table_text = "\n".join(
            _clean_text(cell.text)
            for row in table.rows
            for cell in row.cells
            if _clean_text(cell.text)
        )
        if any(
            ANSWER_MARKER_PATTERN.match(line)
            or SOURCE_ANALYSIS_MARKER_PATTERN.match(line)
            for line in table_text.splitlines()
        ):
            raise UnsupportedTemplateError(
                f"{source_path.name}：第 {table_index + 1} 个表格含答案或解析标记"
            )

    marker_indices: list[int] = []
    for paragraph_index, paragraph in enumerate(document.paragraphs):
        text = _clean_text(paragraph.text)
        if ANSWER_MARKER_PATTERN.match(text):
            marker_indices.append(paragraph_index)
        elif UNKNOWN_ANSWER_MARKER_PATTERN.match(text):
            raise UnsupportedTemplateError(
                f"{source_path.name}：第 {paragraph_index + 1} 段使用未确认答案标记"
            )
        elif SOURCE_ANALYSIS_MARKER_PATTERN.match(text):
            raise UnsupportedTemplateError(
                f"{source_path.name}：第 {paragraph_index + 1} 段出现未确认的独立解析标记"
            )
    if not marker_indices:
        raise UnsupportedTemplateError(f"{source_path.name}：未发现行首“答案”标记")

    return [
        _parse_answer_occurrence(document, source_path, marker_index)
        for marker_index in marker_indices
    ]


def _is_translation_unit(question_unit) -> bool:
    return any("翻译语句" in text for text in question_unit.stem_blocks)


def _occurrence_allocations(source_path: Path, question_units) -> list[int]:
    allocations = [
        max(1, len(question_unit.subquestions))
        if _is_translation_unit(question_unit)
        else 1
        for question_unit in question_units
    ]
    if source_path.name == DOUBLE_MICRO_ANSWER_FILENAME:
        if not allocations or allocations[0] != 1:
            raise UnsupportedTemplateError(
                f"{source_path.name}：双篇微点答案白名单前置条件已变化"
            )
        allocations[0] = 2
    return allocations


def _output_question_ids(source_path: Path, question_units) -> list[str]:
    source_ids = [str(question_unit.question_id) for question_unit in question_units]
    if source_path.name != SUWU_FILENAME:
        return source_ids
    if source_ids != SUWU_SOURCE_IDS:
        raise UnsupportedTemplateError(
            f"{source_path.name}：白名单源题号已变化，实际 {source_ids}，预期 {SUWU_SOURCE_IDS}"
        )
    return list(SUWU_OUTPUT_IDS)


def _extract_titles(
    document: Document,
    source_path: Path,
    first_answer_index: int,
) -> list[CleanTitle]:
    titles = [
        CleanTitle(text=text, source_index=index)
        for index, paragraph in enumerate(document.paragraphs)
        if (text := _clean_text(paragraph.text)) and PURE_TITLE_PATTERN.fullmatch(text)
    ]
    filename_title = _filename_title(source_path)
    if not any(
        title.text == filename_title and title.source_index < first_answer_index
        for title in titles
    ):
        titles.append(CleanTitle(text=filename_title, source_index=-1))
    titles.sort(key=lambda title: title.source_index)
    return titles


def parse_document(
    source_path: str | Path,
    question_path: str | Path,
    *,
    relative_path: str | Path | None = None,
) -> CleanDocument:
    source_path = Path(source_path)
    question_path = Path(question_path)
    relative_path = Path(relative_path) if relative_path is not None else Path(source_path.name)
    try:
        document = Document(source_path)
    except Exception as exc:
        raise UnsupportedTemplateError(f"{source_path.name}：无法打开答案 DOCX") from exc

    try:
        question_units = build_question_units_from_docx(question_path)
    except Exception as exc:
        raise UnsupportedTemplateError(f"{question_path.name}：无法解析配对题目 DOCX") from exc
    if not question_units:
        raise UnsupportedTemplateError(f"{question_path.name}：未识别到 F1 逻辑题块")

    occurrences = _extract_answer_occurrences(document, source_path)
    allocations = _occurrence_allocations(source_path, question_units)
    if sum(allocations) != len(occurrences):
        raise UnsupportedTemplateError(
            f"{source_path.name}：答案标记无法按 F1 题块分组，"
            f"需要 {sum(allocations)} 个，实际 {len(occurrences)} 个"
        )

    output_question_ids = _output_question_ids(source_path, question_units)
    blocks: list[CleanBlock] = []
    cursor = 0
    for question_unit, question_id, allocation in zip(
        question_units,
        output_question_ids,
        allocations,
    ):
        block_occurrences = occurrences[cursor : cursor + allocation]
        cursor += allocation
        blocks.append(
            CleanBlock(
                question_id=question_id,
                source_index=block_occurrences[0].source_index,
                occurrences=block_occurrences,
                is_translation=_is_translation_unit(question_unit),
            )
        )

    titles = _extract_titles(
        document,
        source_path,
        first_answer_index=occurrences[0].source_index,
    )
    clean_document = CleanDocument(
        source_path=source_path,
        question_path=question_path,
        relative_path=relative_path,
        titles=titles,
        blocks=blocks,
        source_answer_marker_count=len(occurrences),
    )
    _validate_known_glyph_counts(clean_document)
    return clean_document


def _validate_known_glyph_counts(clean_document: CleanDocument) -> None:
    expected_counts = EXPECTED_KNOWN_GLYPH_COUNTS.get(clean_document.source_path.name)
    if not expected_counts:
        return
    text = "\n".join(
        line
        for block in clean_document.blocks
        for occurrence in block.occurrences
        for line in occurrence.answer_lines + occurrence.analysis_lines
    )
    for glyph, expected_count in expected_counts.items():
        actual_count = text.count(glyph)
        if actual_count != expected_count:
            raise UnsupportedTemplateError(
                f"{clean_document.source_path.name}：图片字“{glyph}”恢复数量异常，"
                f"实际 {actual_count}，预期 {expected_count}"
            )


def count_score_points(analysis_text: str) -> int:
    payload = re.sub(r"^\s*得分点\s*[：:]\s*", "", analysis_text or "").strip()
    if not payload:
        return 0
    if re.search(r"[；;]", payload):
        return len([part for part in re.split(r"[；;]", payload) if part.strip()])
    if PURE_SCORE_KEYWORD_LIST_PATTERN.fullmatch(payload):
        return len([part for part in re.split(r"[、，,]", payload) if part.strip()])
    return 1


def _validate_full_batch(clean_documents: list[CleanDocument]) -> None:
    if len(clean_documents) != EXPECTED_DOCUMENT_COUNT:
        return
    block_count = sum(len(clean_document.blocks) for clean_document in clean_documents)
    marker_count = sum(
        clean_document.source_answer_marker_count for clean_document in clean_documents
    )
    translation_blocks = [
        block
        for clean_document in clean_documents
        for block in clean_document.blocks
        if block.is_translation
    ]
    continuation_count = sum(
        max(0, len(occurrence.answer_lines) - 1)
        for clean_document in clean_documents
        for block in clean_document.blocks
        for occurrence in block.occurrences
    )
    score_point_count = sum(
        count_score_points(analysis_line)
        for block in translation_blocks
        for occurrence in block.occurrences
        for analysis_line in occurrence.analysis_lines
    )
    actual = (
        block_count,
        marker_count,
        len(translation_blocks),
        sum(len(block.occurrences) for block in translation_blocks),
        continuation_count,
        score_point_count,
    )
    expected = (
        EXPECTED_BLOCK_COUNT,
        EXPECTED_ANSWER_MARKER_COUNT,
        EXPECTED_TRANSLATION_BLOCK_COUNT,
        EXPECTED_TRANSLATION_OCCURRENCE_COUNT,
        EXPECTED_NONEMPTY_CONTINUATION_COUNT,
        EXPECTED_SCORE_POINT_COUNT,
    )
    if actual != expected:
        raise UnsupportedTemplateError(
            "整批结构统计与只读基线不一致，"
            f"实际(题块/标记/翻译块/翻译标记/续行/得分点)={actual}，预期={expected}"
        )


def preflight_source_pairs(source_pairs: list[SourcePair]) -> list[CleanDocument]:
    clean_documents: list[CleanDocument] = []
    errors: list[str] = []
    for pair in source_pairs:
        try:
            clean_documents.append(
                parse_document(
                    pair.answer_path,
                    pair.question_path,
                    relative_path=pair.relative_path,
                )
            )
        except Exception as exc:
            errors.append(str(exc))
    if errors:
        details = "\n".join(f"- {error}" for error in errors)
        raise UnsupportedTemplateError(
            "发现无法确定清洗方式的文档，已停止整批清洗：\n" + details
        )
    _validate_full_batch(clean_documents)
    return clean_documents


def render_clean_lines(clean_document: CleanDocument) -> list[str]:
    events: list[tuple[int, int, CleanTitle | CleanBlock]] = []
    events.extend((title.source_index, 0, title) for title in clean_document.titles)
    events.extend((block.source_index, 1, block) for block in clean_document.blocks)
    events.sort(key=lambda item: (item[0], item[1]))

    lines: list[str] = []
    for _, _, event in events:
        if isinstance(event, CleanTitle):
            if not lines or lines[-1] != event.text:
                lines.append(event.text)
            continue

        lines.append(f"{event.question_id}．")
        multiple = len(event.occurrences) > 1
        has_any_analysis = any(
            occurrence.analysis_lines for occurrence in event.occurrences
        )

        if event.is_translation:
            for occurrence_index, occurrence in enumerate(event.occurrences, start=1):
                prefix = f"({occurrence_index})"
                lines.append(f"{prefix}答案：{occurrence.answer_lines[0]}")
                lines.extend(occurrence.answer_lines[1:])

            lines.append("解析：")
            for occurrence_index, occurrence in enumerate(event.occurrences, start=1):
                prefix = f"({occurrence_index})"
                if occurrence.analysis_lines:
                    lines.append(f"{prefix}{occurrence.analysis_lines[0]}")
                    lines.extend(occurrence.analysis_lines[1:])
                else:
                    lines.append(f"{prefix} ")
            continue

        for occurrence_index, occurrence in enumerate(event.occurrences, start=1):
            prefix = f"({occurrence_index})" if multiple else ""
            lines.append(f"{prefix}答案：{occurrence.answer_lines[0]}")
            lines.extend(occurrence.answer_lines[1:])

            if multiple and has_any_analysis:
                if occurrence.analysis_lines:
                    lines.append(f"{prefix}解析：{occurrence.analysis_lines[0]}")
                    lines.extend(occurrence.analysis_lines[1:])
                else:
                    lines.append(f"{prefix}解析： ")

        if not multiple:
            occurrence = event.occurrences[0]
            if occurrence.analysis_lines:
                lines.append(f"解析：{occurrence.analysis_lines[0]}")
                lines.extend(occurrence.analysis_lines[1:])
            else:
                lines.append("解析： ")
        elif not has_any_analysis:
            lines.append("解析： ")
    return lines


def _set_run_font(run, *, bold: bool = False) -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")


def write_clean_docx(
    clean_document: CleanDocument,
    output_path: str | Path,
) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    title_texts = {title.text for title in clean_document.titles}
    for line in render_clean_lines(clean_document):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        is_empty_analysis_placeholder = line == "解析： "
        paragraph.paragraph_format.space_after = Pt(
            0 if is_empty_analysis_placeholder else 6
        )
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.paragraph_format.line_spacing = (
            1.0 if is_empty_analysis_placeholder else 1.25
        )
        run = paragraph.add_run(line)
        _set_run_font(run, bold=line in title_texts)

    if document.paragraphs and not document.paragraphs[0].text:
        first = document.paragraphs[0]._element
        first.getparent().remove(first)
    document.save(output_path)
    return output_path


def validate_clean_docx(
    output_path: str | Path,
    expected_document: CleanDocument,
) -> None:
    output_path = Path(output_path)
    try:
        document = Document(output_path)
    except Exception as exc:
        raise ValueError(f"{output_path.name}：清洗结果无法重新打开") from exc

    actual_lines = [
        _clean_text(paragraph.text)
        for paragraph in document.paragraphs
        if _clean_text(paragraph.text)
    ]
    expected_lines = [
        _clean_text(line)
        for line in render_clean_lines(expected_document)
        if _clean_text(line)
    ]
    if actual_lines != expected_lines:
        raise ValueError(f"{output_path.name}：写入后文本与预期不一致")

    actual_question_ids = [
        match.group(1)
        for line in actual_lines
        if (match := OUTPUT_QUESTION_LINE_PATTERN.fullmatch(line))
    ]
    expected_question_ids = [block.question_id for block in expected_document.blocks]
    if actual_question_ids != expected_question_ids:
        raise ValueError(
            f"{output_path.name}：题号复查失败，实际 {actual_question_ids}，"
            f"预期 {expected_question_ids}"
        )

    answer_count = sum(bool(OUTPUT_ANSWER_LINE_PATTERN.match(line)) for line in actual_lines)
    if answer_count != expected_document.source_answer_marker_count:
        raise ValueError(
            f"{output_path.name}：答案标记数异常，实际 {answer_count}，"
            f"预期 {expected_document.source_answer_marker_count}"
        )
    if any("得分点" in line for line in actual_lines if "答案：" in line):
        raise ValueError(f"{output_path.name}：得分点仍残留在答案行")
    if not any(OUTPUT_ANALYSIS_LINE_PATTERN.match(line) for line in actual_lines):
        raise ValueError(f"{output_path.name}：缺少解析标记")

    leaked_headings = [
        line
        for line in actual_lines
        if SECTION_HEADING_PATTERN.match(line) or CHINESE_SUBHEADING_PATTERN.match(line)
    ]
    if leaked_headings:
        raise ValueError(f"{output_path.name}：残留题型说明：{leaked_headings[:3]}")


def _output_relative_path(clean_document: CleanDocument) -> Path:
    return clean_document.relative_path.with_name(
        f"{clean_document.relative_path.stem}_已清洗.docx"
    )


def clean_batch(
    answer_dir: str | Path = DEFAULT_ANSWER_DIR,
    question_dir: str | Path = DEFAULT_QUESTION_DIR,
    output_dir: str | Path | None = None,
    *,
    preflight_only: bool = False,
) -> list[Path]:
    answer_dir = Path(answer_dir)
    question_dir = Path(question_dir)
    output_dir = Path(output_dir) if output_dir is not None else answer_dir

    source_pairs = discover_source_pairs(answer_dir, question_dir)
    clean_documents = preflight_source_pairs(source_pairs)
    if preflight_only:
        return []

    output_dir.mkdir(parents=True, exist_ok=True)
    temp_root = Path(
        tempfile.mkdtemp(prefix=".文言文答案清洗临时_", dir=output_dir)
    )
    temp_outputs: list[tuple[Path, Path, CleanDocument]] = []
    try:
        for clean_document in clean_documents:
            relative_output = _output_relative_path(clean_document)
            temp_path = temp_root / relative_output
            final_path = output_dir / relative_output
            write_clean_docx(clean_document, temp_path)
            validate_clean_docx(temp_path, clean_document)
            temp_outputs.append((temp_path, final_path, clean_document))

        for temp_path, final_path, _ in temp_outputs:
            final_path.parent.mkdir(parents=True, exist_ok=True)
            temp_path.replace(final_path)
        output_paths = [final_path for _, final_path, _ in temp_outputs]
        if output_dir.resolve() == answer_dir.resolve():
            paths_by_relative_parent: dict[Path, list[Path]] = {}
            for _, final_path, clean_document in temp_outputs:
                paths_by_relative_parent.setdefault(
                    clean_document.relative_path.parent,
                    [],
                ).append(final_path)
            for relative_parent, group_paths in paths_by_relative_parent.items():
                refresh_review_statuses(
                    group_paths,
                    question_dir=question_dir / relative_parent,
                )
        return output_paths
    finally:
        shutil.rmtree(temp_root, ignore_errors=True)


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="清洗众美高三语文文言文答案：按配对题目 F1 题块输出题号、答案和解析。"
    )
    parser.add_argument(
        "answer_dir",
        nargs="?",
        default=str(DEFAULT_ANSWER_DIR),
        help="原始文言文答案 DOCX 根目录",
    )
    parser.add_argument(
        "--question-dir",
        default=str(DEFAULT_QUESTION_DIR),
        help="配对文言文题目 DOCX 根目录",
    )
    parser.add_argument(
        "--output-dir",
        help="输出根目录；默认写回答案根目录并保留册别子目录",
    )
    parser.add_argument(
        "--preflight-only",
        action="store_true",
        help="只做28份整批只读预检，不写清洗文件",
    )
    return parser


def main() -> None:
    args = build_arg_parser().parse_args()
    source_pairs = discover_source_pairs(args.answer_dir, args.question_dir)
    clean_documents = preflight_source_pairs(source_pairs)
    print(f"模板预检通过：{len(clean_documents)} 份文档")
    print(
        "F1题块/源答案标记："
        f"{sum(len(document.blocks) for document in clean_documents)}/"
        f"{sum(document.source_answer_marker_count for document in clean_documents)}"
    )
    if args.preflight_only:
        return

    output_paths = clean_batch(
        args.answer_dir,
        args.question_dir,
        args.output_dir,
    )
    print(f"清洗完成：{len(output_paths)} 份文档")
    for output_path in output_paths:
        print(f"- {output_path}")


if __name__ == "__main__":
    main()
