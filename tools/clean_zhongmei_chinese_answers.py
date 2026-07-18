from __future__ import annotations

import argparse
import re
import shutil
import sys
import tempfile
from dataclasses import dataclass, field
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from tools.zhongmei_review_status import infer_question_dir, refresh_review_statuses
from tools.merge_zhongmei_reading_answer_groups import (
    merge_document_in_place,
    validate_reading_groups_premerged,
)


DEFAULT_SOURCE_DIR = Path(
    r"D:\墨痕教育题目\众美-高三-语文\答案\对点练案答案"
)
SOURCE_FILE_PATTERN = re.compile(r"^对点练案(\d+)\.docx$")
QUESTION_PATTERN = re.compile(r"^\s*(\d{1,3})\s*[．.]\s*")
SUBQUESTION_PATTERN = re.compile(r"^\s*[（(](\d+)[）)]")
ANSWER_MARKER_PATTERN = re.compile(r"^\s*答案(?:[\s　]*[：:]\s*|[\s　]+|$)")
ANALYSIS_MARKER_PATTERN = re.compile(r"^\s*解析(?:[\s　]*[：:]\s*|[\s　]+|$)")
LEVEL_ONE_HEADING_PATTERN = re.compile(r"^\s*[一二三四五六七八九十百]+、\S+")
LEVEL_TWO_HEADING_PATTERN = re.compile(r"^\s*[（(][一二三四五六七八九十百]+[）)]\S+")
READING_HEADING_PATTERN = re.compile(r"^\s*阅读下面.*完成(?:文后|后面)?题目[。.]?\s*$")


class UnsupportedTemplateError(ValueError):
    """源文档不符合已确认的众美对点练模板。"""


@dataclass
class AnswerOccurrence:
    label: str | None = None
    answer_lines: list[str] = field(default_factory=list)
    analysis_lines: list[str] = field(default_factory=list)


@dataclass
class CleanQuestion:
    number: int
    source_index: int
    occurrences: list[AnswerOccurrence]


@dataclass
class CleanHeading:
    text: str
    source_index: int


@dataclass
class CleanDocument:
    source_path: Path
    headings: list[CleanHeading]
    questions: list[CleanQuestion]


def _clean_text(text: str) -> str:
    return (
        (text or "")
        .replace("\r", "")
        .replace("\n", "")
        .replace("\x07", "")
        .replace("\u00a0", " ")
        .strip()
    )


def _is_heading(text: str) -> bool:
    return bool(
        LEVEL_ONE_HEADING_PATTERN.match(text)
        or LEVEL_TWO_HEADING_PATTERN.match(text)
        or READING_HEADING_PATTERN.match(text)
    )


def _marker_payload(pattern: re.Pattern[str], text: str) -> str:
    match = pattern.match(text)
    if not match:
        return ""
    return text[match.end() :].strip()


def _next_marker_kind(
    paragraphs: list[str],
    start_index: int,
    end_index: int,
) -> str | None:
    for index in range(start_index + 1, end_index):
        text = paragraphs[index]
        if ANSWER_MARKER_PATTERN.match(text):
            return "answer"
        if ANALYSIS_MARKER_PATTERN.match(text):
            return "analysis"
    return None


def _find_answer_associations(
    paragraphs: list[str], source_path: Path
) -> list[tuple[int, int, int]]:
    question_candidates: list[tuple[int, int]] = []
    answer_indices: list[int] = []

    for index, text in enumerate(paragraphs):
        question_match = QUESTION_PATTERN.match(text)
        if question_match:
            question_candidates.append((index, int(question_match.group(1))))
        if ANSWER_MARKER_PATTERN.match(text):
            answer_indices.append(index)

    if not answer_indices:
        raise UnsupportedTemplateError(f"{source_path.name}：未发现行首“答案”标记")

    associations: list[tuple[int, int, int]] = []
    candidate_cursor = 0
    latest_candidate: tuple[int, int] | None = None
    for answer_index in answer_indices:
        while (
            candidate_cursor < len(question_candidates)
            and question_candidates[candidate_cursor][0] < answer_index
        ):
            latest_candidate = question_candidates[candidate_cursor]
            candidate_cursor += 1
        if latest_candidate is None:
            raise UnsupportedTemplateError(
                f"{source_path.name}：第 {answer_index + 1} 段的答案前没有可识别题号"
            )
        question_index, question_number = latest_candidate
        associations.append((answer_index, question_index, question_number))

    sequence: list[int] = []
    question_indices: dict[int, int] = {}
    for _, question_index, question_number in associations:
        if not sequence or sequence[-1] != question_number:
            sequence.append(question_number)
        previous_index = question_indices.setdefault(question_number, question_index)
        if previous_index != question_index:
            raise UnsupportedTemplateError(
                f"{source_path.name}：题号 {question_number} 对应到多个题干，无法确定边界"
            )

    expected = list(range(1, len(sequence) + 1))
    if sequence != expected:
        raise UnsupportedTemplateError(
            f"{source_path.name}：答案对应题号不是从 1 开始连续排列，实际为 {sequence}"
        )

    return associations


def _parse_question_block(
    paragraphs: list[str],
    question_number: int,
    start_index: int,
    end_index: int,
    source_path: Path,
) -> CleanQuestion:
    occurrences: list[AnswerOccurrence] = []
    current: AnswerOccurrence | None = None
    state = "seeking"
    pending_label: str | None = None

    for index in range(start_index + 1, end_index):
        text = paragraphs[index]
        if not text:
            continue
        if _is_heading(text):
            current = None
            state = "seeking"
            pending_label = None
            continue

        subquestion_match = SUBQUESTION_PATTERN.match(text)
        if subquestion_match:
            next_marker = _next_marker_kind(paragraphs, index, end_index)
            if next_marker == "answer":
                pending_label = subquestion_match.group(1)
                state = "seeking"
                current = None
                continue

        if ANSWER_MARKER_PATTERN.match(text):
            current = AnswerOccurrence(label=pending_label)
            pending_label = None
            payload = _marker_payload(ANSWER_MARKER_PATTERN, text)
            if payload:
                current.answer_lines.append(payload)
            occurrences.append(current)
            state = "answer"
            continue

        if ANALYSIS_MARKER_PATTERN.match(text):
            if current is None:
                raise UnsupportedTemplateError(
                    f"{source_path.name}：第 {index + 1} 段出现解析，但当前题没有对应答案"
                )
            payload = _marker_payload(ANALYSIS_MARKER_PATTERN, text)
            if payload:
                current.analysis_lines.append(payload)
            state = "analysis"
            continue

        if current is None:
            continue
        if state == "answer":
            current.answer_lines.append(text)
        elif state == "analysis":
            current.analysis_lines.append(text)

    if not occurrences:
        raise UnsupportedTemplateError(f"{source_path.name}：第 {question_number} 题未提取到答案")
    for occurrence_index, occurrence in enumerate(occurrences, 1):
        if not occurrence.answer_lines:
            raise UnsupportedTemplateError(
                f"{source_path.name}：第 {question_number} 题第 {occurrence_index} 个答案为空"
            )

    if len(occurrences) > 1:
        labels = [occurrence.label for occurrence in occurrences]
        if any(label is None for label in labels):
            raise UnsupportedTemplateError(
                f"{source_path.name}：第 {question_number} 题有多个答案，但小问序号不完整"
            )
        if len(set(labels)) != len(labels):
            raise UnsupportedTemplateError(
                f"{source_path.name}：第 {question_number} 题的小问序号重复，实际为 {labels}"
            )

    return CleanQuestion(
        number=question_number,
        source_index=start_index,
        occurrences=occurrences,
    )


def parse_document(source_path: str | Path) -> CleanDocument:
    source_path = Path(source_path)
    doc = Document(source_path)
    paragraphs = [_clean_text(paragraph.text) for paragraph in doc.paragraphs]

    for table_index, table in enumerate(doc.tables, 1):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = _clean_text(paragraph.text)
                    if ANSWER_MARKER_PATTERN.match(text) or ANALYSIS_MARKER_PATTERN.match(text):
                        raise UnsupportedTemplateError(
                            f"{source_path.name}：第 {table_index} 个表格中含答案或解析标记，"
                            "当前模板不确定如何清洗"
                        )

    associations = _find_answer_associations(paragraphs, source_path)
    first_association_by_question: dict[int, tuple[int, int]] = {}
    for answer_index, question_index, question_number in associations:
        first_association_by_question.setdefault(
            question_number, (answer_index, question_index)
        )

    question_starts = [
        first_association_by_question[number][1]
        for number in sorted(first_association_by_question)
    ]
    questions: list[CleanQuestion] = []
    for offset, question_number in enumerate(sorted(first_association_by_question)):
        start_index = first_association_by_question[question_number][1]
        end_index = (
            question_starts[offset + 1]
            if offset + 1 < len(question_starts)
            else len(paragraphs)
        )
        questions.append(
            _parse_question_block(
                paragraphs,
                question_number,
                start_index,
                end_index,
                source_path,
            )
        )

    headings = [
        CleanHeading(text=text, source_index=index)
        for index, text in enumerate(paragraphs)
        if text and _is_heading(text)
    ]

    return CleanDocument(
        source_path=source_path,
        headings=headings,
        questions=questions,
    )


def render_clean_lines(clean_document: CleanDocument) -> list[str]:
    events: list[tuple[int, int, CleanHeading | CleanQuestion]] = []
    events.extend((heading.source_index, 0, heading) for heading in clean_document.headings)
    events.extend((question.source_index, 1, question) for question in clean_document.questions)
    events.sort(key=lambda item: (item[0], item[1]))

    lines: list[str] = []
    for _, _, event in events:
        if isinstance(event, CleanHeading):
            lines.append(event.text)
            continue

        lines.append(f"{event.number}．")
        multiple = len(event.occurrences) > 1
        has_any_analysis = any(occurrence.analysis_lines for occurrence in event.occurrences)

        for occurrence in event.occurrences:
            label_prefix = f"({occurrence.label})" if multiple else ""
            lines.append(f"{label_prefix}答案：{occurrence.answer_lines[0]}")
            lines.extend(occurrence.answer_lines[1:])

            if multiple and has_any_analysis:
                analysis_prefix = f"({occurrence.label})解析："
                if occurrence.analysis_lines:
                    lines.append(analysis_prefix + occurrence.analysis_lines[0])
                    lines.extend(occurrence.analysis_lines[1:])
                else:
                    lines.append(analysis_prefix + " ")

        if not multiple:
            occurrence = event.occurrences[0]
            if occurrence.analysis_lines:
                lines.append("解析：" + occurrence.analysis_lines[0])
                lines.extend(occurrence.analysis_lines[1:])
            else:
                lines.append("解析： ")
        elif not has_any_analysis:
            lines.append("解析： ")

    return lines


def _set_run_font(run, *, bold: bool = False) -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")


def write_clean_docx(clean_document: CleanDocument, output_path: str | Path) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    heading_texts = {heading.text for heading in clean_document.headings}
    for line in render_clean_lines(clean_document):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.paragraph_format.line_spacing = 1.25
        run = paragraph.add_run(line)
        _set_run_font(run, bold=line in heading_texts)

    if doc.paragraphs and not doc.paragraphs[0].text:
        first = doc.paragraphs[0]._element
        first.getparent().remove(first)

    doc.save(output_path)
    return output_path


def validate_clean_docx(
    output_path: str | Path,
    expected_document: CleanDocument,
) -> None:
    output_path = Path(output_path)
    doc = Document(output_path)
    actual_lines = [_clean_text(paragraph.text) for paragraph in doc.paragraphs if _clean_text(paragraph.text)]
    expected_lines = [_clean_text(line) for line in render_clean_lines(expected_document) if _clean_text(line)]
    if actual_lines != expected_lines:
        raise ValueError(f"{output_path.name}：写入后文本与预期不一致")

    question_lines = [line for line in actual_lines if re.fullmatch(r"\d+．", line)]
    expected_question_lines = [f"{question.number}．" for question in expected_document.questions]
    if question_lines != expected_question_lines:
        raise ValueError(
            f"{output_path.name}：题号复查失败，实际 {question_lines}，预期 {expected_question_lines}"
        )

    analysis_count = sum(
        1
        for line in actual_lines
        if line.startswith("解析：") or re.match(r"^\(\d+\)解析：", line)
    )
    if analysis_count < len(expected_document.questions):
        raise ValueError(
            f"{output_path.name}：解析标记不足，实际 {analysis_count}，"
            f"题目数 {len(expected_document.questions)}"
        )

    forbidden_fragments = ("(分值：", "A．", "B．", "C．", "D．", "答：___")
    leaked = [
        line
        for line in actual_lines
        if any(fragment in line for fragment in forbidden_fragments)
        and not (line.startswith("答案：") or line.startswith("解析：") or "答案：" in line or "解析：" in line)
    ]
    if leaked:
        raise ValueError(f"{output_path.name}：疑似残留题干或选项：{leaked[:3]}")


def discover_source_files(source_dir: str | Path) -> list[Path]:
    source_dir = Path(source_dir)
    files: list[tuple[int, Path]] = []
    for path in source_dir.glob("*.docx"):
        match = SOURCE_FILE_PATTERN.match(path.name)
        if match:
            files.append((int(match.group(1)), path))
    files.sort(key=lambda item: item[0])
    if not files:
        raise FileNotFoundError(f"未在 {source_dir} 找到 对点练案N.docx")
    return [path for _, path in files]


def preflight_source_files(source_files: list[Path]) -> list[CleanDocument]:
    clean_documents: list[CleanDocument] = []
    errors: list[str] = []
    for source_path in source_files:
        try:
            clean_documents.append(parse_document(source_path))
        except Exception as exc:
            errors.append(str(exc))
    if errors:
        details = "\n".join(f"- {error}" for error in errors)
        raise UnsupportedTemplateError(
            "发现无法确定清洗方式的文档，已停止整批清洗：\n" + details
        )
    return clean_documents


def clean_batch(
    source_dir: str | Path = DEFAULT_SOURCE_DIR,
    output_dir: str | Path | None = None,
    *,
    preflight_only: bool = False,
) -> list[Path]:
    source_dir = Path(source_dir)
    output_dir = Path(output_dir) if output_dir else source_dir
    source_files = discover_source_files(source_dir)
    clean_documents = preflight_source_files(source_files)
    if preflight_only:
        return []

    output_dir.mkdir(parents=True, exist_ok=True)
    temp_dir = Path(tempfile.mkdtemp(prefix=".对点练清洗临时_", dir=output_dir))
    temp_outputs: list[tuple[Path, Path]] = []
    try:
        for clean_document in clean_documents:
            final_path = output_dir / f"{clean_document.source_path.stem}_已清洗.docx"
            temp_path = temp_dir / final_path.name
            write_clean_docx(clean_document, temp_path)
            validate_clean_docx(temp_path, clean_document)
            question_path = infer_question_dir(final_path) / f"{clean_document.source_path.stem}.docx"
            if question_path.exists():
                merge_document_in_place(temp_path, question_path=question_path)
                validate_reading_groups_premerged(temp_path, question_path=question_path)
            temp_outputs.append((temp_path, final_path))

        for temp_path, final_path in temp_outputs:
            temp_path.replace(final_path)
        output_paths = [final_path for _, final_path in temp_outputs]
        if output_paths and infer_question_dir(output_paths[0]).is_dir():
            refresh_review_statuses(output_paths)
        return output_paths
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="清洗众美高三语文对点练案：仅保留标题、题号、答案和解析。"
    )
    parser.add_argument(
        "source_dir",
        nargs="?",
        default=str(DEFAULT_SOURCE_DIR),
        help="原始对点练案 DOCX 目录",
    )
    parser.add_argument(
        "--output-dir",
        help="输出目录；默认与源文件同目录，文件名追加 _已清洗",
    )
    parser.add_argument(
        "--preflight-only",
        action="store_true",
        help="只做模板预检，不写入任何清洗文件",
    )
    return parser


def main() -> None:
    args = build_arg_parser().parse_args()
    source_files = discover_source_files(args.source_dir)
    clean_documents = preflight_source_files(source_files)
    print(f"模板预检通过：{len(clean_documents)} 份文档")
    print(
        "题目总数："
        + str(sum(len(clean_document.questions) for clean_document in clean_documents))
    )
    if args.preflight_only:
        return

    output_paths = clean_batch(args.source_dir, args.output_dir)
    print(f"清洗完成：{len(output_paths)} 份文档")
    for output_path in output_paths:
        print(f"- {output_path}")


if __name__ == "__main__":
    main()
