from __future__ import annotations

import argparse
from copy import deepcopy
from dataclasses import asdict, dataclass
from hashlib import sha256
import json
import posixpath
from pathlib import Path
import re
import sys
import tempfile
from zipfile import ZIP_DEFLATED, BadZipFile, ZipFile

from docx import Document
from lxml import etree as ET

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from shared_core import (
    build_answer_units_from_docx,
    build_question_units_from_docx,
    build_review_report,
    export_review_report,
    initialize_review_status,
    map_answers,
    update_review_status,
)
from shared_core.cli_output import configure_utf8_stdio
from shared_core.subject_overlay import (
    FUTURE_BIOLOGY_STRUCTURE_PATTERNS,
    get_subject_overlay,
)


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
V_NS = "urn:schemas-microsoft-com:vml"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PR_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
W = f"{{{W_NS}}}"
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
DOCUMENT_XML = "word/document.xml"
DOCUMENT_RELS = "word/_rels/document.xml.rels"
INPUT_RELATIVE_ROOT = Path("答案") / "按课时截取"
OUTPUT_RELATIVE_ROOT = Path("答案") / "已清洗"
ANSWER_DIRECTORIES = ("选必一答案", "选必二答案")
MANIFEST_FILENAME = "FutureBiologyAnswerCleanManifest.json"

QUESTION_PATTERN = re.compile(r"^\s*(\d+)\s*[．.、]")
ANSWER_MARKER_PATTERN = re.compile(
    r"^\s*(?:【答案】|答案)(?:(?:\s*[：:]\s*)|\s+|(?=$))"
)
ANALYSIS_MARKER_PATTERN = re.compile(
    r"^\s*(?:【解析】|解析)(?:(?:\s*[：:]\s*)|\s+|(?=$))"
)
OPTION_PATTERN = re.compile(r"^\s*[A-DＡ-Ｄ]\s*[．.、]")
STRUCTURE_PATTERNS = tuple(re.compile(pattern) for pattern in FUTURE_BIOLOGY_STRUCTURE_PATTERNS)

RICH_CONTENT_TAGS = {
    f"{{{M_NS}}}oMath",
    f"{{{M_NS}}}oMathPara",
    f"{{{W_NS}}}drawing",
    f"{{{W_NS}}}pict",
    f"{{{W_NS}}}object",
    f"{{{A_NS}}}blip",
    f"{{{V_NS}}}imagedata",
}

ET.register_namespace("w", W_NS)


@dataclass(frozen=True)
class CleanPreflight:
    source_path: Path
    relative_path: Path
    source_sha256: str
    question_count: int
    missing_analysis_count: int
    excluded_structure_count: int
    excluded_media_count: int
    retained_media_count: int


@dataclass(frozen=True)
class CleanResult:
    source_path: Path
    output_path: Path
    review_report_path: Path
    question_count: int
    inserted_analysis_placeholders: int
    excluded_structure_count: int
    excluded_media_count: int
    retained_media_count: int
    blocking_issue_count: int
    source_sha256: str
    output_sha256: str


@dataclass(frozen=True)
class _AnswerBlock:
    question_id: str
    reference_paragraph: object
    question_children: tuple[object, ...]
    answer_children: tuple[object, ...]
    analysis_children: tuple[object, ...]
    excluded_structure_count: int
    excluded_media_count: int


def _sha256_file(path: Path) -> str:
    digest = sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _is_lock_file(path: Path) -> bool:
    return path.name.startswith(("~$", ".~", "~"))


def discover_trimmed_answer_documents(project_root: str | Path) -> list[Path]:
    input_root = Path(project_root) / INPUT_RELATIVE_ROOT
    missing = [name for name in ANSWER_DIRECTORIES if not (input_root / name).is_dir()]
    if missing:
        raise FileNotFoundError(
            "缺少按课时截取答案目录: "
            + "、".join(str(input_root / name) for name in missing)
        )

    paths: list[Path] = []
    for directory_name in ANSWER_DIRECTORIES:
        paths.extend(
            path
            for path in sorted((input_root / directory_name).glob("*.docx"))
            if not _is_lock_file(path) and not path.stem.endswith("_已清洗")
        )
    return paths


def _read_package(path: Path) -> tuple[ET._Element, dict[str, str], dict[str, bytes]]:
    try:
        with ZipFile(path, "r") as package:
            bad_member = package.testzip()
            if bad_member:
                raise ValueError(f"DOCX 压缩包损坏: {path.name} -> {bad_member}")
            root = ET.fromstring(package.read(DOCUMENT_XML))
            relationships: dict[str, str] = {}
            if DOCUMENT_RELS in package.namelist():
                rel_root = ET.fromstring(package.read(DOCUMENT_RELS))
                for relationship in rel_root.findall(f"{{{PR_NS}}}Relationship"):
                    if relationship.get("TargetMode") == "External":
                        continue
                    relationship_id = relationship.get("Id")
                    target = relationship.get("Target")
                    if relationship_id and target:
                        relationships[relationship_id] = target
            media_payloads = {
                name: package.read(name)
                for name in package.namelist()
                if name.startswith("word/media/")
            }
            return root, relationships, media_payloads
    except BadZipFile as exc:
        raise ValueError(f"不是有效的 DOCX 文件: {path}") from exc
    except KeyError as exc:
        raise ValueError(f"DOCX 缺少 {DOCUMENT_XML}: {path}") from exc


def _paragraph_text(paragraph) -> str:
    return "".join(node.text or "" for node in paragraph.iter(f"{W}t"))


def _set_text(node, value: str) -> None:
    node.text = value
    if value.startswith(" ") or value.endswith(" "):
        node.set(XML_SPACE, "preserve")
    else:
        node.attrib.pop(XML_SPACE, None)


def _replace_text_prefix(paragraph, prefix_length: int, replacement: str) -> None:
    text_nodes = list(paragraph.iter(f"{W}t"))
    if not text_nodes:
        raise ValueError("答案或解析标记所在段落没有文本节点")

    cursor = 0
    replacement_written = False
    for node in text_nodes:
        value = node.text or ""
        start = cursor
        end = start + len(value)
        cursor = end

        if end <= prefix_length:
            if not replacement_written:
                _set_text(node, replacement)
                replacement_written = True
            else:
                _set_text(node, "")
            continue
        if start < prefix_length:
            suffix = value[prefix_length - start :]
            _set_text(node, ("" if replacement_written else replacement) + suffix)
            replacement_written = True
            continue
        if not replacement_written:
            _set_text(node, replacement + value)
            replacement_written = True

    if not replacement_written:
        raise ValueError("答案或解析标记替换失败")


def _normalize_marker_run_format(paragraph, marker: str) -> None:
    marker_node = next(
        (
            node
            for node in paragraph.iter(f"{W}t")
            if (node.text or "").startswith(marker)
        ),
        None,
    )
    if marker_node is None:
        raise ValueError(f"标准化后的段落缺少标记: {marker}")

    run = marker_node.getparent()
    while run is not None and run.tag != f"{W}r":
        run = run.getparent()
    if run is None:
        raise ValueError(f"标记不在普通文本运行中，无法设置可见格式: {marker}")

    run_properties = run.find(f"{W}rPr")
    if run_properties is None:
        run_properties = ET.Element(f"{W}rPr")
        run.insert(0, run_properties)
    for hidden_tag in ("vanish", "webHidden", "specVanish"):
        for hidden_node in list(run_properties.findall(f"{W}{hidden_tag}")):
            run_properties.remove(hidden_node)

    run_fonts = run_properties.find(f"{W}rFonts")
    if run_fonts is None:
        run_fonts = ET.SubElement(run_properties, f"{W}rFonts")
    run_fonts.set(f"{W}ascii", "Times New Roman")
    run_fonts.set(f"{W}hAnsi", "Times New Roman")
    run_fonts.set(f"{W}eastAsia", "宋体")
    run_fonts.set(f"{W}cs", "Times New Roman")
    run_fonts.set(f"{W}hint", "eastAsia")

    color = run_properties.find(f"{W}color")
    if color is None:
        color = ET.SubElement(run_properties, f"{W}color")
    color.set(f"{W}val", "000000")
    color.attrib.pop(f"{W}themeColor", None)
    color.attrib.pop(f"{W}themeTint", None)
    color.attrib.pop(f"{W}themeShade", None)


def _new_text_paragraph(reference_paragraph, value: str):
    paragraph = ET.Element(f"{W}p")
    paragraph_properties = reference_paragraph.find(f"{W}pPr")
    if paragraph_properties is not None:
        paragraph.append(deepcopy(paragraph_properties))

    run = ET.SubElement(paragraph, f"{W}r")
    reference_run_properties = reference_paragraph.find(f".//{W}rPr")
    if reference_run_properties is not None:
        run.append(deepcopy(reference_run_properties))
    text = ET.SubElement(run, f"{W}t")
    _set_text(text, value)
    if value.startswith(("答案：", "解析：")):
        _normalize_marker_run_format(paragraph, value[:3])
    return paragraph


def _relationship_package_path(target: str) -> str:
    return posixpath.normpath(posixpath.join("word", target.lstrip("/")))


def _child_media_hashes(
    child,
    relationships: dict[str, str],
    media_payloads: dict[str, bytes],
) -> tuple[str, ...]:
    relationship_ids = child.xpath(
        './/*[local-name()="blip"]/@*[local-name()="embed"]'
        ' | .//*[local-name()="imagedata"]/@*[local-name()="id"]'
    )
    hashes: list[str] = []
    for relationship_id in relationship_ids:
        target = relationships.get(str(relationship_id))
        if not target:
            continue
        payload = media_payloads.get(_relationship_package_path(target))
        if payload is not None:
            hashes.append(sha256(payload).hexdigest())
    return tuple(hashes)


def _has_rich_content(children: tuple[object, ...] | list[object]) -> bool:
    return any(
        node.tag in RICH_CONTENT_TAGS
        for child in children
        for node in child.iter()
    )


def _is_empty_plain_paragraph(child) -> bool:
    return (
        child.tag == f"{W}p"
        and not _paragraph_text(child).strip()
        and not _has_rich_content([child])
    )


def _trim_trailing_empty(children: list[object]) -> list[object]:
    while children and _is_empty_plain_paragraph(children[-1]):
        children.pop()
    return children


def _is_structure_text(text: str) -> bool:
    return any(pattern.match(text or "") for pattern in STRUCTURE_PATTERNS)


def _excluded_media_hashes() -> set[str]:
    overlay = get_subject_overlay("future_biology")
    if overlay is None:
        raise RuntimeError("未注册 future_biology 学科覆盖层")
    return {
        digest
        for digests in overlay.excluded_media_sha256_by_role.values()
        for digest in digests
    }


def _extract_answer_blocks(
    root: ET._Element,
    relationships: dict[str, str],
    media_payloads: dict[str, bytes],
    source_name: str,
) -> list[_AnswerBlock]:
    body = root.find(f"{W}body")
    if body is None:
        raise ValueError(f"文档缺少正文节点: {source_name}")
    children = list(body)
    question_positions: list[tuple[int, str, object]] = []
    for position, child in enumerate(children):
        if child.tag != f"{W}p":
            continue
        match = QUESTION_PATTERN.match(_paragraph_text(child))
        if match:
            question_positions.append((position, match.group(1), child))

    if not question_positions:
        raise ValueError(f"{source_name} 未识别到顶层题号")
    question_ids = [int(question_id) for _, question_id, _ in question_positions]
    if question_ids != list(range(1, len(question_ids) + 1)):
        raise ValueError(f"{source_name} 顶层题号不是从 1 连续排列: {question_ids}")

    excluded_hashes = _excluded_media_hashes()
    blocks: list[_AnswerBlock] = []
    for index, (question_position, question_id, question_paragraph) in enumerate(
        question_positions
    ):
        next_question_position = (
            question_positions[index + 1][0]
            if index + 1 < len(question_positions)
            else len(children)
        )
        if children and children[-1].tag == f"{W}sectPr":
            next_question_position = min(next_question_position, len(children) - 1)

        answer_positions = [
            position
            for position in range(question_position + 1, next_question_position)
            if children[position].tag == f"{W}p"
            and ANSWER_MARKER_PATTERN.match(_paragraph_text(children[position]))
        ]
        if len(answer_positions) != 1:
            raise ValueError(
                f"{source_name} 第 {question_id} 题答案标记数量异常: {len(answer_positions)}"
            )
        answer_position = answer_positions[0]

        payload_boundary = next_question_position
        excluded_structure_count = 0
        excluded_media_count = 0
        for position in range(answer_position + 1, next_question_position):
            child = children[position]
            text = _paragraph_text(child).strip() if child.tag == f"{W}p" else ""
            media_hashes = _child_media_hashes(child, relationships, media_payloads)
            if _is_structure_text(text):
                payload_boundary = position
                excluded_structure_count = 1
                break
            if any(digest in excluded_hashes for digest in media_hashes):
                payload_boundary = position
                excluded_media_count = 1
                break

        for position in range(payload_boundary + 1, next_question_position):
            child = children[position]
            text = _paragraph_text(child).strip() if child.tag == f"{W}p" else ""
            media_hashes = _child_media_hashes(child, relationships, media_payloads)
            excluded_structure_count += int(_is_structure_text(text))
            excluded_media_count += sum(
                digest in excluded_hashes for digest in media_hashes
            )

        analysis_positions = [
            position
            for position in range(answer_position + 1, payload_boundary)
            if children[position].tag == f"{W}p"
            and ANALYSIS_MARKER_PATTERN.match(_paragraph_text(children[position]))
        ]
        if len(analysis_positions) > 1:
            raise ValueError(
                f"{source_name} 第 {question_id} 题解析标记数量异常: {len(analysis_positions)}"
            )
        analysis_position = analysis_positions[0] if analysis_positions else None
        answer_end = analysis_position if analysis_position is not None else payload_boundary

        answer_children = _trim_trailing_empty(
            list(children[answer_position:answer_end])
        )
        analysis_children = _trim_trailing_empty(
            list(children[analysis_position:payload_boundary])
            if analysis_position is not None
            else []
        )
        if not answer_children:
            raise ValueError(f"{source_name} 第 {question_id} 题答案内容为空")
        first_answer = answer_children[0]
        if first_answer.tag != f"{W}p" or not ANSWER_MARKER_PATTERN.match(
            _paragraph_text(first_answer)
        ):
            raise ValueError(f"{source_name} 第 {question_id} 题答案块首段异常")

        blocks.append(
            _AnswerBlock(
                question_id=question_id,
                reference_paragraph=question_paragraph,
                question_children=tuple(
                    children[question_position:answer_position]
                ),
                answer_children=tuple(answer_children),
                analysis_children=tuple(analysis_children),
                excluded_structure_count=excluded_structure_count,
                excluded_media_count=excluded_media_count,
            )
        )
    return blocks


def _normalize_payload_children(
    children: tuple[object, ...],
    pattern: re.Pattern[str],
    marker: str,
) -> list[object]:
    normalized = [deepcopy(child) for child in children]
    first = normalized[0]
    if first.tag != f"{W}p":
        raise ValueError(f"{marker}块首节点不是段落")
    match = pattern.match(_paragraph_text(first))
    if not match:
        raise ValueError(f"{marker}块首段缺少标记")
    _replace_text_prefix(first, match.end(), marker)
    _normalize_marker_run_format(first, marker)

    retained = [first]
    for child in normalized[1:]:
        if (
            child.tag != f"{W}p"
            or _has_rich_content([child])
            or not _paragraph_text(child).strip()
        ):
            retained.append(child)
            continue

        separator_run = ET.Element(f"{W}r")
        separator_text = ET.SubElement(separator_run, f"{W}t")
        _set_text(separator_text, " ")
        first.append(separator_run)
        for paragraph_child in list(child):
            if paragraph_child.tag != f"{W}pPr":
                first.append(paragraph_child)
    return retained


def _render_clean_body(root: ET._Element, blocks: list[_AnswerBlock]) -> bytes:
    body = root.find(f"{W}body")
    if body is None:
        raise ValueError("文档缺少正文节点")
    section_properties = body.find(f"{W}sectPr")
    for child in list(body):
        if child is not section_properties:
            body.remove(child)

    rendered: list[object] = []
    for block in blocks:
        rendered.append(
            _new_text_paragraph(block.reference_paragraph, f"{block.question_id}．")
        )
        rendered.extend(
            _normalize_payload_children(
                block.answer_children,
                ANSWER_MARKER_PATTERN,
                "答案：",
            )
        )
        if block.analysis_children:
            rendered.extend(
                _normalize_payload_children(
                    block.analysis_children,
                    ANALYSIS_MARKER_PATTERN,
                    "解析：",
                )
            )
        else:
            rendered.append(
                _new_text_paragraph(block.reference_paragraph, "解析： ")
            )

    for child in rendered:
        if section_properties is None:
            body.append(child)
        else:
            body.insert(list(body).index(section_properties), child)
    return ET.tostring(
        root,
        encoding="utf-8",
        xml_declaration=True,
        standalone=True,
    )


def _render_question_only_body(root: ET._Element, blocks: list[_AnswerBlock]) -> bytes:
    body = root.find(f"{W}body")
    if body is None:
        raise ValueError("文档缺少正文节点")
    section_properties = body.find(f"{W}sectPr")
    for child in list(body):
        if child is not section_properties:
            body.remove(child)

    for block in blocks:
        for child in block.question_children:
            cloned = deepcopy(child)
            if section_properties is None:
                body.append(cloned)
            else:
                body.insert(list(body).index(section_properties), cloned)
    return ET.tostring(
        root,
        encoding="utf-8",
        xml_declaration=True,
        standalone=True,
    )


def _build_question_units_for_review(source_path: Path):
    root, relationships, media_payloads = _read_package(source_path)
    blocks = _extract_answer_blocks(
        root,
        relationships,
        media_payloads,
        source_path.name,
    )
    question_xml = _render_question_only_body(root, blocks)
    with tempfile.TemporaryDirectory(prefix="future_biology_questions_") as temp_dir:
        question_path = Path(temp_dir) / source_path.name
        _write_docx_with_document_xml(
            source_path,
            question_path,
            question_xml,
            overwrite=False,
        )
        return build_question_units_from_docx(question_path, grade_hint="高二")


def _write_docx_with_document_xml(
    source_path: Path,
    output_path: Path,
    document_xml: bytes,
    *,
    overwrite: bool,
) -> None:
    if source_path.resolve() == output_path.resolve():
        raise ValueError("清洗输出不能覆盖截取答案或原答案")
    if output_path.exists() and not overwrite:
        raise FileExistsError(f"输出已存在，未覆盖: {output_path}")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    with tempfile.NamedTemporaryFile(
        prefix=f"{output_path.stem}_",
        suffix=".docx",
        dir=output_path.parent,
        delete=False,
    ) as temporary_file:
        temporary_path = Path(temporary_file.name)

    try:
        with ZipFile(source_path, "r") as source_zip, ZipFile(
            temporary_path,
            "w",
            compression=ZIP_DEFLATED,
        ) as target_zip:
            for info in source_zip.infolist():
                payload = (
                    document_xml
                    if info.filename == DOCUMENT_XML
                    else source_zip.read(info.filename)
                )
                target_zip.writestr(info, payload)
        with ZipFile(temporary_path, "r") as check_zip:
            bad_member = check_zip.testzip()
            if bad_member:
                raise ValueError(f"清洗后的 DOCX 压缩包损坏: {bad_member}")
        Document(temporary_path)
        temporary_path.replace(output_path)
    finally:
        if temporary_path.exists():
            temporary_path.unlink()


def _count_retained_media(blocks: list[_AnswerBlock]) -> int:
    return sum(
        1
        for block in blocks
        for child in (*block.answer_children, *block.analysis_children)
        for node in child.iter()
        if node.tag in {f"{{{A_NS}}}blip", f"{{{V_NS}}}imagedata"}
    )


def preflight_answer_document(
    source_path: str | Path,
    *,
    input_root: str | Path | None = None,
) -> CleanPreflight:
    source = Path(source_path)
    if not source.is_file():
        raise FileNotFoundError(f"未找到截取答案: {source}")
    source_hash = _sha256_file(source)
    root, relationships, media_payloads = _read_package(source)
    blocks = _extract_answer_blocks(
        root,
        relationships,
        media_payloads,
        source.name,
    )
    if _sha256_file(source) != source_hash:
        raise RuntimeError(f"预检过程中源答案发生变化: {source}")
    relative_path = (
        source.relative_to(Path(input_root)) if input_root else Path(source.name)
    )
    return CleanPreflight(
        source_path=source,
        relative_path=relative_path,
        source_sha256=source_hash,
        question_count=len(blocks),
        missing_analysis_count=sum(not block.analysis_children for block in blocks),
        excluded_structure_count=sum(block.excluded_structure_count for block in blocks),
        excluded_media_count=sum(block.excluded_media_count for block in blocks),
        retained_media_count=_count_retained_media(blocks),
    )


def preflight_clean_batch(
    project_root: str | Path,
    *,
    expected_source_count: int | None = None,
    expected_question_count: int | None = None,
) -> list[CleanPreflight]:
    root = Path(project_root)
    input_root = root / INPUT_RELATIVE_ROOT
    sources = discover_trimmed_answer_documents(root)
    if expected_source_count is not None and len(sources) != expected_source_count:
        raise ValueError(
            f"截取答案数量不符合预期: 预期 {expected_source_count}，实际 {len(sources)}"
        )
    preflights = [
        preflight_answer_document(source, input_root=input_root) for source in sources
    ]
    total_questions = sum(item.question_count for item in preflights)
    if expected_question_count is not None and total_questions != expected_question_count:
        raise ValueError(
            f"答案题数不符合预期: 预期 {expected_question_count}，实际 {total_questions}"
        )
    return preflights


def clean_answer_document(
    source_path: str | Path,
    output_path: str | Path,
    *,
    overwrite: bool = False,
) -> tuple[CleanPreflight, str]:
    source = Path(source_path)
    output = Path(output_path)
    preflight = preflight_answer_document(source)
    root, relationships, media_payloads = _read_package(source)
    blocks = _extract_answer_blocks(
        root,
        relationships,
        media_payloads,
        source.name,
    )
    document_xml = _render_clean_body(root, blocks)
    _write_docx_with_document_xml(
        source,
        output,
        document_xml,
        overwrite=overwrite,
    )

    output_doc = Document(output)
    output_texts = [paragraph.text.strip() for paragraph in output_doc.paragraphs]
    answer_count = sum(bool(ANSWER_MARKER_PATTERN.match(text)) for text in output_texts)
    analysis_count = sum(bool(ANALYSIS_MARKER_PATTERN.match(text)) for text in output_texts)
    if len(output_doc.tables) or answer_count != len(blocks) or analysis_count != len(blocks):
        output.unlink(missing_ok=True)
        raise ValueError(
            f"{source.name} 清洗结果结构异常: 表格 {len(output_doc.tables)}，"
            f"答案 {answer_count}，解析 {analysis_count}，题数 {len(blocks)}"
        )
    if any(OPTION_PATTERN.match(text) or _is_structure_text(text) for text in output_texts):
        output.unlink(missing_ok=True)
        raise ValueError(f"{source.name} 清洗结果仍含题干选项或结构标题")

    answer_units = build_answer_units_from_docx(output, preserve_source_positions=True)
    if len(answer_units) != len(blocks):
        output.unlink(missing_ok=True)
        raise ValueError(
            f"{source.name} 清洗后答案块数量异常: 期望 {len(blocks)}，实际 {len(answer_units)}"
        )
    if _sha256_file(source) != preflight.source_sha256:
        output.unlink(missing_ok=True)
        raise RuntimeError(f"清洗过程中源答案发生变化: {source}")
    return preflight, _sha256_file(output)


def _write_manifest(
    results: list[CleanResult],
    report_root: Path,
) -> Path:
    report_root.mkdir(parents=True, exist_ok=True)
    manifest_path = report_root / MANIFEST_FILENAME
    payload = {
        "schema_version": "1.0",
        "project": "未来-高二-生物",
        "source_immutable": True,
        "document_count": len(results),
        "question_count": sum(result.question_count for result in results),
        "inserted_analysis_placeholders": sum(
            result.inserted_analysis_placeholders for result in results
        ),
        "excluded_structure_count": sum(
            result.excluded_structure_count for result in results
        ),
        "excluded_media_count": sum(
            result.excluded_media_count for result in results
        ),
        "retained_media_count": sum(result.retained_media_count for result in results),
        "approved_document_count": sum(
            result.blocking_issue_count == 0 for result in results
        ),
        "results": [
            {
                **asdict(result),
                "source_path": str(result.source_path.resolve()),
                "output_path": str(result.output_path.resolve()),
                "review_report_path": str(result.review_report_path.resolve()),
            }
            for result in results
        ],
    }
    manifest_path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    json.loads(manifest_path.read_text(encoding="utf-8"))
    return manifest_path


def clean_answer_batch(
    project_root: str | Path,
    *,
    output_dir: str | Path | None = None,
    report_dir: str | Path | None = None,
    expected_source_count: int | None = None,
    expected_question_count: int | None = None,
    overwrite: bool = False,
) -> tuple[list[CleanResult], Path]:
    root = Path(project_root)
    input_root = root / INPUT_RELATIVE_ROOT
    output_root = Path(output_dir) if output_dir else root / OUTPUT_RELATIVE_ROOT
    report_root = Path(report_dir) if report_dir else root / "答案" / "审核清单"
    preflights = preflight_clean_batch(
        root,
        expected_source_count=expected_source_count,
        expected_question_count=expected_question_count,
    )
    planned_outputs = [
        output_root
        / item.relative_path.parent
        / f"{item.source_path.stem}_已清洗.docx"
        for item in preflights
    ]
    existing = [path for path in planned_outputs if path.exists()]
    if existing and not overwrite:
        raise FileExistsError(
            f"已有 {len(existing)} 份清洗结果；如需重建请显式使用 --overwrite"
        )

    results: list[CleanResult] = []
    for preflight, output_path in zip(preflights, planned_outputs):
        refreshed, output_hash = clean_answer_document(
            preflight.source_path,
            output_path,
            overwrite=overwrite,
        )
        question_units = _build_question_units_for_review(preflight.source_path)
        answer_units = build_answer_units_from_docx(
            output_path,
            preserve_source_positions=True,
        )
        for answer_unit in answer_units:
            answer_unit.metadata["allow_answer_defined_subquestions"] = True
        mapped_units = map_answers(question_units, answer_units)
        review = build_review_report(output_path.name, question_units, mapped_units)
        review_report_path = (
            report_root
            / preflight.relative_path.parent
            / f"{output_path.stem}_审核清单.md"
        )
        export_review_report(review, review_report_path)
        initialize_review_status(
            output_path,
            report_path=str(review_report_path),
            report=review,
        )
        blocking_issue_count = sum(
            issue.severity == "error" for issue in review.issues
        )
        update_review_status(
            output_path,
            status="approved" if blocking_issue_count == 0 else "rejected",
            reviewer="system",
            note=(
                "未来高二生物答案自动检查通过"
                if blocking_issue_count == 0
                else f"自动检查发现 {blocking_issue_count} 个高风险问题"
            ),
        )
        results.append(
            CleanResult(
                source_path=preflight.source_path,
                output_path=output_path,
                review_report_path=review_report_path,
                question_count=refreshed.question_count,
                inserted_analysis_placeholders=refreshed.missing_analysis_count,
                excluded_structure_count=refreshed.excluded_structure_count,
                excluded_media_count=refreshed.excluded_media_count,
                retained_media_count=refreshed.retained_media_count,
                blocking_issue_count=blocking_issue_count,
                source_sha256=refreshed.source_sha256,
                output_sha256=output_hash,
            )
        )
    return results, _write_manifest(results, report_root)


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="清洗未来高二生物按课时答案：删除题干选项，标准化题号、答案和解析。"
    )
    parser.add_argument("project_root", help="未来-高二-生物项目根目录")
    parser.add_argument("--output-dir", help="清洗结果目录；默认 答案/已清洗")
    parser.add_argument("--report-dir", help="审核清单和清洗清单目录；默认 答案/审核清单")
    parser.add_argument("--preflight-only", action="store_true", help="只做整批只读预检")
    parser.add_argument("--overwrite", action="store_true", help="显式覆盖既有清洗结果")
    parser.add_argument("--expected-source-count", type=int)
    parser.add_argument("--expected-question-count", type=int)
    return parser


def main() -> int:
    configure_utf8_stdio()
    args = _build_parser().parse_args()
    preflights = preflight_clean_batch(
        args.project_root,
        expected_source_count=args.expected_source_count,
        expected_question_count=args.expected_question_count,
    )
    print(f"预检文件: {len(preflights)}")
    print(f"预检题数: {sum(item.question_count for item in preflights)}")
    print(f"需补空解析: {sum(item.missing_analysis_count for item in preflights)}")
    print(f"排除结构标题: {sum(item.excluded_structure_count for item in preflights)}")
    print(f"排除装饰媒体: {sum(item.excluded_media_count for item in preflights)}")
    print(f"保留答案/解析媒体: {sum(item.retained_media_count for item in preflights)}")
    if args.preflight_only:
        return 0

    results, manifest_path = clean_answer_batch(
        args.project_root,
        output_dir=args.output_dir,
        report_dir=args.report_dir,
        expected_source_count=args.expected_source_count,
        expected_question_count=args.expected_question_count,
        overwrite=args.overwrite,
    )
    approved = sum(result.blocking_issue_count == 0 for result in results)
    print(f"清洗文件: {len(results)}")
    print(f"答案总数: {sum(result.question_count for result in results)}")
    print(f"自动检查通过: {approved}/{len(results)}")
    print(f"清洗清单: {manifest_path.resolve()}")
    return 0 if approved == len(results) else 1


if __name__ == "__main__":
    raise SystemExit(main())
