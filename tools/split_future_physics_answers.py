from __future__ import annotations

import argparse
import re
from copy import deepcopy
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT


HEADING_PATTERN = re.compile(r"^(课时跟踪检测\([^)]+\))\s*$")
QUESTION_FILE_HEADING_PATTERN = re.compile(r"^(课时跟踪检测\([^)]+\))")


def _clean_text(text: str) -> str:
    return (text or "").replace("\r", "").replace("\n", "").replace("\x07", "").strip()


def _clear_doc_body(doc_obj: Document) -> None:
    body = doc_obj._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def extract_sections_from_docx(source_path: str | Path) -> list[dict]:
    doc = Document(source_path)
    sections: list[dict] = []
    current_section: dict | None = None

    for paragraph_index, paragraph in enumerate(doc.paragraphs):
        text = _clean_text(paragraph.text)
        if not text and current_section is None:
            continue

        heading_match = HEADING_PATTERN.match(text)
        if heading_match:
            if current_section is not None and current_section["paragraphs"]:
                sections.append(current_section)
            current_section = {
                "heading": heading_match.group(1),
                "paragraphs": [heading_match.group(1)],
                "paragraph_indices": [paragraph_index],
            }
            continue

        if current_section is None:
            continue

        current_section["paragraphs"].append(text)
        current_section["paragraph_indices"].append(paragraph_index)

    if current_section is not None and current_section["paragraphs"]:
        sections.append(current_section)

    return sections


def build_question_heading_map(question_dir: str | Path) -> dict[str, str]:
    mapping: dict[str, str] = {}
    for question_path in sorted(Path(question_dir).glob("*.docx")):
        match = QUESTION_FILE_HEADING_PATTERN.match(question_path.stem)
        if not match:
            continue
        mapping[match.group(1)] = question_path.stem
    return mapping


def _dedupe_output_stem(base_stem: str, seen_stems: dict[str, int]) -> str:
    current_count = seen_stems.get(base_stem, 0)
    seen_stems[base_stem] = current_count + 1
    if current_count == 0:
        return base_stem
    return f"{base_stem}-{current_count + 1}"


def write_section_docx(paragraphs: list[str], output_path: str | Path) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _clear_doc_body(doc)
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(output_path)
    return output_path


def _copy_related_media(source_doc: Document, target_doc: Document, copied_element) -> None:
    rel_attrs = (qn("r:embed"), qn("r:link"), qn("r:id"))
    for node in copied_element.iter():
        for attr_name in rel_attrs:
            old_rid = node.get(attr_name)
            if not old_rid or old_rid not in source_doc.part.rels:
                continue
            rel = source_doc.part.rels[old_rid]
            if rel.reltype != RT.IMAGE:
                continue
            new_rid, _ = target_doc.part.get_or_add_image(BytesIO(rel.target_part.blob))
            node.set(attr_name, new_rid)


def _append_copied_paragraph(source_doc: Document, target_doc: Document, paragraph_index: int) -> None:
    copied_element = deepcopy(source_doc.paragraphs[paragraph_index]._element)
    _copy_related_media(source_doc, target_doc, copied_element)
    target_doc._element.body.insert_element_before(copied_element, "w:sectPr")


def write_rich_section_docx(
    source_doc: Document,
    paragraph_indices: list[int],
    output_path: str | Path,
) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    target_doc = Document()
    _clear_doc_body(target_doc)
    for paragraph_index in paragraph_indices:
        _append_copied_paragraph(source_doc, target_doc, paragraph_index)
    target_doc.save(output_path)
    return output_path


def split_answer_docx(
    source_path: str | Path,
    *,
    question_dir: str | Path | None = None,
    output_dir: str | Path | None = None,
) -> list[Path]:
    source_path = Path(source_path)
    question_dir = Path(question_dir) if question_dir else source_path.parent.parent
    output_dir = Path(output_dir) if output_dir else source_path.parent / "按课时拆分"

    source_doc = Document(source_path)
    sections = extract_sections_from_docx(source_path)
    question_heading_map = build_question_heading_map(question_dir)
    seen_stems: dict[str, int] = {}
    output_paths: list[Path] = []

    for section in sections:
        heading = section["heading"]
        base_stem = question_heading_map.get(heading, heading)
        output_stem = _dedupe_output_stem(f"{base_stem}-答案", seen_stems)
        output_path = output_dir / f"{output_stem}.docx"
        output_paths.append(
            write_rich_section_docx(
                source_doc,
                section["paragraph_indices"],
                output_path,
            )
        )

    return output_paths


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="按课时拆分未来高二物理总答案。")
    parser.add_argument("source", help="总答案 docx 路径")
    parser.add_argument(
        "--question-dir",
        help="题目目录，用于把拆分后的答案命名为对应题目文件名",
    )
    parser.add_argument(
        "--output-dir",
        help="拆分输出目录，默认在总答案同级下创建 按课时拆分",
    )
    return parser


def main() -> None:
    parser = build_arg_parser()
    args = parser.parse_args()

    output_paths = split_answer_docx(
        args.source,
        question_dir=args.question_dir,
        output_dir=args.output_dir,
    )

    print(f"已拆分 {len(output_paths)} 份答案：")
    for output_path in output_paths[:10]:
        print(f"- {output_path}")
    if len(output_paths) > 10:
        print(f"... 其余 {len(output_paths) - 10} 份已写入同目录")


if __name__ == "__main__":
    main()
