# coding: utf-8
from __future__ import annotations

import argparse
import re
import shutil
import sys
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from shared_core import (
    build_answer_units_from_docx,
    build_question_units_from_docx,
    map_answers,
)
from shared_core.answer_core import infer_grouped_question_ids
from tools.zhongmei_review_status import derive_question_path, refresh_review_statuses


DEFAULT_SOURCE_DIR = Path(
    r"D:\墨痕教育题目\众美-高三-语文\答案\对点练案答案"
)
CLEANED_FILE_PATTERN = re.compile(r"^对点练案(\d+)_已清洗\.docx$")
UNSUPPORTED_XML_TAGS = {
    qn("w:drawing"),
    qn("w:object"),
    qn("m:oMath"),
    qn("m:oMathPara"),
    qn("w:hyperlink"),
    qn("w:bookmarkStart"),
    qn("w:bookmarkEnd"),
}


@dataclass(frozen=True)
class ReadingGroup:
    question_id: str
    grouped_question_ids: tuple[str, ...]
    source_span: tuple[int, int]
    answer_lines: tuple[str, ...]
    analysis_lines: tuple[str, ...]


def discover_cleaned_documents(source_dir: str | Path) -> list[Path]:
    source_dir = Path(source_dir)
    files = [
        path
        for path in source_dir.glob("*_已清洗.docx")
        if not path.name.startswith("~$") and CLEANED_FILE_PATTERN.match(path.name)
    ]
    return sorted(files, key=lambda path: int(CLEANED_FILE_PATTERN.match(path.name).group(1)))


def _office_lock_path(docx_path: Path) -> Path:
    return docx_path.with_name(f"~${docx_path.name[2:]}")


def _reading_group_questions(question_path: str | Path):
    return [
        question
        for question in build_question_units_from_docx(question_path)
        if len(infer_grouped_question_ids(question)) > 1
    ]


def validate_reading_groups_premerged(
    docx_path: str | Path,
    *,
    question_path: str | Path | None = None,
) -> int:
    docx_path = Path(docx_path)
    question_path = Path(question_path) if question_path else derive_question_path(docx_path)
    questions = build_question_units_from_docx(question_path)
    answers = map_answers(questions, build_answer_units_from_docx(docx_path))
    answer_by_qid = {unit.question_id: unit for unit in answers}
    group_count = 0

    for question in questions:
        grouped_ids = tuple(infer_grouped_question_ids(question))
        if len(grouped_ids) <= 1:
            continue
        group_count += 1
        answer = answer_by_qid.get(question.question_id)
        if not answer:
            raise ValueError(f"{docx_path.name} 第{question.question_id}题：未找到阅读组答案")
        if answer.metadata.get("mapping_method") != "material_group_premerged":
            raise ValueError(f"{docx_path.name} 第{question.question_id}题：阅读组尚未永久合并")
        if len(answer.answer_items) != len(grouped_ids):
            raise ValueError(
                f"{docx_path.name} 第{question.question_id}题：合并后答案数量与阅读组题号数量不一致"
            )
        if len(answer.analysis_items) != len(grouped_ids):
            raise ValueError(
                f"{docx_path.name} 第{question.question_id}题：合并后解析数量与阅读组题号数量不一致"
            )

    return group_count


def find_legacy_reading_groups(
    docx_path: str | Path,
    *,
    question_path: str | Path | None = None,
) -> list[ReadingGroup]:
    docx_path = Path(docx_path)
    question_path = Path(question_path) if question_path else derive_question_path(docx_path)
    questions = build_question_units_from_docx(question_path)
    raw_answers = build_answer_units_from_docx(docx_path)
    answers = map_answers(questions, raw_answers)
    answer_by_qid = {unit.question_id: unit for unit in answers}
    raw_answer_by_qid = {unit.question_id: unit for unit in raw_answers}
    groups: list[ReadingGroup] = []

    for question in questions:
        grouped_ids = tuple(infer_grouped_question_ids(question))
        if len(grouped_ids) <= 1:
            continue
        answer = answer_by_qid.get(question.question_id)
        if not answer or not answer.metadata.get("is_material_group"):
            raise ValueError(f"{docx_path.name} 第{question.question_id}题：未找到完整阅读组答案")
        if answer.metadata.get("mapping_method") == "material_group_premerged":
            continue
        if len(answer.answer_items) != len(grouped_ids):
            raise ValueError(
                f"{docx_path.name} 第{question.question_id}题：答案数量与阅读组题号数量不一致"
            )

        analysis_texts: list[str] = []
        for grouped_id in grouped_ids:
            raw_answer = raw_answer_by_qid.get(grouped_id)
            if raw_answer is None:
                raise ValueError(f"{docx_path.name} 第{grouped_id}题：未找到原始答案块")
            analysis_texts.append(
                "\n".join(item.text.strip() for item in raw_answer.analysis_items).strip()
            )

        groups.append(
            ReadingGroup(
                question_id=question.question_id,
                grouped_question_ids=grouped_ids,
                source_span=(max(1, answer.source_span[0] - 1), answer.source_span[1]),
                answer_lines=tuple(
                    f"({index})答案：{item.text.strip()}"
                    for index, item in enumerate(answer.answer_items, 1)
                ),
                analysis_lines=tuple(
                    f"({index}){analysis_text}"
                    for index, analysis_text in enumerate(analysis_texts, 1)
                ),
            )
        )

    groups.sort(key=lambda group: group.source_span[0])
    for previous, current in zip(groups, groups[1:]):
        if previous.source_span[1] >= current.source_span[0]:
            raise ValueError(f"{docx_path.name}：阅读组段落范围发生重叠")
    return groups


def _assert_supported_span(paragraphs, group: ReadingGroup, docx_name: str) -> None:
    start = group.source_span[0] - 1
    end = group.source_span[1]
    for paragraph in paragraphs[start:end]:
        if any(element.tag in UNSUPPORTED_XML_TAGS for element in paragraph._element.iter()):
            raise ValueError(
                f"{docx_name} 第{group.question_id}题含图片、公式、超链接或书签，已停止改写"
            )


def _paragraph_shell(template, text: str):
    paragraph = OxmlElement("w:p")
    paragraph_properties = template._element.find(qn("w:pPr"))
    if paragraph_properties is not None:
        paragraph.append(deepcopy(paragraph_properties))

    run = OxmlElement("w:r")
    run_properties = template._element.find(".//" + qn("w:rPr"))
    if run_properties is not None:
        run.append(deepcopy(run_properties))

    chunks = text.split("\n")
    for index, chunk in enumerate(chunks):
        if index:
            run.append(OxmlElement("w:br"))
        text_node = OxmlElement("w:t")
        if chunk.startswith(" ") or chunk.endswith(" "):
            text_node.set(qn("xml:space"), "preserve")
        text_node.text = chunk
        run.append(text_node)
    paragraph.append(run)
    return paragraph


def _replace_group(document: Document, group: ReadingGroup) -> None:
    paragraphs = document.paragraphs
    start = group.source_span[0] - 1
    end = group.source_span[1]
    old_paragraphs = paragraphs[start:end]
    if not old_paragraphs:
        raise ValueError(f"第{group.question_id}题：未找到待合并段落")

    question_template = old_paragraphs[0]
    answer_template = next(
        (paragraph for paragraph in old_paragraphs[1:] if "答案" in paragraph.text),
        question_template,
    )
    analysis_template = next(
        (paragraph for paragraph in old_paragraphs[1:] if paragraph.text.strip().startswith("解析")),
        answer_template,
    )
    new_paragraphs = [
        _paragraph_shell(question_template, f"{group.question_id}．"),
        *(_paragraph_shell(answer_template, line) for line in group.answer_lines),
        _paragraph_shell(analysis_template, "解析："),
        *(_paragraph_shell(analysis_template, line) for line in group.analysis_lines),
    ]

    anchor = old_paragraphs[0]._element
    for paragraph in new_paragraphs:
        anchor.addprevious(paragraph)
    for paragraph in old_paragraphs:
        paragraph._element.getparent().remove(paragraph._element)


def merge_document_in_place(
    docx_path: str | Path,
    *,
    question_path: str | Path | None = None,
) -> int:
    docx_path = Path(docx_path)
    question_path = Path(question_path) if question_path else derive_question_path(docx_path)
    groups = find_legacy_reading_groups(docx_path, question_path=question_path)
    if not groups:
        return 0

    document = Document(docx_path)
    for group in groups:
        _assert_supported_span(document.paragraphs, group, docx_path.name)
    for group in reversed(groups):
        _replace_group(document, group)

    temp_path = docx_path.with_name(f".{docx_path.stem}_合并中.docx")
    try:
        document.save(temp_path)
        Document(temp_path)
        validate_reading_groups_premerged(temp_path, question_path=question_path)
        temp_path.replace(docx_path)
    finally:
        if temp_path.exists():
            temp_path.unlink()
    return len(groups)


def merge_batch(
    source_dir: str | Path,
    backup_dir: str | Path,
    *,
    expected_files: int = 44,
    expected_groups: int = 40,
) -> tuple[int, int]:
    source_dir = Path(source_dir)
    backup_dir = Path(backup_dir)
    files = discover_cleaned_documents(source_dir)
    if len(files) != expected_files:
        raise ValueError(f"已清洗文档应为 {expected_files} 份，实际为 {len(files)} 份")

    question_paths = {path: derive_question_path(path) for path in files}
    missing_questions = [path for path in question_paths.values() if not path.exists()]
    if missing_questions:
        names = "、".join(path.name for path in missing_questions)
        raise FileNotFoundError(f"缺少配对题目文档：{names}")

    total_group_count = sum(
        len(_reading_group_questions(question_path))
        for question_path in question_paths.values()
    )
    if total_group_count != expected_groups:
        raise ValueError(
            f"完整阅读组应为 {expected_groups} 组，实际为 {total_group_count} 组"
        )

    groups_by_path = {
        path: find_legacy_reading_groups(path, question_path=question_paths[path])
        for path in files
    }
    legacy_group_count = sum(len(groups) for groups in groups_by_path.values())

    changed_files = [path for path, groups in groups_by_path.items() if groups]
    locked_files = [path for path in changed_files if _office_lock_path(path).exists()]
    if locked_files:
        names = "、".join(path.name for path in locked_files)
        raise PermissionError(f"以下文档正在被 WPS/Word 占用，请先关闭：{names}")

    if changed_files:
        backup_dir.mkdir(parents=True, exist_ok=False)
        for path in changed_files:
            shutil.copy2(path, backup_dir / path.name)

    modified_files: list[Path] = []
    try:
        merged_count = 0
        for path in changed_files:
            merged_count += merge_document_in_place(
                path,
                question_path=question_paths[path],
            )
            modified_files.append(path)
        if merged_count != legacy_group_count:
            raise ValueError(
                f"实际合并 {merged_count} 组，与预检的 {legacy_group_count} 组不一致"
            )
        verified_group_count = sum(
            validate_reading_groups_premerged(
                path,
                question_path=question_paths[path],
            )
            for path in files
        )
        if verified_group_count != expected_groups:
            raise ValueError(
                f"合并后完整阅读组应为 {expected_groups} 组，实际为 {verified_group_count} 组"
            )
        refresh_review_statuses(files)
    except Exception:
        for path in modified_files:
            backup = backup_dir / path.name
            if backup.exists():
                shutil.copy2(backup, path)
        raise

    return len(changed_files), merged_count


def main() -> int:
    parser = argparse.ArgumentParser(description="合并众美语文完整阅读题的逐题答案与解析")
    parser.add_argument("--source-dir", type=Path, default=DEFAULT_SOURCE_DIR)
    parser.add_argument("--backup-dir", type=Path, required=True)
    parser.add_argument("--expected-files", type=int, default=44)
    parser.add_argument("--expected-groups", type=int, default=40)
    args = parser.parse_args()

    changed_files, merged_groups = merge_batch(
        args.source_dir,
        args.backup_dir,
        expected_files=args.expected_files,
        expected_groups=args.expected_groups,
    )
    print(f"已更新 {changed_files} 份文档，合并 {merged_groups} 个完整阅读组。")
    print(f"临时备份：{args.backup_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
